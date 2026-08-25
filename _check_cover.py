"""
检查教学设计封面、表0课程信息表、表1教学要求的结构和内容
"""
from docx import Document
from docx.oxml.ns import qn

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)

print("=== 封面段落（P0-P40）===")
for i, p in enumerate(doc.paragraphs[:45]):
    text = p.text.strip()
    if text:
        runs_info = ""
        if p.runs:
            r = p.runs[0]
            fn = r.font.name or ""
            sz = r.font.size.pt if r.font.size else ""
            bold = r.font.bold
            runs_info = f"[{fn},{sz},bold={bold}]"
        print(f"  P{i}: {runs_info} {text[:80]}")

print(f"\n=== 表0（课程信息表）===")
t0 = doc.tables[0]
print(f"{len(t0.rows)}行 x {len(t0.rows[0].cells)}列")
for ri in range(len(t0.rows)):
    cells = t0.rows[ri].cells
    texts = [c.text.strip()[:30] for c in cells]
    print(f"  R{ri}: {texts}")

print(f"\n=== 表1（教学要求/课程内容设计）===")
t1 = doc.tables[1]
print(f"{len(t1.rows)}行 x {len(t1.rows[0].cells)}列")
for ri in range(len(t1.rows)):
    cells = t1.rows[ri].cells
    texts = [c.text.strip()[:30] for c in cells]
    print(f"  R{ri}: {texts}")

# 检查表5第一个任务的教学要求部分
print(f"\n=== 表5第一个任务R9-R14（教学要求）===")
t5 = doc.tables[5]
for ri in range(9, min(15, len(t5.rows))):
    cells = t5.rows[ri].cells
    texts = []
    for ci in range(min(10, len(cells))):
        t = cells[ci].text.strip()[:40]
        texts.append(f"[{ci}]{t}")
    print(f"  R{ri}: {' | '.join(texts)}")
