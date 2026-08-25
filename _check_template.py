from docx import Document

fp = r'原始资料\模板\模板5：教学设计\模板5：教学设计 模板（2023-2024）.docx'
doc = Document(fp)

# 找到第一个单元教学设计表格
print(f'总表格数: {len(doc.tables)}')
# Table 5是第一个单元教学设计表
t = doc.tables[5]
print(f'\n=== Template Table 5: {len(t.rows)} rows, {len(t.columns)} cols ===')
for ri in range(len(t.rows)):
    texts = [c.text.strip()[:40] for c in t.rows[ri].cells]
    print(f'  R{ri}: {texts}')

t6 = doc.tables[6]
print(f'\n=== Template Table 6: {len(t6.rows)} rows, {len(t6.columns)} cols ===')
for ri in range(len(t6.rows)):
    texts = [c.text.strip()[:40] for c in t6.rows[ri].cells]
    print(f'  R{ri}: {texts}')
