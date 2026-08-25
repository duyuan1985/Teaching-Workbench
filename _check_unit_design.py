from docx import Document

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

# 找到"教学设计·基本信息"段落位置
for pi, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if '教学设计' in txt and '基本信息' in txt:
        print(f'P{pi}: {txt}')

# 查看后续表格
print(f'\n总表格数: {len(doc.tables)}')
# Table 5开始是单元教学设计
for ti in range(5, min(8, len(doc.tables))):
    t = doc.tables[ti]
    print(f'\n=== Table {ti}: {len(t.rows)} rows, {len(t.columns)} cols ===')
    for ri in range(len(t.rows)):
        texts = [c.text.strip()[:40] for c in t.rows[ri].cells]
        print(f'  R{ri}: {texts}')
