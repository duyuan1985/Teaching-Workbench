"""
修复所有教学组织表格的行高：将固定高度(exact)改为自动高度(auto)
让行高随内容自动增加
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

# 处理所有29个教学组织表格（偶数索引：6, 8, 10, ..., 62）
table_indices = list(range(6, 63, 2))

for ti in table_indices:
    t = doc.tables[ti]
    
    # R3-R10都需要改为auto
    for ri in range(3, 11):
        tr = t.rows[ri]._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        
        # 找到或创建trHeight
        trHeight = trPr.find(qn('w:trHeight'))
        if trHeight is None:
            trHeight = OxmlElement('w:trHeight')
            trPr.append(trHeight)
        
        # 设置为auto，让行高随内容自动调整
        trHeight.set(qn('w:val'), '0')
        trHeight.set(qn('w:hRule'), 'auto')
    
    print(f'Table {ti}: R3-R10 行高改为auto')

doc.save(fp)
print(f'\n保存完成')

# 验证
doc2 = Document(fp)
t = doc2.tables[6]
for ri in range(3, 11):
    tr = t.rows[ri]._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is not None:
        trHeight = trPr.find(qn('w:trHeight'))
        if trHeight is not None:
            val = trHeight.get(qn('w:val'))
            hRule = trHeight.get(qn('w:hRule'))
            print(f'R{ri}: val={val} hRule={hRule}')
