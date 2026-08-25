"""
修复表2合并单元格：
1. vMerge改为正确值（第一行restart，后续行continue）
2. 清空continue行的文字
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)
t2 = doc.tables[2]

# 确定合并范围（从之前的数据）
# (1,3,'1'), (4,6,'2'), (7,9,'3'), (10,12,'4'), (13,15,'5'), (16,18,'6'), (19,21,'7'), (22,24,'8'), (25,26,'9')
merge_groups = [
    (1, 3, '1', '初识数据分析'),
    (4, 6, '2', 'Excel数据分析工具'),
    (7, 9, '3', 'Numpy数学运算库'),
    (10, 12, '4', 'Pandas数据分析库'),
    (13, 15, '5', 'SciPy科学计算库'),
    (16, 18, '6', 'Sklearn数据统计基础'),
    (19, 21, '7', 'Sklearn数据统计进阶'),
    (22, 24, '8', 'Seaborn可视化分析库'),
    (25, 26, '9', '综合评价与课程总结'),
]

for start, end, seq_val, name_val in merge_groups:
    # col0: 编号
    # 第一行：vMerge=restart，写编号
    tc0_start = t2.cell(start, 0)._tc
    tcPr0 = tc0_start.find(qn('w:tcPr'))
    if tcPr0 is not None:
        vm = tcPr0.find(qn('w:vMerge'))
        if vm is not None:
            tcPr0.remove(vm)
        vm_new = OxmlElement('w:vMerge')
        vm_new.set(qn('w:val'), 'restart')
        tcPr0.append(vm_new)
    
    # 清空第一行文字后重新写入
    cell0 = t2.cell(start, 0)
    for p in cell0.paragraphs:
        for r in p.runs:
            r.text = ''
    if cell0.paragraphs and cell0.paragraphs[0].runs:
        cell0.paragraphs[0].runs[0].text = seq_val
    else:
        cell0.paragraphs[0].add_run(seq_val)
    
    # col1: 名称，同样处理
    tc1_start = t2.cell(start, 1)._tc
    tcPr1 = tc1_start.find(qn('w:tcPr'))
    if tcPr1 is not None:
        vm = tcPr1.find(qn('w:vMerge'))
        if vm is not None:
            tcPr1.remove(vm)
        vm_new = OxmlElement('w:vMerge')
        vm_new.set(qn('w:val'), 'restart')
        tcPr1.append(vm_new)
    
    cell1 = t2.cell(start, 1)
    for p in cell1.paragraphs:
        for r in p.runs:
            r.text = ''
    if cell1.paragraphs and cell1.paragraphs[0].runs:
        cell1.paragraphs[0].runs[0].text = name_val
    else:
        cell1.paragraphs[0].add_run(name_val)
    
    # 后续行：vMerge=continue，清空文字
    for ri in range(start + 1, end + 1):
        for ci in [0, 1]:
            tc = t2.cell(ri, ci)._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                vm = tcPr.find(qn('w:vMerge'))
                if vm is not None:
                    tcPr.remove(vm)
                vm_new = OxmlElement('w:vMerge')
                # continue不需要val属性
                tcPr.append(vm_new)
            
            # 清空文字
            cell = t2.cell(ri, ci)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.text = ''

print("修复完成，验证:")
doc.save(fp)

doc2 = Document(fp)
t2 = doc2.tables[2]
for ri in range(min(10, len(t2.rows))):
    cells = t2.rows[ri].cells
    texts = [c.text.strip()[:20] for c in cells[:3]]
    print(f"  R{ri}: {texts}")
