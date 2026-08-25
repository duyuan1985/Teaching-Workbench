"""
教学设计全面审计脚本
检查：空格、模板残留、重复内容、表格结构、内容完整度
"""
import re
from docx import Document
from docx.oxml.ns import qn
from collections import Counter

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)

print(f"段落总数: {len(doc.paragraphs)}")
print(f"表格总数: {len(doc.tables)}")

# ============================================================
# 1. 段落审计
# ============================================================
print(f"\n{'='*60}")
print("一、段落审计")
print(f"{'='*60}")

para_issues = []
residue_patterns = [r'×+', r'XX+', r'XXX', r'（注：', r'\(注：', r'（建议', r'【模板', r'（模板', r'体例之一', r'巩固知识']
residue_found = []

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    # 空段落
    if not text and i > 5:
        # 只记录连续空段落
        if i > 0 and doc.paragraphs[i-1].text.strip():
            para_issues.append(f"P{i}: 空段落（前一段有内容）")
    # 模板残留
    for pat in residue_patterns:
        if re.search(pat, text):
            residue_found.append(f"P{i}: {text[:80]}")
            break
    # 检查字体格式
    if p.runs:
        r = p.runs[0]
        # 检查是否有firstLineChars异常
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                flc = ind.get(qn('w:firstLineChars'))
                if flc and int(flc) > 10:
                    para_issues.append(f"P{i}: 异常缩进firstLineChars={flc} | {text[:40]}")
        # 检查spacing exact
        if pPr is not None:
            sp = pPr.find(qn('w:spacing'))
            if sp is not None:
                rule = sp.get(qn('w:lineRule'))
                if rule == 'exact':
                    val = sp.get(qn('w:line'))
                    if val and int(val) < 400:
                        para_issues.append(f"P{i}: 固定行距过小exact={val} | {text[:40]}")

print(f"段落问题: {len(para_issues)}个")
for iss in para_issues[:20]:
    print(f"  {iss}")

print(f"\n模板残留: {len(residue_found)}处")
for r in residue_found[:20]:
    print(f"  {r}")

# ============================================================
# 2. 表格结构审计
# ============================================================
print(f"\n{'='*60}")
print("二、表格结构审计")
print(f"{'='*60}")

# 教学设计模板结构：表0-4是整体设计，表5+6是第1个任务，表7+8是第2个任务...
# 表5,7,9...是基本信息（15行x20列）
# 表6,8,10...是教学组织（11行x6列）

task_tables = []
for ti in range(5, len(doc.tables), 2):
    if ti + 1 < len(doc.tables):
        t_basic = doc.tables[ti]
        t_org = doc.tables[ti + 1]
        task_idx = (ti - 5) // 2
        task_tables.append((task_idx, ti, t_basic, t_org))

print(f"任务表格对: {len(task_tables)}个")

# ============================================================
# 3. 基本信息表（表5,7,9...）审计
# ============================================================
print(f"\n{'='*60}")
print("三、基本信息表审计（每任务1个，15行x20列）")
print(f"{'='*60}")

basic_issues = []
for task_idx, ti, t, _ in task_tables:
    if len(t.rows) < 15:
        basic_issues.append(f"任务{task_idx+1} 表{ti}: 只有{len(t.rows)}行（应15行）")
        continue
    
    # R0: 周次/课时/班级
    r0 = t.rows[0]
    week = r0.cells[2].text.strip()
    hours = r0.cells[4].text.strip()
    cls = r0.cells[6].text.strip()
    
    # R1: 教师/日期
    r1 = t.rows[1]
    teacher = r1.cells[2].text.strip()
    date = r1.cells[6].text.strip()
    
    # R2: 课程类型/教学环境
    r2 = t.rows[2]
    ctype = r2.cells[2].text.strip()
    env = r2.cells[6].text.strip()
    
    # R3: 课题
    r3 = t.rows[3] if len(t.rows) > 3 else None
    title = r3.cells[2].text.strip() if r3 else ""
    
    # R4: 教学目标
    r4 = t.rows[4] if len(t.rows) > 4 else None
    goal = r4.cells[2].text.strip()[:50] if r4 else ""
    
    # R5: 重点
    r5 = t.rows[5] if len(t.rows) > 5 else None
    key_point = r5.cells[2].text.strip()[:30] if r5 else ""
    
    # R6: 难点
    r6 = t.rows[6] if len(t.rows) > 6 else None
    diff_point = r6.cells[2].text.strip()[:30] if r6 else ""
    
    # R7: 教学方法
    r7 = t.rows[7] if len(t.rows) > 7 else None
    method = r7.cells[2].text.strip()[:30] if r7 else ""
    
    # R8-14: 德育/板书/作业等
    r8_text = t.rows[8].cells[2].text.strip()[:30] if len(t.rows) > 8 else ""
    
    # 检查空值
    if not week: basic_issues.append(f"任务{task_idx+1}: 周次空")
    if not teacher: basic_issues.append(f"任务{task_idx+1}: 教师空")
    if not date: basic_issues.append(f"任务{task_idx+1}: 日期空")
    if not title: basic_issues.append(f"任务{task_idx+1}: 课题空")
    if not goal: basic_issues.append(f"任务{task_idx+1}: 教学目标空")
    if not key_point: basic_issues.append(f"任务{task_idx+1}: 重点空")
    if not diff_point: basic_issues.append(f"任务{task_idx+1}: 难点空")
    
    # 检查模板残留
    for ri in range(min(15, len(t.rows))):
        for ci in range(len(t.rows[ri].cells)):
            ct = t.rows[ri].cells[ci].text.strip()
            if '×' in ct or 'XX' in ct or 'XXX' in ct:
                basic_issues.append(f"任务{task_idx+1} R{ri}C{ci}: 模板残留 '{ct[:30]}'")
            if '（注：' in ct or '（建议' in ct:
                basic_issues.append(f"任务{task_idx+1} R{ri}C{ci}: 模板说明 '{ct[:30]}'")
    
    if task_idx < 3 or task_idx >= len(task_tables) - 2:
        print(f"  任务{task_idx+1} (表{ti}): 周{week} 师{teacher} 日{date}")
        print(f"    课题: {title[:40]}")
        print(f"    目标: {goal[:40]}")
        print(f"    重点: {key_point[:30]} 难点: {diff_point[:30]}")
        print(f"    方法: {method[:30]}")
        print(f"    R8: {r8_text}")

print(f"\n基本信息问题: {len(basic_issues)}个")
for iss in basic_issues[:30]:
    print(f"  {iss}")

# ============================================================
# 4. 教学组织表（表6,8,10...）审计
# ============================================================
print(f"\n{'='*60}")
print("四、教学组织表审计（每任务1个，11行x6列）")
print(f"{'='*60}")

org_issues = []

# 收集所有小结、作业、反思内容，检查重复
all_summaries = []
all_homeworks = []
all_reflections = []
all_teaching_contents = []

for task_idx, _, _, t in task_tables:
    if len(t.rows) < 11:
        org_issues.append(f"任务{task_idx+1}: 教学组织表只有{len(t.rows)}行（应11行）")
        continue
    
    # 教学组织表结构：
    # R0: 教学场景设计（col0标题, col1-3内容）
    # R1: 教学资源准备
    # R2: 教学过程（表头：教学步骤/教法学法/达成目标/时间）
    # R3-8: 教学过程内容（6行）
    # R9: 课堂小结
    # R10: 课后作业/教学反思
    
    # 检查每行内容
    for ri in range(min(11, len(t.rows))):
        cells = t.rows[ri].cells
        for ci in range(len(cells)):
            ct = cells[ci].text.strip()
            if not ct and ri > 1:
                # 空格子（可能是合并单元格）
                pass
            elif ct:
                # 收集内容用于重复检查
                if '小结' in cells[0].text or '课堂小结' in cells[0].text:
                    all_summaries.append((task_idx, ct[:50]))
                elif '作业' in cells[0].text or '课后作业' in cells[0].text:
                    all_homeworks.append((task_idx, ct[:50]))
                elif '反思' in cells[0].text or '教学反思' in cells[0].text:
                    all_reflections.append((task_idx, ct[:50]))
            
            # 模板残留
            if '×' in ct or 'XX' in ct:
                if ri > 1:
                    org_issues.append(f"任务{task_idx+1} R{ri}C{ci}: 模板残留 '{ct[:30]}'")
    
    # 打印前2个和最后1个任务的详细内容
    if task_idx < 2 or task_idx == len(task_tables) - 1:
        print(f"\n  任务{task_idx+1} 教学组织表:")
        for ri in range(min(11, len(t.rows))):
            cells = t.rows[ri].cells
            texts = [c.text.strip()[:40] for c in cells]
            print(f"    R{ri}: {texts}")

# ============================================================
# 5. 重复内容检查
# ============================================================
print(f"\n{'='*60}")
print("五、重复内容检查")
print(f"{'='*60}")

def check_duplicates(items, name):
    """检查重复项"""
    texts = [t for _, t in items]
    counter = Counter(texts)
    dups = {t: c for t, c in counter.items() if c > 1}
    unique = len(set(texts))
    total = len(texts)
    print(f"\n{name}: {total}条，唯一{unique}条，重复{len(dups)}种")
    if dups:
        for t, c in list(dups.items())[:10]:
            tasks = [idx+1 for idx, txt in items if txt == t]
            print(f"  重复{c}次（任务{tasks}）: {t[:50]}")
    return unique, total

su_u, su_t = check_duplicates(all_summaries, "课堂小结")
hw_u, hw_t = check_duplicates(all_homeworks, "课后作业")
rf_u, rf_t = check_duplicates(all_reflections, "教学反思")

# ============================================================
# 6. 表格空格率
# ============================================================
print(f"\n{'='*60}")
print("六、表格空格率")
print(f"{'='*60}")

total_cells = 0
empty_cells = 0
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            total_cells += 1
            if not cell.text.strip():
                empty_cells += 1

print(f"总格子: {total_cells}, 空格子: {empty_cells}, 空格率: {empty_cells/total_cells*100:.1f}%")

# ============================================================
# 7. 教学过程内容详细度检查
# ============================================================
print(f"\n{'='*60}")
print("七、教学过程内容详细度")
print(f"{'='*60}")

for task_idx, _, _, t in task_tables[:5]:
    if len(t.rows) >= 9:
        # R3是教学过程第一行内容
        r3 = t.rows[3]
        content = r3.cells[2].text.strip() if len(r3.cells) > 2 else ""
        content_len = len(content)
        print(f"  任务{task_idx+1} 教学过程R3: {content_len}字 | {content[:60]}")

# ============================================================
# 8. 汇总
# ============================================================
print(f"\n{'='*60}")
print("八、审计汇总")
print(f"{'='*60}")
print(f"段落问题: {len(para_issues)}个")
print(f"模板残留: {len(residue_found)}处")
print(f"基本信息问题: {len(basic_issues)}个")
print(f"教学组织问题: {len(org_issues)}个")
print(f"课堂小结: {su_u}/{su_t}唯一")
print(f"课后作业: {hw_u}/{hw_t}唯一")
print(f"教学反思: {rf_u}/{rf_t}唯一")
print(f"空格率: {empty_cells}/{total_cells} = {empty_cells/total_cells*100:.1f}%")
