from docx import Document
from docx.oxml.ns import qn

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

# 检查所有教学组织表格（偶数索引）的R3-R10行高
table_indices = list(range(6, 63, 2))
all_ok = True

for ti in table_indices:
    t = doc.tables[ti]
    for ri in range(3, 11):
        tr = t.rows[ri]._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is not None:
            trHeight = trPr.find(qn('w:trHeight'))
            if trHeight is not None:
                hRule = trHeight.get(qn('w:hRule'))
                if hRule != 'auto':
                    print(f'Table {ti} R{ri}: hRule={hRule} (需要修复)')
                    all_ok = False

if all_ok:
    print('所有29个表格的R3-R10行高已全部为auto')

# 同时检查基本信息表格（奇数索引）是否有固定行高
print('\n=== 检查基本信息表格 ===')
odd_indices = list(range(5, 63, 2))
for ti in odd_indices[:3]:
    t = doc.tables[ti]
    for ri in range(len(t.rows)):
        tr = t.rows[ri]._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is not None:
            trHeight = trPr.find(qn('w:trHeight'))
            if trHeight is not None:
                val = trHeight.get(qn('w:val'))
                hRule = trHeight.get(qn('w:hRule'))
                if hRule == 'exact':
                    print(f'Table {ti} R{ri}: val={val} hRule={hRule}')
