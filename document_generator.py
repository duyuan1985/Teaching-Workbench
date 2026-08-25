import json
import os
import re
import shutil
import tempfile
from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from lxml import etree

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.table import Table
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

import store
import template_filler
import enhanced_content_author
from assessment_scheme import component_text, get_scheme
from course_standard_model import _subscenarios
from document_utils import (
    _set_paragraph,
    _set_cell,
    _set_underlined_field,
    _set_body_font,
    _set_cell_body_font,
    _set_cell_small_font,
    _set_table_font,
    _replace_prefix,
    _term_label,
    _plan_term_label,
)


def _get_template_contract(template_file_id):
    """获取模板分析契约（表格角色等信息）"""
    rows = store.rows(
        "SELECT contract_json FROM template_analyses WHERE template_file_id=?",
        (template_file_id,)
    )
    if not rows:
        return None
    import json
    return json.loads(rows[0]["contract_json"])


def _find_table_by_role(document, contract, role_name):
    """
    根据表格角色定位表格索引。
    返回表格对象和索引，找不到返回 (None, -1)
    """
    if not contract or "tables" not in contract:
        return None, -1
    for table_info in contract["tables"]:
        if table_info["role"] == role_name:
            idx = table_info["index"]
            if idx < len(document.tables):
                return document.tables[idx], idx
    return None, -1


def _get_template_file_id(offering_id, document_type):
    """获取模板文件ID"""
    rows = store.rows(
        "SELECT id FROM template_files WHERE offering_id=? AND document_type=?",
        (offering_id, document_type)
    )
    return rows[0]["id"] if rows else None


def _set_merged_cell(cell, text):
    _set_paragraph(cell.paragraphs[0], text)
    for extra in list(cell.paragraphs[1:]):
        _remove(extra._p)


def _numbered(items):
    return "\n".join(f"{index}. {item}" for index, item in enumerate(items, 1))


def _numbered_compact(items):
    return "\n".join(f"{index}.{item}" for index, item in enumerate(items, 1))


def _task_detail(task):
    """Remove the chapter prefix from a task title when the table has a chapter column."""
    title = str(task.get("title") or "").strip()
    chapter = str(task.get("chapter") or "").strip()
    if chapter and title.startswith(chapter):
        title = title[len(chapter):].lstrip(" ：:、")
    return title or chapter


def _even_sample(items, count):
    if len(items) <= count:
        return list(items)
    return [items[round(index * (len(items) - 1) / (count - 1))] for index in range(count)]


# ============================================================
# 富文本填充：支持 [[图片:]] 和 [[源码:]] 标记
# ============================================================

def _find_resource_file(offering_id, filename):
    """根据文件名在资源库中查找完整路径"""
    rows = store.rows(
        "SELECT file_path, resource_type FROM resource_items WHERE offering_id=? AND file_path LIKE ? LIMIT 5",
        (offering_id, f"%{filename}"),
    )
    for row in rows:
        if Path(row["file_path"]).name == filename:
            return row["file_path"], row["resource_type"]
    return None, None


def _insert_image_into_cell(cell, image_path, width_inches=3.0):
    """在单元格末尾插入一张图片（居中）"""
    # 获取或创建最后一个段落
    if cell.paragraphs and cell.paragraphs[-1].text.strip():
        para = cell.add_paragraph()
    else:
        para = cell.paragraphs[-1] if cell.paragraphs else cell.add_paragraph()
    
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    try:
        run.add_picture(image_path, width=Inches(width_inches))
    except Exception:
        run.text = f"[图片: {Path(image_path).name}]"
    
    # 图片说明
    caption = cell.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption.add_run(f"图：{Path(image_path).stem}")
    cap_run.font.size = Pt(8)
    cap_run.font.name = "宋体"
    rpr = cap_run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")


def _insert_code_into_cell(cell, code_path, max_lines=30):
    """在单元格中插入源码片段（带背景的代码样式）"""
    try:
        with open(code_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()
    except Exception:
        return
    
    # 截取行数
    if len(lines) > max_lines:
        lines = lines[:max_lines]
        lines.append("...\n")
    
    code_text = "".join(lines).rstrip()
    
    # 添加代码段落
    para = cell.add_paragraph()
    para.paragraph_format.left_indent = Pt(12)
    run = para.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), "Consolas")
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")
    
    # 文件名标注
    caption = cell.add_paragraph()
    cap_run = caption.add_run(f"代码：{Path(code_path).name}")
    cap_run.font.size = Pt(8)
    cap_run.font.name = "宋体"
    cap_run.font.italic = True
    rpr = cap_run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")


def _fill_cell_rich(cell, text, offering_id=None, image_width=2.8):
    """
    富文本填充单元格：解析 [[图片:xxx.png]] 和 [[源码:xxx.py]] 标记，
    将文本+图片+代码混合填入单元格。
    
    Args:
        cell: Word 表格单元格
        text: 包含标记的文本
        offering_id: 课程实例ID（用于查找资源文件）
        image_width: 图片宽度（英寸）
    """
    # 清空单元格现有内容
    for para in list(cell.paragraphs):
        p = para._p
        p.getparent().remove(p)
    
    if not text:
        cell.add_paragraph("")
        return
    
    # 按行分割处理
    lines = text.split("\n")
    current_para = cell.add_paragraph()
    
    for line in lines:
        stripped = line.strip()
        
        # 检查是否是图片标记
        img_match = re.fullmatch(r'\[\[图片:(.+?)\]\]', stripped)
        if img_match and offering_id:
            filename = img_match.group(1).strip()
            img_path, _ = _find_resource_file(offering_id, filename)
            if img_path and Path(img_path).exists():
                _insert_image_into_cell(cell, img_path, width_inches=image_width)
                current_para = cell.add_paragraph()
            else:
                # 找不到图片，保留标记文字
                if current_para.text or len(cell.paragraphs) > 1:
                    current_para = cell.add_paragraph()
                current_para.add_run(stripped)
            continue
        
        # 检查是否是源码标记
        code_match = re.fullmatch(r'\[\[源码:(.+?)\]\]', stripped)
        if code_match and offering_id:
            filename = code_match.group(1).strip()
            code_path, _ = _find_resource_file(offering_id, filename)
            if code_path and Path(code_path).exists():
                _insert_code_into_cell(cell, code_path)
                current_para = cell.add_paragraph()
            else:
                if current_para.text or len(cell.paragraphs) > 1:
                    current_para = cell.add_paragraph()
                current_para.add_run(stripped)
            continue
        
        # 普通文本行
        if not stripped:
            # 空行
            if current_para.text:
                current_para = cell.add_paragraph()
        else:
            if current_para.text:
                current_para.add_run("\n" + stripped)
            else:
                current_para.add_run(stripped)
    
    # 设置整体字体（普通文本段落）
    for para in cell.paragraphs:
        for run in para.runs:
            # 跳过已经设置过特殊字体的（代码块等）
            if run.font.name and run.font.name in ("Consolas", "Courier New"):
                continue
            if not run.font.name or run.font.name == "宋体":
                run.font.name = "宋体"
                run.font.size = Pt(9) if run.font.size is None else run.font.size
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = rpr.makeelement(qn("w:rFonts"), {})
                    rpr.insert(0, rfonts)
                rfonts.set(qn("w:eastAsia"), "宋体")


def _drafts(offering_id, document_type):
    rows = store.rows(
        "SELECT section_key,repeat_key,content_json FROM authored_sections WHERE offering_id=? AND document_type=?",
        (offering_id, document_type),
    )
    return {(row["section_key"], row["repeat_key"]): json.loads(row["content_json"]) for row in rows}


def _insert_after(paragraph, text, model=None):
    new_p = deepcopy((model or paragraph)._p)
    for node in list(new_p):
        if node.tag.endswith("}r") or node.tag.endswith("}hyperlink"):
            new_p.remove(node)
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    _set_paragraph(result, text)
    _apply_body_font(result)
    return result


def _apply_body_font(paragraph, font_name="仿宋_GB2312", size_pt=12):
    """按模板正文规范（小四仿宋）统一新增段落的字体。"""
    for run in paragraph.runs:
        if not run.text.strip():
            continue
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


def _remove(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _study_semester(teaching_class, term):
    """Convert cohort and academic term to the student's numbered semester."""
    class_str = str(teaching_class or "")
    class_match = re.search(r"(?<!\d)(\d{2})\d{4}(?!\d)", class_str)
    year_match = re.search(r"(?<!\d)(20\d{2})(?!\d)", class_str)
    term_match = re.fullmatch(r"(\d{4})-(\d{4})-([12])", str(term or ""))
    if not term_match or not (class_match or year_match):
        return _term_label(term)
    cohort_year = 2000 + int(class_match.group(1)) if class_match else int(year_match.group(1))
    semester_number = (int(term_match.group(1)) - cohort_year) * 2 + int(term_match.group(3))
    chinese = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六", 7: "七", 8: "八"}
    if semester_number <= 0:
        return _term_label(term)
    return f"第{chinese.get(semester_number, semester_number)}学期"


def _fill_standard_2025(offering_id, template, output):
    offering = store.rows("SELECT * FROM offerings WHERE id=?", (offering_id,))[0]
    tasks = store.rows("SELECT * FROM tasks WHERE offering_id=? ORDER BY seq", (offering_id,))
    drafts = _drafts(offering_id, "课程标准")
    # 支持增量填充
    output_path = Path(output)
    if not output_path.exists():
        shutil.copy2(template, output)
    document = Document(output)
    
    # 获取模板分析契约
    template_file_id = _get_template_file_id(offering_id, "课程标准")
    contract = _get_template_contract(template_file_id) if template_file_id else None
    _replace_prefix(document, "《**********》课程标准", f"《{offering['course_name']}》课程标准")
    _replace_prefix(document, "《**********》（居中", f"《{offering['course_name']}》")
    cover_fields = (
        ("适用专业：", offering["major"]),
        ("编制单位：", f"{offering['major']}教研室"),
        ("合作单位：", store.get_setting("partner_company", "天津滨海迅腾科技集团有限公司")),
    )
    for prefix, value in cover_fields:
        paragraph = next((p for p in document.paragraphs if p.text.strip().startswith(prefix)), None)
        if paragraph:
            _set_underlined_field(paragraph, prefix, value)
    generated_paragraphs = []
    info = document.tables[0]
    theory = sum(int(t["theory_hours"] or 0) for t in tasks)
    practice = sum(int(t["practice_hours"] or 0) for t in tasks)
    # The template uses irregular horizontal merges. These indexes point to
    # the blank value cells to the right of each label, not the label merge.
    values = {(0, 2): offering["course_name"], (0, 7): store.get_setting("department", "经济贸易系"), (1, 2): offering["course_code"],
              (1, 6): _study_semester(offering.get("teaching_class"), offering["term"]), (2, 2): offering.get("course_nature") or "必修课", (2, 6): offering["course_type"],
              (3, 1): offering.get("assessment_type") or "期末考核", (3, 4): offering.get("assessment_method") or "实操", (3, 8): offering["credits"],
              (4, 1): offering["total_hours"], (4, 4): theory, (4, 8): practice}
    for (row, col), value in values.items():
        _set_cell(info.rows[row].cells[col], value)
        _set_cell_body_font(info.rows[row].cells[col])
    nature = list(drafts.get(("course_nature", ""), []))
    design = drafts.get(("course_design", ""), [])
    goals = drafts.get(("course_goals", ""), {})
    nature_word = (offering.get("course_nature") or "必修课").removesuffix("课")
    project_names = list(dict.fromkeys(t["chapter"] for t in tasks if t["chapter"]))
    project_text = "、".join(project_names[:6])
    prerequisite = offering.get("prerequisite_courses") or "相关专业基础课程"
    followup = offering.get("followup_courses") or "相关综合实践课程"
    nature_content = [
        f"《{offering['course_name']}》是{offering['major']}专业开设的一门{nature_word}{offering['course_type']}。本课程旨在使学生系统掌握{project_text}所涉及的基础知识、核心方法和规范化项目实施流程，形成运用专业工具解决实际业务问题的能力，为本专业技术应用、业务实施和项目协作等岗位能力培养提供支撑。",
        f"课程以{project_text}等教材项目为载体，将知识学习、技能训练、成果检查和规范交付贯穿教学全过程，突出实践性、职业性和综合性，在专业课程体系中承担由基础操作能力向综合项目实施能力过渡的作用。",
        f"本课程以《{prerequisite}》形成的知识与操作能力为学习基础，并为《{followup}》中的综合任务实施奠定知识、技术、规范意识和协作基础。",
    ]
    task_content = [
        f"课程服务{offering['major']}专业数字化业务实施、数据处理与分析、项目运营及技术支持等相关岗位。通过{project_text}等任务训练，使学生掌握课程核心知识和工具使用方法，能够完成需求分析、方案制定、操作实现、运行检查、问题修正和成果交付。",
        "课程同步培养学生遵守技术标准和操作规范的意识，将数据安全、个人信息保护、知识产权、诚信责任、劳动精神、团队协作和持续改进融入项目过程，促进知识、能力、思政素养和职业素质协调发展。",
    ]
    content_design = [
        f"课程内容依据课程目标和相关岗位工作过程进行重构，以{project_text}等项目为主线，将知识点、技能点和职业规范转化为递进任务，按照“理解原理—示范操作—技术练习—项目实施—检查评价—迭代改进”的路径组织{offering['total_hours']}学时教学。",
        "内容设计融入课程思政、劳动教育、安全教育和课证赛融通要求，突出规范操作、数据与网络安全、知识产权和诚信交付；结合行业发展补充新标准、新技术、新工艺和新方法，以真实或仿真成果作为学习产出，体现课程标准的职业性、先进性和可评价性。",
    ]
    teaching_design = list(design) or [
        "课程采用理实一体、项目教学和任务驱动模式，以学生为主体、成果为导向组织教学。教师通过案例导入、问题引导、技术分析和操作示范明确任务，学生通过独立练习、小组协作、运行调试、成果展示和复盘改进达成目标。",
        "教学实施引入企业工作流程和质量要求，综合运用课堂观察、任务检查、阶段成果、学生互评和综合作品评价，兼顾过程评价、结果评价与增值评价，实现知识技能培养、职业素养形成和思想价值引领相统一。",
    ]
    for heading, content in (("1.课程性质", nature_content), ("2.课程任务", task_content), ("1.内容设计", content_design), ("2.教学设计", teaching_design)):
        paragraph = next((p for p in document.paragraphs if p.text.strip() == heading), None)
        if paragraph:
            cursor = paragraph
            for text in content:
                cursor = _insert_after(cursor, text)
                generated_paragraphs.append(cursor)
    knowledge_clause = "，并".join(item.rstrip("。；") for item in goals.get("知识目标", [])[:2])
    ability_clause = "，并能".join(item.removeprefix("具有").removeprefix("能够").rstrip("。；") for item in goals.get("能力目标", [])[:2])
    quality_clause = "，".join(item.removeprefix("在").rstrip("。；") for item in (goals.get("思政目标", [])[:1] + goals.get("素质目标", [])[:2]))
    overall_goal = (
        f"通过本课程学习，学生能够{knowledge_clause}；能够{ability_clause}；"
        f"在项目学习与实践中{quality_clause}，达到{offering['major']}专业对知识、能力和素质协调发展的培养要求。"
    )
    overall_heading = next((p for p in document.paragraphs if p.text.strip().startswith("（一）总体目标")), None)
    if overall_heading:
        paragraph = _insert_after(overall_heading, overall_goal)
        generated_paragraphs.append(paragraph)
    substantive_tasks = [task for task in tasks if "综合评价与课程总结" not in task["chapter"]]
    goal_tasks = _even_sample(substantive_tasks, 9)
    verbs = ("了解", "熟悉", "掌握")
    knowledge_items = []
    ability_items = []
    for index, task in enumerate(goal_tasks):
        focus = task["title"].split("：", 1)[-1]
        chapter = task["chapter"]
        verb = verbs[index % len(verbs)]
        descriptions = {
            "了解": "的基本概念、主要作用和应用场景",
            "熟悉": "的基本规则、操作流程和使用规范",
            "掌握": "的关键方法、实现步骤和配合关系",
        }
        knowledge_items.append(f"{verb}{focus}{descriptions[verb]}")
        ability_items.append(f"能根据任务要求运用{focus}完成“{chapter}”相关实践任务，并进行运行检查、问题修正和规范提交")
    quality_items = [
        "树立正确的政治方向和价值观，认识数字技术服务经济社会发展的责任",
        "形成诚实守信、严谨负责、尊重事实和规范提交的职业道德",
        "具备遵守技术标准、操作规范和质量要求的职业意识",
        "提高网络安全、数据安全、个人信息保护和风险防范意识",
        "形成尊重知识产权、软件许可和数字资源版权的良好习惯",
        "具备团队分工、沟通反馈、互助协作和客观评价的职业素养",
        "发展自主学习、独立思考、创新实践、自我管理和持续改进能力",
    ]
    _set_cell(document.tables[2].rows[1].cells[0], _numbered_compact(quality_items))
    _set_cell(document.tables[2].rows[1].cells[1], _numbered_compact(knowledge_items))
    _set_cell(document.tables[2].rows[1].cells[2], _numbered_compact(ability_items))
    for cell in document.tables[2].rows[1].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "仿宋_GB2312"
                run.font.size = Pt(9)
                run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    structure = document.tables[3]
    sample = deepcopy(structure.rows[1]._tr)
    for row in list(structure.rows[1:]): _remove(row._tr)
    for task in tasks:
        structure._tbl.append(deepcopy(sample))
        for col, value in enumerate((task["seq"], task["chapter"], _task_detail(task), task["hours"])):
            _set_cell(structure.rows[-1].cells[col], value)
            _set_cell_small_font(structure.rows[-1].cells[col])
    requirements = document.tables[4]
    sample = deepcopy(requirements.rows[1]._tr)
    for row in list(requirements.rows[1:]): _remove(row._tr)
    for task in tasks:
        requirements._tbl.append(deepcopy(sample))
        values = (task["chapter"], _task_detail(task), task["knowledge_goal"], task["ability_goal"], task["ideological_goal"], "源代码、运行结果、任务报告或项目作品")
        for col, value in enumerate(values):
            _set_cell(requirements.rows[-1].cells[col], value)
            _set_cell_small_font(requirements.rows[-1].cells[col])
    scheme = get_scheme(offering_id)
    scheme_text = component_text(offering_id)
    assessment_data = ((f"过程性评价（{scheme['process_total']:g}%）", scheme_text, "教师、学生", "考勤记录、课堂表现记录、学习通作业成绩及其他过程证据"),
                       (f"终结性评价（{scheme['final_total']:g}%）", "期末综合作品、演示与答辩", "教师、学生", "依据成果质量、运行效果、规范性和答辩表现评分"),
                       ("增值评价（不单独计分）", "学习进步、规范意识、协作创新、安全版权", "教师、学生", "作为过程性或终结性评价的观察依据"),
                       ("综合评价", "技能竞赛、职业证书、创新成果和社会服务", "加分项", "成果证明与社会评价"))
    for row, row_values in zip(document.tables[5].rows[1:], assessment_data):
        for col, value in enumerate(row_values):
            _set_cell(row.cells[col], value)
            _set_cell_small_font(row.cells[col])
    _set_table_font(document.tables[3], size=9)
    _set_table_font(document.tables[4], size=9)
    _set_table_font(document.tables[5], size=9)
    structure_heading = next((p for p in document.paragraphs if p.text.strip() == "（一）课程结构"), None)
    if structure_heading:
        structure_text = (
            f"课程对接{offering['major']}专业相关岗位任务和职业能力要求，以教材项目为基础重构为{len(project_names)}个模块、{len(tasks)}个教学任务，"
            f"共{offering['total_hours']}学时，其中理论{theory}学时、实践{practice}学时。课程由基础认知与工具使用逐步过渡到综合任务实施，"
            "融入行业标准、技术规范、技能竞赛评价要求以及新技术、新工艺和新方法，形成岗课赛证融通、理实一体的课程结构。"
        )
        paragraph = _insert_after(structure_heading, structure_text)
        generated_paragraphs.append(paragraph)
    assessment_heading = next((p for p in document.paragraphs if p.text.strip() == "四、课程考核与评价"), None)
    if assessment_heading:
        paragraph = _insert_after(assessment_heading, "课程建立过程性评价、结果评价、增值评价和综合评价相结合的多元评价体系，重点考查知识理解、操作过程、项目成果、问题解决、规范安全、协作创新和学习进步。评价由教师评价、学生自评互评及必要的行业企业评价共同构成，以评价结果促进学生持续改进。")
        generated_paragraphs.append(paragraph)
    implementation_sections = (
        ("1.教学模式（或策略）", "坚持学生中心、产教融合、工学结合、德技并修和成果导向，实施理实一体、项目引领、任务驱动的教学模式。以工作过程组织课堂，以可运行、可展示、可评价的项目成果检验学习效果，并根据学生基础实施分层指导。"),
        ("2.教学方法", "本课程根据课程内容、学生基础和项目特点，以学定教，综合采用案例分析、启发引导、分组讨论、任务教学、项目教学、操作演示和成果评价等方法，兼顾教学效果、课堂可操作性和学生差异。"),
        ("（1）教授方法", "根据课程内容和学生特点，综合采用案例教学法、问题引导法、任务驱动法、项目教学法、操作演示法和分组讨论法。教师通过案例呈现、概念分析、技术示范、巡视指导和共性问题讲评，引导学生主动思考、规范操作并及时修正问题。案例教学法是以典型案例为载体，引导学生从现象观察、问题分析到方案比较；问题引导法是以递进问题组织认知过程，促进学生主动思考；任务驱动法是将学习目标分解为可操作、可检查的任务；项目教学法是围绕完整成果组织需求分析、实施、测试和交付；操作演示法通过教师示范关键步骤和易错点帮助学生形成规范动作；分组讨论法通过角色分工、交流互评和共同决策培养协作能力。"),
        ("（2）学习方法", "引导学生采用自主学习、观察分析、模仿练习、协同合作、探究验证和复盘改进等方法。学生通过查阅教材与数字资源、研读示例、动手实践、记录问题、自评互评和成果迭代，提高自主学习、问题解决、沟通协作和创新实践能力。"),
        ("1.授课教室", "授课教室应配备多媒体教学设备、稳定网络和教学管理平台，能够开展课件展示、案例讲解、操作演示、课堂互动、作品展示和过程评价。"),
        ("（1）校内实训基地条件要求", "校内实训场所应按班级规模配置满足课程项目运行要求的计算机、网络环境和相关软件，支持教师演示、学生独立操作、文件存储、成果提交及数据备份；设备与账号管理应符合网络安全、数据安全和软件版权要求。"),
        ("（2）校外实训基地条件要求", "校外实训基地应具有与课程相关的真实业务任务、规范工作流程和必要软硬件条件，能够安排企业人员参与任务指导、过程检查和成果评价，并保障实训期间的人员、设备、数据和信息安全。"),
    )
    for heading, text in implementation_sections:
        anchor = next((p for p in document.paragraphs if p.text.strip() == heading), None)
        if anchor:
            paragraph = _insert_after(anchor, text)
            generated_paragraphs.append(paragraph)
    teacher_intro = "任课教师应强化课程思政意识，将思想政治教育融入专业教学全过程；持续学习行业新标准、新技术、新工艺和新方法，具备扎实专业基础、项目实践能力、课程设计能力以及安全、创新、版权和数据保护意识。"
    teacher_sections = (
        ("（二）教师基本要求", [teacher_intro]),
        ("1.专任教师", [f"专任教师应具备{offering['course_name']}相关专业知识，熟悉本课程教材项目、岗位工作流程和质量标准，能够使用课程所需软件与数字工具完成示范、指导、调试和评价；应具有较强的实践能力、教学设计能力、课堂组织能力和因材施教能力，能够及时解决学生在操作、运行、测试和成果交付中的问题。"]),
        ("2.行业导师/企业教师", ["行业导师或企业教师应具有相关岗位实践经历，熟悉行业工作流程、技术规范和质量要求，能够提供真实或仿真任务、指导项目实施、评价成果质量，并将企业安全、协作和交付要求融入教学。"]),
        ("3.兼职教师", ["兼职教师应具备与课程相关的专业技能和职业素养，能够承担专题讲座、案例指导、实践教学或成果评价，遵守学校教学管理要求和学生数据、知识产权保护规定。"]),
    )
    resources = drafts.get(("course_resources", ""), {})
    resource_sections = (
        ("（四）教学资源", ["教学资源建设应服务课程目标和项目任务，符合高职教育教学要求，保持内容的职业性、实践性、先进性和可更新性。"]),
        ("1.教材选用与开发", [f"本课程选用《{offering['textbook_version']}》作为主要教材，教材应符合高职教育要求，内容以项目和工作过程为载体，能够支撑知识学习、技能训练和成果评价。教材使用中应结合{offering['major']}专业岗位需求补充行业标准、软件版本、工具方法、数据安全、知识产权和课程思政内容，并根据技术发展和教学反馈持续更新。"]),
        ("2.书籍配备", ["配备与课程相关的专业基础、技术规范、岗位标准和项目实践参考书，满足教师备课、学生拓展阅读和项目问题解决需要；书籍版本应优先选择正规出版物并及时更新。"]),
        ("3.数字资源", ["建设并使用教材配套课件、实训指导书、任务单、操作演示视频、示例代码、项目素材、习题库、评价量表和优秀作品库。充分利用国家职业教育智慧教育平台、学习通、国家及省级在线精品课程资源和专业技术网站，开展线上线下混合式教学；数字资源应来源合法、标注清晰，并落实网络安全、数据安全、个人信息保护和数字版权要求。"]),
    )
    for heading, content in (*teacher_sections, *resource_sections):
        paragraph = next((p for p in document.paragraphs if p.text.strip() == heading), None)
        if paragraph:
            cursor = paragraph
            for text in content:
                if text:
                    cursor = _insert_after(cursor, text)
                    generated_paragraphs.append(cursor)
    for paragraph in generated_paragraphs:
        _set_body_font(paragraph)
    clean_headings = {
        "（一）总体目标": "（一）总体目标",
        "（二）具体要求": "（二）具体要求",
        "（二）课程内容及要求": "（二）课程内容及要求",
        "五、课程实施与保障": "五、课程实施与保障",
        "（一）教学要求": "（一）教学要求",
        "（三）教学设施": "（三）教学设施",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        for prefix, replacement in clean_headings.items():
            if text.startswith(prefix) and text != replacement:
                _set_paragraph(paragraph, replacement)
                break
    placeholder_prefixes = (
        "注：", "（注", "(注", "说明：", "（说明", "(说明", "示例一", "示例二", "示例：",
        "（针对", "(针对", "（要求", "(要求", "(教学组织", "（教学组织", "(表格中", "（表格中",
        "文字+表格", "评价建议", "根据本课程的教学目标要求", "XXXX教学法",
        "担任本课程的主讲教师除了", "要求在行业/企业", "（提出选用教材", "（课程教学需要",
    )
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if text.startswith(placeholder_prefixes) or "填写说明" in text or "××" in text or "……" in text:
            _remove(paragraph._p)
    document.save(output)


def _replace_xml_text(docx_path, replacements):
    docx_path = Path(docx_path)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, dir=docx_path.parent) as handle:
        temp_path = Path(handle.name)
    with ZipFile(docx_path, "r") as source, ZipFile(temp_path, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                nodes = root.xpath("//*[local-name()='t']")
                for old, new in replacements.items():
                    while True:
                        full = "".join(node.text or "" for node in nodes)
                        start = full.find(old)
                        if start < 0:
                            break
                        end = start + len(old)
                        cursor = 0
                        touched = []
                        for node in nodes:
                            length = len(node.text or "")
                            if cursor < end and cursor + length > start:
                                touched.append(node)
                            cursor += length
                        if not touched:
                            break
                        touched[0].text = new
                        for node in touched[1:]:
                            node.text = ""
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            target.writestr(item, data)
    temp_path.replace(docx_path)


def _replace_labeled_paragraphs(docx_path, replacements, underline_values=False):
    """Replace labels that live in text boxes/shapes, which python-docx cannot expose."""
    docx_path = Path(docx_path)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, dir=docx_path.parent) as handle:
        temp_path = Path(handle.name)
    with ZipFile(docx_path, "r") as source, ZipFile(temp_path, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                for paragraph in root.xpath("//*[local-name()='p']"):
                    nodes = paragraph.xpath(".//*[local-name()='t']")
                    full = "".join(node.text or "" for node in nodes)
                    for label, value in replacements.items():
                        if label not in full:
                            continue
                        first = nodes[0] if nodes else None
                        if first is None:
                            continue
                        first.text = f"{label}{value}"
                        for node in nodes[1:]:
                            node.text = ""
                        if underline_values:
                            run = first.getparent()
                            rpr = run.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
                            if rpr is None:
                                rpr = etree.SubElement(run, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
                            underline = rpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}u")
                            if underline is None:
                                underline = etree.SubElement(rpr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}u")
                            underline.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "single")
                        break
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            target.writestr(item, data)
    temp_path.replace(docx_path)


def _fill_xml_goal_blocks(docx_path, blocks):
    """Fill goal lines inside Word text boxes/shapes, invisible to python-docx."""
    docx_path = Path(docx_path)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, dir=docx_path.parent) as handle:
        temp_path = Path(handle.name)
    wns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with ZipFile(docx_path, "r") as source, ZipFile(temp_path, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                paragraphs = root.xpath("//*[local-name()='p']")
                texts = ["".join(node.text or "" for node in p.xpath(".//*[local-name()='t']")) for p in paragraphs]
                for marker, values in blocks.items():
                    for index, text in enumerate(texts):
                        if marker not in text:
                            continue
                        nodes = paragraphs[index].xpath(".//*[local-name()='t']")
                        if nodes:
                            nodes[0].text = text + "\n" + "\n".join(f"{offset}. {value}" for offset, value in enumerate(values, 1))
                            for node in nodes[1:]:
                                node.text = ""
                        break
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            target.writestr(item, data)
    temp_path.replace(docx_path)


def _underline_labeled_paragraphs(docx_path, labels):
    docx_path = Path(docx_path)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, dir=docx_path.parent) as handle:
        temp_path = Path(handle.name)
    wns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with ZipFile(docx_path, "r") as source, ZipFile(temp_path, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                for paragraph in root.xpath("//*[local-name()='p']"):
                    full = "".join(node.text or "" for node in paragraph.xpath(".//*[local-name()='t']"))
                    if not any(label in full for label in labels):
                        continue
                    for run in paragraph.xpath(".//*[local-name()='r']"):
                        rpr = run.find(f"{{{wns}}}rPr")
                        if rpr is None:
                            rpr = etree.SubElement(run, f"{{{wns}}}rPr")
                        underline = rpr.find(f"{{{wns}}}u")
                        if underline is None:
                            underline = etree.SubElement(rpr, f"{{{wns}}}u")
                        underline.set(f"{{{wns}}}val", "single")
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            target.writestr(item, data)
    temp_path.replace(docx_path)


def _set_tc_vmerge(tc, mode):
    """设置单元格纵向合并标记。mode: 'restart' 开始合并，'continue' 延续合并。"""
    tcPr = tc.get_or_add_tcPr()
    vm = tcPr.get_or_add_vMerge()
    if mode == "restart":
        vm.set(qn("w:val"), "restart")
    elif vm.get(qn("w:val")) is not None:
        del vm.attrib[qn("w:val")]


def _fill_standard(offering_id, template, output):
    if any(p.text.strip() == "一、课程概述" for p in Document(template).paragraphs):
        return _fill_standard_2025(offering_id, template, output)
    offering = store.rows("SELECT * FROM offerings WHERE id=?", (offering_id,))[0]
    units = store.rows("SELECT * FROM curriculum_units WHERE offering_id=? AND approval_status='已确认' AND review_action<>'删除' ORDER BY seq", (offering_id,))
    drafts = _drafts(offering_id, "课程标准")
    model_row = store.rows("SELECT model_json FROM course_content_models WHERE offering_id=?", (offering_id,))
    semantic_model = json.loads(model_row[0]["model_json"]) if model_row else {}
    scenario_tools = [t for t in semantic_model.get("tools_technology", []) if isinstance(t, str) and t.strip()]
    tools_line = "、".join(scenario_tools[:3]) if scenario_tools else "课程配套软件工具"
    # 支持增量填充：如果输出文件已存在（如已由槽位填充器处理过），则直接打开
    output_path = Path(output)
    if not output_path.exists():
        shutil.copy2(template, output)
    document = Document(output)

    # 获取模板分析契约，用于表格角色定位
    template_file_id = _get_template_file_id(offering_id, "课程标准")
    contract = _get_template_contract(template_file_id) if template_file_id else None
    _replace_prefix(document, "《**********》（居中", f"《{offering['course_name']}》")
    _replace_prefix(document, "适用专业：", f"适用专业：{offering['major']}")
    _replace_prefix(document, "编制单位：", f"编制单位：{store.get_setting('department', '经济贸易系')}")
    _replace_prefix(document, "合作单位：", f"合作单位：{store.get_setting('partner_company', '')}" if store.get_setting("partner_company") else "")
    _replace_prefix(document, "《**********》课程标准", f"《{offering['course_name']}》课程标准")
    _replace_prefix(document, "课程名称及课程编号：", f"课程名称及课程编号：{offering['course_name']}（{offering['course_code']}）")
    _replace_prefix(document, "课程类型：", f"课程类型：{offering['course_type']}")
    credits = int(offering["credits"]) if float(offering["credits"]).is_integer() else offering["credits"]
    _replace_prefix(document, "学时学分：", f"学时学分：{offering['total_hours']}学时（{credits}学分）")
    _replace_prefix(document, "开设学期：", f"开设学期：{_study_semester(offering.get('teaching_class'), offering['term'])}")
    notes = [p for p in document.paragraphs if p.text.strip().startswith(("（注", "(注", "（正文", "（要求"))]
    nature_heading = next(p for p in document.paragraphs if p.text.strip().startswith("一、课程性质"))
    cursor = nature_heading
    project_names = list(dict.fromkeys(u["project_title"] for u in units if u["project_title"]))
    project_text = "、".join(project_names[:6])
    nature_content = [t for t in drafts.get(("course_nature", ""), []) if str(t).strip()]
    if not nature_content:
        nature_content = [
            f"《{offering['course_name']}》是{offering['major']}专业开设的一门{offering['course_type']}，旨在使学生系统掌握{project_text}所涉及的基础知识、核心方法和规范化的项目实施流程，形成运用专业工具解决实际业务问题的能力，为专业岗位能力培养提供支撑。",
            f"课程以{project_text}等教材项目为载体，将知识学习、技能训练、成果检查和规范交付贯穿教学全过程，突出实践性、职业性和综合性；前导课程为相关专业基础课程，后续支撑综合实践课程与岗位实习。",
        ]
    for text in nature_content:
        cursor = _insert_after(cursor, text, document.paragraphs[29])
    goals = drafts.get(("course_goals", ""), {})
    if not isinstance(goals, dict) or not any(goals.get(k) for k in ("知识目标", "能力目标", "思政目标", "素质目标")):
        knowledge_items = [str(k) for k in semantic_model.get("knowledge_system", []) if str(k).strip()]
        ability_items = [
            str(p.get("expected_outcome") or "").strip()
            for p in semantic_model.get("projects", [])
            if str(p.get("expected_outcome") or "").strip()
        ]
        goals = {
            "知识目标": [f"掌握{item}" for item in knowledge_items[:8]] or [f"掌握{offering['course_name']}的基础知识与核心方法"],
            "能力目标": ability_items[:8] or ["能够运用课程工具完成项目任务并规范交付成果"],
            "思政目标": ["坚持诚信分析、数据安全与职业道德", "树立规范意识、责任意识和社会责任感"],
            "素质目标": ["形成团队协作、沟通反馈和持续改进习惯", "具备自主学习、独立思考和问题解决能力"],
        }
    for col, key in enumerate(("知识目标", "能力目标", "思政目标", "素质目标")):
        _set_cell(document.tables[0].rows[1].cells[col], _numbered(goals.get(key, [])))
    design_heading = next(p for p in document.paragraphs if p.text.strip().startswith("1、本课程设计的总体思路"))
    cursor = design_heading
    design_content = [t for t in drafts.get(("course_design", ""), []) if str(t).strip()]
    if not design_content:
        design_content = [
            f"本课程依据专业人才培养方案和{offering['major']}专业岗位能力要求设计，以{project_text}等项目为主线，将知识点、技能点和职业规范转化为递进的学习任务，按照“任务分析—方案设计—实施操作—检查评价—迭代改进”的路径组织{offering['total_hours']}学时教学。",
            "课程设计融入课程思政、劳动教育和数据安全、知识产权要求，结合行业发展补充新标准、新技术、新工艺和新方法，以真实或仿真项目成果作为学习产出，体现课程标准的职业性、实践性、先进性和可评价性。",
        ]
    for text in design_content:
        cursor = _insert_after(cursor, text, document.paragraphs[29])
    content_table = document.tables[1]
    template_row = deepcopy(content_table.rows[1]._tr)
    for row in list(content_table.rows[1:-1]):
        _remove(row._tr)
    scenarios = []
    for unit in units:
        for sub in _subscenarios(unit):
            scenarios.append((unit, sub))
    for unit, sub in scenarios:
        row_xml = deepcopy(template_row)
        content_table.rows[-1]._tr.addprevious(row_xml)
        row = content_table.rows[-2]
        for index, value in enumerate((unit["seq"], unit["project_title"], sub["seq"], sub["title"], sub["hours"], f"{unit['suggested_hours']}（理论{int(unit['suggested_hours'])//2}+实践{int(unit['suggested_hours'])-int(unit['suggested_hours'])//2}）")):
            _set_cell(row.cells[index], value)
    _set_cell(content_table.rows[-1].cells[-1], offering["total_hours"])
    _set_table_font(content_table, "仿宋_GB2312", 9)
    assessment = drafts.get(("assessment", ""), {})
    if not isinstance(assessment, dict):
        assessment = {}
    assessment_table = document.tables[2]
    scheme = get_scheme(offering_id)
    _set_cell(assessment_table.rows[0].cells[2], f"过程性考核（{scheme['process_total']:g}%）")
    _set_cell(assessment_table.rows[0].cells[4], f"终结性考核（{scheme['final_total']:g}%）")
    process_eval = assessment.get("过程性评价", {})
    if not isinstance(process_eval, dict):
        process_eval = {}
    selected = [str(p) for p in process_eval.get("抽取项目", []) if str(p).strip()]
    if not selected:
        selected = [
            p["title"] for p in semantic_model.get("projects", [])
            if isinstance(p, dict) and p.get("title") and p["title"] != "综合评价与课程总结"
        ][:4]
    row_index = 2
    # 模板 c3 表头为"权重"；槽位填充可能将其误写为其他值，这里恢复
    _set_cell(assessment_table.rows[1].cells[3], "权重")
    # 清除数据区原有纵向合并，避免合并单元格解析导致写入丢失或清空穿透
    for row in assessment_table.rows[2:]:
        for tc in row._tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            vm = tcPr.find(qn("w:vMerge")) if tcPr is not None else None
            if vm is not None:
                tcPr.remove(vm)
    for project in selected:
        for task_name, fraction in (("知识技能", 0.4), ("任务实施", 0.6)):
            if row_index >= len(assessment_table.rows):
                break
            methods = []
            percents = []
            for component in scheme["components"]:
                value = (float(component["weight"]) / scheme["process_total"] * 25 * fraction) if scheme["process_total"] else 0
                methods.append(str(component["component_name"]))
                percents.append(f"{value:g}%")
            _set_cell(assessment_table.rows[row_index].cells[0], f"{project}（25%）")
            _set_cell(assessment_table.rows[row_index].cells[1], task_name)
            _set_cell(assessment_table.rows[row_index].cells[2], "\n".join(methods))
            _set_cell(assessment_table.rows[row_index].cells[3], "\n".join(percents))
            _set_cell(assessment_table.rows[row_index].cells[4], f"课程综合作品、成果展示与答辩（{scheme['final_total']:g}%）")
            row_index += 1
    # 删除未使用的多余数据行（模板预留行多于实际项目数）
    for row in list(assessment_table.rows[row_index:]):
        _remove(row._tr)
    # 重建纵向合并：学习情境列按项目两行合并，终结性考核列覆盖整个数据区
    for ri in range(2, row_index, 2):
        _set_tc_vmerge(assessment_table.rows[ri]._tr.findall(qn("w:tc"))[0], "restart")
        _set_tc_vmerge(assessment_table.rows[ri + 1]._tr.findall(qn("w:tc"))[0], "continue")
    if row_index > 2:
        for ri in range(2, len(assessment_table.rows)):
            tcs = assessment_table.rows[ri]._tr.findall(qn("w:tc"))
            if len(tcs) > 4:
                _set_tc_vmerge(tcs[4], "restart" if ri == 2 else "continue")
    _set_table_font(assessment_table, "仿宋_GB2312", 9)
    scenario_template = document.tables[3]
    pristine_scenario_xml = deepcopy(scenario_template._tbl)
    anchor = scenario_template._tbl
    scenarios = list(drafts.get(("learning_scenarios", ""), []))
    if not scenarios:
        scenarios = [
            {"项目": p.get("title", ""), "学时": p.get("hours", ""), "知识技能": p.get("knowledge_skills", [])}
            for p in semantic_model.get("projects", [])
            if isinstance(p, dict) and p.get("title")
        ]
    for index, project in enumerate(scenarios):
        table_xml = scenario_template._tbl if index == 0 else deepcopy(pristine_scenario_xml)
        if index:
            anchor.addnext(table_xml)
            anchor = table_xml
        table = Table(table_xml, scenario_template._parent)
        skills = [s for s in (project.get("知识技能") or []) if str(s).strip()]
        _set_cell(table.rows[0].cells[0], f"课程：{offering['course_name']}")
        _set_cell(table.rows[0].cells[2], "学分")
        _set_cell(table.rows[0].cells[3], credits)
        _set_cell(table.rows[0].cells[5], f"总学时：{offering['total_hours']}")
        _set_cell(table.rows[1].cells[0], f"学习情境{index+1}：{project['项目']}")
        _set_cell(table.rows[1].cells[5], f"学时：{project['学时']}")
        _set_cell(table.rows[3].cells[0], f"知识目标：\n1.理解{'、'.join(skills[:2])}\n2.掌握{'、'.join(skills[2:4] or skills[:2])}\n能力目标：\n1.完成项目任务与成果\n2.检查修改并解决问题\n3.规范整理和提交成果\n思政目标：\n1.遵守数据安全与规范要求\n2.坚持诚信、规范和责任意识\n素质目标：\n1.形成协作、检查和持续改进习惯")
        _set_cell(table.rows[3].cells[1], "主要内容：\n" + _numbered(skills) + "\n项目流程：任务分析、方案设计、实施操作、检查测试、迭代优化、成果交付。")
        _set_cell(table.rows[3].cells[4], "教学方法：\n1.项目教学\n2.任务驱动\n3.演示练习\n建议：\n1.使用教材PPT和项目素材\n2.实施分层指导\n3.依据评价清单改进")
        _set_cell(table.rows[5].cells[0], f"多媒体教学设备\n教材配套PPT\n项目数据与素材\n{tools_line}")
        _set_cell(table.rows[5].cells[1], "评价内容：知识理解、操作过程、项目成果、规范安全、创新与协作。\n评价方法：课堂观察、练习检查、阶段作品、同伴互评和成果评价。")
        _set_cell(table.rows[5].cells[4], f"通过本学习情境掌握{project['项目']}所需知识技能，提升项目实施、问题解决和规范交付能力。")
        unit = units[index] if index < len(units) else None
        subs = _subscenarios(unit) if unit is not None else []
        block_rows = [deepcopy(table.rows[row]._tr) for row in range(6, 11)]
        for sub_index in range(1, len(subs)):
            for row_xml in block_rows:
                table._tbl.append(deepcopy(row_xml))
        if not subs:
            for row in range(6, 11):
                for cell in table.rows[row].cells:
                    _set_cell(cell, "")
        for sub_index, sub in enumerate(subs):
            base = 6 + sub_index * 5
            _set_cell(table.rows[base].cells[0], f"学习子情境{index+1}.{sub_index+1}：{sub['title']}")
            _set_cell(table.rows[base].cells[5], f"学时：{sub['hours']}")
            _set_cell(table.rows[base+2].cells[0], f"教学目标：\n1.理解{sub['title']}的基本规则\n2.能够完成对应练习或项目实施\n思政目标：\n1.遵守技术规范与版权要求\n2.坚持诚信操作和责任意识\n素质目标：\n1.形成主动实践和及时检查习惯")
            _set_cell(table.rows[base+2].cells[1], f"1.{sub['title']}的概念、规则或实施要求\n2.教材示例与关键操作\n3.常见错误及检查方法")
            _set_cell(table.rows[base+2].cells[4], "教学方法：\n1.提问引导\n2.操作演示\n3.技术练习\n建议：\n1.对照示例操作\n2.记录错误现象\n3.完成课堂验收")
            _set_cell(table.rows[base+4].cells[0], f"多媒体设备、教材PPT、项目数据与素材、{tools_line}")
            _set_cell(table.rows[base+4].cells[1], "评价知识理解、操作规范、练习或阶段成果、安全版权和改进情况；采用课堂观察、操作检查、成果评价与反馈修改。")
            _set_cell(table.rows[base+4].cells[4], f"掌握{sub['title']}并能够用于“{project['项目']}”项目。")
    implementation = [t for t in drafts.get(("teacher_requirements", ""), []) if str(t).strip()]
    resource = drafts.get(("course_resources", ""), {})
    if not isinstance(resource, dict):
        resource = {}
    if not implementation:
        knowledge_items = [str(k) for k in semantic_model.get("knowledge_system", []) if str(k).strip()]
        work_process = [str(w) for w in semantic_model.get("work_process", []) if str(w).strip()]
        teaching_methods = [str(m) for m in semantic_model.get("teaching_methods", []) if str(m).strip()]
        implementation = [
            "任课教师应坚持正确政治方向，强化课程思政意识，将职业道德、数据安全意识、社会责任和诚信评价融入专业教学全过程。",
            f"任课教师应具备扎实的{offering['course_name']}专业基础，熟悉{'、'.join(knowledge_items[:10])}，能够依据教材项目解释原理、示范操作并处理学生实践中的典型问题。",
            f"任课教师应能够熟练使用{tools_line}开展项目实践、运行检查和成果评价，持续关注课程涉及的标准、技术和工具更新。",
            f"任课教师应能按照{'；'.join(work_process)}将教材项目转化为教学任务，实施{'、'.join(teaching_methods)}，并兼顾数据安全、知识产权、创新意识和分层指导。",
        ]
    textbook = resource.get("教材") or offering["textbook_version"]
    materials_text = resource.get("教学资料") or "使用教材配套课件、实训文档、项目数据与素材、任务单和评价清单；教学时按项目映射调用，保持内容与当前学期教材一致。"
    development_text = resource.get("开发利用") or "围绕各项目建设课件要点、操作演示、任务单、评价清单和优秀作品库；资源更新应记录来源、适用项目和版本，并检查版权、个人信息和运行安全。"
    for heading, content in (("（一）教师知识素质要求", implementation), ("（二）教材编写与选用", [f"本课程选用《{textbook}》作为主要教材。", f"教材编写与选用应符合高职{offering['major']}专业教学要求，内容以项目和工作过程为载体，能够支撑知识学习、技能训练和成果评价，并结合技术发展和教学反馈持续更新。"]), ("（三）课程资源的开发与利用", [materials_text, development_text])):
        paragraph = next(p for p in document.paragraphs if p.text.strip() == heading)
        cursor = paragraph
        for text in content:
            cursor = _insert_after(cursor, text, document.paragraphs[29])
    _replace_prefix(document, "根据本课程的教学目标要求和课程特点", "根据课程项目特点和学生学习基础，综合采用项目教学法、任务驱动法、案例教学法、演示教学法、分层练习法和小组合作法。教师以项目成果明确学习目标，以任务单组织知识学习和技术训练，通过示范、练习、检查、反馈和迭代修改促进学生掌握知识技能并形成规范交付能力。")
    _replace_prefix(document, "XXXX教学法是", "项目教学法以完整项目成果为载体组织教学；任务驱动法将项目分解为可操作、可检查的学习任务；案例与演示教学用于分析关键技术和操作过程；分层练习与合作学习兼顾学生差异并培养沟通协作能力。")
    _replace_prefix(document, "《×××》课程考核评价", f"《{offering['course_name']}》课程考核评价")
    _replace_prefix(document, "注：情境（子情境）", "")
    book_match = re.match(r"(.+?)[（(](.+?)[)）]$", str(offering["textbook_version"] or "").strip())
    book_title = book_match.group(1) if book_match else str(offering["textbook_version"] or "").strip()
    book_publisher = f"，{book_match.group(2)}。" if book_match else "。"
    reference_books = [
        f"（1）《{book_title}》{book_publisher}",
        "（2）本课程教材配套的实训指导书、任务单与项目数据集。",
        f"（3）与{offering['course_name']}相关的行业标准、岗位规范与典型案例汇编。",
    ]
    tools_text = "".join(scenario_tools)
    reference_sites = ["（1）国家职业教育智慧教育平台：https://vocational.smartedu.cn/"]
    site_index = 2
    if "Python" in tools_text or "pandas" in tools_text:
        reference_sites.append(f"（{site_index}）Python 官方文档：https://docs.python.org/zh-cn/3/")
        site_index += 1
    if "pandas" in tools_text:
        reference_sites.append(f"（{site_index}）pandas 官方文档：https://pandas.pydata.org/docs/")
        site_index += 1
    if "电子表格" in tools_text or "Excel" in tools_text:
        reference_sites.append(f"（{site_index}）Microsoft Excel 帮助与培训：https://support.microsoft.com/zh-cn/excel")
        site_index += 1
    journals = ["（1）《计算机教育》", "（2）《现代教育技术》", "（3）《软件导刊》", "（4）《电脑知识与技术》"]
    for prefix, items in (("1.参考书", reference_books), ("2.期刊", journals), ("3.网站", reference_sites)):
        heading = next((p for p in document.paragraphs if p.text.strip().startswith(prefix)), None)
        if heading is None:
            continue
        cursor = heading
        for item in items:
            cursor = _insert_after(cursor, item, document.paragraphs[29])
    for paragraph in notes:
        _set_paragraph(paragraph, "")
    document.save(output)


def _fill_plan(offering_id, template, output):
    offering = store.rows("SELECT * FROM offerings WHERE id=?", (offering_id,))[0]
    tasks = store.rows("SELECT * FROM tasks WHERE offering_id=? ORDER BY seq", (offering_id,))
    sessions = store.rows("SELECT * FROM sessions WHERE offering_id=? AND status='已确认' ORDER BY lesson_date,id", (offering_id,))
    # 支持增量填充
    output_path = Path(output)
    if not output_path.exists():
        shutil.copy2(template, output)
    document = Document(output)
    
    # 获取模板分析契约
    template_file_id = _get_template_file_id(offering_id, "授课计划")
    contract = _get_template_contract(template_file_id) if template_file_id else None
    term_label = _term_label(offering["term"])
    plan_term = _plan_term_label(offering["term"])
    _replace_prefix(document, "20XX——20XX学年第", term_label)
    _replace_prefix(document, "《……》教案", f"《{offering['course_name']}》教案")
    _replace_prefix(document, "主讲教师：", f"主讲教师：{offering.get('teacher_name', '') or store.get_setting('teacher_name', '')}        授课班级：{offering['teaching_class']}")
    first_date = next((s["lesson_date"] for s in sessions if s["lesson_date"]), "")
    _replace_prefix(document, "20 年 月 日", first_date[:7].replace("-", "年") + "月" if first_date else "")
    schedule_title = f"山西林业职业技术学院{plan_term}《{offering['course_name']}》课程教学日程表"
    title_paragraph = next((p for p in document.paragraphs if "课程教学日程表" in p.text or "20XX" in p.text and "山西林业职业技术学院" in p.text), None)
    if title_paragraph:
        _set_paragraph(title_paragraph, schedule_title)
    theory = sum(int(t["theory_hours"] or 0) for t in tasks)
    practice = sum(int(t["practice_hours"] or 0) for t in tasks)
    first_class = str(offering.get("teaching_class") or "").replace("；", ";").split(";")[0].strip()
    _replace_prefix(document, "专业：", f"专业：{offering['major']}  班级：{first_class}  总学时：{offering['total_hours']}   （其中：课堂教学：{theory}    实验实习：{practice}）")
    _replace_prefix(document, "20  年  月  日", f"{(first_date or '')[:4]}  年  {(first_date or '')[5:7]}  月  {(first_date or '')[8:10]}  日" if first_date else "")
    table = document.tables[0]
    template_row = deepcopy(table.rows[2]._tr)
    for row in list(table.rows[2:]):
        _remove(row._tr)
    for task in tasks:
        session = next((s for s in sessions if s["lesson_date"] == task["lesson_date"]), {})
        detail = _task_detail(task)
        skills = [item.strip() for item in re.split(r"[、；;\n]+", detail) if item.strip()]
        if detail == task["title"] or skills == [task["chapter"]]:
            skills = []
        theory_lines = [task["chapter"], *skills]
        start_row = len(table.rows)
        for line in theory_lines:
            table._tbl.append(deepcopy(template_row))
            _set_cell(table.rows[-1].cells[3], line)
        theory_end_row = len(table.rows) - 1
        table._tbl.append(deepcopy(template_row))
        practice_row = len(table.rows) - 1
        _set_cell(table.rows[practice_row].cells[6], f"任务实施：{task['chapter']}")
        _set_cell(table.rows[practice_row].cells[7], task["practice_hours"])
        _set_cell(table.rows[practice_row].cells[8], session.get("classroom", ""))
        _set_cell(table.rows[practice_row].cells[9], "")
        for col, value in {0: task["seq"], 1: task["week_no"] or "", 2: task["lesson_date"]}.items():
            cell = table.cell(start_row, col)
            cell = cell.merge(table.cell(practice_row, col))
            _set_merged_cell(cell, value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for col, value in {4: task["theory_hours"], 5: session.get("classroom", "")}.items():
            cell = table.cell(start_row, col)
            if theory_end_row > start_row:
                cell = cell.merge(table.cell(theory_end_row, col))
            _set_merged_cell(cell, value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # 获取已采纳的内容更新，融入授课计划说明
    from content_updater import format_updates_for_prompt
    accepted_updates = store.get_accepted_updates(offering_id)
    update_note = ""
    if accepted_updates:
        update_topics = "、".join(u["topic"] for u in accepted_updates[:5])
        update_note = f"本课程已融入{len(accepted_updates)}项内容更新（{update_topics}等），教学中将结合行业新标准、新技术对教材内容进行补充和优化。"
    
    _set_cell(document.tables[1].cell(0, 0), f"本授课计划依据{offering['term']}教学安排、学校校历和《{offering['textbook_version']}》编制。课程采用理实一体、项目化和任务驱动教学，共完成{offering['total_hours']}学时。教学内容结合课程所属行业的新标准、新技术、新工艺和新方法，按照需求分析、方案设计、任务实施、检查测试、迭代优化和成果交付的工作流程组织。{update_note}因节假日产生的学时调整，以学校校历、调课和补课记录为准。")
    _set_cell(document.tables[2].cell(0, 0), f"课程德育主线：将诚信学习、劳动纪律、数据安全、数字版权、团队协作和精益求精贯穿《{offering['course_name']}》各项任务。\n实施方式：结合教材项目和项目任务，在需求分析、代码编写、运行调试、成果评价与答辩中适时渗透。\n评价方式：通过课堂观察、过程记录、项目成果、自评互评和综合评价考查学生的责任意识、规范意识、安全意识与协作表现。")
    document.save(output)


def _fill_design_table_basic(table, task, offering, classroom):
    task_detail = _task_detail(task)
    values = {(0, 2): task["week_no"] or "", (0, 4): task["hours"], (0, 6): task.get("_class_label") or offering["teaching_class"], (1, 2): offering.get("teacher_name", "") or store.get_setting("teacher_name", ""), (1, 6): task.get("_date_label") or task["lesson_date"], (2, 2): "理实一体", (2, 6): "√多媒体教室", (2, 10): f"√实训室（{str(classroom).replace('市场营销理实一体化教室（804）', '804教室')}）", (2, 13): "□企业", (2, 16): "□实习基地", (3, 2): task_detail}
    for (row, col), value in values.items():
        _set_cell(table.rows[row].cells[col], value)
    for row, key in ((4, "knowledge_goal"), (5, "ability_goal"), (6, "ideological_goal"), (7, "quality_goal")):
        _set_cell(table.rows[row].cells[2], task[key])
    _set_cell(table.rows[8].cells[2], f"教材以“{task['chapter']}”为载体组织本次知识与技能，围绕{task_detail}开展概念理解、示范操作、分步练习和结果检查。学生已具备前序基础，但在综合运用、错误定位和规范提交方面仍需通过分层指导与及时反馈提升。坚持学生主体、做中学和成果导向，将技术规范、版权安全与职业责任自然融入活动。")
    skills = [x.strip() for x in task_detail.split("、") if x.strip()][:4]
    for index in range(4):
        _set_cell(table.rows[11 + index].cells[2], f"{index+1}. {skills[index] if index < len(skills) else '项目综合实施与质量检查'}")


def _fill_design_table_org(table, task, activity, refs, offering_id=None):
    _set_cell(table.rows[0].cells[2], f"多媒体实训室；围绕“{task['chapter']}”组织效果观察、技术分析、操作练习、项目实施和成果评价。")
    _set_cell(table.rows[1].cells[2], "教材配套PPT、项目实训文档、源代码、图片/音视频素材、Visual Studio Code、浏览器及开发者工具。\n资源依据：" + "；".join(Path(x).name for x in refs if "\\" in x)[:500])
    _set_cell(table.rows[3].cells[2], activity["教学导入"])
    _set_cell(table.rows[3].cells[3], "案例导入、提问引导；学生观察、回忆、讨论并提出初步方案。")
    _set_cell(table.rows[3].cells[4], "激发学习兴趣，建立前后知识联系，培养观察、分析和判断能力。")
    _set_cell(table.rows[3].cells[5], "10分钟")
    _set_cell(table.rows[4].cells[2], "教师归纳学生观察与回答，明确本次教学任务、知识能力目标、操作步骤和成果验收标准。")
    _set_cell(table.rows[4].cells[3], "总结、归纳；学生思考、记录并确认任务要求。")
    _set_cell(table.rows[4].cells[4], "明确教学任务与目标，形成可执行的学习路径。")
    _set_cell(table.rows[4].cells[5], "5分钟")
    methods = {
        "任务1": "问题导向、案例分析、讲授归纳；观察、讨论、辨析练习",
        "任务2": "示范教学、源码研读、操作练习；模仿、调试、错误复盘",
        "任务3": "项目教学、任务驱动、合作学习；实践、互评、迭代改进",
    }
    teacher_actions = {
        "任务1": "展示目标效果并提出递进问题；结合PPT讲解概念、规则、适用场景和易错点；组织正误案例辨析并归纳结论。",
        "任务2": "打开教材源码定位关键结构、样式或参数；分步演示编写、运行、检查元素和错误定位；巡视并针对共性问题再次示范。",
        "任务3": "发布项目任务单和评价清单；指导学生按需求分析、制作、测试、优化顺序实施；组织成果展示、互评和现场修改。",
    }
    student_actions = {
        "任务1": "观察效果、回答问题并记录规则；比较正误示例，标注错误、说明依据并形成技术方案。",
        "任务2": "研读源码并同步输入、运行、观察和修改；使用开发者工具定位问题，保存运行截图和错误修改记录。",
        "任务3": "独立或协作完成阶段成果；开展运行检查、自评互评和迭代修改，提交源文件、截图及问题记录。",
    }
    goal_lines = {
        "任务1": "理解本次核心概念、规则和适用场景，能够据需求选择技术方案。",
        "任务2": "掌握关键代码与操作流程，能够运行、检查并排除常见错误。",
        "任务3": "完成可运行的阶段成果，达到规范制作、测试优化和诚信交付要求。",
    }
    for row, key in ((5, "任务1"), (6, "任务2"), (7, "任务3")):
        # 任务内容列（列2）使用富文本填充，支持图片和源码
        content = activity[key]
        if offering_id and ("[[图片:" in content or "[[源码:" in content):
            _fill_cell_rich(table.rows[row].cells[2], content, offering_id=offering_id, image_width=2.5)
        else:
            _set_cell(table.rows[row].cells[2], content)
        _set_cell(table.rows[row].cells[3], methods[key] + "。教师：" + teacher_actions[key] + "\n学生：" + student_actions[key])
        _set_cell(table.rows[row].cells[4], goal_lines[key])
        # The school example fixes these rows to a short sample height. Generated
        # activities are longer, so remove that limit and let Word grow/split them.
        tr_pr = table.rows[row]._tr.get_or_add_trPr()
        for child in list(tr_pr):
            if child.tag.endswith("}trHeight") or child.tag.endswith("}cantSplit"):
                tr_pr.remove(child)
    for row, minutes in ((5, "45分钟"), (6, "50分钟"), (7, "55分钟")):
        _set_cell(table.rows[row].cells[5], minutes)
    _set_cell(table.rows[8].cells[2], activity["课堂小结"])
    _set_cell(table.rows[8].cells[5], "15分钟")
    _set_cell(table.rows[9].cells[2], activity["课后作业"])
    _set_cell(table.rows[9].cells[5], "30分钟")
    _set_cell(table.rows[10].cells[2], activity["教学反思"])


def _fill_design(offering_id, template, output, week_no=None):
    offering = store.rows("SELECT * FROM offerings WHERE id=?", (offering_id,))[0]
    if week_no is not None:
        tasks = store.rows(
            "SELECT * FROM tasks WHERE offering_id=? AND week_no=? ORDER BY seq",
            (offering_id, week_no),
        )
        if not tasks:
            raise ValueError(f"第{week_no}周没有教学任务，无法生成该周单元教学设计。")
    else:
        tasks = store.rows("SELECT * FROM tasks WHERE offering_id=? ORDER BY seq", (offering_id,))
    sessions = store.rows("SELECT * FROM sessions WHERE offering_id=? AND status='已确认'", (offering_id,))

    # 增强模式：自动生成增强版单元教学设计（周次模式下仅处理该周任务）
    template_file_id = _get_template_file_id(offering_id, "教学设计")
    if enhanced_content_author.is_enhanced_mode_enabled():
        print("[增强模式] 检查单元教学设计内容...")
        enhanced_content_author.generate_and_save_unit_designs(
            offering_id=offering_id,
            task_ids=[t["id"] for t in tasks] if week_no is not None else None,
            force_regenerate=False,
        )
    
    drafts = _drafts(offering_id, "教学设计")
    # 支持增量填充
    output_path = Path(output)
    if not output_path.exists():
        shutil.copy2(template, output)
    document = Document(output)
    
    # 获取模板分析契约
    template_file_id = _get_template_file_id(offering_id, "教学设计")
    contract = _get_template_contract(template_file_id) if template_file_id else None
    goals = drafts.get(("course_goals", ""), {})
    replacements = {"《 》课程整体教学设计": f"《{offering['course_name']}》课程整体教学设计"}
    for paragraph in document.paragraphs:
        if paragraph.text.strip() in replacements:
            _set_paragraph(paragraph, replacements[paragraph.text.strip()])
    standard_drafts = _drafts(offering_id, "课程标准")
    course_nature = standard_drafts.get(("course_nature", ""), [])
    for heading, content in (
        ("二、课程定位", course_nature),
        ("三、教学设计的理念", ["坚持立德树人、学生中心、成果导向和持续改进，以教材真实项目为载体，将知识学习、技术训练、职业素养和价值引领统一到项目实施全过程。", "采用理实一体、项目教学和任务驱动，按照需求分析、设计、编码、测试、优化和交付组织教学，并通过分层练习、合作学习和过程性评价支持学生达成目标。"]),
        ("六、课程教学实施条件", ["配备多媒体实训室、可运行教材项目的计算机、Visual Studio Code、主流浏览器及开发者工具；网络和素材使用符合安全、版权及个人信息保护要求。"]),
        ("七、教学资源", [f"教材：{offering['textbook_version']}", "教学资料：教材配套PPT、实训文档、项目源代码、图片及音视频素材、任务单、评价清单和拓展案例。"]),
        ("八、需要说明的其他问题", ["（1）课程按每次4学时组织理实一体教学。", "（2）因节假日产生的缺课按学校校历和补课安排执行。", "（3）课程资源更新须记录来源、版本和适用项目，不使用历史课程文档替代本学期内容分析。"]),
    ):
        paragraph = next((p for p in document.paragraphs if p.text.strip() == heading), None)
        if paragraph:
            cursor = paragraph
            for text in content:
                cursor = _insert_after(cursor, text)
    goal_starts = {"1、认知目标：": "知识目标", "2、能力目标：": "能力目标", "3、思政目标：": "思政目标", "4、素质目标：": "素质目标"}
    for index, paragraph in enumerate(document.paragraphs):
        key = goal_starts.get(paragraph.text.strip())
        if key:
            for offset, text in enumerate(goals.get(key, [])[:4], 1):
                target = index + offset
                if target < len(document.paragraphs):
                    _set_paragraph(document.paragraphs[target], f"{offset}. {text}")
    for marker, values in (("教学模式：", ["①理实一体教学", "②项目化教学"]), ("教学方法：", ["①任务驱动、案例分析与演示操作", "②分层练习、合作学习与成果评价"])):
        index = next((i for i, p in enumerate(document.paragraphs) if p.text.strip() == marker), None)
        if index is not None:
            for offset, text in enumerate(values, 1):
                _set_paragraph(document.paragraphs[index + offset], text)
    for paragraph in document.paragraphs:
        if paragraph.text.strip() in {"（1）", "（2）", "（3）", "①", "②", "③", "④"}:
            _set_paragraph(paragraph, "")
    info = document.tables[0]
    first_lesson_date = next((s["lesson_date"] for s in sessions if s.get("lesson_date")), "")
    for row, col, value in ((0, 1, offering["course_name"]), (0, 3, offering["course_code"]), (0, 5, store.get_setting("department", "经济贸易系")), (1, 1, first_lesson_date or offering["term"]), (1, 3, offering.get("teacher_name", "") or store.get_setting("teacher_name", "")), (2, 1, offering["course_type"]), (2, 3, offering["total_hours"]), (2, 5, offering["credits"]), (3, 1, offering["major"]), (4, 1, offering.get("prerequisite_courses", "")), (4, 4, offering.get("followup_courses", ""))):
        _set_cell(info.rows[row].cells[col], value)
    _set_table_font(info, size=10.5)
    content = document.tables[1]
    units = store.rows("SELECT project_title,suggested_hours FROM curriculum_units WHERE offering_id=? AND approval_status='已确认' AND review_action<>'删除' ORDER BY seq", (offering_id,))
    sample_content_row = deepcopy(content.rows[1]._tr)
    for row in list(content.rows[1:]):
        _remove(row._tr)
    for unit in units:
        content._tbl.append(deepcopy(sample_content_row))
        _set_cell(content.rows[-1].cells[0], unit["project_title"])
        _set_cell(content.rows[-1].cells[1], unit["suggested_hours"])
    content._tbl.append(deepcopy(sample_content_row))
    _set_cell(content.rows[-1].cells[0], "合计")
    _set_cell(content.rows[-1].cells[1], offering["total_hours"])
    _set_table_font(content, size=9)
    # Fill the template's vocational ability project table instead of leaving its sample rows blank.
    ability_projects = document.tables[2]
    project_rows = store.rows(
        "SELECT project_title,suggested_hours FROM curriculum_units WHERE offering_id=? "
        "AND approval_status='已确认' AND review_action<>'删除' ORDER BY seq",
        (offering_id,),
    )
    ability_template = deepcopy(ability_projects.rows[1]._tr)
    for row in list(ability_projects.rows[1:]):
        _remove(row._tr)
    for index, unit in enumerate(project_rows, 1):
        ability_projects._tbl.append(deepcopy(ability_template))
        row = ability_projects.rows[-1]
        values = (
            index,
            f"{index}.1",
            unit["project_title"],
            "分析任务需求，完成页面或功能制作、运行调试和规范交付",
            "教材项目知识、关键技术与质量标准",
            "项目任务书、源代码、运行截图和评价清单",
            f"完成可运行成果（{unit['suggested_hours']}学时）",
        )
        for col, value in enumerate(values):
            if col < len(row.cells):
                _set_cell(row.cells[col], value)
    _set_table_font(ability_projects, size=9)
    progress = document.tables[3]
    progress_template = deepcopy(progress.rows[2]._tr)
    for row in list(progress.rows[2:]):
        _remove(row._tr)
    # 教学进度表是课程级内容，周次模式下仍展示全部任务
    progress_tasks = store.rows("SELECT * FROM tasks WHERE offering_id=? ORDER BY seq", (offering_id,)) if week_no is not None else tasks
    for task in progress_tasks:
        progress._tbl.append(deepcopy(progress_template))
        row = progress.rows[-1]
        task_detail = _task_detail(task)
        if len(row.cells) == 9:
            values = (task["seq"], task["hours"], task_detail, task["ability_goal"], task["seq"], task["knowledge_goal"],
                      f"{task['ideological_goal']}\n{task['quality_goal']}", "项目教学、任务驱动、演示练习", "阶段成果与评价清单")
        else:
            values = (task["seq"], task["hours"], task_detail, task["ability_goal"], task["seq"], task["knowledge_goal"],
                      task["ideological_goal"], task["quality_goal"], "项目教学、任务驱动、演示练习", "阶段成果与评价清单")
        for col, value in enumerate(values):
            if col < len(row.cells):
                _set_cell(row.cells[col], value)
    _set_table_font(progress, size=9)
    first_heading = next((p for p in document.paragraphs if p.text.strip() == "九、第一节课设计梗概"), None)
    if first_heading:
        outline = f"第一节课以《{offering['course_name']}》课程任务和学习成果导入，介绍课程内容、学习要求、平台资源和考核方式；围绕“{_task_detail(progress_tasks[0]) if progress_tasks else offering['course_name']}”开展案例提问、概念分析、操作演示和分步练习，学生完成首个可运行成果并依据检查清单进行自评。课堂强调规范命名、软件与数字资源版权、数据安全、诚信学习和按时提交，最后布置复习与完善任务。"
        _insert_after(first_heading, outline)
    cutoff = next(p for p in document.paragraphs if p.text.strip() in {"九、单元教学设计", "十、单元教学设计"})
    body = document._body._element
    reached = False
    for child in list(body):
        if child is cutoff._p:
            reached = True
            continue
        if reached and not child.tag.endswith("}sectPr"):
            body.remove(child)
    # Template style 2: tabular teaching-design scheme. It provides separate
    # columns for teacher/student activity, methods, outcomes and platform.
    basic_xml = deepcopy(Document(template).tables[-2]._tbl)
    org_xml = deepcopy(Document(template).tables[-1]._tbl)
    title_style = next((p for p in Document(template).paragraphs if p.text.strip() == "教学设计·基本信息"), None)
    design_rows = {key[1]: value for key, value in drafts.items() if key[0] == "unit_design"}
    for task_index, task in enumerate(tasks):
        task = dict(task)
        if task_index > 0:
            document.add_page_break()
        p1 = document.add_paragraph("教学设计·基本信息")
        if title_style and title_style._p.pPr is not None:
            p1._p.insert(0, deepcopy(title_style._p.pPr))
        document._body._element.insert(-1, deepcopy(basic_xml))
        basic = document.tables[-1]
        same_week = [s for s in sessions if s.get("week_no") == task.get("week_no") or s.get("lesson_date") == task.get("lesson_date")]
        classroom = next((s["classroom"] for s in same_week), "多媒体实训室")
        task["_class_label"] = "\n".join(dict.fromkeys(s.get("class_name", "") for s in same_week if s.get("class_name"))) or offering.get("teaching_class", "")
        task["_date_label"] = "\n".join(dict.fromkeys(s.get("lesson_date", "") for s in same_week if s.get("lesson_date"))) or task.get("lesson_date", "")
        _fill_design_table_basic(basic, task, offering, classroom)
        p2 = document.add_paragraph("教学设计·教学组织")
        if title_style and title_style._p.pPr is not None:
            p2._p.insert(0, deepcopy(title_style._p.pPr))
        document._body._element.insert(-1, deepcopy(org_xml))
        org = document.tables[-1]
        draft = design_rows[str(task["seq"])]
        refs = json.loads(task["resource_refs"] or "[]")
        _fill_design_table_org(org, task, draft["教学组织"], refs, offering_id=offering_id)
    for table in document.tables:
        seen = set()
        for row in table.rows:
            for cell in row.cells:
                marker = id(cell._tc)
                if marker in seen:
                    continue
                seen.add(marker)
                text = cell.text
                if "……" in text or "××" in text:
                    _set_cell(cell, text.replace("学习……\n理解……\n掌握……", "理解知识\n掌握方法\n完成操作").replace("……", "相关要求").replace("××分钟", "20分钟").replace("××", "线下"))
    # The cover fields use equal-width underlines with centered values.
    for label, value in (("课程名称：", offering["course_name"]), ("班    级：", offering["teaching_class"]), ("教    材：", offering["textbook_version"]), ("授课教师：", store.get_setting("teacher_name", "杜媛"))):
        paragraph = next((p for p in document.paragraphs if p.text.strip().startswith(label)), None)
        if paragraph:
            _set_underlined_field(paragraph, label, value, field_width=30)
    document.save(output)
    _replace_xml_text(output, {"《   》课程整体教学设计": f"《{offering['course_name']}》课程整体教学设计", "《  》课程整体教学设计": f"《{offering['course_name']}》课程整体教学设计"})
    _fill_xml_goal_blocks(output, {"素质目标：": goals.get("素质目标", [])[:4], "知识目标：": goals.get("知识目标", [])[:4], "能力目标：": goals.get("能力目标", [])[:4], "思政目标：": goals.get("思政目标", [])[:4]})
    _underline_labeled_paragraphs(output, ("课程名称：", "班    级：", "教    材：", "授课教师："))
    _replace_xml_text(output, {
        "           学年  第    学期": _term_label(offering["term"]),
        "课程名称：\u3000\u3000          \u3000\u3000": f"课程名称：{offering['course_name']}       ",
        "班    级：\u3000\u3000\u3000      \u3000\u3000\u3000": f"班    级：{offering['teaching_class']}       ",
        "教    材：\u3000\u3000         \u3000 \u3000": f"教    材：{offering['textbook_version']}       ",
        "授课教师：\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000": f"授课教师：{offering.get('teacher_name', '') or store.get_setting('teacher_name', '')}",
    })


def generate_documents(offering_id, output_dir, document_types=None, week_no=None):
    all_types = ("课程标准", "授课计划", "教学设计")
    if document_types:
        unknown = [dt for dt in document_types if dt not in all_types]
        if unknown:
            raise ValueError(f"不支持的文档类型：{'、'.join(unknown)}")
        selected = [dt for dt in all_types if dt in document_types]
        if not selected:
            raise ValueError("未选择任何要生成的文档。")
    else:
        selected = list(all_types)
    if week_no is not None and "教学设计" not in selected:
        raise ValueError("仅在选择生成教学设计时才能指定周次。")

    templates = {row["document_type"]: row["template_path"] for row in store.rows("SELECT document_type,template_path FROM template_files WHERE offering_id=?", (offering_id,))}
    offering = store.rows("SELECT * FROM offerings WHERE id=?", (offering_id,))[0]
    teacher_name = offering.get("teacher_name", "") or store.get_setting("teacher_name", "")
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # 检查模板文件是否存在
    for document_type in selected:
        template_path = templates.get(document_type)
        if not template_path:
            raise ValueError(f"缺少模板文件：{document_type}，请先上传对应模板。")
        if not Path(template_path).exists():
            raise ValueError(f"模板文件不存在：{document_type}，路径={template_path}")

    design_suffix = f" 第{week_no}周" if week_no is not None else ""
    names = {
        "课程标准": output_dir / f"{offering['term']}《{offering['course_name']}》课程标准 {teacher_name}.docx",
        "授课计划": output_dir / f"{offering['term']}《{offering['course_name']}》授课计划 {teacher_name}.docx",
        "教学设计": output_dir / f"{offering['term']}《{offering['course_name']}》课程教学设计{design_suffix} {teacher_name}.docx",
    }
    names = {dt: names[dt] for dt in selected}

    # 第一步：生成到临时文件，避免覆盖用户正在打开的Word文档
    import tempfile
    fill_stats = {}
    temp_paths = {}
    for document_type in selected:
        temp_fd, temp_name = tempfile.mkstemp(suffix=".docx", dir=str(output_dir))
        temp_path = Path(temp_name)
        os.close(temp_fd)
        try:
            filled, skipped = template_filler.fill_simple_fields(
                templates[document_type], temp_path, offering_id
            )
            fill_stats[document_type] = (filled, skipped)
        except Exception:
            fill_stats[document_type] = (0, -1)
            shutil.copy2(templates[document_type], temp_path)
        temp_paths[document_type] = temp_path

    # 第二步：填充复杂结构化内容到临时文件
    if "课程标准" in selected:
        _fill_standard(offering_id, temp_paths["课程标准"], temp_paths["课程标准"])
    if "授课计划" in selected:
        _fill_plan(offering_id, temp_paths["授课计划"], temp_paths["授课计划"])
    if "教学设计" in selected:
        _fill_design(offering_id, temp_paths["教学设计"], temp_paths["教学设计"], week_no=week_no)

    # 第三步：将临时文件移动到最终位置（若目标被Word占用则生成编号副本）
    for document_type, final_path in names.items():
        temp_file = temp_paths[document_type]
        try:
            if final_path.exists():
                final_path.unlink()
            temp_file.rename(final_path)
        except (PermissionError, OSError):
            candidate = final_path.with_name(f"{final_path.stem}（新生成）{final_path.suffix}")
            serial = 2
            while candidate.exists():
                candidate = final_path.with_name(f"{final_path.stem}（新生成{serial}）{final_path.suffix}")
                serial += 1
            temp_file.rename(candidate)
            names[document_type] = candidate
        except Exception:
            temp_file.unlink(missing_ok=True)
            raise

    with store.connect() as db:
        for document_type, path in names.items():
            # 周次模式的教学设计单独记录，不覆盖整本教学设计的生成记录
            record_type = f"教学设计（第{week_no}周）" if week_no is not None and document_type == "教学设计" else document_type
            template_id = store.rows("SELECT id FROM template_files WHERE offering_id=? AND document_type=?", (offering_id, document_type))[0]["id"]
            db.execute(
                "INSERT INTO generated_documents (offering_id,template_file_id,document_type,output_path,generation_status,structural_check,visual_check,notes,generated_at) "
                "VALUES (?,?,?,?,?,'待检查','待检查','由空白模板和当前原始资料重新生成',CURRENT_TIMESTAMP) "
                "ON CONFLICT(offering_id,document_type) DO UPDATE SET template_file_id=excluded.template_file_id,output_path=excluded.output_path,generation_status=excluded.generation_status,structural_check='待检查',visual_check='待检查',notes=excluded.notes,generated_at=CURRENT_TIMESTAMP",
                (offering_id, template_id, record_type, str(path.resolve()), "已生成"),
            )
        db.commit()
    return names


def generate_offering_documents(offering_id, document_types=None, week_no=None):
    offering = store.rows("SELECT * FROM offerings WHERE id=?", (offering_id,))[0]
    root = Path(store.get_setting("output_root", Path(__file__).parent / "生成结果"))
    output_dir = root / offering["term"] / offering["course_name"]

    if offering.get("offering_kind") == "实训课程":
        return generate_training_documents(offering_id, output_dir)

    issues = store.rows("SELECT severity,message FROM quality_issues WHERE offering_id=? AND severity='错误'", (offering_id,))
    if issues:
        raise ValueError("生成基础仍有错误，不能套版：" + "；".join(item["message"] for item in issues[:3]))
    return generate_documents(offering_id, output_dir, document_types=document_types, week_no=week_no)


def generate_training_documents(offering_id, output_dir):
    """实训课程：只生成实训资料文件（模板8）"""
    from training_materials import generate_training_materials, TEMPLATE as T8_TEMPLATE
    offering = store.rows("SELECT * FROM offerings WHERE id=?", (offering_id,))[0]

    if not T8_TEMPLATE.exists():
        raise ValueError(f"实训资料模板不存在：{T8_TEMPLATE}")

    # 确保有一条 template_files 记录指向模板8
    existing = store.rows(
        "SELECT id FROM template_files WHERE offering_id=? AND document_type='实训资料'",
        (offering_id,),
    )
    if existing:
        template_file_id = existing[0]["id"]
        with store.connect() as db:
            db.execute(
                "UPDATE template_files SET template_path=? WHERE id=?",
                (str(T8_TEMPLATE), template_file_id),
            )
            db.commit()
    else:
        with store.connect() as db:
            cursor = db.execute(
                "INSERT INTO template_files (offering_id,document_type,template_name,template_path,notes) VALUES (?,?,?,?,?)",
                (offering_id, "实训资料", "模板8：《XXX》实训资料.docx", str(T8_TEMPLATE), "实训资料模板（模板8）"),
            )
            template_file_id = cursor.lastrowid
            db.commit()

    output_path = generate_training_materials(offering_id, str(output_dir))
    final_path = Path(output_path)

    with store.connect() as db:
        db.execute(
            "DELETE FROM generated_documents WHERE offering_id=? AND document_type NOT IN ('实训资料')",
            (offering_id,),
        )
        db.execute(
            "INSERT INTO generated_documents (offering_id,template_file_id,document_type,output_path,generation_status,structural_check,visual_check,notes,generated_at) "
            "VALUES (?,?,?,?,'已生成','待检查','待检查','由实训资料模板（模板8）生成',CURRENT_TIMESTAMP) "
            "ON CONFLICT(offering_id,document_type) DO UPDATE SET template_file_id=excluded.template_file_id,output_path=excluded.output_path,generation_status=excluded.generation_status,structural_check='待检查',visual_check='待检查',notes=excluded.notes,generated_at=CURRENT_TIMESTAMP",
            (offering_id, template_file_id, "实训资料", str(final_path.resolve())),
        )
        db.commit()
    return {"实训资料": str(final_path.resolve())}

# Unified full-document entry: UI and CLI use the enhanced generators.
_legacy_generate_offering_documents = generate_offering_documents


def generate_offering_documents(offering_id, document_types=None, week_no=None):
    offering = store.rows("SELECT * FROM offerings WHERE id=?", (offering_id,))[0]
    if offering.get("offering_kind") == "实训课程" or week_no is not None:
        return _legacy_generate_offering_documents(
            offering_id, document_types=document_types, week_no=week_no
        )

    issues = store.rows(
        "SELECT severity,message FROM quality_issues WHERE offering_id=? AND severity='错误'",
        (offering_id,),
    )
    if issues:
        raise ValueError("生成基础仍有错误，不能套版：" + "；".join(item["message"] for item in issues[:3]))

    type_map = {"课程标准": "standard", "授课计划": "plan", "教学设计": "design"}
    selected = document_types or list(type_map)
    unknown = [item for item in selected if item not in type_map]
    if unknown:
        raise ValueError(f"不支持的文档类型：{'、'.join(unknown)}")

    root = Path(store.get_setting("output_root", Path(__file__).parent / "生成结果"))
    output_dir = root / offering["term"] / offering["course_name"]
    from generate import generate_all
    generated = generate_all(
        offering_id,
        [type_map[item] for item in selected],
        str(output_dir),
    )

    reverse_map = {value: key for key, value in type_map.items()}
    names = {}
    errors = []
    for short_type, result in generated.items():
        document_type = reverse_map[short_type]
        if result.get("error"):
            errors.append(f"{document_type}：{result['error']}")
        elif result.get("path"):
            names[document_type] = Path(result["path"])
    if errors:
        raise ValueError("；".join(errors))

    with store.connect() as db:
        for document_type, path in names.items():
            template = db.execute(
                "SELECT id FROM template_files WHERE offering_id=? AND document_type=? ORDER BY id LIMIT 1",
                (offering_id, document_type),
            ).fetchone()
            if not template:
                raise ValueError(f"缺少模板登记：{document_type}")
            db.execute(
                "INSERT INTO generated_documents "
                "(offering_id,template_file_id,document_type,output_path,generation_status,structural_check,visual_check,notes,generated_at) "
                "VALUES (?,?,?,?,?,'已检查','待检查','界面与命令行统一增强生成器生成',CURRENT_TIMESTAMP) "
                "ON CONFLICT(offering_id,document_type) DO UPDATE SET "
                "template_file_id=excluded.template_file_id,output_path=excluded.output_path,"
                "generation_status=excluded.generation_status,structural_check=excluded.structural_check,"
                "visual_check=excluded.visual_check,notes=excluded.notes,generated_at=CURRENT_TIMESTAMP",
                (offering_id, template["id"], document_type, str(path.resolve()), "已生成"),
            )
        db.commit()
    return names
