import store
store.initialize()

offerings = store.rows("SELECT id, course_name, course_type, course_nature, offering_kind, term FROM offerings ORDER BY id")
print("=== offerings ===")
for o in offerings:
    print(f'  {o["id"]}: {o["course_name"]} | type={o["course_type"]} | nature={o["course_nature"]} | kind={o["offering_kind"]} | term={o["term"]}')

print()
print("=== template_files document_type values ===")
types = store.rows("SELECT DISTINCT document_type FROM template_files ORDER BY document_type")
for t in types:
    print(f'  {t["document_type"]}')

print()
print("=== Check template8 path ===")
import os
t8 = r"e:\开发\AIGC\教学档案工作台\原始资料\模板\模板8：《XXX》实训资料.docx"
print(f"  Exists: {os.path.exists(t8)}")

from docx import Document
doc = Document(t8)
print(f"  Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
for i, p in enumerate(doc.paragraphs[:20]):
    if p.text.strip():
        print(f"  P{i}: {p.text.strip()[:80]}")
print()
for ti, table in enumerate(doc.tables):
    print(f"  Table {ti}: {len(table.rows)}x{len(table.columns)}")
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip()[:30] for c in row.cells]
        print(f"    Row {ri}: {' | '.join(cells)}")
