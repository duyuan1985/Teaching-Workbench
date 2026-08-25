"""
修复表2（职业能力训练项目设计）：
1. 相关支撑知识：每个子目标展开写多个①②③，每个不一样
2. 训练方式手段及步骤：每个子目标展开写多个①②③，每个不一样
3. 结果（可展示）：每个子目标不一样
4. 合并第一列（编号）和第二列（能力训练项目名称）
"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from copy import deepcopy

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)
t2 = doc.tables[2]

# 每个子目标的详细内容（基于PPT和源码）
capacity_data = {
    # 单元1：初识数据分析
    ("1", "1.1 认识数据分析"): {
        "knowledge": "①数据分析的基本概念、内涵与分类\n②描述性分析、诊断性分析、预测性分析与指导性分析的区别\n③数据分析在电商行业的应用场景与价值\n④数据分析的基本流程：明确目标→数据采集→数据清洗→数据分析→数据可视化→报告撰写",
        "method": "①案例教学法：展示农产品电商平台用户行为分析案例\n②讨论法：学生分组讨论数据分析在电商中的应用场景\n③讲授法：讲解数据分析基本概念和分类\n④任务驱动法：布置电商数据分析需求分析任务",
        "result": "完成电商数据分析需求分析报告1份，包含分析目标、数据来源、分析方法选择"
    },
    ("1", "1.2 数据分析指标"): {
        "knowledge": "①电商核心数据分析指标体系：GMV、转化率、客单价、复购率、退货率\n②流量指标：UV、PV、跳失率、停留时长\n③用户行为指标：点击率、加购率、收藏率\n④商品分析指标：动销率、售罄率、毛利率",
        "method": "①案例分析法：分析某电商平台真实指标数据\n②练习法：学生计算给定数据集的各指标值\n③小组讨论法：讨论各指标间的关系和影响因素\n④演示法：教师演示Excel数据透视表计算指标",
        "result": "提交电商数据分析指标计算表1份，含8个核心指标的计算过程和结果"
    },
    ("1", "1.3 数据分析方法理论"): {
        "knowledge": "①对比分析法：同比、环比、定基比的概念与计算\n②分组分析法：按维度分组统计\n③交叉分析法：多维度交叉分析\n④时间序列分析法：趋势分析、季节性分析\n⑤漏斗分析法：用户转化路径分析",
        "method": "①讲授法：讲解各分析方法的理论基础\n②案例演示法：用电商数据演示每种方法\n③实操练习法：学生用Excel练习各分析方法\n④小组互评法：各组交换分析结果互相点评",
        "result": "完成电商数据多维度分析报告1份，至少运用3种分析方法"
    },
    # 单元2：Excel数据分析工具
    ("2", "2.1 Excel概述"): {
        "knowledge": "①Excel在数据分析中的定位与优势\n②数据分析常用功能模块：数据透视表、图表、函数\n③Excel数据分析插件安装与配置\n④数据导入与格式化处理",
        "method": "①演示法：教师演示Excel数据分析功能\n②练习法：学生安装数据分析插件并导入数据\n③任务驱动法：完成电商数据导入和格式化任务\n④个别指导法：教师巡回指导操作问题",
        "result": "完成Excel数据分析环境配置，成功导入电商数据集并格式化"
    },
    ("2", "2.2 函数相关概念"): {
        "knowledge": "①常用统计函数：SUM、AVERAGE、COUNT、MAX、MIN\n②条件函数：SUMIF、COUNTIF、AVERAGEIF\n③查找函数：VLOOKUP、HLOOKUP、INDEX、MATCH\n④文本函数：LEFT、RIGHT、MID、CONCATENATE",
        "method": "①讲授法：讲解各函数语法和参数\n②演示法：用电商数据演示函数应用\n③练习法：学生完成10道函数练习题\n④互评法：同桌交换检查函数结果",
        "result": "提交Excel函数练习作业1份，含统计、条件、查找、文本4类函数各3题"
    },
    ("2", "2.3 Excel函数"): {
        "knowledge": "①数据透视表创建与布局设置\n②数据透视图制作\n③切片器与时间轴使用\n④数据透视表分组与计算字段",
        "method": "①项目驱动法：以电商销售数据为项目\n②演示法：演示数据透视表全流程操作\n③实操练习法：学生独立完成数据透视表\n④成果展示法：学生展示分析结果并讲解",
        "result": "完成电商销售数据透视分析报告1份，含3个维度的交叉分析和数据透视图"
    },
    # 单元3：Numpy数学运算库
    ("3", "3.1 位运算函数"): {
        "knowledge": "①NumPy库概述与安装（pip install numpy）\n②ndarray数组对象：创建、属性（shape、dtype、ndim）\n③位运算函数：bitwise_and、bitwise_or、bitwise_xor、bitwise_not\n④位运算在数据加密和权限控制中的应用",
        "method": "①环境搭建法：学生安装Anaconda和Jupyter\n②讲授法：讲解ndarray和位运算原理\n③演示法：演示位运算函数代码和运行结果\n④练习法：学生编写位运算代码练习",
        "result": "提交Jupyter Notebook文件1份，含NumPy数组创建和位运算5个代码单元格"
    },
    ("3", "3.2 取反"): {
        "knowledge": "① bitwise_not取反运算原理与二进制表示\n②逻辑取反与按位取反的区别\n③取反运算在数据掩码中的应用\n④数组元素条件筛选与布尔索引",
        "method": "①讲授法：讲解取反运算二进制原理\n②演示法：演示取反运算和掩码应用代码\n③练习法：学生编写数据掩码筛选代码\n④小组讨论法：讨论取反运算的应用场景",
        "result": "完成数据掩码筛选练习代码，能正确筛选满足条件的数据元素"
    },
    ("3", "3.3 数学函数"): {
        "knowledge": "①统计函数：np.mean、np.median、np.std、np.var、np.percentile\n②三角函数：np.sin、np.cos、np.tan\n③指数对数函数：np.exp、np.log、np.log10\n④数组运算：广播机制、矩阵运算（np.dot、np.matmul）",
        "method": "①讲授法：讲解各数学函数的用途和参数\n②演示法：用电商数据演示统计函数计算\n③练习法：学生完成电商数据统计指标计算\n④项目法：用数学函数实现数据标准化",
        "result": "提交电商数据统计分析代码1份，计算均值、标准差、百分位数等5项指标"
    },
    # 单元4：Pandas数据分析库
    ("4", "4.1 统计函数"): {
        "knowledge": "①Pandas库概述与DataFrame数据结构\n②数据读取：pd.read_csv、pd.read_excel、pd.read_sql\n③统计函数：df.describe()、df.sum()、df.mean()、df.count()\n④分组统计：df.groupby()与agg()聚合",
        "method": "①环境搭建法：确认pandas安装和导入\n②演示法：演示数据读取和统计函数\n③练习法：学生完成电商数据读取和统计分析\n④对比法：对比Pandas与Excel统计的效率",
        "result": "完成电商数据集的Pandas统计分析，输出describe()统计摘要和分组统计结果"
    },
    ("4", "4.2 标准化数据"): {
        "knowledge": "①数据标准化概念：Min-Max标准化、Z-score标准化\n②Min-Max标准化公式与实现：(x-min)/(max-min)\n③Z-score标准化公式与实现：(x-mean)/std\n④标准化在机器学习预处理中的作用",
        "method": "①讲授法：讲解标准化原理和公式\n②演示法：用Pandas实现两种标准化\n③练习法：学生对电商数据进行标准化处理\n④讨论法：讨论不同标准化方法的适用场景",
        "result": "提交标准化处理代码和结果对比表，含Min-Max和Z-score两种方法的输出"
    },
    ("4", "4.3 标准差标准化数据"): {
        "knowledge": "①标准差标准化的数学原理与公式推导\n②Pandas实现标准差标准化：df.transform()\n③异常值检测与处理：3σ原则、IQR方法\n④数据清洗：缺失值填充df.fillna()、重复值删除df.drop_duplicates()",
        "method": "①项目驱动法：以电商数据清洗为项目\n②演示法：演示异常值检测和清洗流程\n③练习法：学生完成完整数据清洗流程\n④互评法：学生交换检查清洗结果的正确性",
        "result": "提交数据清洗报告1份，含异常值检测、缺失值处理、重复值删除的代码和结果"
    },
    # 单元5：SciPy科学计算库
    ("5", "5.1 SciPy简介与安装"): {
        "knowledge": "①SciPy库概述与子模块介绍（cluster、stats、optimize、integrate）\n②SciPy安装与导入：pip install scipy\n③SciPy与NumPy的关系\n④SciPy在科学计算和数据分析中的应用场景",
        "method": "①环境搭建法：安装SciPy并验证导入\n②讲授法：讲解SciPy子模块功能\n③演示法：演示SciPy基本用法\n④练习法：学生完成SciPy安装和基本操作",
        "result": "成功安装SciPy并完成基本导入测试，提交环境配置截图和测试代码"
    },
    ("5", "5.2 cluster模块"): {
        "knowledge": "①scipy.cluster模块概述与层次聚类\n②hierarchy.linkage()层次聚类链接函数\n③hierarchy.dendrogram()树状图绘制\n④fcluster()聚类结果切割与标签获取",
        "method": "①讲授法：讲解层次聚类原理\n②演示法：演示电商用户聚类代码\n③练习法：学生对电商用户行为数据进行聚类\n④可视化法：绘制树状图和聚类结果散点图",
        "result": "提交电商用户聚类分析代码1份，含树状图和3类聚类标签结果"
    },
    ("5", "5.3 SciPy常用模块"): {
        "knowledge": "①scipy.stats：概率分布、假设检验（ttest_ind、chi2_contingency）\n②scipy.optimize：函数优化与曲线拟合（curve_fit）\n③scipy.integrate：数值积分\n④scipy.spatial：距离计算（distance.euclidean、distance.cdist）",
        "method": "①讲授法：讲解各模块功能和应用场景\n②演示法：用电商数据演示假设检验和曲线拟合\n③练习法：学生完成A/B测试假设检验\n④讨论法：讨论统计检验在电商决策中的应用",
        "result": "完成电商A/B测试假设检验代码1份，输出t统计量、p值和检验结论"
    },
    # 单元6：Sklearn数据统计基础
    ("6", "6.1 sklearn简介及安装"): {
        "knowledge": "①scikit-learn库概述与机器学习流程\n②sklearn安装与导入：pip install scikit-learn\n③sklearn数据集：load_iris、load_boston、load_digits\n④sklearn基本API：fit()、predict()、transform()",
        "method": "①环境搭建法：安装scikit-learn并验证\n②讲授法：讲解机器学习基本流程\n③演示法：演示iris数据集加载和基本操作\n④练习法：学生完成sklearn环境搭建",
        "result": "成功安装sklearn，完成iris数据集加载和基本操作测试代码"
    },
    ("6", "6.2 sklearn安装"): {
        "knowledge": "①sklearn依赖库管理：NumPy、SciPy、joblib、threadpoolctl\n②sklearn模块结构：datasets、model_selection、preprocessing、metrics\n③模型保存与加载：joblib.dump()、joblib.load()\n④超参数概念与网格搜索GridSearchCV",
        "method": "①讲授法：讲解sklearn模块结构和依赖关系\n②演示法：演示模型训练、保存和加载流程\n③练习法：学生完成模型训练和保存\n④对比法：对比不同超参数的模型效果",
        "result": "提交模型训练、保存、加载的完整代码1份，含超参数调优记录"
    },
    ("6", "6.3 数据集"): {
        "knowledge": "①train_test_split：训练集与测试集划分\n②cross_val_score：交叉验证\n③StandardScaler：数据标准化预处理\n④LabelEncoder：标签编码与OneHotEncoder独热编码",
        "method": "①讲授法：讲解数据集划分和预处理原理\n②演示法：演示电商数据集划分和标准化\n③练习法：学生对电商数据进行预处理\n④讨论法：讨论不同划分比例对模型的影响",
        "result": "完成电商数据集的预处理Pipeline代码，含划分、标准化、编码3个步骤"
    },
    # 单元7：Sklearn数据统计进阶
    ("7", "7.1 分类模型"): {
        "knowledge": "①逻辑回归LogisticRegression：二分类与多分类\n②决策树DecisionTreeClassifier：信息增益与基尼系数\n③随机森林RandomForestClassifier：集成学习\n④分类模型评估：accuracy_score、precision、recall、f1_score",
        "method": "①讲授法：讲解各分类算法原理\n②演示法：用电商用户购买预测数据演示\n③练习法：学生训练3种分类模型并比较效果\n④评估法：用混淆矩阵和分类报告评估",
        "result": "提交3种分类模型的训练代码和评估报告，含准确率、精确率、召回率对比表"
    },
    ("7", "7.2 回归模型"): {
        "knowledge": "①线性回归LinearRegression：最小二乘法\n②岭回归Ridge：L2正则化\n③Lasso回归：L1正则化与特征选择\n④回归模型评估：MSE、RMSE、R2_score",
        "method": "①讲授法：讲解回归分析原理和正则化\n②演示法：用电商销售额预测数据演示\n③练习法：学生训练3种回归模型并比较\n④可视化法：绘制回归拟合曲线和残差图",
        "result": "提交3种回归模型的训练代码和评估报告，含MSE和R2对比及拟合曲线图"
    },
    ("7", "7.3 聚类模型"): {
        "knowledge": "①KMeans聚类：K值选择与肘部法则\n②DBSCAN密度聚类：eps与min_samples参数\n③层次聚类AgglomerativeClustering\n④聚类评估：silhouette_score轮廓系数、inertia_",
        "method": "①讲授法：讲解各聚类算法原理和适用场景\n②演示法：用电商用户分群数据演示\n③练习法：学生对电商用户进行聚类分析\n④可视化法：绘制聚类结果散点图和肘部图",
        "result": "提交电商用户聚类分析报告1份，含3种算法对比和最优K值选择过程"
    },
    # 单元8：Seaborn可视化分析库
    ("8", "8.1 seaborn概述"): {
        "knowledge": "①Seaborn库概述与Matplotlib的关系\n②Seaborn安装与导入：pip install seaborn\n③Seaborn主题与样式设置：set_theme()、set_style()\n④Seaborn调色板：color_palette()、set_palette()",
        "method": "①环境搭建法：安装Seaborn并验证\n②讲授法：讲解Seaborn与Matplotlib区别\n③演示法：演示不同主题和调色板效果\n④练习法：学生设置不同风格绘制图表",
        "result": "成功安装Seaborn，完成3种主题和3种调色板的图表样式设置代码"
    },
    ("8", "8.2 关系图"): {
        "knowledge": "①relplot()关系图：散点图与线图\n②scatterplot()散点图：hue、size、style参数\n③lineplot()线图：置信区间与聚合\n④regplot()回归图：拟合线与置信区间",
        "method": "①讲授法：讲解各关系图函数的参数和用法\n②演示法：用电商销售数据演示4种关系图\n③练习法：学生用电商数据绘制4种关系图\n④美化法：调整图表标题、坐标轴、图例样式",
        "result": "提交电商数据关系图可视化代码1份，含4种图表类型的展示效果"
    },
    ("8", "8.3 分类图"): {
        "knowledge": "①catplot()分类图：kind参数选择\n②boxplot()箱线图：数据分布与异常值展示\n③violinplot()小提琴图：密度分布\n④barplot()与countplot()：分类统计与计数",
        "method": "①讲授法：讲解各分类图的适用场景\n②演示法：用电商商品分类数据演示4种分类图\n③练习法：学生对电商数据进行分类可视化\n④互评法：学生展示图表并互相点评美观度",
        "result": "提交电商分类数据可视化报告1份，含4种分类图表和分析说明"
    },
    # 单元9：综合评价与课程总结
    ("9", "9.1 课程成果汇报与评价"): {
        "knowledge": "①数据分析项目完整流程回顾：数据采集→清洗→分析→建模→可视化→报告\n②综合项目成果评价指标：完整性、准确性、创新性、规范性\n③项目答辩PPT制作要点与演讲技巧\n④同行评价与教师评价标准",
        "method": "①成果展示法：各组展示综合项目成果\n②答辩法：学生进行项目答辩并回答提问\n③互评法：各组之间按标准互相评分\n④教师点评法：教师对每组项目进行综合点评",
        "result": "提交综合项目报告1份、答辩PPT 1份，完成项目答辩和互评"
    },
    ("9", "9.2 课程总结与复习"): {
        "knowledge": "①课程知识点体系梳理：数据分析基础→Excel→Python(Numpy/Pandas/SciPy)→机器学习(Sklearn)→可视化(Seaborn)\n②重点难点回顾：数据清洗、特征工程、模型选择与调优\n③行业新技术展望：大语言模型辅助分析、AIGC图表生成、自动洞察\n④课程证书与竞赛对接：1+X证书备考要点、技能大赛要点",
        "method": "①总结归纳法：教师梳理课程知识体系\n②小组讨论法：学生讨论重点难点和易错点\n③模拟测试法：进行课程综合模拟测试\n④反馈法：收集学生课程学习反馈和改进建议",
        "result": "完成课程知识体系思维导图1份、综合模拟测试卷1份"
    },
}

# 填充表2
print("填充表2...")
filled = 0
for ri in range(1, len(t2.rows)):
    cells = t2.rows[ri].cells
    if len(cells) < 6:
        continue
    
    seq = cells[0].text.strip()
    sub_goal = cells[2].text.strip()
    
    key = (seq, sub_goal)
    if key in capacity_data:
        data = capacity_data[key]
        
        # col3: 相关支撑知识
        for p in cells[3].paragraphs:
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = data['knowledge']
            else:
                p.add_run(data['knowledge'])
        
        # col4: 训练方式手段及步骤
        for p in cells[4].paragraphs:
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = data['method']
            else:
                p.add_run(data['method'])
        
        # col5: 结果（可展示）
        for p in cells[5].paragraphs:
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = data['result']
            else:
                p.add_run(data['result'])
        
        filled += 1
    else:
        print(f"  未找到匹配: {key}")

print(f"  填充了{filled}/{len(t2.rows)-1}行")

# 合并第一列和第二列（相同编号的行合并）
print("合并单元格...")
# 收集需要合并的行范围
merge_ranges = []
current_seq = None
start_row = 1
for ri in range(1, len(t2.rows)):
    seq = t2.rows[ri].cells[0].text.strip()
    if seq != current_seq:
        if current_seq is not None and start_row < ri - 1:
            merge_ranges.append((start_row, ri - 1, current_seq))
        current_seq = seq
        start_row = ri
# 最后一组
if current_seq is not None and start_row < len(t2.rows) - 1:
    merge_ranges.append((start_row, len(t2.rows) - 1, current_seq))

print(f"  合并范围: {merge_ranges}")

# 执行合并
for start, end, seq in merge_ranges:
    # 合并col0
    cell_a = t2.cell(start, 0)
    cell_b = t2.cell(end, 0)
    merged = cell_a.merge(cell_b)
    # 合并col1
    cell_a = t2.cell(start, 1)
    cell_b = t2.cell(end, 1)
    merged = cell_a.merge(cell_b)

print("  合并完成")

doc.save(fp)
print(f"\n保存完成: {fp}")
