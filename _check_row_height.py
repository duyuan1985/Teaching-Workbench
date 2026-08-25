from docx import Document
from docx.oxml.ns import qn

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

# 检查Table 6的R3-R7行高设置
t = doc.tables[6]
for ri in range(3, 8):
    tr = t.rows[ri]._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is not None:
        trHeight = trPr.find(qn('w:trHeight'))
        if trHeight is not None:
            val = trHeight.get(qn('w:val'))
            hRule = trHeight.get(qn('w:hRule'))
            print(f'R{ri}: trHeight val={val} hRule={hRule}')
        else:
            print(f'R{ri}: 无trHeight')
    else:
        print(f'R{ri}: 无trPr')

# 检查模板中的行高设置
tmpl_fp = r'原始资料\模板\模板5：教学设计\模板5：教学设计 模板（2023-2024）.docx'
tmpl_doc = Document(tmpl_fp)
tmpl_t = tmpl_doc.tables[6]
print('\n=== 模板 ===')
for ri in range(3, 8):
    tr = tmpl_t.rows[ri]._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is not None:
        trHeight = trPr.find(qn('w:trHeight'))
        if trHeight is not None:
            val = trHeight.get(qn('w:val'))
            hRule = trHeight.get(qn('w:hRule'))
            print(f'R{ri}: trHeight val={val} hRule={hRule}')
        else:
            print(f'R{ri}: 无trHeight')
    else:
        print(f'R{ri}: 无trPr')
