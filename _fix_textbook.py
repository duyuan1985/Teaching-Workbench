"""
统一教学设计和课程标准中的教材描述
统一表述：《大数据分析方法项目实战》，天津大学出版社。由天津滨海迅腾科技集团有限公司编写，天津大学出版社出版。该教材为十三五职业教育国家规划教材，以项目引领、任务驱动方式组织教学内容，内容涵盖数据采集、清洗、分析、可视化全流程，配有丰富的电商案例和实训项目，符合高职高专学生的认知规律和理实一体化教学要求。
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 统一后的教材描述
unified_textbook_desc = '《大数据分析方法项目实战》，天津大学出版社。由天津滨海迅腾科技集团有限公司编写，天津大学出版社出版。该教材为十三五职业教育国家规划教材，以项目引领、任务驱动方式组织教学内容，内容涵盖数据采集、清洗、分析、可视化全流程，配有丰富的电商案例和实训项目，符合高职高专学生的认知规律和理实一体化教学要求。'

# ============================================================
# 1. 修改教学设计 P98
# ============================================================
fp_d = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc_d = Document(fp_d)

p98 = doc_d.paragraphs[98]
print(f'教学设计 修改前: {p98.text.strip()[:100]}')

# 获取格式
font_name = '仿宋_GB2312'
font_size_val = '21'
if p98.runs:
    r = p98.runs[0]
    if r.font.name:
        font_name = r.font.name
    if r.font.size:
        font_size_val = str(int(r.font.size.pt * 2))

# 清除原有run
for r in p98.runs:
    r._element.getparent().remove(r._element)

# 添加新run
run = p98.add_run(unified_textbook_desc)
run.font.name = font_name
run.font.size = None  # 用段落默认
run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

doc_d.save(fp_d)
print(f'教学设计 修改后: {p98.text.strip()[:100]}')

# ============================================================
# 2. 修改课程标准 P48
# ============================================================
fp_s = r'生成结果\精修版\2023-2024-2《商务数据分析》课程标准 杜媛.docx'
doc_s = Document(fp_s)

p48 = doc_s.paragraphs[48]
print(f'\n课程标准 修改前: {p48.text.strip()[:100]}')

# 获取格式
font_name_s = '仿宋'
font_size_val_s = '21'
if p48.runs:
    r = p48.runs[0]
    if r.font.name:
        font_name_s = r.font.name
    if r.font.size:
        font_size_val_s = str(int(r.font.size.pt * 2))

# 课程标准P48包含两段内容（教材选用+教材编写要求），需要只替换第一段
# 拆分文本
old_text = p48.text
# 找到"（2）教材编写要求"的位置
split_marker = '（2）教材编写要求'
if split_marker in old_text:
    part1_old = old_text[:old_text.index(split_marker)].strip()
    part2 = old_text[old_text.index(split_marker):]
    print(f'  原第一段: {part1_old[:100]}')
    print(f'  原第二段: {part2[:100]}')
    
    # 新的第一段
    new_part1 = f'（1）教材选用：本课程选用{unified_textbook_desc}'
    new_full_text = new_part1 + '\n' + part2
else:
    new_full_text = f'（1）教材选用：本课程选用{unified_textbook_desc}'

# 清除原有run
for r in p48.runs:
    r._element.getparent().remove(r._element)

# 添加新run（可能包含换行）
lines = new_full_text.split('\n')
for li, line in enumerate(lines):
    if li > 0:
        # 用换行符
        pass  # python-docx add_run 不支持\n，需要用段落
    run = p48.add_run(line)
    run.font.name = font_name_s
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_s)

doc_s.save(fp_s)
print(f'\n课程标准 修改后: {p48.text.strip()[:200]}')

# ============================================================
# 验证
# ============================================================
print('\n=== 验证 ===')
doc_d2 = Document(fp_d)
for pi, p in enumerate(doc_d2.paragraphs):
    txt = p.text.strip()
    if '大数据分析方法' in txt and '天津大学' in txt:
        print(f'教学设计 P{pi}: {txt[:300]}')

doc_s2 = Document(fp_s)
for pi, p in enumerate(doc_s2.paragraphs):
    txt = p.text.strip()
    if '教材选用' in txt:
        print(f'课程标准 P{pi}: {txt[:300]}')
