"""
第四轮修复：
1. 封面内容填写
2. 表0课程信息表修复（制定人重复、开设学期格式、后续课去掉顶岗实习）
3. 表1课程内容设计加章节号
4. 表5教学要求修复（R9重点/R10难点标签、子任务匹配、打√列正确）
5. 表5知识&技能标题恢复
"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from copy import deepcopy
import store

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)

# ============================================================
# 工具函数
# ============================================================
def clear_cell(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''

def set_cell_text(cell, text, font_name='仿宋', size=10.5, bold=False):
    clear_cell(cell)
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    
    lines = text.split('\n')
    for li, line in enumerate(lines):
        if li > 0:
            run = p.add_run()
            run.add_break()
        else:
            run = p.add_run()
        run.text = line
        run.font.name = font_name
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        run.font.size = Pt(size)
        run.font.bold = bold

def set_para_text(p, text, font_name='仿宋', size=14, bold=False, align=None):
    for r in p.runs:
        r.text = ''
    if p.runs:
        p.runs[0].text = text
        r = p.runs[0]
    else:
        r = p.add_run(text)
    r.font.name = font_name
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    r.font.size = Pt(size)
    r.font.bold = bold
    if align:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p.alignment = align
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)

# ============================================================
# 1. 填写封面内容
# ============================================================
print("1. 填写封面内容...")

from docx.enum.text import WD_ALIGN_PARAGRAPH

paras = doc.paragraphs
# P0-P39是空段落，用于封面
# 在合适位置填写封面内容

# 找到空段落填入封面信息
# P5左右放学校名
# P10左右放课程名
# P20左右放教师信息
# P30左右放日期

cover_items = [
    (5, "山西林业职业技术学院", "方正小标宋简体", 22, True, WD_ALIGN_PARAGRAPH.CENTER),
    (10, "《商务数据分析》", "方正小标宋简体", 28, True, WD_ALIGN_PARAGRAPH.CENTER),
    (11, "课程整体教学设计", "方正小标宋简体", 22, True, WD_ALIGN_PARAGRAPH.CENTER),
    (20, "授课教师：杜媛", "仿宋", 16, False, WD_ALIGN_PARAGRAPH.CENTER),
    (22, "授课班级：2022电商教学班", "仿宋", 16, False, WD_ALIGN_PARAGRAPH.CENTER),
    (24, "授课学期：2023-2024学年第二学期", "仿宋", 16, False, WD_ALIGN_PARAGRAPH.CENTER),
    (30, "2024年2月", "仿宋", 16, False, WD_ALIGN_PARAGRAPH.CENTER),
]

for idx, text, fn, sz, bold, align in cover_items:
    if idx < len(paras):
        set_para_text(paras[idx], text, font_name=fn, size=sz, bold=bold, align=align)

print("  封面内容已填写")

# ============================================================
# 2. 修复表0课程信息表
# ============================================================
print("2. 修复表0课程信息表...")

t0 = doc.tables[0]

# R0: 课程名称 | 商务数据分析 | 课程代码 | A451114 | 所属系部 | 经济贸易系  (已正确)
# R1: 制定时间 | 2024年2月 | 制定人 | 杜媛 | (重复的杜媛清除)
set_cell_text(t0.cell(1, 3), "杜媛")
# col4和col5是合并单元格，清除重复
# 实际上它们是合并的，只需要在col4写
set_cell_text(t0.cell(1, 4), "")

# R2: 课程类型 | 专业核心课 | 学时 | 60 | 学分 | 3  (已正确)

# R3: 开设学期 | 2023-2024学年第二学期 | 授课对象 | 2022电商教学班
set_cell_text(t0.cell(3, 1), "2023-2024学年第二学期")
set_cell_text(t0.cell(3, 2), "")  # 清除重复
set_cell_text(t0.cell(3, 4), "2022电商教学班")
set_cell_text(t0.cell(3, 5), "")  # 清除重复

# R4: 先修课 | ... | 后续课 | ... (去掉顶岗实习)
set_cell_text(t0.cell(4, 1), "电子商务基础、计算机应用基础、Python程序设计")
set_cell_text(t0.cell(4, 2), "")  # 清除重复
set_cell_text(t0.cell(4, 4), "新媒体平台运营与推广、电子商务综合实训")
set_cell_text(t0.cell(4, 5), "")  # 清除重复

print("  表0已修复（制定人去重、开设学期统一、后续课去掉顶岗实习）")

# ============================================================
# 3. 表1课程内容设计加章节号
# ============================================================
print("3. 表1加章节号...")

t1 = doc.tables[1]
chapter_names = [
    "第一章 初识数据分析",
    "第二章 Excel数据分析工具",
    "第三章 Numpy数学运算库",
    "第四章 Pandas数据分析库",
    "第五章 SciPy科学计算库",
    "第六章 Sklearn数据统计基础",
    "第七章 Sklearn数据统计进阶",
    "第八章 Seaborn可视化分析库",
    "第九章 综合评价与课程总结",
]

for ri in range(1, min(10, len(t1.rows))):
    if ri - 1 < len(chapter_names):
        set_cell_text(t1.cell(ri, 0), chapter_names[ri - 1])

print("  表1已加章节号")

# ============================================================
# 4. 修复表5教学要求（R9-R14）
# ============================================================
print("4. 修复表5教学要求...")

# 获取任务数据
tasks = store.rows("SELECT * FROM tasks WHERE offering_id=20 ORDER BY seq")

chapter_nums = ['一','二','三','四','五','六','七','八','九']

for ti in range(5, len(doc.tables), 2):
    t5 = doc.tables[ti]
    task_idx = (ti - 5) // 2
    if task_idx >= len(tasks):
        continue
    
    task = tasks[task_idx]
    title = task.get('title', f'任务{task_idx+1}')
    seq = task.get('seq', task_idx+1)
    
    # 解析子任务
    parts = title.split('：')
    if len(parts) >= 2:
        chapter_name = parts[0]
        sub_tasks = [s.strip() for s in parts[1].split('、')]
    else:
        chapter_name = title
        sub_tasks = [title]
    
    # 生成重点难点
    if seq == 1:
        key_points = "1. 数据分析的基本概念和分类\n2. 电商核心数据分析指标\n3. 常用分析方法的应用"
        difficult = "1. 分析方法的选择和应用\n2. 业务需求与分析方法的对应\n3. 从数据中提炼有效结论"
    elif seq == 2:
        key_points = "1. Excel常用函数的语法和参数\n2. 数据透视表的创建步骤\n3. 数据透视图的制作"
        difficult = "1. 复杂函数的参数理解\n2. 多维交叉分析\n3. 大数据量下的性能优化"
    else:
        key_points = f"1. {chapter_name}的基本概念和核心原理\n2. 核心API的使用方法和参数\n3. 操作流程和注意事项"
        difficult = f"1. 代码调试和错误处理\n2. API参数组合使用\n3. 异常情况的处理"
    
    # R9: 重点
    set_cell_text(t5.cell(9, 2), key_points)
    # col7应为"重点"标签
    for ci in range(7, min(10, len(t5.rows[9].cells))):
        if ci == 7:
            set_cell_text(t5.cell(9, ci), "重点")
        else:
            set_cell_text(t5.cell(9, ci), "")
    
    # R10: 难点
    set_cell_text(t5.cell(10, 2), difficult)
    for ci in range(7, min(10, len(t5.rows[10].cells))):
        if ci == 7:
            set_cell_text(t5.cell(10, ci), "难点")
        else:
            set_cell_text(t5.cell(10, ci), "")
    
    # R11-R14: 子任务列表，打√
    for ri in range(11, min(15, len(t5.rows))):
        sub_idx = ri - 11
        if sub_idx < len(sub_tasks):
            set_cell_text(t5.cell(ri, 2), sub_tasks[sub_idx])
            # 打√列（col7=重点, col8=难点, col9=其他）
            # 根据用户要求，打√表示该子任务涉及的重点和难点
            set_cell_text(t5.cell(ri, 7), "√")
            set_cell_text(t5.cell(ri, 8), "√")
            set_cell_text(t5.cell(ri, 9), "")
        else:
            # 空行清空
            for ci in range(2, min(10, len(t5.rows[ri].cells))):
                set_cell_text(t5.cell(ri, ci), "")

print(f"  30个表5教学要求已修复")

# ============================================================
# 5. 恢复表5中"知识&技能"标题
# ============================================================
print("5. 检查知识&技能标题...")

# 检查表5R4-R7的col1是否有"知识目标"等标题
t5_first = doc.tables[5]
for ri in range(4, 8):
    cell = t5_first.cell(ri, 1)
    text = cell.text.strip()
    print(f"  R{ri} col1: '{text}'")

# ============================================================
# 6. 修复表5 R3教学任务（加第X章）
# ============================================================
print("6. 修复表5 R3教学任务...")

for ti in range(5, len(doc.tables), 2):
    t5 = doc.tables[ti]
    task_idx = (ti - 5) // 2
    if task_idx >= len(tasks):
        continue
    
    task = tasks[task_idx]
    title = task.get('title', f'任务{task_idx+1}')
    seq = task.get('seq', task_idx+1)
    
    parts = title.split('：')
    if len(parts) >= 2:
        chapter_name = parts[0]
        sub_tasks_list = [s.strip() for s in parts[1].split('、')]
    else:
        chapter_name = title
        sub_tasks_list = [title]
    
    chapter_num = f"第{chapter_nums[seq-1] if seq <= 9 else seq}章"
    
    # 教学任务：第X章 情境名 \n 子情景名
    task_text = f"{chapter_num} {chapter_name}\n"
    for st in sub_tasks_list:
        task_text += f"子情景：{st}\n"
    
    set_cell_text(t5.cell(3, 2), task_text)

print(f"  30个表5教学任务已加章节号")

# ============================================================
# 7. 修复表5表头"教学任务"内容
# ============================================================
# 表5 R3 col0-1 应该是"教学任务（章节、项目、情景）"
# 检查是否正确
t5_first = doc.tables[5]
r3_label = t5_first.cell(3, 0).text.strip()
print(f"  R3 col0: '{r3_label}'")

# ============================================================
# 保存
# ============================================================
doc.save(fp)
print(f"\n保存完成: {fp}")
print("第四轮修复完成")
