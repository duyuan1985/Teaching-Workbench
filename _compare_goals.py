"""
详细对比课程目标和教学方法
"""
from docx import Document

fp_design = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
fp_standard = r'生成结果\精修版\2023-2024-2《商务数据分析》课程标准 杜媛.docx'

doc_d = Document(fp_design)
doc_s = Document(fp_standard)

# ============================================================
# 课程标准目标表 Table 0
# ============================================================
print('='*80)
print('课程标准 目标表 (Table 0)')
print('='*80)
t0_s = doc_s.tables[0]
for ri in range(len(t0_s.rows)):
    for ci in range(len(t0_s.columns)):
        txt = t0_s.cell(ri, ci).text.strip()
        if txt:
            print(f'  R{ri}C{ci} [{t0_s.cell(0,ci).text.strip()[:10]}]:')
            for line in txt.split('\n'):
                if line.strip():
                    print(f'    {line.strip()[:100]}')

# ============================================================
# 教学设计目标段落
# ============================================================
print('\n' + '='*80)
print('教学设计 课程目标')
print('='*80)
in_goals = False
for p in doc_d.paragraphs:
    txt = p.text.strip()
    if '（一）课程目标设计' in txt:
        in_goals = True
    if in_goals:
        print(f'  {txt[:150]}')
    if in_goals and '（二）' in txt:
        break

# ============================================================
# 教学方法对比
# ============================================================
print('\n' + '='*80)
print('教学方法详细对比')
print('='*80)

print('\n--- 教学设计 ---')
in_method = False
for p in doc_d.paragraphs:
    txt = p.text.strip()
    if '（四）课程教学模式和教学方法设计' in txt:
        in_method = True
    if in_method:
        print(f'  {txt[:200]}')
    if in_method and '（五）' in txt:
        break

print('\n--- 课程标准 ---')
in_method = False
for p in doc_s.paragraphs:
    txt = p.text.strip()
    if '3、教学方法描述' in txt:
        in_method = True
    if in_method:
        print(f'  {txt[:200]}')
    if in_method and '4.' in txt:
        break
