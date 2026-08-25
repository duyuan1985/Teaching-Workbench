"""
修复教学设计文档中的考核方案表（Table 4）
直接操作XML避免python-docx cell()索引问题
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from copy import deepcopy

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)
t4 = doc.tables[4]

print(f"修复前: {len(t4.rows)} rows")
for ri in range(len(t4.rows)):
    texts = [c.text.strip()[:40] for c in t4.rows[ri].cells]
    print(f"  R{ri}: {texts}")

tbl = t4._tbl
rows = tbl.findall(qn('w:tr'))

# ============================================================
# Step 1: 清除所有合并属性（vMerge, gridSpan），重建基本结构
# ============================================================
for tr in rows:
    for tc in tr.findall(qn('w:tc')):
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            # 删除 vMerge
            for vm in tcPr.findall(qn('w:vMerge')):
                tcPr.remove(vm)
            # 删除 gridSpan（恢复为独立单元格）
            for gs in tcPr.findall(qn('w:gridSpan')):
                tcPr.remove(gs)

# ============================================================
# Step 2: 删除多余的汇总行（R1）
# ============================================================
row_to_delete = rows[1]
tbl.remove(row_to_delete)
rows = tbl.findall(qn('w:tr'))
print(f"\n删除R1后: {len(rows)} rows")

# ============================================================
# Step 3: 确保每行有正确的单元格数量（5个）
# ============================================================
for ri, tr in enumerate(rows):
    tcs = tr.findall(qn('w:tc'))
    # 如果单元格数量不对，需要调整
    print(f"  Row {ri}: {len(tcs)} cells")

# ============================================================
# Step 4: 更新单元格内容
# ============================================================
def clear_and_set_text(tc, text, font_name='仿宋', font_size='21', bold=False):
    """清除tc中的内容并设置新文本"""
    # 清除所有段落内容
    for p in tc.findall(qn('w:p')):
        # 删除所有run
        for r in p.findall(qn('w:r')):
            p.remove(r)

    # 获取或创建第一个段落
    paras = tc.findall(qn('w:p'))
    if not paras:
        p = OxmlElement('w:p')
        tc.append(p)
    else:
        p = paras[0]
        # 删除多余的段落
        for extra_p in paras[1:]:
            tc.remove(extra_p)

    # 设置段落格式
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p.insert(0, pPr)

    # 居中对齐
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'center')

    # 行距auto
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), 'auto')

    # 创建run
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 字体
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

    # 字号
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), font_size)
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), font_size)
    rPr.append(szCs)

    # 加粗
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)

    r.append(rPr)

    # 文本
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)

    p.append(r)

# 数据内容
data = [
    # R0: header (keep original or rewrite)
    ['项目', '考核方式', '考核标准', '权重', '合计'],
    # R1: 过程性考核 - 签到
    ['商务数据分析', '过程性考核', '签到', '10%', '40%'],
    # R2: 过程性考核 - 课堂表现
    ['', '', '课堂表现', '10%', ''],
    # R3: 过程性考核 - 作业
    ['', '', '作业', '20%', ''],
    # R4: 终结性考核 - 综合作品
    ['综合作品', '终结性考核', '基于真实电商数据集完成完整数据分析项目，含数据采集、清洗、分析、可视化和报告撰写，需答辩展示', '60%', '60%'],
]

rows = tbl.findall(qn('w:tr'))
for ri, tr in enumerate(rows):
    tcs = tr.findall(qn('w:tc'))
    for ci, tc in enumerate(tcs):
        if ri < len(data) and ci < len(data[ri]):
            is_header = (ri == 0)
            clear_and_set_text(tc, data[ri][ci], '仿宋', '21', bold=is_header)

# ============================================================
# Step 5: 应用纵向合并
# ============================================================
def set_vmerge(tc, restart=False):
    """设置纵向合并"""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    # 清除已有的 vMerge
    for vm in tcPr.findall(qn('w:vMerge')):
        tcPr.remove(vm)
    vm = OxmlElement('w:vMerge')
    if restart:
        vm.set(qn('w:val'), 'restart')
    tcPr.append(vm)

rows = tbl.findall(qn('w:tr'))

# 项目列 (col 0): R1-R3 合并
set_vmerge(rows[1].findall(qn('w:tc'))[0], restart=True)
set_vmerge(rows[2].findall(qn('w:tc'))[0], restart=False)
set_vmerge(rows[3].findall(qn('w:tc'))[0], restart=False)

# 考核方式列 (col 1): R1-R3 合并
set_vmerge(rows[1].findall(qn('w:tc'))[1], restart=True)
set_vmerge(rows[2].findall(qn('w:tc'))[1], restart=False)
set_vmerge(rows[3].findall(qn('w:tc'))[1], restart=False)

# 合计列 (col 4): R1-R3 合并
set_vmerge(rows[1].findall(qn('w:tc'))[4], restart=True)
set_vmerge(rows[2].findall(qn('w:tc'))[4], restart=False)
set_vmerge(rows[3].findall(qn('w:tc'))[4], restart=False)

# ============================================================
# 保存
# ============================================================
doc.save(fp)
print(f"\n修复完成，保存到: {fp}")

# 验证
doc2 = Document(fp)
t4_2 = doc2.tables[4]
print(f"\n验证: {len(t4_2.rows)} rows, {len(t4_2.columns)} cols")
for ri in range(len(t4_2.rows)):
    cells = t4_2.rows[ri].cells
    texts = [c.text.strip()[:50] for c in cells]
    print(f"  R{ri}: {texts}")
