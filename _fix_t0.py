"""
修复表0合并单元格内容
R1: 制定人 col3-5合并 → 写col3
R3: 开设学期 col1-2合并 → 写col1；授课对象 col4-5合并 → 写col4
R4: 先修课 col1-2合并 → 写col1；后续课 col4-5合并 → 写col4
"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)
t0 = doc.tables[0]

def set_cell_text(cell, text, font_name='仿宋', size=10.5, bold=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    
    lines = text.split('\n')
    for li, line in enumerate(lines):
        if li > 0:
            run = p.add_run()
            run.add_break()
        else:
            run = p.add_run()
        run.text = line
        run.font.name = font_name
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        run.font.size = Pt(size)
        run.font.bold = bold

# R1: 制定人 → col3 (合并col3-5)
set_cell_text(t0.cell(1, 3), "杜媛")

# R3: 开设学期 → col1 (合并col1-2)
set_cell_text(t0.cell(3, 1), "2023-2024学年第二学期")

# R3: 授课对象 → col4 (合并col4-5)
set_cell_text(t0.cell(3, 4), "2022电商教学班")

# R4: 先修课 → col1 (合并col1-2)
set_cell_text(t0.cell(4, 1), "电子商务基础、计算机应用基础、Python程序设计")

# R4: 后续课 → col4 (合并col4-5)
set_cell_text(t0.cell(4, 4), "新媒体平台运营与推广、电子商务综合实训")

doc.save(fp)
print("表0合并单元格内容已修复")

# 验证
doc2 = Document(fp)
t0 = doc2.tables[0]
for ri in range(len(t0.rows)):
    cells = t0.rows[ri].cells
    texts = [c.text.strip()[:35] for c in cells]
    print(f"  R{ri}: {texts}")
