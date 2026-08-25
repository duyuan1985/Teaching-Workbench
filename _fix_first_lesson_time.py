"""
调整第一节课各环节时间，总时长控制在90分钟
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

# 找到P110（总述）和P111-P118（各环节）
# 原时间：15+20+25+20+10+10+10+10=120min
# 新时间：10+15+20+15+5+10+5+10=90min
new_lines = [
    "第一节课为课程导入课（2学时，90分钟），主要内容包括：",
    "（1）课程介绍（10分钟）：介绍课程性质——农村电子商务专业核心课，总学时60学时（理论30+实践30），3学分。讲解课程教学目标（知识目标、能力目标、思政目标、素质目标），说明过程性考核（40%）与终结性考核（60%）的具体标准和权重分配，介绍超星学习通平台使用方法及课程资源构成。",
    "（2）行业认知（15分钟）：通过招聘网站数据展示电商行业数据分析岗位需求量、薪资水平（初级8-12K、中级12-20K、高级20K+）和职业发展路径（数据分析师→数据产品经理→数据总监），分析岗位核心技能要求，激发学生学习兴趣和职业认同感。",
    "（3）案例分析（20分钟）：展示农产品电商平台真实数据分析案例——某农产品电商平台的用户购买行为分析，包含数据采集、数据清洗、数据分析（RFM模型用户分群）、数据可视化和报告撰写全流程，让学生直观感受数据分析在商务决策中的实际价值。",
    "（4）工具介绍（15分钟）：现场演示Python数据分析环境搭建，包括Anaconda安装、Jupyter Notebook启动和基本操作；同时演示Excel高级功能（数据透视表、图表制作），让学生初步了解课程将使用的两大工具体系。",
    "（5）思政融入（5分钟）：讲解《数据安全法》《个人信息保护法》对数据分析的合规要求，通过\u201c大数据杀熟\u201d反面案例引导学生树立数据伦理意识，强调数据分析师应坚持真实分析、诚信报告的职业操守。",
    "（6）分组与任务分配（10分钟）：按4-5人一组将学生分组，每组分配本学期综合项目的初始方向（农产品电商、生鲜配送、农村物流等不同场景），发放项目任务书，明确团队分工要求，建立小组协作机制。",
    "（7）课堂互动（5分钟）：通过超星学习通发起随堂投票和抢答，检验学生对课程框架和考核方式的理解程度，收集学生疑问并现场解答。",
    "（8）课后任务布置（10分钟）：要求学生完成以下任务：①注册超星学习通账号并加入课程班级；②预习第一章微课视频；③安装Anaconda环境并测试Jupyter Notebook运行；④阅读教材第一章节内容；⑤各小组讨论确定综合项目的数据集方向，下次课提交项目选题报告。",
]

# 找到P110元素
p110 = doc.paragraphs[110]
p110_elem = p110._element

# 获取格式
font_name = '仿宋'
font_size_val = '28'
if p110.runs:
    r = p110.runs[0]
    if r.font.name:
        font_name = r.font.name
    if r.font.size:
        font_size_val = str(int(r.font.size.pt * 2))

# 收集P110-P118的所有元素
paras_to_update = []
for pi in range(110, 119):
    paras_to_update.append(doc.paragraphs[pi]._element)

# 替换已有段落内容
for i, p_elem in enumerate(paras_to_update):
    # 清除原有run
    for r in p_elem.findall(qn('w:r')):
        p_elem.remove(r)
    # 添加新run
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), font_size_val)
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), font_size_val)
    rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = new_lines[i]
    r.append(t)
    p_elem.append(r)

doc.save(fp)
print("保存完成")

# 验证
doc2 = Document(fp)
total = 0
for pi, p in enumerate(doc2.paragraphs):
    txt = p.text.strip()
    if '九、第一节课' in txt:
        in_section = True
    if in_section and txt:
        print(f'P{pi}: {txt[:120]}')
        # 提取时间
        import re
        m = re.search(r'(\d+)分钟', txt)
        if m:
            total += int(m.group(1))
    if in_section and '十、' in txt:
        break
print(f'\n总时长: {total}分钟')
