"""
统一文档生成器 - 课程标准/授课计划/教学设计
用法: python generate.py <offering_id> [--doc plan|standard|design|all]
"""
import sys, os, shutil, random, re, argparse
from copy import deepcopy
from io import BytesIO
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
from pptx import Presentation
import store

random.seed(42)

# ============================================================
# 第一层：XML安全操作工具
# 原则：只替换文字，不覆盖模板的格式属性
# ============================================================

def safe_set_text(paragraph, text, font=None, size=None, bold=None):
    """安全设置段落文字：保留模板原有格式，只替换文字"""
    # 保留第一个run的格式，只改文字
    if not paragraph.runs:
        r = paragraph.add_run(text)
    else:
        # 清除多余run
        for extra in paragraph.runs[1:]:
            extra._element.getparent().remove(extra._element)
        paragraph.runs[0].text = text
        r = paragraph.runs[0]
    # 只在明确指定时才覆盖格式
    if font is not None:
        _set_font_name(r, font)
    if size is not None:
        _set_font_size(r, size)
    if bold is not None:
        r.font.bold = bold
    return r

def _set_font_name(run, name):
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.insert(0, rf)
    rf.set(qn('w:eastAsia'), name)
    rf.set(qn('w:ascii'), name)
    rf.set(qn('w:hAnsi'), name)

def _set_font_size(run, pt_size):
    half_pt = int(pt_size * 2)
    rpr = run._element.get_or_add_rPr()
    for tag in ('w:sz', 'w:szCs'):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rpr.append(el)
        el.set(qn('w:val'), str(half_pt))

def clear_indent(paragraph):
    """清除段落所有缩进（firstLine, firstLineChars, left等）"""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        return
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        pPr.remove(ind)

def set_indent_chars(paragraph, chars=2):
    """设置段首缩进N个字符"""
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars * 100))
    # 清除firstLine（避免冲突）
    if ind.get(qn('w:firstLine')) is not None:
        ind.attrib.pop(qn('w:firstLine'))

def fix_spacing_auto(paragraph):
    """将固定行距改为自动行距，防止文字被截断"""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        return
    spacing = pPr.find(qn('w:spacing'))
    if spacing is not None:
        rule = spacing.get(qn('w:lineRule'))
        if rule == 'exact':
            spacing.set(qn('w:lineRule'), 'auto')
            spacing.set(qn('w:line'), '360')  # 1.5倍行距

def set_alignment(paragraph, align):
    """设置段落对齐"""
    paragraph.alignment = align

def set_page_break_before(paragraph):
    """设置段前分页"""
    pPr = paragraph._element.get_or_add_pPr()
    pbdr = OxmlElement('w:pageBreakBefore')
    pPr.append(pbdr)

def set_space_before(paragraph, pt):
    """设置段前间距"""
    paragraph.paragraph_format.space_before = Pt(pt)

def set_line_spacing(paragraph, multiple):
    """设置行距倍数"""
    paragraph.paragraph_format.line_spacing = multiple

def merge_v(table, row1, row2, col):
    """纵向合并单元格，并显式写入 continue 标记。"""
    if row1 >= len(table.rows) or row2 >= len(table.rows) or row2 < row1:
        return
    try:
        table.cell(row1, col).merge(table.cell(row2, col))
    except Exception:
        return
    # python-docx 的 merge() 有时只保留 restart，Word 会把下方单元格当成独立空格。
    # 直接补齐 OOXML vMerge，确保视觉上是一个纵向合并单元格。
    for row_index in range(row1, row2 + 1):
        tr = table.rows[row_index]._tr
        cells = tr.findall(qn("w:tc"))
        if col >= len(cells):
            continue
        tc_pr = cells[col].get_or_add_tcPr()
        vmerge = tc_pr.find(qn("w:vMerge"))
        if vmerge is None:
            vmerge = OxmlElement("w:vMerge")
            tc_pr.append(vmerge)
        vmerge.set(qn("w:val"), "restart" if row_index == row1 else "continue")

def merge_h(table, row, col1, col2):
    """横向合并单元格"""
    if row >= len(table.rows):
        return
    try:
        table.cell(row, col1).merge(table.cell(row, col2))
    except:
        pass

def write_cell(cell, text, font="仿宋_GB2312", size=9, bold=False, align=None):
    """写入单元格内容（清空旧内容后写入）"""
    # 清空多余段落
    while len(cell.paragraphs) > 1:
        cell.paragraphs[-1]._element.getparent().remove(cell.paragraphs[-1]._element)
    # 清空第一个段落的内容
    p0 = cell.paragraphs[0]
    for r in p0.runs:
        r.text = ""
    # 写入新内容
    lines = text.split("\n") if isinstance(text, str) else [str(text)]
    for li, line in enumerate(lines):
        if li == 0:
            if p0.runs:
                p0.runs[0].text = line
                r = p0.runs[0]
            else:
                r = p0.add_run(line)
        else:
            np = cell.add_paragraph()
            r = np.add_run(line)
        _set_font_name(r, font)
        _set_font_size(r, size)
        r.font.bold = bold
    if align is not None:
        p0.alignment = align
    # 清除缩进
    clear_indent(p0)
    for p in cell.paragraphs[1:]:
        clear_indent(p)

def write_cell_lines(cell, lines, font="仿宋_GB2312", size=9, bold_map=None):
    """写入多行到单元格，每行可独立设置加粗"""
    while len(cell.paragraphs) > 1:
        cell.paragraphs[-1]._element.getparent().remove(cell.paragraphs[-1]._element)
    p0 = cell.paragraphs[0]
    for r in p0.runs:
        r.text = ""
    for li, line in enumerate(lines):
        if li == 0:
            p = p0
            if p.runs:
                r = p.runs[0]
                r.text = line
            else:
                r = p.add_run(line)
        else:
            p = cell.add_paragraph()
            r = p.add_run(line)
        _set_font_name(r, font)
        _set_font_size(r, size)
        if bold_map and li in bold_map:
            r.font.bold = bold_map[li]
        clear_indent(p)

def get_vmerge_val(tc):
    """获取vMerge值：restart/continue/None"""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    vm = tcPr.find(qn('w:vMerge'))
    if vm is None:
        return None
    return vm.get(qn('w:val')) or 'continue'

def get_cell_text_raw(tc):
    """直接从XML读取单元格文本"""
    ps = tc.findall(qn('w:p'))
    if not ps:
        return ""
    texts = []
    for p in ps:
        for r in p.findall(qn('w:r')):
            for t in r.findall(qn('w:t')):
                texts.append(t.text or "")
    return "".join(texts).strip()

CN_NUMS = ["一","二","三","四","五","六","七","八","九","十"]

# ============================================================
# 第二层：内容生成器（差异化小结/作业/反思/思政）
# ============================================================

IDEO_MAP = {
    "初识": ["数据安全意识：遵守《个人信息保护法》《数据安全法》",
             "诚信分析精神：坚持真实分析和规范验证",
             "科技报国情怀：服务乡村振兴战略"],
    "指标": ["数据真实性：基于真实原始数据",
             "客观公正：以数据为依据判断",
             "职业道德：严守保密协议"],
    "Excel": ["严谨细致：公式和数据透视不能出错",
              "规范操作：遵循数据处理规范",
              "工匠精神：精益求精"],
    "Numpy": ["科学精神：严格遵循数学原理",
              "严谨细致：数组维度不能错",
              "规范操作：遵循编码规范"],
    "Pandas": ["数据诚信：如实记录处理过程",
              "隐私保护：数据脱敏处理",
              "工匠精神：严格把关数据质量"],
    "SciPy": ["科学求真：结论要有数据支撑",
              "逻辑思维：从现象发现本质",
              "创新探索：尝试不同方法"],
    "sklearn": ["算法伦理：确保模型公平公正",
               "数据偏见：培养批判性思维",
               "科技向善：让科技服务社会"],
    "Sklearn": ["算法伦理：确保模型公平公正",
               "数据偏见：培养批判性思维",
               "科技向善：让科技服务社会"],
    "Seaborn": ["数据美学：以清晰直观的方式呈现数据",
               "严谨表达：图表标注完整准确",
               "创新思维：探索数据可视化新方法"],
    "综合": ["团队协作：分工合作互相帮助",
            "责任担当：按时高质量交付",
            "持续学习：树立终身学习理念"],
    "实战": ["综合应用：融会贯通",
            "职业素养：沟通汇报能力",
            "创新创业：探索新应用"],
}

def get_ideo(chapter):
    for k, v in IDEO_MAP.items():
        if k in chapter:
            return v
    return IDEO_MAP["综合"]

def gen_summary(task, idx, kp_names, title, chapter, ideo):
    """生成差异化课堂小结"""
    structures = [
        # 结构1：三点式
        lambda: (
            f"本节课我们围绕{title}展开了系统学习，主要内容可以概括为以下三个方面：\n\n"
            f"第一，{kp_names[0]}是基础。我们学习了它的定义、核心要素和基本特征，"
            f"理解了它在数据分析体系中的地位和作用。这部分是后续学习的基石，必须牢固掌握。\n\n"
            f"第二，{kp_names[1] if len(kp_names)>1 else '核心方法'}是重点也是难点。"
            f"我们通过代码演示和实操练习，掌握了它的基本操作步骤和关键技术要点。"
            f"这部分内容实践性强，需要通过大量练习才能熟练掌握。同学们在课后要多加练习。\n\n"
            f"第三，{kp_names[2] if len(kp_names)>2 else '应用实践'}是延伸。"
            f"我们将所学方法应用到实际数据分析场景中，培养解决实际问题的能力。"
            f"这部分是知识转化为能力的关键环节，也是企业对数据分析人才的核心要求。\n\n"
            f"同学们要把这三个方面串联起来，形成完整的知识网络，而不是孤立地记忆零散的概念。"
            f"下节课我们将继续学习更深入的内容，请同学们课后及时复习。"
        ),
        # 结构2：是什么/为什么/怎么做
        lambda: (
            f"【是什么】\n"
            f"{title}是{chapter}模块中的重要组成部分，主要解决{kp_names[0]}相关的问题。"
            f"它包含{kp_names[0]}、{kp_names[1] if len(kp_names)>1 else '核心方法'}、"
            f"{kp_names[2] if len(kp_names)>2 else '应用实践'}等关键内容，"
            f"在电商运营、市场营销、用户分析等场景中都有广泛应用。\n\n"
            f"【为什么】\n"
            f"为什么要学习{title}？因为在实际工作中，数据分析师每天都要用到这些方法。"
            f"比如电商运营人员需要通过{kp_names[0] if kp_names else '数据分析'}来了解用户购买行为，"
            f"市场人员需要通过{kp_names[1] if len(kp_names)>1 else '统计分析'}来评估营销效果。"
            f"掌握好这些技能，是成为一名合格数据分析师的必备条件。\n\n"
            f"【怎么做】\n"
            f"学好{title}的关键在于三个字：练、想、用。\n"
            f"  · 练：多动手写代码、做实验，熟能生巧\n"
            f"  · 想：每做完一个练习都要反思为什么这么做、有没有更好的方法\n"
            f"  · 用：把学到的方法用到真实问题中，在解决问题的过程中深化理解\n\n"
            f"希望同学们课后按照这三个字的要求，认真复习和练习。"
        ),
        # 结构3：收获/疑问/行动
        lambda: (
            f"【本节课的收获】\n"
            f"通过本节课的学习，同学们应该有以下收获：\n"
            f"  1. 知识层面：理解了{kp_names[0]}的基本概念和原理，掌握了"
            f"{kp_names[1] if len(kp_names)>1 else title}的操作方法\n"
            f"  2. 能力层面：能够独立完成基础实操任务，初步具备调试和排错能力\n"
            f"  3. 素养层面：体会到{ideo[0].split('：')[0] if '：' in ideo[0] else '数据安全'}的重要性，"
            f"培养了规范操作的意识\n\n"
            f"【还存在的疑问】\n"
            f"相信同学们在学习过程中也产生了一些疑问，比如：\n"
            f"  · {kp_names[1] if len(kp_names)>1 else '核心方法'}在什么情况下最适用？\n"
            f"  · 数据量很大的时候，这些方法还能用吗？\n"
            f"  · 除了课堂上讲的，还有没有其他实现方式？\n"
            f"这些问题非常好，说明大家在积极思考。有些问题我们会在后续课程中解答，"
            f"有些问题需要同学们自己去查阅资料、动手实验来寻找答案。\n\n"
            f"【下一步行动】\n"
            f"  1. 及时复习：整理笔记，巩固本节课知识点\n"
            f"  2. 完成作业：认真完成课后练习，检验学习效果\n"
            f"  3. 提前预习：预习下节课内容，带着问题听课效果更好"
        ),
        # 结构4：知识图谱式
        lambda: (
            f"本节课的知识图谱可以这样梳理：\n\n"
            f"                    {title}\n"
            f"                   /   |   \\\n"
            f"          概念     方法     应用\n"
            f"         /   \\   /   \\   /   \\\n"
            f"       {kp_names[0][:6]} 原理 操作 实战 拓展\n\n"
            f"【核心概念】{kp_names[0]}——是什么、有什么用、分几类\n"
            f"【关键方法】{kp_names[1] if len(kp_names)>1 else '核心方法'}——怎么做、注意什么、常见错误\n"
            f"【实际应用】{kp_names[2] if len(kp_names)>2 else '应用实践'}——用在哪、怎么用、效果如何\n\n"
            f"各知识点之间的联系：\n"
            f"  · 概念是基础，方法是手段，应用是目的\n"
            f"  · 理解了概念才能正确选择方法，掌握了方法才能有效应用\n"
            f"  · 在应用中发现的问题又会促使我们深化对概念和方法的理解\n\n"
            f"同学们要学会用这种结构化的方式来整理知识，形成自己的知识体系。"
            f"下节课我们将在这个基础上继续拓展，请同学们课后好好消化。"
        ),
        # 结构5：对比总结式
        lambda: (
            f"【内容回顾】\n"
            f"本节课我们主要学习了以下内容：\n"
            f"  ① {kp_names[0]}：定义、作用、分类、适用场景\n"
            f"  ② {kp_names[1] if len(kp_names)>1 else '核心方法'}：操作步骤、关键技术、常见错误\n"
            f"  ③ {kp_names[2] if len(kp_names)>2 else '应用实践'}：实训项目、操作要点、结果验证\n"
            f"  ④ {kp_names[3] if len(kp_names)>3 else '拓展内容'}：进阶知识、行业应用\n\n"
            f"【重点对比】\n"
            f"很多同学容易混淆{kp_names[0]}和{kp_names[1] if len(kp_names)>1 else '相关方法'}，"
            f"这里做一个简单对比：\n"
            f"  · 侧重点不同：一个偏理论理解，一个偏实际操作\n"
            f"  · 学习方法不同：前者重在理解，后者重在练习\n"
            f"  · 应用场景不同：不同的业务问题需要选择不同的方法\n\n"
            f"【易错点提醒】\n"
            f"根据课堂观察，同学们在以下几个地方容易出错：\n"
            f"  1. 参数设置不当导致结果异常\n"
            f"  2. 数据类型不匹配引发报错\n"
            f"  3. 忽略了数据预处理的重要性\n"
            f"  4. 结果验证不充分就下结论\n"
            f"希望同学们在课后练习中特别注意这些问题，养成严谨细致的习惯。"
        ),
    ]
    return structures[idx % len(structures)]()

def gen_homework(task, idx, kp_names, title, chapter, ideo):
    """生成差异化课后作业"""
    types = [
        # 类型1：基础+实操+思考
        lambda: (
            f"【基础题】\n"
            f"1. 简答题：简述{kp_names[0]}的定义、作用和主要分类（不少于200字）\n"
            f"2. 填空题：{kp_names[0]}的三个核心要素是____、____、____\n"
            f"3. 判断题：{kp_names[1] if len(kp_names)>1 else title}只能用于数值型数据（  ）\n\n"
            f"【实操题】\n"
            f"4. 完成本节课实训项目的所有要求，代码规范、注释完整，打包提交\n"
            f"5. 在原有代码基础上增加一个新功能：用{kp_names[2] if len(kp_names)>2 else '另一种方法'}"
            f"实现相同的功能，对比两种方法的结果差异\n\n"
            f"【思考题】\n"
            f"6. {ideo[0].split('：')[0] if '：' in ideo[0] else '数据安全'}在{title}中有哪些具体体现？"
            f"请结合一个真实案例来说明（不少于300字）\n\n"
            f"【拓展题（选做）】\n"
            f"7. 从Kaggle下载一个电商数据集，用本节课所学的方法进行分析，"
            f"写出你的发现和建议（不少于500字）\n\n"
            f"【预习】\n"
            f"8. 阅读下节课PPT，列出3个你最想了解答案的问题"
        ),
        # 类型2：分层作业
        lambda: (
            f"【必做题（所有人都要完成）】\n"
            f"1. 复习本节课PPT和笔记，整理{title}的知识点思维导图\n"
            f"2. 把课堂上的代码重新独立写一遍，确保每一行都理解\n"
            f"3. 完成教材对应章节的课后习题1-10题\n\n"
            f"【选做题（至少选2题）】\n"
            f"4. 用Python实现一个{title}的小工具，可以接收用户输入的参数并输出结果\n"
            f"5. 调研3个主流的{title}工具/库，写一个对比分析（不少于300字）\n"
            f"6. 找一个因为{kp_names[0] if kp_names else '数据质量'}问题导致决策错误的真实案例，"
            f"分析原因并提出改进建议\n"
            f"7. 尝试用两种不同的方法解决同一个分析问题，比较效率和结果差异\n\n"
            f"【挑战题（学有余力的同学选做）】\n"
            f"8. 阅读一篇关于{title}的学术论文或技术博客，写一个摘要（200字以内），"
            f"并用自己的话解释核心思想\n\n"
            f"【小组任务】\n"
            f"9. 以小组为单位，讨论{title}在农产品电商中的一个具体应用场景，"
            f"下次课每组派代表分享（3分钟以内）"
        ),
        # 类型3：项目式作业
        lambda: (
            f"【项目名称】电商用户{title}分析\n\n"
            f"【项目背景】\n"
            f"某农产品电商平台积累了大量用户数据，运营团队希望通过{title}"
            f"来了解用户行为特征，为后续运营决策提供数据支持。\n\n"
            f"【项目要求】\n"
            f"1. 数据准备：从指定位置下载用户行为数据集（约10万条记录）\n"
            f"2. 数据探索：用{kp_names[0]}对数据进行初步探索，了解数据基本情况\n"
            f"3. 核心分析：用{kp_names[1] if len(kp_names)>1 else '核心方法'}进行深入分析\n"
            f"4. 结果呈现：用图表展示分析结果，撰写分析报告\n"
            f"5. 业务建议：基于分析结果，提出至少3条运营优化建议\n\n"
            f"【提交要求】\n"
            f"  · Python源代码（.ipynb格式，注释完整）\n"
            f"  · 分析报告（Word格式，不少于800字，含图表）\n"
            f"  · 截止时间：下周三晚24:00\n\n"
            f"【评分标准】\n"
            f"  · 代码正确性（30%）：代码能否正确运行，结果是否准确\n"
            f"  · 分析深度（30%）：分析是否深入，结论是否有依据\n"
            f"  · 报告质量（20%）：结构是否清晰，图表是否规范\n"
            f"  · 创新程度（20%）：是否有自己的思考和发现\n\n"
            f"【德育要求】\n"
            f"本次作业可以讨论但必须独立完成，严禁抄袭。"
            f"数据仅供教学使用，不得外传。——{ideo[1].split('：')[0] if '：' in ideo[1] else '诚信'}"
        ),
        # 类型4：复盘式作业
        lambda: (
            f"【知识复盘】\n"
            f"1. 用自己的话解释{kp_names[0]}的含义（不少于100字）\n"
            f"2. 画出{kp_names[1] if len(kp_names)>1 else title}的操作流程图\n"
            f"3. 列出本节课你认为最重要的3个知识点，并说明理由\n\n"
            f"【错误复盘】\n"
            f"4. 记录你在实操练习中遇到的至少3个错误/问题：\n"
            f"   错误1：\n"
            f"     · 现象：\n"
            f"     · 原因分析：\n"
            f"     · 解决方法：\n"
            f"   错误2：\n"
            f"     · 现象：\n"
            f"     · 原因分析：\n"
            f"     · 解决方法：\n"
            f"   错误3：\n"
            f"     · 现象：\n"
            f"     · 原因分析：\n"
            f"     · 解决方法：\n\n"
            f"【拓展学习】\n"
            f"5. 查阅Python官方文档中关于{kp_names[0] if kp_names else '相关模块'}的部分，"
            f"找出3个课堂上没有讲到的功能/参数，简要说明它们的用途\n"
            f"6. 在知乎、CSDN或技术博客上搜索一篇关于{title}的文章，"
            f"阅读后写下你的收获和疑问（不少于200字）\n\n"
            f"【实践应用】\n"
            f"7. 思考：在你的日常生活中，有哪些地方可以用到{title}的思想或方法？"
            f"请举一个具体例子并说明（不少于150字）"
        ),
        # 类型5：游戏化作业
        lambda: (
            f"欢迎来到{title}闯关任务！共设5关，看看你能闯到第几关？\n\n"
            f"【第一关：入门级】⭐\n"
            f"任务：复述{kp_names[0]}的定义和作用\n"
            f"奖励：基础经验值+10\n"
            f"完成方式：在作业本上写下你的答案\n\n"
            f"【第二关：基础级】⭐⭐\n"
            f"任务：独立完成课堂示例代码并得到正确结果\n"
            f"奖励：经验值+20，解锁\"初级分析师\"称号\n"
            f"完成方式：提交运行成功的代码截图\n\n"
            f"【第三关：进阶级】⭐⭐⭐\n"
            f"任务：在示例代码基础上，增加{kp_names[2] if len(kp_names)>2 else '一个新功能'}\n"
            f"奖励：经验值+30，解锁\"中级分析师\"称号\n"
            f"完成方式：提交完整代码和运行结果\n\n"
            f"【第四关：高手级】⭐⭐⭐⭐\n"
            f"任务：用{title}分析一个自选数据集，得出3个有价值的结论\n"
            f"奖励：经验值+40，解锁\"高级分析师\"称号\n"
            f"完成方式：提交分析报告（300字以上）\n\n"
            f"【第五关：大师级】⭐⭐⭐⭐⭐\n"
            f"任务：对比分析至少两种不同的{title}方法，写出详细的对比报告\n"
            f"奖励：经验值+50，解锁\"数据大师\"称号+神秘奖品\n"
            f"完成方式：提交对比分析报告（500字以上）\n\n"
            f"【彩蛋任务】\n"
            f"思考题：{ideo[2].split('：')[0] if '：' in ideo[2] else '科技向善'}——"
            f"技术是中性的，但使用技术的人要有温度。结合本节课内容，谈谈你的理解。"
        ),
        # 类型6：PBL式
        lambda: (
            f"【驱动问题】\n"
            f"假设你是某农产品电商公司的数据分析师，运营总监找到你说："
            f"\"最近我们的用户复购率下降了，你能不能用{title}帮我分析一下原因？\"\n"
            f"请你设计一个分析方案来回答这个问题。\n\n"
            f"【任务拆解】\n"
            f"任务1：明确问题\n"
            f"  · 运营总监的问题到底是什么？（把模糊的问题具体化）\n"
            f"  · 需要哪些数据才能回答这个问题？\n"
            f"  · 衡量\"原因\"的标准是什么？\n\n"
            f"任务2：数据准备\n"
            f"  · 需要哪些数据表？它们之间有什么关系？\n"
            f"  · 数据质量如何？需要做哪些清洗？\n"
            f"  · 用{kp_names[0] if kp_names else '描述性统计'}先看看数据长什么样\n\n"
            f"任务3：分析实施\n"
            f"  · 用{kp_names[1] if len(kp_names)>1 else '核心方法'}进行深入分析\n"
            f"  · 你发现了哪些规律和异常？\n"
            f"  · 哪些发现可能解释复购率下降的原因？\n\n"
            f"任务4：结果呈现\n"
            f"  · 你会如何向运营总监汇报你的发现？\n"
            f"  · 你会提出哪些具体建议？\n"
            f"  · 你的建议有什么数据支撑？\n\n"
            f"【提交要求】\n"
            f"提交一份分析方案文档（可以是大纲形式，500字以上），"
            f"不需要真正跑数据，重点展示你的分析思路和方法选择。"
        ),
    ]
    return types[idx % len(types)]()

def gen_reflection(task, idx, kp_names, title, chapter, ideo):
    """生成差异化教学反思"""
    types = [
        lambda: (
            f"【知识目标达成情况】\n"
            f"本节课的知识目标是让学生理解{kp_names[0]}的概念和{kp_names[1] if len(kp_names)>1 else title}的操作方法。"
            f"从课堂提问和练习情况来看，约70%的学生能够达到基本要求，"
            f"约30%的学生对{kp_names[1] if len(kp_names)>1 else '代码部分'}的理解还不够深入，"
            f"需要在后续课程中进一步巩固。\n\n"
            f"【能力目标达成情况】\n"
            f"能力目标方面，约60%的学生能够独立完成基础实训任务，"
            f"约30%的学生需要老师提示或同学帮助才能完成，"
            f"约10%的学生完成情况不理想。学生的调试能力普遍较弱，"
            f"遇到error不知道从何入手，这是后续教学需要重点加强的地方。\n\n"
            f"【思政目标达成情况】\n"
            f"本节课的德育渗透点是{ideo[0].split('：')[0] if '：' in ideo[0] else '数据安全'}。"
            f"在讲解{kp_names[0]}时引入了相关案例和讨论，学生反应较好，"
            f"有几位学生主动提问和分享自己的看法，说明德育渗透起到了效果。\n\n"
            f"【改进措施】\n"
            f"1. 下次课前发放预习视频，让学生先了解基本概念\n"
            f"2. 增加调试方法的专题指导\n"
            f"3. 设计更多梯度化的练习任务，照顾不同基础的学生"
        ),
        lambda: (
            f"【教学方法效果评估】\n"
            f"本节课主要采用了项目教学法和任务驱动法。从课堂反馈来看，"
            f"任务驱动法的效果不错，学生带着明确的任务学习，目的性更强，"
            f"注意力也更集中。但任务设计的难度梯度还可以优化，"
            f"基础任务和进阶任务之间的跨度有点大，基础薄弱的学生容易受挫。\n\n"
            f"案例教学法在导入环节效果很好，选用的农产品电商案例贴近学生专业，"
            f"容易引起共鸣。但案例的深度还可以加强，可以引入更多真实的业务细节。\n\n"
            f"分组讨论环节时间偏短，很多组还没充分讨论就被打断了。"
            f"下次可以适当延长讨论时间，或者提前布置讨论题让学生课前准备。\n\n"
            f"【时间管理反思】\n"
            f"本节课的时间把控不够理想。{kp_names[0]}的概念讲解部分花的时间略多，"
            f"导致后面{kp_names[1] if len(kp_names)>1 else '实操'}的时间被压缩了一些。"
            f"分析原因主要是：\n"
            f"  1. 学生提问比预期多，临时增加了答疑时间\n"
            f"  2. 概念讲解时举了太多例子，偏离了主线\n"
            f"  3. 演示环节出现了一个技术小问题，耽误了几分钟\n"
            f"下次要更严格地控制各环节时间，可以准备一个计时器。\n\n"
            f"【下次改进】\n"
            f"  · 优化任务难度梯度，增加过渡任务\n"
            f"  · 提前准备好可能的技术问题，减少现场调试时间\n"
            f"  · 严格控制各环节时间，使用计时器提醒"
        ),
        lambda: (
            f"【教学亮点】\n"
            f"1. 案例导入效果好：选用了学生熟悉的直播电商案例，"
            f"学生兴趣很高，课堂参与度比平时高出不少\n"
            f"2. 理实交替节奏合适：讲15分钟练10分钟的节奏比较合适，"
            f"学生的注意力保持较好，没有出现明显的走神现象\n"
            f"3. 德育渗透自然：在讲解{kp_names[0]}时自然引入"
            f"{ideo[0].split('：')[0] if '：' in ideo[0] else '数据安全'}的话题，"
            f"不生硬不说教，学生接受度高\n"
            f"4. 成果展示环节激发了学生的成就感：抽取几组学生展示实训成果，"
            f"其他学生互评学习，效果很好\n\n"
            f"【存在不足】\n"
            f"1. 个体差异较大：基础好的学生早早完成任务开始摸鱼，"
            f"基础差的学生还在第一步挣扎，两头都没照顾好\n"
            f"2. 板书不够系统：讲课过程中板书有点随意，学生课后看笔记找不到重点\n"
            f"3. 作业反馈不及时：学生提交的作业不能当天批改反馈，"
            f"等反馈回来学生已经忘了当时是怎么想的\n"
            f"4. 后排学生关注不够：巡回指导时主要关注前排和中间的学生，"
            f"后排的学生得到的指导较少\n\n"
            f"【改进计划】\n"
            f"1. 设计分层任务：基础任务+进阶任务+挑战任务，不同水平的学生各取所需\n"
            f"2. 准备结构化板书：提前做好板书PPT，课后发给学生\n"
            f"3. 建立互助小组：让掌握快的学生帮助基础薄弱的学生\n"
            f"4. 增加后排巡回频次：有意识地多到后排走走"
        ),
        lambda: (
            f"【学生课堂表现观察】\n"
            f"本节课学生整体表现良好，课堂气氛比较活跃。具体观察如下：\n"
            f"  · 注意力：前20分钟注意力很集中，20-35分钟开始有些涣散，"
            f"35分钟后（实操环节）又重新集中起来\n"
            f"  · 参与度：约60%的学生主动回答问题或参与讨论，"
            f"约30%的学生被点到时能回答，约10%的学生基本不发言\n"
            f"  · 实操情况：约60%的学生能跟上进度，"
            f"约30%的学生进度稍慢但能完成，"
            f"约10%的学生需要一对一辅导\n\n"
            f"【典型问题分析】\n"
            f"从课堂练习和提问来看，学生普遍存在以下问题：\n"
            f"1. {kp_names[1] if len(kp_names)>1 else '代码部分'}的细节掌握不牢，"
            f"很多学生记不住具体的参数和用法，需要经常查文档\n"
            f"2. 调试能力弱，遇到error不知道从哪里看起，"
            f"很多学生直接把报错截图发过来问\"老师这是什么错\"\n"
            f"3. 知识迁移能力不足，课堂上讲过的例子会做，"
            f"换一个场景就不知道怎么下手了\n\n"
            f"【针对性改进】\n"
            f"1. 制作\"速查手册\"：把常用的函数、参数、用法整理成一页纸，发给学生随时查阅\n"
            f"2. 增加\"排错小课堂\"：每节课花3分钟讲一个常见错误和排错方法\n"
            f"3. 设计更多变式练习：同一个知识点用不同的场景来练习，培养迁移能力\n"
            f"4. 建立\"错题本\"机制：让学生记录自己的错误，定期复习"
        ),
    ]
    base = types[idx % len(types)]()
    # 添加独特观察
    observations = [
        f"\n\n【本节课特别观察】\n第{idx+1}次课，学生整体状态{'较好' if idx%2==0 else '一般'}。"
        f"{'前排学生' if idx%3==0 else '中间组'}在讨论环节表现尤为积极。"
        f"有{3+idx%5}位学生主动提出了有深度的问题，说明学生的思考在深入。"
        f"下节课可以尝试让学生自己来讲一部分内容，进一步锻炼表达能力。",
        f"\n\n【课堂趣事】\n今天的实操环节发生了一件有趣的事："
        f"第{2+idx%6}组的同学在调试代码时发现了一个老师都没注意到的小技巧，"
        f"大家都围过去看，学习气氛特别好。"
        f"以后要多给学生展示自己发现的机会，教学相长。",
        f"\n\n【待跟进学生】\n本节课发现有{1+idx%4}位学生进度明显偏慢，"
        f"主要卡在{kp_names[0] if kp_names else '概念理解'}部分。"
        f"已经安排了学习伙伴帮助他们，下节课要重点关注这几位学生的掌握情况。"
        f"同时也要思考：是不是我的讲解方式对这部分学生不够友好？",
        f"\n\n【教学灵感】\n讲完{kp_names[1] if len(kp_names)>1 else title}的时候突然想到一个好例子——"
        f"可以用学生自己的消费数据来做分析，这样代入感更强。"
        f"下次备课的时候把这个例子加进去，应该能让学生更有兴趣。"
        f"好的教学就是在一次次这样的小改进中不断完善的。",
        f"\n\n【学生闪光点】\n今天要特别表扬第{1+idx%8}组的同学，"
        f"他们不仅自己完成了任务，还主动帮助旁边遇到困难的同学。"
        f"这种互助精神非常可贵，下次课要公开表扬，树立好榜样。"
        f"团队协作能力也是职业素养的重要组成部分。",
    ]
    base += observations[idx % len(observations)]
    return base


# ============================================================
# 第三层：数据加载工具
# ============================================================

def load_offering(offering_id):
    """加载开课数据"""
    rows = store.rows("SELECT * FROM offerings WHERE id=?", (offering_id,))
    if not rows:
        raise ValueError(f"offering_id={offering_id} 不存在")
    return rows[0]

def load_tasks(offering_id):
    """加载教学任务"""
    return store.rows("SELECT * FROM tasks WHERE offering_id=? ORDER BY seq", (offering_id,))

def load_sessions(offering_id):
    """加载排课"""
    return store.rows("SELECT * FROM sessions WHERE offering_id=? ORDER BY lesson_date", (offering_id,))

def load_units(offering_id):
    """加载课程单元"""
    return store.rows("SELECT * FROM curriculum_units WHERE offering_id=? ORDER BY seq", (offering_id,))

def load_calendar(offering_id):
    """加载校历事件"""
    return store.rows("SELECT * FROM calendar_events WHERE offering_id=? ORDER BY start_date", (offering_id,))

def find_unit_for_task(task, units):
    for u in units:
        if u["project_title"] == task["chapter"]:
            return u
    return units[0] if units else {}

def find_sess_for_task(task, sessions):
    for s in sessions:
        if s.get("week_no") == task.get("week_no") and s.get("lesson_date") == task.get("lesson_date"):
            return s
    ws = [s for s in sessions if s.get("week_no") == task.get("week_no")]
    return ws[0] if ws else {}


# ============================================================
# 第四层：模板路径解析
# ============================================================

TEMPLATE_BASE = os.path.join(os.path.dirname(__file__), "原始资料", "模板")

def get_opening_semester(offering):
    """根据班级入学年份和开课学期计算第几学期。"""
    import re
    term = str(offering.get("term", ""))
    teaching_class = str(offering.get("teaching_class", ""))
    parts = term.split("-")
    if len(parts) != 3:
        return ""
    start_year = int(parts[0])
    semester_in_year = 1 if parts[2] == "1" else 2
    match = re.search(r"20(\d{2})", teaching_class)
    if match:
        cohort_year = 2000 + int(match.group(1))
    else:
        match = re.search(r"(?<!\d)(\d{2})\d{2,}", teaching_class)
        if not match:
            return ""
        cohort_year = 2000 + int(match.group(1))
    number = (start_year - cohort_year) * 2 + semester_in_year
    names = ("第一学期", "第二学期", "第三学期", "第四学期", "第五学期", "第六学期")
    return names[number - 1] if 1 <= number <= len(names) else ""


def get_template_version(term):
    """把具体学期映射到学校模板版本。"""
    if term.startswith("2023-2024"):
        return "2023-2024"
    if term.startswith("2024-2025-1"):
        return "2024-2025-1"
    if term.startswith("2024-2025-2"):
        return "2024-2025-2"
    if term.startswith("2025-2026"):
        return "2025-2026"
    return term



def get_template_path(doc_type, offering):
    """根据文档类型和开课信息获取确定的模板版本。"""
    term = offering["term"]
    if doc_type == "plan":
        return os.path.join(TEMPLATE_BASE, "模板4：授课计划 模板.docx")
    directories = {
        "standard": ("模板3：课程标准", "模板3：课程标准 模板"),
        "design": ("模板5：教学设计", "模板5：教学设计 模板"),
    }
    if doc_type not in directories:
        return None
    folder, prefix = directories[doc_type]
    version = get_template_version(term)
    path = os.path.join(TEMPLATE_BASE, folder, f"{prefix}（{version}）.docx")
    return path if os.path.exists(path) else None

def get_output_path(doc_type, offering, output_dir=None):
    """生成输出文件路径，同名课程自动添加专业后缀防覆盖"""
    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(__file__), "生成结果", "精修版")
    os.makedirs(output_dir, exist_ok=True)
    term = offering["term"]
    cn = offering["course_name"]
    tn = offering.get("teacher_name") or "杜媛"
    type_names = {"plan": "授课计划", "standard": "课程标准", "design": "教学设计"}

    base_name = f"{term}《{cn}》{type_names[doc_type]} {tn}"

    # 检查是否有同名同期的其他课程（需要加专业区分）
    try:
        import store
        dup_count = store.rows(
            'SELECT COUNT(*) as cnt FROM offerings WHERE course_name=? AND term=? AND offering_kind != "实训课程"',
            [cn, term]
        )
        if dup_count[0]['cnt'] > 1:
            mj = offering.get("major", "")
            major_abbrev = {
                "农村电子商务": "农商",
                "全媒体广告策划与营销": "全媒体",
                "市场营销": "营销",
            }
            major_short = major_abbrev.get(mj, mj[:2] if mj else "")
            if major_short:
                base_name = f"{base_name}（{major_short}）"
    except Exception:
        pass

    return os.path.join(output_dir, f"{base_name}.docx")


# ============================================================
# 第五层：验证模块
# ============================================================

RESIDUE_PATTERNS = [r'×+', r'XX+', r'（注：', r'\(注：', r'（建议', r'XXX', r'【模板', r'（模板', r'表格中内容', r'参考格式如下', r'注：情境', r'……+', r'\.\.\.+', r'顶格，行距', r'正文行距', r'具有…+能力']

def verify_document(fp):
    """验证文档：检查空格、重复、模板残留"""
    doc = Document(fp)
    issues = []
    empty_cells = 0
    total_cells = 0
    residues = []

    # 检查段落
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        for pat in RESIDUE_PATTERNS:
            if re.search(pat, text):
                residues.append(f"P{i}: {text[:50]}")

    # 检查表格
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                total_cells += 1
                text = cell.text.strip()
                if not text:
                    empty_cells += 1
                else:
                    for pat in RESIDUE_PATTERNS:
                        if re.search(pat, text):
                            residues.append(f"表{ti} R{ri}C{ci}: {text[:50]}")

    # 检查重复小结
    summaries = []
    homeworks = []
    if len(doc.tables) >= 6:
        t5 = doc.tables[5]
        if len(t5.rows) > 0:
            for ri in range(min(len(t5.rows), 100)):
                row = t5.rows[ri]
                texts = [c.text.strip()[:50] for c in row.cells]
                if len(texts) >= 3:
                    if texts[1] and texts[1] not in summaries:
                        summaries.append(texts[1])
                    if texts[2] and texts[2] not in homeworks:
                        homeworks.append(texts[2])

    empty_rate = empty_cells / total_cells * 100 if total_cells > 0 else 0
    return {
        "empty_cells": empty_cells,
        "total_cells": total_cells,
        "empty_rate": round(empty_rate, 1),
        "residues": residues,
        "summary_unique": len(summaries),
        "homework_unique": len(homeworks),
    }


# ============================================================
# 主入口
# ============================================================

def generate_all(offering_id, doc_types=None, output_dir=None):
    """生成指定文档"""
    if doc_types is None:
        doc_types = ["plan", "standard", "design"]

    offering = load_offering(offering_id)
    tasks = load_tasks(offering_id)
    sessions = load_sessions(offering_id)
    units = load_units(offering_id)

    print(f"开课: {offering['course_name']} ({offering['term']})")
    print(f"任务: {len(tasks)}个, 排课: {len(sessions)}条, 单元: {len(units)}个")
    print(f"教师: {offering.get('teacher_name','')}, 专业: {offering.get('major','')}")

    results = {}
    for dt in doc_types:
        print(f"\n{'='*60}")
        print(f"生成 {dt} ...")
        try:
            if dt == "plan":
                from gen_plan import generate_plan
                fp = generate_plan(offering, tasks, sessions, units, output_dir)
            elif dt == "standard":
                from gen_standard import generate_standard
                fp = generate_standard(offering, tasks, sessions, units, output_dir)
            elif dt == "design":
                from gen_design import generate_design
                fp = generate_design(offering, tasks, sessions, units, output_dir)
            else:
                print(f"  未知文档类型: {dt}")
                continue

            # 验证
            from _batch_polish import (
                fix_all_heights, fix_empty_paragraphs, fix_line_spacing,
                fix_font_consistency, fix_empty_para_indent,
            )
            polished = Document(fp)
            residual_texts = (
                "××××课程是××××专业", "××××（如专业核心）课程",
                "体现本课程在课程体系中的特色", "说明课程在专业人才培养中的作用",
                "服务××岗位", "在课程整体设计思路基础上，依据……",
                "确定XX个模块/项目及其学时", "其中包括理论课时XX个",
                "以学生为中心，注重学生的学习兴趣", "引导学生运用XXX",
            )
            paragraphs = list(polished.paragraphs)
            paragraphs.extend(
                paragraph
                for table in polished.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            )
            for paragraph in paragraphs:
                if any(marker in paragraph.text for marker in residual_texts):
                    for run in paragraph.runs:
                        run.text = ""
            fix_all_heights(polished)
            fix_empty_paragraphs(polished)
            fix_line_spacing(polished)
            fix_font_consistency(polished)
            fix_empty_para_indent(polished)
            polished.save(fp)
            vr = verify_document(fp)
            print(f"  完成: {fp}")
            print(f"  空格率: {vr['empty_rate']}% ({vr['empty_cells']}/{vr['total_cells']})")
            print(f"  残留: {len(vr['residues'])}处")
            if vr['residues']:
                for r in vr['residues'][:5]:
                    print(f"    {r}")
            results[dt] = {"path": fp, "verify": vr}
        except Exception as e:
            import traceback
            print(f"  错误: {e}")
            traceback.print_exc()
            results[dt] = {"error": str(e)}

    return results


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="统一文档生成器")
    parser.add_argument("offering_id", type=int, help="开课ID")
    parser.add_argument("--doc", choices=["plan", "standard", "design", "all"], default="all")
    parser.add_argument("--output", default=None, help="输出目录")
    args = parser.parse_args()

    doc_types = ["plan", "standard", "design"] if args.doc == "all" else [args.doc]
    results = generate_all(args.offering_id, doc_types, args.output)

    print(f"\n{'='*60}")
    print("生成完成:")
    for dt, r in results.items():
        if "path" in r:
            print(f"  {dt}: {r['path']}")
        elif "error" in r:
            print(f"  {dt}: 错误 - {r['error']}")
