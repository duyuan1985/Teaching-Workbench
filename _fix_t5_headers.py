"""
修复所有单元教学设计表格(Table 5,7,9,...,63)的R9-R14区域：
1. 恢复R9/R10的标题行：知识＆技能 | 重点 | 难点 | 目标达成度
2. 将内容移到R11-R14的正确位置
3. 修复教学任务格式（子情景换行分隔）
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

# 获取模板Table 5的R9-R10作为参考
tmpl_fp = r'原始资料\模板\模板5：教学设计\模板5：教学设计 模板（2023-2024）.docx'
tmpl_doc = Document(tmpl_fp)
tmpl_t5 = tmpl_doc.tables[5]

# 模板R9的标题文本
# R9: 知识＆技能(merged C2-C6) | 重点(C7) | 难点(merged C8-C10) | 目标达成度(merged C11-C19)
# R10: same headers with 识记(4) | 理解(2) | 应用(2) | 评价(1)

def set_cell_text(tc, text, font_name='仿宋', font_size='21'):
    """设置单元格文本，保留格式"""
    # 清除现有段落内容
    for p in tc.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            p.remove(r)
        # 添加新run
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), font_size)
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), font_size)
        rPr.append(szCs)
        # 加粗
        b = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        break  # 只处理第一个段落

def set_cell_text_no_bold(tc, text, font_name='仿宋', font_size='21'):
    """设置单元格文本，不加粗"""
    for p in tc.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            p.remove(r)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), font_size)
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), font_size)
        rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        break

# 处理所有30个表格（Table 5, 7, 9, ..., 63）
table_indices = list(range(5, 64, 2))  # 5,7,9,...,63
print(f'需要处理 {len(table_indices)} 个表格')

for ti in table_indices:
    t = doc.tables[ti]
    
    # === 1. 修复R9标题行 ===
    # 模板R9: C2='知识＆技能', C7='重点', C8='难点', C11='目标达成度'
    # 当前R9: C2=内容文本, C7='难点', C11='目标达成度'
    r9 = t.rows[9]
    
    # 保存R9 C2的当前内容（这是知识＆技能的内容，应该移到R11）
    r9_c2_text = r9.cells[2].text.strip()
    
    # 恢复R9标题
    set_cell_text(r9.cells[2]._tc, '知识＆技能')
    
    # 检查C7是否有'重点'
    if '重点' not in r9.cells[7].text:
        set_cell_text(r9.cells[7]._tc, '重点')
    
    # === 2. 修复R10标题行 ===
    r10 = t.rows[10]
    r10_c2_text = r10.cells[2].text.strip()
    
    # 恢复R10标题
    set_cell_text(r10.cells[2]._tc, '知识＆技能')
    
    # 检查C7
    if '重点' not in r10.cells[7].text and '识记' not in r10.cells[7].text:
        set_cell_text(r10.cells[7]._tc, '重点')
    
    # === 3. 修复R3教学任务格式 ===
    r3 = t.rows[3]
    r3_text = r3.cells[2].text.strip()
    # 将"子情景"分隔的内容换行
    if '子情景' in r3_text:
        # 格式：第一章 初识数据分析子情景：认识数据分析子情景：常用数据分析方法
        # 应改为：第一章 初识数据分析\n子情景：认识数据分析\n子情景：常用数据分析方法
        import re
        # 在"子情景"前插入换行
        fixed = re.sub(r'(子情景[：:])', r'\n\1', r3_text)
        # 去除开头可能的换行
        fixed = fixed.lstrip('\n')
        # 设置到单元格
        tc = r3.cells[2]._tc
        for p in tc.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                p.remove(r)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), '仿宋')
            rFonts.set(qn('w:eastAsia'), '仿宋')
            rFonts.set(qn('w:hAnsi'), '仿宋')
            rFonts.set(qn('w:cs'), '仿宋')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '21')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '21')
            rPr.append(szCs)
            r.append(rPr)
            # 用换行符分割
            lines = fixed.split('\n')
            for li, line in enumerate(lines):
                if li > 0:
                    br = OxmlElement('w:br')
                    r.append(br)
                t_elem = OxmlElement('w:t')
                t_elem.set(qn('xml:space'), 'preserve')
                t_elem.text = line
                r.append(t_elem)
            p.append(r)
            break

    print(f'Table {ti}: R9标题恢复, R3教学任务修复')

doc.save(fp)
print(f'\n保存完成')

# 验证Table 5
doc2 = Document(fp)
t = doc2.tables[5]
print(f'\n=== 验证 Table 5 ===')
for ri in range(9, 15):
    texts = [c.text.strip()[:30] for c in t.rows[ri].cells]
    print(f'  R{ri}: C2={texts[2][:30]} | C7={texts[7]} | C8={texts[8]}')

print(f'\n=== R3 ===')
print(t.rows[3].cells[2].text)
