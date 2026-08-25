"""
模板填充引擎 - 基于 template_slots 的通用模板填充

提供通用的模板填充能力，根据模板分析生成的槽位（template_slots）
自动定位并填充内容。主要处理"事实填写"类的简单字段，复杂的结构化内容
仍由各专用生成器处理。
"""

import json
import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

import store


# ============================================================
# 字体工具
# ============================================================

def _extract_run_font(run):
    """从 run 中提取字体信息（含东亚字体）"""
    font_name = ""
    size_pt = None
    bold = None

    if run.font.name:
        font_name = run.font.name
    if run.font.size:
        size_pt = run.font.size / 12700  # EMU → pt
    if run.font.bold is not None:
        bold = run.font.bold

    rpr = run._element.find(qn("w:rPr"))
    if rpr is not None:
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            ea = rfonts.get(qn("w:eastAsia"))
            if ea:
                font_name = ea

    return {"name": font_name, "size_pt": size_pt, "bold": bold}


def _apply_run_font(run, font_info):
    """将字体信息应用到 run"""
    name = font_info.get("name", "")
    size_pt = font_info.get("size_pt")
    bold = font_info.get("bold")

    if name:
        run.font.name = name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:eastAsia"), name)
        rfonts.set(qn("w:ascii"), name)
        rfonts.set(qn("w:hAnsi"), name)
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def _get_cell_font(cell):
    """从单元格中提取字体信息（取第一个有文字的 run）"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                return _extract_run_font(run)
    return None


def _get_paragraph_font(paragraph):
    """从段落中提取字体信息（取第一个有文字的 run）"""
    for run in paragraph.runs:
        if run.text.strip():
            return _extract_run_font(run)
    return None


def _set_paragraph_text(paragraph, text, font_info=None):
    """设置段落文本，保留或指定字体格式"""
    if paragraph.runs:
        if font_info is None:
            font_info = _extract_run_font(paragraph.runs[0])
        paragraph.runs[0].text = str(text)
        _apply_run_font(paragraph.runs[0], font_info)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(str(text))
        if font_info:
            _apply_run_font(run, font_info)


def _set_cell_text(cell, text, font_info=None):
    """设置单元格文本，保留或指定字体格式"""
    if not cell.paragraphs:
        return
    if font_info is None:
        font_info = _get_cell_font(cell)
    _set_paragraph_text(cell.paragraphs[0], text, font_info)
    for extra in cell.paragraphs[1:]:
        _set_paragraph_text(extra, "", font_info)


def _set_cell_font(cell, font_info):
    """设置单元格字体（兼容旧接口）"""
    if isinstance(font_info, str):
        font_info = {"name": font_info, "size_pt": 10.5, "bold": None}
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            _apply_run_font(run, font_info)


def _parse_locator(locator):
    """
    解析槽位定位器。
    返回: (location_type, table_index, row_index, col_index, direction)
    """
    if locator.startswith("paragraph:"):
        # paragraph:123
        parts = locator.split(":")
        return ("paragraph", int(parts[1]), None, None, None)
    
    if locator.startswith("table:"):
        # table:0/row:2/col:3/horizontal
        # table:0/col:2/header
        parts = locator.split("/")
        table_idx = int(parts[0].split(":")[1])
        
        row_idx = None
        col_idx = None
        direction = None
        
        for part in parts[1:]:
            if part.startswith("row:"):
                row_idx = int(part.split(":")[1])
            elif part.startswith("col:"):
                col_idx = int(part.split(":")[1])
            elif part in ("horizontal", "vertical"):
                direction = part
            elif part == "header":
                direction = "header"
        
        return ("table", table_idx, row_idx, col_idx, direction)
    
    return ("unknown", None, None, None, None)


def _get_slot_value(field_name, offering, tasks=None):
    """
    根据字段名获取填充值（简单事实类字段）。
    返回: 值字符串，或 None（表示该字段需由专用逻辑处理）
    """
    # 基本信息映射
    simple_mappings = {
        "课程名称": offering.get("course_name", ""),
        "课程基本信息": offering.get("course_name", ""),
        "课程类型": offering.get("course_type", ""),
        "课程性质": offering.get("course_nature", "") or "必修课",
        "课程代码": offering.get("course_code", ""),
        "适用专业": offering.get("major", ""),
        "编制单位": offering.get("major", "") + "教研室",
        "所属系部": offering.get("department", "") or store.get_setting("department", ""),
        "开设学期": offering.get("term", ""),
        "授课教师": offering.get("teacher_name", "") or store.get_setting("teacher_name", ""),
        "实训教师": "",
        "授课班级": offering.get("teaching_class", ""),
    }
    
    if field_name in simple_mappings:
        return simple_mappings[field_name]
    
    # 学时学分特殊处理
    if field_name == "学时学分":
        credits = offering.get("credits", 0)
        if float(credits).is_integer():
            credits = int(credits)
        return "%d学时（%s学分）" % (offering.get("total_hours", 0), credits)

    # 考核类型/考核方式
    if field_name in ("考核类型", "考核方式"):
        return offering.get("assessment_type", "") or store.get_setting("assessment_type", "")
    # 学时构成
    if field_name == "总学时":
        return str(offering.get("total_hours", ""))
    if field_name == "理论学时":
        return str(offering.get("lecture_hours", 0))
    if field_name == "实践学时":
        return str(offering.get("practice_hours", 0))
    if field_name == "学分":
        credits = offering.get("credits", 0)
        return str(int(credits)) if float(credits).is_integer() else str(credits)
    
    if field_name == "任务学时":
            task_hours = store.rows(
                "SELECT SUM(hours) as total FROM tasks WHERE offering_id=?",
                (offering.get("id"),)
            )
            if task_hours and task_hours[0]["total"]:
                return str(task_hours[0]["total"])
            return str(offering.get("total_hours", ""))
    
    # 需要专用逻辑处理的字段返回 None
    complex_fields = (
        "教学任务", "知识目标", "能力目标", "思政目标", "素质目标",
        "课程设计总体思路", "课程内容与学时", "考核评价",
        "教师知识能力要求", "教材选用", "教学资源",
        "教材学情分析及教育理念", "教学场景设计",
        "教学活动流程", "教法学法", "达成目标",
        "课堂小结", "课后作业", "教学反思",
        "教学时间", "周次", "授课日期", "教学环境",
    )
    if field_name in complex_fields:
        return None
    
    return None


def fill_simple_fields(template_path, output_path, offering_id):
    """
    填充模板中的简单事实字段。
    
    Args:
        template_path: 模板文件路径
        output_path: 输出文件路径
        offering_id: 课程实例ID
    
    Returns:
        filled_count: 成功填充的字段数
        skipped_count: 跳过的字段数（需专用逻辑处理）
    """
    offering = store.rows("SELECT * FROM offerings WHERE id=?", (offering_id,))
    if not offering:
        raise ValueError("课程实例不存在。")
    offering = offering[0]
    
    # 获取模板文件ID
    template_files = store.rows(
        "SELECT id FROM template_files WHERE offering_id=? AND template_path=?",
        (offering_id, str(Path(template_path).resolve()))
    )
    if not template_files:
        # 尝试模糊匹配
        template_files = store.rows(
            "SELECT id FROM template_files WHERE offering_id=? ORDER BY id",
            (offering_id,)
        )
        if not template_files:
            return 0, 0
    
    template_file_id = template_files[0]["id"]
    slots = store.rows(
        "SELECT * FROM template_slots WHERE template_file_id=? ORDER BY slot_key",
        (template_file_id,)
    )
    
    if not slots:
        return 0, 0
    
    # 复制模板
    shutil.copy2(template_path, output_path)
    document = Document(output_path)
    
    filled_count = 0
    skipped_count = 0
    
    for slot in slots:
        field_name = slot["field_name"]
        locator = slot["locator"]
        content_kind = slot["content_kind"]

        # 获取值
        value = _get_slot_value(field_name, offering)

        if value is None:
            skipped_count += 1
            continue

        # 解析定位器
        loc_type, table_idx, row_idx, col_idx, direction = _parse_locator(locator)

        try:
            if loc_type == "paragraph":
                # 段落级填充
                if row_idx is not None and row_idx < len(document.paragraphs):
                    paragraph = document.paragraphs[row_idx]
                    # 提取原有字体
                    font_info = _get_paragraph_font(paragraph)
                    # 对于带标签的段落，替换标签后的内容
                    text = paragraph.text.strip()
                    # 尝试替换"标签：值"格式
                    if "：" in text or ":" in text:
                        separator = "：" if "：" in text else ":"
                        label = text.split(separator)[0] + separator
                        _set_paragraph_text(paragraph, label + str(value), font_info)
                    else:
                        _set_paragraph_text(paragraph, str(value), font_info)
                    filled_count += 1

            elif loc_type == "table":
                # 表格级填充
                if table_idx is not None and table_idx < len(document.tables):
                    table = document.tables[table_idx]

                    if direction == "horizontal" and row_idx is not None and col_idx is not None:
                        # 水平标签-值对：从标签列读取字体，应用到值列
                        if row_idx < len(table.rows):
                            cells = table.rows[row_idx].cells
                            if col_idx < len(cells):
                                # 从标签单元格读取字体
                                label_font = None
                                if col_idx > 0 and cells[col_idx - 1].text.strip():
                                    label_font = _get_cell_font(cells[col_idx - 1])
                                if label_font is None:
                                    label_font = _get_cell_font(cells[col_idx])
                                _set_cell_text(cells[col_idx], str(value), label_font)
                                filled_count += 1

                    elif direction == "vertical" and col_idx is not None:
                        # 垂直表头：从表头单元格读取字体，应用到值单元格
                        if len(table.rows) > 1 and col_idx < len(table.rows[1].cells):
                            header_font = _get_cell_font(table.rows[0].cells[col_idx])
                            _set_cell_text(table.rows[1].cells[col_idx], str(value), header_font)
                            filled_count += 1
        except Exception as e:
            # 单个槽位填充失败不影响整体
            skipped_count += 1
            continue
    
    # 清理模板中的说明文字和占位符。
    # 顺序：先填标题占位符（防止含***的标题段被当占位内容删除），
    # 再做段落级清理（纯指令段整段删除），最后内联清理（保留标题、只删括号说明）。
    _fill_title_placeholders(document, offering["course_name"])
    clean_placeholders(document)
    clean_inline_instructions(document)
    _remove_punct_only_paragraphs(document)

    document.save(output_path)
    return filled_count, skipped_count


def fill_paragraph_after(document, heading_text, content_lines, style_paragraph=None):
    """
    在指定标题后插入内容段落。
    
    Args:
        document: Document 对象
        heading_text: 标题文本（精确匹配）
        content_lines: 要插入的内容行列表
        style_paragraph: 样式参考段落
    """
    heading = next((p for p in document.paragraphs if p.text.strip() == heading_text), None)
    if not heading:
        return 0
    
    cursor = heading
    inserted = 0
    for text in content_lines:
        if not text:
            continue
        new_p = deepcopy(heading._p) if style_paragraph is None else deepcopy(style_paragraph._p)
        # 清除文本
        for r in new_p.findall(qn("w:r")):
            new_p.remove(r)
        # 添加新文本
        new_run = heading.add_run("")._r if heading.runs else None
        if new_run is None:
            new_run = heading.add_run("")._r
        # 实际上用 python-docx 的方式更简单
        cursor = _insert_after_p(cursor, text, document)
        inserted += 1
    
    return inserted


def _insert_after_p(paragraph, text, document):
    """在段落后插入新段落（内部工具）"""
    new_p = paragraph._p.addnext(deepcopy(paragraph._p))
    # 清除所有 run
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    # 添加新的 run
    new_run = paragraph._p.makeelement(qn("w:r"), {})
    new_t = paragraph._p.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
    new_t.text = str(text)
    new_run.append(new_t)
    new_p.append(new_run)
    
    # 返回新段落对象
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


# 说明性括号内的关键词（出现任意一个即判定为说明文字）
_INSTRUCTION_KEYWORDS = (
    "注", "建议", "说明", "体现", "如", "若", "不限于", "针对", "要求",
    "此部分", "文字+表格", "理论课为", "实践类", "基于", "即合作",
    "可展示", "填写", "居中", "第一人为", "结合", "简要", "包括",
    "根据课程", "需体现", "需", "应", "可", "以", "描述",
    "保留", "删除", "增减", "选择", "比较", "决定",
    # 排版格式指令（如"(顶格，行距23，黑体、小四号，下同)"）
    "顶格", "行距", "黑体", "宋体", "仿宋", "楷体", "加粗",
    "段后", "段前", "缩进", "下同", "字号", "居左", "居右",
)

# 不应被清理的括号内容（版本号、章节编号、权重等）
_KEEP_PATTERNS = (
    re.compile(r"^\d{4}版?$"),  # 2024版
    re.compile(r"^[一二三四五六七八九十]+$"),  # （一）（二）
    re.compile(r"^\d+%$"),  # 30%
    re.compile(r"^\d+分$"),  # 30分
    re.compile(r"^必修$|^选修$"),  # 必修/选修
)


_PUNCT_ONLY = re.compile(r"^[\s。、，；：！？．.,;:!?·]+$")


def _remove_punct_only_paragraphs(document):
    """删除内联清理后只剩标点的段落（如"（注：…）。（可以多个）"清完括号后残余的"。"）。"""
    removed = 0
    for paragraph in list(document.paragraphs):
        text = paragraph.text
        if text and text.strip() and _PUNCT_ONLY.match(text.strip()):
            paragraph._p.getparent().remove(paragraph._p)
            removed += 1
    return removed


def _fill_title_placeholders(document, course_name):
    """把模板中的标题占位符替换为实际课程名。

    必须在 clean_placeholders 之前执行：含 "**********"/"××" 的段落会被
    占位清理整段删除，导致后续 _replace_prefix 无段可填、文档丢失标题。
    """
    replacements = {
        "《**********》": f"《{course_name}》",
        "《×××》": f"《{course_name}》",
        "《xxx》": f"《{course_name}》",
        "《XXX》": f"《{course_name}》",
    }
    replaced = 0

    def _apply(paragraph):
        nonlocal replaced
        text = paragraph.text
        if not text:
            return
        new_text = text
        for old, new in replacements.items():
            if old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != text:
            font_info = _get_paragraph_font(paragraph)
            _set_paragraph_text(paragraph, new_text, font_info)
            replaced += 1

    for paragraph in document.paragraphs:
        _apply(paragraph)
    for table in document.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                for paragraph in cell.paragraphs:
                    _apply(paragraph)
    return replaced


def _is_instruction_paren(text):
    """判断括号内文字是否为说明性内容"""
    text = text.strip()
    if not text:
        return False
    # 检查是否匹配保留模式
    for pattern in _KEEP_PATTERNS:
        if pattern.match(text):
            return False
    # 检查是否包含说明关键词
    for kw in _INSTRUCTION_KEYWORDS:
        if kw in text:
            return True
    # 检查是否为纯占位符（×××, ***, XX等）
    if re.fullmatch(r"[\s×*\u2606\u2605X×]+", text):
        return True
    return False


def clean_inline_instructions(document):
    """
    清理段落中的内嵌说明文字（括号内的提示）。
    保留标题/标签部分，只移除说明性括号内容。

    例如：
      【教学场景设计】（理论课为××教室...）→ 【教学场景设计】
      适用专业：     （不限于一个）→ 适用专业：
      《**********》（居中，与专业...一致）→ 《**********》
    """
    cleaned = 0
    for paragraph in document.paragraphs:
        text = paragraph.text
        if not text or not text.strip():
            continue

        # 提取原有字体信息
        font_info = _get_paragraph_font(paragraph)

        # 查找所有括号对
        new_text = text
        changed = False

        # 处理中文括号和英文括号
        for pair in [("（", "）"), ("(", ")")]:
            open_ch, close_ch = pair
            # 逐个查找括号对
            search_start = 0
            while True:
                open_idx = new_text.find(open_ch, search_start)
                if open_idx == -1:
                    break
                close_idx = new_text.find(close_ch, open_idx + 1)
                if close_idx == -1:
                    break
                inner = new_text[open_idx + 1:close_idx]
                if _is_instruction_paren(inner):
                    # 移除说明括号及其前后多余空格
                    before = new_text[:open_idx].rstrip()
                    after = new_text[close_idx + 1:].lstrip()
                    new_text = before + after
                    changed = True
                    search_start = len(before)
                else:
                    search_start = close_idx + 1

        if changed:
            new_text = new_text.rstrip()
            if new_text != text.strip():
                _set_paragraph_text(paragraph, new_text, font_info)
                cleaned += 1

    # 同样处理表格单元格
    for table in document.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if not text or not text.strip():
                        continue
                    font_info = _get_paragraph_font(paragraph)
                    new_text = text
                    changed = False
                    for pair in [("（", "）"), ("(", ")")]:
                        open_ch, close_ch = pair
                        search_start = 0
                        while True:
                            open_idx = new_text.find(open_ch, search_start)
                            if open_idx == -1:
                                break
                            close_idx = new_text.find(close_ch, open_idx + 1)
                            if close_idx == -1:
                                break
                            inner = new_text[open_idx + 1:close_idx]
                            if _is_instruction_paren(inner):
                                before = new_text[:open_idx].rstrip()
                                after = new_text[close_idx + 1:].lstrip()
                                new_text = before + after
                                changed = True
                                search_start = len(before)
                            else:
                                search_start = close_idx + 1
                    if changed:
                        new_text = new_text.rstrip()
                        if new_text != text.strip():
                            _set_paragraph_text(paragraph, new_text, font_info)
                            cleaned += 1

    return cleaned


def clean_placeholders(document):
    """
    清理模板中的占位内容（注释、说明、示例等）。

    Returns:
        removed_count: 移除的段落数
    """
    placeholder_prefixes = (
        "注：", "（注", "(注", "说明：", "（说明", "(说明",
        "示例一", "示例二", "示例：", "示例",
        "（针对", "(针对", "（要求", "(要求",
        "（教学组织", "(教学组织", "（表格中", "(表格中",
        "（基于", "(基于", "（即合作", "评价建议",
        "文字+表格", "XXXX教学法", "××",
        "……", "填写说明", "参考格式",
    )

    removed = 0
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if (any(text.startswith(p) for p in placeholder_prefixes) or
            "填写说明" in text or "××" in text or "……" in text or
            "**********" in text):
            paragraph._p.getparent().remove(paragraph._p)
            removed += 1

    return removed


def get_template_fill_data(offering_id, document_type):
    """
    获取模板填充所需的所有数据。
    这是一个统一的数据收集接口，供各文档生成器使用。
    """
    offering = store.rows("SELECT * FROM offerings WHERE id=?", (offering_id,))[0]
    tasks = store.rows("SELECT * FROM tasks WHERE offering_id=? ORDER BY seq", (offering_id,))
    units = store.rows(
        "SELECT * FROM curriculum_units WHERE offering_id=? AND approval_status='已确认' AND review_action<>'删除' ORDER BY seq",
        (offering_id,)
    )
    
    theory_hours = sum(int(t.get("theory_hours") or 0) for t in tasks)
    practice_hours = sum(int(t.get("practice_hours") or 0) for t in tasks)
    
    return {
        "offering": offering,
        "tasks": tasks,
        "units": units,
        "theory_hours": theory_hours,
        "practice_hours": practice_hours,
        "total_hours": offering.get("total_hours", 0),
        "project_names": list(dict.fromkeys(t["chapter"] for t in tasks if t["chapter"])),
    }
