"""
设置表5和表6所有单元格格式为仿宋五号（10.5pt）
"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)

count = 0
for ti in range(5, len(doc.tables), 2):
    for t in [doc.tables[ti], doc.tables[ti+1] if ti+1 < len(doc.tables) else None]:
        if t is None:
            continue
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = '仿宋'
                        rPr = r._element.get_or_add_rPr()
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is None:
                            rFonts = rPr.makeelement(qn('w:rFonts'), {})
                            rPr.insert(0, rFonts)
                        rFonts.set(qn('w:eastAsia'), '仿宋')
                        rFonts.set(qn('w:ascii'), '仿宋')
                        rFonts.set(qn('w:hAnsi'), '仿宋')
                        r.font.size = Pt(10.5)
                        # 清除缩进
                        pPr = p._element.find(qn('w:pPr'))
                        if pPr is not None:
                            ind = pPr.find(qn('w:ind'))
                            if ind is not None:
                                pPr.remove(ind)
                        count += 1

print(f"设置了{count}个run的格式为仿宋五号")
doc.save(fp)
print(f"保存完成: {fp}")
