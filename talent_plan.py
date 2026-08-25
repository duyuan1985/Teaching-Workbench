"""人才培养方案解析：提取培养目标、培养规格、课程设置表，供内容生成引用。

支持 docx（python-docx）与 pdf（pdfplumber）。方案按专业名匹配开课；
课程信息按课程名查找（代码/类别/学分/学时）。
"""

import re
from pathlib import Path

import config
import store

PLAN_DIR_NAME = "人才培养方案"
CODE_RE = re.compile(r"^[A-Z]\d{4,6}$")
CATEGORY_KEYWORDS = ("公共基础", "专业基础", "专业核心", "专业拓展", "专业（技能）", "综合实践", "人文素质")
CATEGORY_MAP = {
    "公共基础课程": "公共基础课",
    "专业基础课程": "专业基础课",
    "专业核心课程": "专业核心课",
    "专业拓展课程": "专业拓展课",
    "综合实践课程": "综合实践课",
    "人文素质拓展课程": "人文素质拓展课",
}
GOAL_MARKERS = ("（一）培养目标", "(一)培养目标")
SPEC_START_MARKERS = ("（二）培养规格", "(二)培养规格")
SPEC_END_MARKERS = ("六、课程设置及要求", "六、课程设置及要求", "七、", "六、课程设置")
MAJOR_RE = re.compile(r"专业名称[：:]\s*([\u4e00-\u9fa5A-Za-z]+)")

ORIENTATION_HEADERS = {
    "industries": ("对应行业",),
    "job_categories": ("主要职业类别",),
    "job_positions": ("主要岗位群", "主要岗位群或技术领域举例"),
    "certificates": ("职业资格证书", "职业资格或技能等级证书"),
}


JOB_TITLE_RE = re.compile(r"[\u4e00-\u9fa5]{2,5}(?:主管|专员|经理|设计师|店长|工程师|师|员)")


def _split_items(text):
    """拆分岗位/类别列表：顿号/分号/冒号优先（PDF 断行在词内无标点），否则按换行。"""
    text = (text or "").strip()
    if not text:
        return []
    if any(sep in text for sep in "、；;：:"):
        parts = re.split(r"[、；;：:]", text)
    else:
        parts = text.splitlines()
    items = []
    for part in parts:
        clean = re.sub(r"\s+", "", part)
        clean = clean.rstrip("等。；;")
        if not clean or clean == "等":
            continue
        # 源文档个别岗位连写无分隔（如“渠道销售主管直播销售员”），按职务后缀拆开
        if len(clean) >= 10 and re.fullmatch(r"[\u4e00-\u9fa5]+", clean):
            titles = JOB_TITLE_RE.findall(clean)
            if titles and "".join(titles) == clean:
                items.extend(titles)
                continue
        items.append(clean)
    return items


def _extract_orientation(rows_cells):
    """从职业面向表（表头+数据行）提取行业/职业类别/岗位群/证书。"""
    header = [_norm(c) for c in rows_cells[0]]
    if not any("主要岗位群" in h for h in header):
        return None
    data_row = rows_cells[1] if len(rows_cells) > 1 else []
    orientation = {}
    for field, keys in ORIENTATION_HEADERS.items():
        col = next((i for i, h in enumerate(header) if any(k in h for k in keys)), None)
        if col is None or col >= len(data_row):
            continue
        items = _split_items(data_row[col])
        if items:
            orientation[field] = items
    return orientation or None


def _norm(text):
    return re.sub(r"\s+", "", text or "")


def _category_of(text):
    normed = _norm(text)
    for keyword in CATEGORY_KEYWORDS:
        if keyword in normed:
            for full, short in CATEGORY_MAP.items():
                if _norm(full) == normed or (keyword in full and keyword in normed):
                    return short if _norm(full) == normed else short
            return normed
    return ""


def _dedup_adjacent(texts):
    out = []
    for t in texts:
        if not out or out[-1] != t:
            out.append(t)
    return out


def _semester_numbers(cells, semester_cols):
    """按表头学期列索引提取各学期周学时：{学期号: 周学时}。"""
    result = {}
    for sem_no, col in semester_cols.items():
        if col < len(cells):
            value = _norm(cells[col])
            if re.match(r"^[\d.]+$", value):
                result[sem_no] = int(float(value))
    return result


def _parse_course_row(cells, category, semester_cols=None):
    """从一行单元格提取课程信息（先邻位去重消除水平合并重复）。"""
    texts = [(c or "").strip() for c in cells]
    code_idx = next((i for i, t in enumerate(texts) if CODE_RE.match(_norm(t))), None)
    if code_idx is None:
        return None
    deduped = _dedup_adjacent(texts[code_idx + 1:])
    name = ""
    name_end = 0
    for i, t in enumerate(deduped):
        clean = _norm(t)
        if clean and not re.match(r"^[\d.]+$", clean) and "合计" not in clean:
            name = clean
            name_end = i
            break
    if not name or "合计" in name:
        return None
    numbers = []
    for t in deduped[name_end + 1:]:
        clean = _norm(t)
        if re.match(r"^[\d.]+$", clean):
            numbers.append(float(clean))
        elif clean and not re.match(r"^[√*]+$", clean):
            break
    credits = int(numbers[0]) if numbers else None
    hours = int(numbers[1]) if len(numbers) > 1 else None
    semesters = _semester_numbers(texts, semester_cols or {})
    return {
        "name": name, "code": _norm(texts[code_idx]), "category": category,
        "credits": credits, "hours": hours, "semesters": semesters,
    }


def _row_cells(row):
    if hasattr(row, "cells"):
        return [c.text for c in row.cells]
    return list(row)


def _find_semester_columns(rows):
    """在表10表头行中定位学期列（一~六）的物理列索引。"""
    for row in rows[:6]:
        cells = [_norm(c) for c in _row_cells(row)]
        cols = {}
        for i, cell in enumerate(cells):
            if cell in ("一", "二", "三", "四", "五", "六"):
                cols["一二三四五六".index(cell) + 1] = i
        if len(cols) >= 4:
            return cols
    return {}


def _parse_docx(path):
    from docx import Document

    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    def section_between(start_markers, end_markers):
        start = next((i for i, t in enumerate(paragraphs) if _norm(t) in { _norm(m) for m in start_markers }), None)
        if start is None:
            return ""
        end = next((i for i in range(start + 1, len(paragraphs)) if _norm(paragraphs[i]) in {_norm(m) for m in end_markers} or any(_norm(paragraphs[i]).startswith(_norm(m)) for m in end_markers)), None)
        return "\n".join(t for t in paragraphs[start + 1:end or len(paragraphs)] if t)

    goals = section_between(GOAL_MARKERS, SPEC_START_MARKERS)
    specs = section_between(SPEC_START_MARKERS, SPEC_END_MARKERS)

    courses = []
    current_category = ""
    orientation = None
    for table in doc.tables:
        if orientation is None:
            rows_cells = [[c.text for c in row.cells] for row in table.rows]
            orientation = _extract_orientation(rows_cells)
        header_cells = [_norm(c.text) for c in table.rows[1].cells] if len(table.rows) > 1 else []
        if not ("课程类别" in header_cells and "课程代码" in header_cells):
            continue
        semester_cols = _find_semester_columns(table.rows)
        for row in table.rows[2:]:
            cells = [c.text for c in row.cells]
            first = _norm(cells[0]) if cells else ""
            if first and any(k in first for k in CATEGORY_KEYWORDS):
                current_category = _category_of(first)
            course = _parse_course_row(cells, current_category, semester_cols)
            if course:
                courses.append(course)
        break
    return goals, specs, courses, orientation


def _parse_pdf(path):
    import pdfplumber

    goals, specs = "", ""
    courses = []
    current_category = ""
    with pdfplumber.open(path) as pdf:
        first_page_text = (pdf.pages[0].extract_text() or "") if pdf.pages else ""
        all_lines = []
        for page in pdf.pages:
            text = page.extract_text() or ""
            all_lines.extend(line.strip() for line in text.splitlines())
        joined = "\n".join(all_lines)

        def between(start_markers, end_markers):
            start = next((joined.find(m) for m in start_markers if joined.find(m) >= 0), -1)
            if start < 0:
                return ""
            rest = joined[start:]
            end = next((rest.find(m, len(start_markers[0])) for m in end_markers if rest.find(m, len(start_markers[0])) >= 0), -1)
            return rest[len(start_markers[0]):end] if end > 0 else rest[len(start_markers[0]):]

        goals = between(GOAL_MARKERS, SPEC_START_MARKERS)
        specs = between(SPEC_START_MARKERS, SPEC_END_MARKERS)

        orientation = None
        for page in pdf.pages:
            for table in page.extract_tables():
                if not table or not table[0]:
                    continue
                if orientation is None and "主要岗位群" in _norm("".join((c or "") for c in table[0])):
                    orientation = _extract_orientation([[(c or "") for c in row] for row in table[:2]])
                flat = _norm("".join((c or "") for c in table[0]))
                has_code = any(CODE_RE.match(_norm(c or "")) for row in table for c in row)
                if "课程类别" not in flat and "课程代码" not in flat and not has_code:
                    continue
                semester_cols = _find_semester_columns(table)
                for row in table:
                    first = _norm(row[0]) if row and row[0] else ""
                    if first and any(k in first for k in CATEGORY_KEYWORDS):
                        current_category = _category_of(first)
                    course = _parse_course_row(row, current_category, semester_cols)
                    if course:
                        courses.append(course)
    seen = set()
    unique = []
    for course in courses:
        if course["code"] not in seen:
            seen.add(course["code"])
            unique.append(course)
    return goals, specs, unique, orientation, first_page_text


def parse_plan_file(path):
    path = Path(path)
    text_head = ""
    if path.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(path)
        text_head = "\n".join(p.text for p in doc.paragraphs[:20])
        goals, specs, courses, orientation = _parse_docx(path)
    else:
        goals, specs, courses, orientation, text_head = _parse_pdf(path)
    major_match = MAJOR_RE.search(text_head) or MAJOR_RE.search(goals + specs)
    major = major_match.group(1) if major_match else ""
    cohort_match = re.search(r"(20\d{2})级", path.stem + text_head[:200])
    cohort = cohort_match.group(1) + "级" if cohort_match else ""
    return {
        "major": major,
        "cohort": cohort,
        "goals": goals.strip(),
        "specs": specs.strip(),
        "orientation": orientation,
        "courses": courses,
        "source_path": str(path),
    }


def sync_talent_plans(directory):
    """扫描人才培养方案目录，解析并入库（按 major+cohort 去重更新）。"""
    root = Path(directory)
    if not root.is_dir():
        return {"added": [], "updated": [], "failed": []}
    added, updated, failed = [], [], []
    with store.connect() as db:
        for path in sorted(root.rglob("*")):
            if not path.is_file() or path.suffix.lower() not in (".docx", ".pdf"):
                continue
            try:
                plan = parse_plan_file(path)
            except Exception as exc:
                failed.append(f"{path.name}: {exc}")
                continue
            if not plan["major"]:
                failed.append(f"{path.name}: 未识别专业名称")
                continue
            existing = db.execute(
                "SELECT id FROM talent_plans WHERE major=? AND cohort=?",
                (plan["major"], plan["cohort"]),
            ).fetchone()
            payload = _payload(plan)
            if existing:
                db.execute(
                    "UPDATE talent_plans SET source_path=?, plan_json=?, updated_at=CURRENT_TIMESTAMP WHERE id=?",
                    (plan["source_path"], payload, existing["id"]),
                )
                updated.append(f"{plan['major']}（{plan['cohort']}）")
            else:
                db.execute(
                    "INSERT INTO talent_plans(major,cohort,source_path,plan_json) VALUES (?,?,?,?)",
                    (plan["major"], plan["cohort"], plan["source_path"], payload),
                )
                added.append(f"{plan['major']}（{plan['cohort']}）")
        db.commit()
    return {"added": added, "updated": updated, "failed": failed}


def _payload(plan):
    import json
    return json.dumps(plan, ensure_ascii=False)


def get_plan_for_major(major):
    """按专业取人才培养方案；无同名方案专业按 MAJOR_PLAN_ALIASES 参照相近专业方案。"""
    actual_major = config.MAJOR_PLAN_ALIASES.get(major, major)
    rows = store.rows("SELECT * FROM talent_plans WHERE major=? ORDER BY cohort DESC", (actual_major,))
    if not rows:
        return None
    import json
    return json.loads(rows[0]["plan_json"])


def find_course(plan, course_name):
    """在方案课程表中查找课程：先精确匹配，再按名称包含（长名优先，避免“实训”误配）。"""
    if not plan:
        return None
    aliases = {"Pyhton程序设计": "Python程序设计"}
    target = _norm(course_name).replace("《", "").replace("》", "")
    target = aliases.get(target, target)
    courses = plan.get("courses", [])
    for course in courses:
        if _norm(course["name"]) == target:
            return course
    candidates = [c for c in courses if target in _norm(c["name"]) or _norm(c["name"]) in target]
    if candidates:
        return max(candidates, key=lambda c: len(_norm(c["name"])))
    return None
