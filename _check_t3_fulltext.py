from docx import Document

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)
t3 = doc.tables[3]

# 检查几个单元格的完整文本
print('=== R3 完整文本 ===')
for ci in range(len(t3.rows[3].cells)):
    txt = t3.cell(3, ci).text.strip()
    print(f'  C{ci} [{len(txt)} chars]: {txt[:200]}')
    if len(txt) > 200:
        print(f'    ... (total {len(txt)} chars)')

print('\n=== R4 完整文本 ===')
for ci in range(len(t3.rows[4].cells)):
    txt = t3.cell(4, ci).text.strip()
    print(f'  C{ci} [{len(txt)} chars]: {txt[:200]}')

print('\n=== R31 完整文本 ===')
for ci in range(len(t3.rows[31].cells)):
    txt = t3.cell(31, ci).text.strip()
    print(f'  C{ci} [{len(txt)} chars]: {txt[:200]}')

# 检查教学方法段落是否有截断
print('\n=== 教学方法段落 ===')
for pi, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if '教学模式' in txt or '教学法' in txt or '教学方法' in txt or '数字技术' in txt:
        print(f'P{pi} [{len(txt)} chars]: {txt[:300]}')
