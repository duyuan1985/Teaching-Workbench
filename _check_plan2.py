"""详细检查授课计划"""
from docx import Document

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》授课计划 杜媛.docx"
doc = Document(fp)

# 封面P16完整内容
for i, p in enumerate(doc.paragraphs[:20]):
    text = p.text.strip()
    if "总学时" in text:
        print(f"P{i} 完整: {text}")

# 表0前5行
t0 = doc.tables[0]
print(f"\n表0: {len(t0.rows)}行 x {len(t0.rows[0].cells)}列")
for ri in range(min(8, len(t0.rows))):
    cells = t0.rows[ri].cells
    texts = [c.text.strip()[:25] for c in cells]
    print(f"  R{ri}: {texts}")

# 表1补充说明
t1 = doc.tables[1]
print(f"\n表1: {len(t1.rows)}行")
cell = t1.rows[0].cells[0]
for pi, p in enumerate(cell.paragraphs[:6]):
    text = p.text.strip()[:60]
    print(f"  P{pi}: {text}")

# 表2德育
t2 = doc.tables[2]
print(f"\n表2: {len(t2.rows)}行")
cell = t2.rows[0].cells[0]
for pi, p in enumerate(cell.paragraphs[:4]):
    text = p.text.strip()[:60]
    # 检查缩进
    from docx.oxml.ns import qn
    pPr = p._element.find(qn('w:pPr'))
    ind_info = ""
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            flc = ind.get(qn('w:firstLineChars'))
            ind_info = f" [firstLineChars={flc}]"
        else:
            ind_info = " [无ind元素]"
    else:
        ind_info = " [无pPr]"
    print(f"  P{pi}: {text}{ind_info}")
