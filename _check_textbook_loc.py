from docx import Document

# 教学设计
fp_d = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc_d = Document(fp_d)
print('=== 教学设计 教材段落 ===')
for pi, p in enumerate(doc_d.paragraphs):
    txt = p.text.strip()
    if '教材' in txt or '天津大学' in txt or '大数据分析方法' in txt:
        print(f'P{pi}: [{len(txt)}] {txt[:300]}')

# 课程标准
fp_s = r'生成结果\精修版\2023-2024-2《商务数据分析》课程标准 杜媛.docx'
doc_s = Document(fp_s)
print('\n=== 课程标准 教材段落 ===')
for pi, p in enumerate(doc_s.paragraphs):
    txt = p.text.strip()
    if '教材' in txt or '天津大学' in txt or '大数据分析方法' in txt:
        print(f'P{pi}: [{len(txt)}] {txt[:300]}')
