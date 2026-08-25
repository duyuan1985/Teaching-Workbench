from docx import Document
from docx.oxml.ns import qn

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)
print(f'总表格数: {len(doc.tables)}')

# 检查Table 62
t = doc.tables[62]
r3_txt = ''
for tc in t.rows[3]._tr.findall(qn('w:tc')):
    txt = ''
    for p in tc.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            for elem in r:
                if elem.tag == qn('w:t'):
                    txt += elem.text or ''
                elif elem.tag == qn('w:br'):
                    txt += '|'
    if txt:
        r3_txt = txt
print(f'Table 62 R3: {r3_txt[:100]}')

# 检查表格类型：奇数索引是大表(20列)，偶数是小表(6列)
for ti in range(60, 63):
    t = doc.tables[ti]
    print(f'Table {ti}: {len(t.rows)} rows, {len(t.columns)} cols')
