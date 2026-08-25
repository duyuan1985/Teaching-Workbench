from docx import Document

# 检查课程标准中课程内容划分和课时分配
fp_s = r'生成结果\精修版\2023-2024-2《商务数据分析》课程标准 杜媛.docx'
doc_s = Document(fp_s)

print('=== 课程标准 课程内容划分 ===')
in_section = False
for pi, p in enumerate(doc_s.paragraphs):
    txt = p.text.strip()
    if '课程内容划分' in txt or '课时分配' in txt:
        in_section = True
    if in_section:
        print(f'P{pi}: {txt[:200]}')
    if in_section and '3、教学方法' in txt:
        break

# 检查课程标准中的学习情境表
print('\n=== 课程标准 学习情境表 ===')
for ti, t in enumerate(doc_s.tables):
    first_cell = t.rows[0].cells[0].text.strip() if t.rows else ''
    if '序号' in first_cell or '学习情境' in first_cell or '情境' in first_cell:
        print(f'\nTable {ti}: rows={len(t.rows)}')
        for ri in range(len(t.rows)):
            texts = [c.text.strip()[:40] for c in t.rows[ri].cells]
            print(f'  R{ri}: {texts}')

# 检查教学设计中的学习情境表
print('\n=== 教学设计 学习情境表 ===')
fp_d = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc_d = Document(fp_d)
for ti, t in enumerate(doc_d.tables):
    first_cell = t.rows[0].cells[0].text.strip() if t.rows else ''
    if '情景' in first_cell or '学时' in first_cell or '情境' in first_cell:
        print(f'\nTable {ti}: rows={len(t.rows)}')
        for ri in range(len(t.rows)):
            texts = [c.text.strip()[:40] for c in t.rows[ri].cells]
            print(f'  R{ri}: {texts}')

# 检查授课计划中的进度表
print('\n=== 授课计划 进度表 ===')
fp_p = r'生成结果\精修版\2023-2024-2《商务数据分析》授课计划 杜媛.docx'
doc_p = Document(fp_p)
for ti, t in enumerate(doc_p.tables):
    first_cell = t.rows[0].cells[0].text.strip() if t.rows else ''
    if '周' in first_cell or '次' in first_cell or '序号' in first_cell:
        print(f'\nTable {ti}: rows={len(t.rows)}, cols={len(t.columns)}')
        for ri in range(min(5, len(t.rows))):
            texts = [c.text.strip()[:20] for c in t.rows[ri].cells]
            print(f'  R{ri}: {texts}')
        if len(t.rows) > 5:
            print(f'  ... ({len(t.rows)} rows total)')
