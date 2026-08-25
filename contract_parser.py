"""模板库契约解析器：目录扫描入库 + 三契约（结构/格式/内容）解析。

三契约对应方案第 4 章：
- 结构契约 structural_json：表格角色/合并规则/循环结构/封面结构
- 格式契约 format_json：槽位字体/字号/对齐/下划线/eastAsia（含封面字段下划线提议）
- 内容契约 content_json：模板指令 + 参考格式结构化 + 官方完善要求（四新条款）

槽位分类 classification：A=事实提取 / B=AI润色 / C=结构生成 / 人工。
"""

import hashlib
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

import store
from template_analyzer import (
    FIELD_SPECS,
    HEADING_RE,
    _cell_merge,
    _cell_text,
    _format,
    _identify_label_value_pairs,
    _match_field,
    _row_unique_cells_with_indices,
    _table_role,
)

KIND_CLASSIFICATION = {
    "事实填写": "A",
    "按班事实填写": "A",
    "按排课填写": "A",
    "可选事实填写": "A",
    "分析撰写": "B",
    "事实与分析": "B",
    "结构生成": "C",
    "按任务填写": "C",
    "按任务撰写": "C",
    "资源映射": "C",
}

# 文件名 → 文档类型识别规则（按顺序匹配，特殊规则在前）
DOC_TYPE_RULES = (
    (("完善要求",), "完善要求"),
    (("课程标准",), "课程标准"),
    (("授课计划",), "授课计划"),
    (("教学设计",), "教学设计"),
    (("成绩分析",), "成绩分析"),
    (("实训资料",), "实训资料"),
    (("温听",), "听课记录"),
    (("通用", ".xlsx"), "听课记录"),
)

VERSION_RE = re.compile(r"（(\d{4}-\d{4}(?:-\d)?)）")

# 封面字段模式：标签 + 冒号 + 空白值区 + 可选（提示）；排除注释类文本
COVER_FIELD_RE = re.compile(r"^([\u4e00-\u9fa5A-Za-z0-9]{2,10})[：:][\s\u3000]*（?.*$")

# 课程名称标题占位：《****》《×××》等
COURSE_TITLE_RE = re.compile(r"《[\*＊×xX＿]{2,}》")
# 括号内的模板提示（要求删除）：支持全角/半角括号，允许多层嵌套一层
PAREN_NOTE_RE = re.compile(r"[（(]([^（）()]*)[）)]")

ASPECT_RE = re.compile(r"([\u4e00-\u9fa5]{2,6})方面[，,]")

# 封面字段填写规则（A 类事实的具体取值口径）
FIELD_FILL_RULES = {
    "编制单位": "按课程所属专业名称＋“教研室”（如：农村电子商务教研室），不填系部名称",
    "适用专业": "按课程所属专业名称填写（与人才培养方案一致）",
    "合作单位": "取系统设置中的合作单位；无合作单位时整行删除",
}
NUMBERED_PLACEHOLDER_RE = re.compile(r"\d+[.、][^。\n]{1,30}……")
DEFINITION_RE = re.compile(r"XXXX[^。]{0,20}是[：:]")

# 格式指令关键词（区别于"注：xxx"类内容指令）
FORMAT_KEYWORD_RE = re.compile(r"行距|字体|仿宋|黑体|宋体|楷体|顶格|加粗|下同|居中|小四|小五|小二|四号|五号|二号")
CONTENT_NOTE_PREFIX = ("（注", "(注", "注：", "注:", "（要求", "(要求", "（说明", "(说明")

# 数字编号子标题（如 1、本课程设计的总体思路）：父节注释为其组说明的标志
SUBSECTION_RE = re.compile(r"^\d+[、.]\s*")
# 中文序号父节（一、二、…）：新父节开始，组说明作用域重置
CHAPTER_RE = re.compile(r"^[一二三四五六七八九十]+、")
# 中文序号子节（（一）（二）…）：同样可能是空节
SUBSECTION_CN_RE = re.compile(r"^（[一二三四五六七八九十]+）")
# 资源映射类空节（参考文献等：内容来自教材资料而非 AI 撰写）
RESOURCE_HEADING_RE = re.compile(r"参考书|期刊|参考书目|参考资料|网站")
# 空节的分类与填写要求（教材相关节 → 资源映射；其余 → AI 撰写）
EMPTY_SECTION_RULES = {
    "教材编写与选用": ("资源映射", "依据教材资料与教材版本信息填写教材编写与选用原则、教材信息，不由 AI 自由撰写"),
    "课程资源的开发与利用": ("资源映射", "依据教材包资料（PPT/大纲/源码/案例）与数字资源清单填写课程资源开发与利用说明"),
}


def _is_pure_format_instruction(text: str) -> bool:
    """判断整段括号文字是否为纯格式指令（如（正文行距23，小四号仿宋字，下同））。"""
    stripped = text.strip()
    match = PAREN_NOTE_RE.match(stripped)
    if not match:
        return False
    inner = match.group(1)
    if stripped.startswith(CONTENT_NOTE_PREFIX):
        return False
    return bool(FORMAT_KEYWORD_RE.search(inner)) and len(inner) <= 60


def _detect_doc_type(filename: str) -> str:
    for keywords, doc_type in DOC_TYPE_RULES:
        if all(k in filename for k in keywords):
            return doc_type
    return "其他"


def _file_hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def scan_templates(directory: str) -> dict:
    """扫描模板目录，登记/刷新模板库（按 file_path 去重，指纹变化更新状态）。"""
    root = Path(directory)
    if not root.is_dir():
        raise ValueError(f"模板目录不存在：{directory}")
    added, updated, skipped = 0, 0, 0
    with store.connect() as db:
        for path in sorted(root.rglob("*")):
            if not path.is_file() or path.suffix.lower() not in (".docx", ".doc", ".xlsx"):
                continue
            digest = _file_hash(path)
            row = db.execute(
                "SELECT id,file_hash FROM template_library WHERE file_path=?", (str(path),)
            ).fetchone()
            detected = _detect_doc_type(path.name)
            version_match = VERSION_RE.search(path.stem)
            version_label = version_match.group(1) if version_match else ""
            if row:
                if row["file_hash"] != digest:
                    db.execute(
                        "UPDATE template_library SET file_hash=?,status='待解析',doc_type=?,updated_at=CURRENT_TIMESTAMP WHERE id=?",
                        (digest, detected, row["id"]),
                    )
                    updated += 1
                else:
                    db.execute(
                        "UPDATE template_library SET doc_type=? WHERE id=? AND doc_type<>?",
                        (detected, row["id"], detected),
                    )
                    skipped += 1
                continue
            db.execute(
                """INSERT INTO template_library (doc_type,name,version_label,file_path,file_hash,file_format,status,notes)
                VALUES (?,?,?,?,?,?,?,?)""",
                (
                    detected,
                    path.stem,
                    version_label,
                    str(path),
                    digest,
                    path.suffix.lower().lstrip("."),
                    "待解析",
                    "",
                ),
            )
            added += 1
        db.commit()
    return {"added": added, "updated": updated, "skipped": skipped}


def _is_body_title(text: str) -> bool:
    """判断标题占位段落是否为正文大标题：占位后除括号提示外还有实际文字（如《***》课程标准）。"""
    match = COURSE_TITLE_RE.search(text)
    if not match:
        return False
    remainder = text.replace(match.group(0), "", 1)
    for note in PAREN_NOTE_RE.findall(remainder):
        remainder = remainder.replace(f"（{note}）", "").replace(f"({note})", "")
    return bool(remainder.strip())


def _cover_fields(document) -> list:
    """提取封面字段：正文大标题（如《***》课程标准）或首个章节标题之前的 标签：空白 段落。

    格式契约提议：值区加下划线、整段居中（用户对课标封面的明确要求）。
    封面裸标题《***》（居中…）不是分界；正文中的标签字段（课程名称及课程编号、课程类型等，
    位于第二页正文标题之后）不属于封面，不下划线。
    """
    fields = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if _is_body_title(text) or HEADING_RE.match(text):
            break
        match = COVER_FIELD_RE.match(text)
        if match:
            fields.append({
                "kind": "paragraph",
                "locator": f"paragraph:{index}",
                "label": match.group(1),
                "hint": re.sub(r"^（|）$|^\\(|\\)$", "", text.split("：", 1)[-1].strip()) if "：" in text else "",
                "proposed_format": {"underline": True, "alignment": "center"},
            })
    if not fields:
        # 封面整体为表格的模板（如教学设计）：扫第一个表格的 标签|空值 对
        for table_index, table in enumerate(document.tables[:2]):
            for row_index, row in enumerate(table.rows):
                for col_index, cell in _row_unique_cells_with_indices(row):
                    text = _cell_text(cell)
                    if not text or not _match_field(text, mode="label"):
                        continue
                    field_name, _, _ = _match_field(text, mode="label")
                    right = row.cells[col_index + 1] if col_index + 1 < len(row.cells) else None
                    if right is None or _cell_text(right):
                        continue
                    fields.append({
                        "kind": "table",
                        "locator": f"table:{table_index}/row:{row_index}/cell:{col_index + 1}",
                        "label": field_name,
                        "hint": "",
                        "proposed_format": {"underline": True, "alignment": "center"},
                    })
    return fields


def _course_title_fields(document) -> list:
    """识别课程名称标题占位（如《**********》（居中，与专业人才培养方案中名称一致））。

    契约规则：占位符替换为《课程名称》；括号内的模板提示（居中说明、字体说明等）删除。
    章节归属：首个占位为封面标题，其后（如《***》课程标准、考核评价标题）为正文标题；
    正文标题页自空行开始（保留模板标题前的空段落）。
    """
    fields = []
    current_heading = ""
    seen_body = False
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if HEADING_RE.search(text) and not text.startswith("（注"):
            current_heading = text.split("(")[0].split("（注")[0].strip()
        match = COURSE_TITLE_RE.search(text)
        if not match:
            continue
        if _is_body_title(text):
            seen_body = True
        page = "正文" if (seen_body or _is_body_title(text)) else "封面"
        notes = PAREN_NOTE_RE.findall(text.replace(match.group(0), "", 1))
        fill_template = text.replace(match.group(0), "《{课程名称}》", 1)
        for note in notes:
            fill_template = fill_template.replace(f"（{note}）", "").replace(f"({note})", "")
        fields.append({
            "locator": f"paragraph:{index}",
            "section": current_heading,
            "placeholder": match.group(0),
            "fill_template": re.sub(r"\s{2,}", " ", fill_template).strip(),
            "clean_instructions": [f"（{n}）" for n in notes],
            "format": _format(paragraph),
            "page": page,
        })
    return fields


def _reference_formats(text: str) -> list:
    """从模板文本中识别参考格式结构（内容契约的核心）。"""
    formats = []
    aspects = ASPECT_RE.findall(text)
    if len(aspects) >= 2 and "……" in text:
        formats.append({"type": "aspect_list", "aspects": aspects, "example": "马列主义方面，……"})
    if NUMBERED_PLACEHOLDER_RE.search(text):
        items = re.findall(r"(\d+[.、][^。\n]{1,30}……)", text)
        formats.append({"type": "numbered_placeholder", "items": items[:8]})
    if DEFINITION_RE.search(text):
        formats.append({"type": "definition_pattern", "example": "XXXX教学法是：（对涉及的教学法简单说明其内涵、特点等）"})
    return formats


def _mandates_for(doc_type: str) -> list:
    """从库内'完善要求'文档提取官方强制条款（四新要求），挂到课标/教学设计契约。"""
    if doc_type not in ("课程标准", "教学设计"):
        return []
    rows = store.rows("SELECT file_path FROM template_library WHERE doc_type='完善要求' AND file_format='docx'")
    mandates = []
    for row in rows:
        try:
            document = Document(row["file_path"])
        except Exception:
            continue
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text and re.match(r"^\d+[、.]", text):
                mandates.append(text)
    return mandates


def _next_is_numbered_subsection(paragraphs, index, lookahead=4) -> bool:
    """判断当前段之后（有限窗口内）首个非空段是否为数字编号子标题（1、2、…）。

    用于区分：父节注释（如"三、课程设计"的（注：…）后跟 1、2、3、4 子节 → 组说明）
    与本节撰写指令（如"一、课程性质"的（注：…）后跟同级"二、课程目标"）。
    """
    checked = 0
    for paragraph in paragraphs[index + 1:]:
        text = paragraph.text.strip()
        if not text:
            continue
        if SUBSECTION_RE.match(text):
            return True
        if HEADING_RE.match(text):
            return False
        checked += 1
        if checked >= lookahead:
            return False
    return False


def _extract_content(document) -> dict:
    """内容契约：模板指令 + 参考格式 + 格式指令 + 章节正文区 + 官方强制条款。"""
    instructions = []
    references = []
    format_instructions = []
    section_bodies = []
    current_heading = ""
    pending_heading = None  # (标题, locator) 等待其后的首个正文段
    last_group_note = ""  # 最近父节组说明（数字子节的撰写依据）
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        is_heading = bool(HEADING_RE.search(text)) and not text.startswith("（注")
        if is_heading:
            # 上一标题为空节（pending 未消费）且是子标题（数字/中文序号）：建空节槽位
            if pending_heading is not None:
                heading_old, locator_old = pending_heading
                if SUBSECTION_RE.match(heading_old) or SUBSECTION_CN_RE.match(heading_old):
                    section_bodies.append({
                        "heading": heading_old,
                        "heading_locator": locator_old,
                        "body_locator": locator_old,
                        "instruction": "",
                        "example": "",
                        "format": {},
                        "type": "empty",
                        "group_note": last_group_note,
                    })
            current_heading = text.split("(")[0].split("（注")[0].strip()
            if CHAPTER_RE.match(current_heading):
                last_group_note = ""  # 新父节开始，旧组说明失效
            pending_heading = (current_heading, f"paragraph:{index}")
            for note in PAREN_NOTE_RE.findall(text):
                if FORMAT_KEYWORD_RE.search(note):
                    format_instructions.append({
                        "locator": f"paragraph:{index}",
                        "scope": "节标题",
                        "section": current_heading,
                        "text": f"（{note}）",
                    })
            continue
        if _is_pure_format_instruction(text):
            inner = PAREN_NOTE_RE.match(text).group(1)
            scope = "全文档正文" if "下同" in inner and "正文" in inner else (
                "后续同类元素" if "下同" in inner else "本处")
            format_instructions.append({
                "locator": f"paragraph:{index}",
                "scope": scope,
                "section": current_heading,
                "text": text,
            })
            continue
        refs = _reference_formats(text)
        if refs:
            references.append({"locator": f"paragraph:{index}", "section": current_heading, "formats": refs})
        is_instruction = text.startswith(CONTENT_NOTE_PREFIX) or "要求：" in text
        if is_instruction:
            instructions.append({"locator": f"paragraph:{index}", "section": current_heading, "text": text})
        # 章节正文区：节标题后首个非格式非标签段落（AI 撰写区）
        # 字号 ≤10.5pt（小五）视为表格注释而非散文正文，跳过（该类节内容由表格槽位承担）
        if pending_heading and not COVER_FIELD_RE.match(text) and not COURSE_TITLE_RE.search(text):
            fmt = _format(paragraph)
            size = fmt.get("font_size_pt")
            if size and size <= 10.5:
                pending_heading = None
                continue
            # 父节注释：内容指令后紧跟数字子标题（如"三、课程设计"注后跟 1、2、3、4）→ 组说明，不建正文槽
            if is_instruction and _next_is_numbered_subsection(document.paragraphs, index):
                last_group_note = text
                pending_heading = None
                continue
            heading, heading_locator = pending_heading
            pending_heading = None
            section_bodies.append({
                "heading": heading,
                "heading_locator": heading_locator,
                "body_locator": f"paragraph:{index}",
                "instruction": text if is_instruction else "",
                "example": "" if is_instruction else text[:120],
                "format": fmt,
                "type": "body",
            })
    # 文末冲刷：最后一个标题若为空子节（如"3.网站"），同样建槽
    if pending_heading is not None:
        heading_old, locator_old = pending_heading
        if SUBSECTION_RE.match(heading_old) or SUBSECTION_CN_RE.match(heading_old):
            section_bodies.append({
                "heading": heading_old,
                "heading_locator": locator_old,
                "body_locator": locator_old,
                "instruction": "",
                "example": "",
                "format": {},
                "type": "empty",
                "group_note": last_group_note,
            })
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                text = _cell_text(cell)
                if not text:
                    continue
                refs = _reference_formats(text)
                if refs:
                    references.append({
                        "locator": f"table:{table_index}/row:{row_index}/cell:{col_index}",
                        "section": _table_role_header(table),
                        "formats": refs,
                    })
                if "要求：" in text or "（要求" in text or "（注：" in text:
                    instructions.append({
                        "locator": f"table:{table_index}/row:{row_index}/cell:{col_index}",
                        "section": _table_role_header(table),
                        "text": text,
                    })
    seen = set()
    unique_instructions = []
    for item in instructions:
        if item["text"] not in seen:
            seen.add(item["text"])
            unique_instructions.append(item)
    return {
        "instructions": unique_instructions,
        "reference_formats": references,
        "format_instructions": format_instructions,
        "section_bodies": section_bodies,
    }


def _table_role_header(table) -> str:
    header = " | ".join(_cell_text(c) for c in table.rows[0].cells) if table.rows else ""
    return _table_role(header, "") if header else "普通表格"


def _document_order(document):
    """按文档正文流（w:p / w:tbl 交错顺序）给段落和表格编号，用于槽位排序。"""
    para_seq, table_seq = {}, {}
    p_i = t_i = seq = 0
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para_seq[p_i] = seq
            p_i += 1
            seq += 1
        elif child.tag == qn("w:tbl"):
            table_seq[t_i] = seq
            t_i += 1
            seq += 1
    return para_seq, table_seq


def _locator_sort_key(locator, para_seq, table_seq):
    """locator → (文档序, 行, 列)；无法解析的排最后。"""
    match = re.match(r"^paragraph:(\d+)", locator)
    if match:
        n = int(match.group(1))
        return (para_seq.get(n, 999999), 0, 0)
    match = re.match(r"^table:(\d+)", locator)
    if match:
        t = int(match.group(1))
        row_m = re.search(r"row:(\d+)", locator)
        col_m = re.search(r"(?:col|cell):(\d+)", locator)
        return (
            table_seq.get(t, 999999),
            int(row_m.group(1)) if row_m else 0,
            int(col_m.group(1)) if col_m else 0,
        )
    return (999999, 0, 0)


def _build_contract_slots(db, contract_id, document, structural, content_extra=None) -> int:
    """构建契约槽位（分类 A/B/C + 格式 + 置信度）。"""
    db.execute("DELETE FROM contract_slots WHERE contract_id=?", (contract_id,))
    slots = {}
    cover_locators = {f["locator"] for f in structural.get("cover", {}).get("fields", [])}

    def add_slot(slot_key, locator, section, field_name, kind, repeat_scope, sources, label_text, fmt, confidence,
                 structure_extra=None, content_req_override=None):
        if slot_key in slots:
            return
        classification = KIND_CLASSIFICATION.get(kind, "B")
        structure = {"locator": locator, "repeat_scope": repeat_scope, "sources": sources}
        if structure_extra:
            structure.update(structure_extra)
        content_req = content_req_override or ""
        if not content_req_override:
            if kind in ("分析撰写", "事实与分析"):
                content_req = "依据已审核教材语义模型与所在专业润色生成"
            elif classification == "C":
                content_req = "依据课程蓝本/任务结构生成"
        slots[slot_key] = (
            contract_id, slot_key, locator, section, field_name, classification,
            json.dumps(structure, ensure_ascii=False),
            json.dumps(fmt, ensure_ascii=False),
            content_req, confidence, 1, 0, "待确认",
        )

    # 封面字段槽位（A 类事实 + 下划线/居中格式契约提议 + 字段级填写规则）
    for field in structural.get("cover", {}).get("fields", []):
        add_slot(
            f"cover:{field['locator']}", field["locator"], "封面",
            field["label"], "事实填写", "单次", "课程实例、系统设置",
            field["label"], dict(field["proposed_format"]), "高",
            content_req_override=FIELD_FILL_RULES.get(field["label"]),
        )

    # 课程名称标题槽位（《****》占位 → 《课程名称》；模板提示删除）
    for title in structural.get("cover", {}).get("title_fields", []):
        req = f"占位符{title['placeholder']}替换为课程名称（生成 {title['fill_template']}）"
        if title["clean_instructions"]:
            req += "；删除模板提示：" + "、".join(title["clean_instructions"])
        if title["page"] == "正文":
            req += "；正文标题页自空行开始，保留标题前空段落"
        title_section = "封面" if title["page"] == "封面" else (title["section"] or "正文")
        add_slot(
            f"{title['locator']}:课程名称", title["locator"], title_section,
            "课程名称", "事实填写", "单次", "课程实例",
            title["placeholder"], title["format"], "高",
            structure_extra={
                "page": title["page"],
                "placeholder": title["placeholder"],
                "fill_template": title["fill_template"],
                "clean_instructions": title["clean_instructions"],
            },
            content_req_override=req,
        )

    # 章节正文槽位（整节 AI 撰写区：课程性质/教学设计思路/教学方法描述等，B 类）
    for body in content_extra.get("section_bodies", []):
        body_type = body.get("type", "body")
        is_resource = bool(RESOURCE_HEADING_RE.search(body["heading"]))
        rule = next((v for k, v in EMPTY_SECTION_RULES.items() if k in body["heading"]), None)
        kind = rule[0] if rule else ("资源映射" if is_resource else "分析撰写")
        if body_type == "empty":
            group_note = body.get("group_note", "")
            if rule:
                req = rule[1]
                confidence = "中"
            elif is_resource:
                req = "模板空节；内容依据教材资料与参考资源映射填写（书目/期刊清单），不由 AI 自由撰写"
                confidence = "中"
            elif group_note:
                req = (f"模板空节，需 AI 撰写；依据父节组说明撰写本节内容，聚焦本节标题，"
                       f"不与其他小节重复。父节组说明：{group_note[:100]}")
                confidence = "中"
            else:
                req = "模板空节，需 AI 撰写；聚焦本节标题撰写，不与其他小节重复"
                confidence = "低"
            fmt = {}
        else:
            instruction = body.get("instruction", "")
            if instruction:
                req = f"按模板指令撰写：{instruction[:90]}"
                confidence = "高"
            else:
                req = f"参考模板示例撰写：{body.get('example', '')[:60]}"
                confidence = "中"
            fmt = body.get("format", {})
        add_slot(
            f"section:{body['body_locator']}", body["body_locator"], body["heading"],
            "章节正文", kind, "单次", "教材语义模型、专业人才培养方案" if not is_resource else "教材资料、参考资源",
            body["heading"], fmt, confidence,
            structure_extra={
                "heading": body["heading"],
                "heading_locator": body["heading_locator"],
                "instruction": body.get("instruction", ""),
                "example": body.get("example", ""),
                "body_type": body_type,
                "group_note": body.get("group_note", ""),
            },
            content_req_override=req,
        )

    # 段落槽位（封面字段已单独建槽，跳过；正文标签区位于首个章节标题之前，归属"正文"）
    current_heading = "正文"
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if f"paragraph:{index}" in cover_locators:
            continue
        if HEADING_RE.search(text) and not text.startswith("（注"):
            current_heading = text.split("(")[0].split("（注")[0].strip()
            continue
        matched = _match_field(text, mode="label") or _match_field(text, mode="heading")
        if not matched:
            continue
        field_name, kind, sources = matched
        add_slot(
            f"paragraph:{index}:{field_name}", f"paragraph:{index}", current_heading,
            field_name, kind, "单次", sources, text, _format(paragraph),
            "高" if _match_field(text, mode="label") else "中",
        )

    # 表格槽位（同表同字段只保留首个槽位：配对逻辑与表头逻辑可能重复命中）
    seen_table_fields = set()
    for table_index, table in enumerate(document.tables):
        meta = structural["tables"][table_index]
        role, repeat_scope = meta["role"], meta["repeat_mode"]
            # 学习情境描述表：首行（课程/学分/总学时）为整表唯一表头，非重复结构
        if role == "学习情境描述表":
            add_slot(
                f"table:{table_index}:课程学分总学时", f"table:{table_index}/row:0/header-only", role,
                "课程基本信息", "事实填写", "单次（整表唯一，不随情境重复）", "课程实例",
                "课程：/学分/总学时", {}, "高",
                structure_extra={
                    "table_header_row": 0,
                    "repeat_note": "首行表头整表唯一；自'学习情境1'行起按情境重复",
                },
            )
            continue
        for pair in _identify_label_value_pairs(table):
            matched = _match_field(pair["label"], mode="label")
            if not matched:
                continue
            # 水平配对要求值单元格为空：表头行相邻标签互配（如 知识目标|能力目标）不是待填槽位
            if pair["direction"] == "horizontal":
                if _cell_text(table.rows[pair["row"]].cells[pair["value_col"]]):
                    continue
            field_name, kind, sources = matched
            # 角色感知命名：考核评价表的"学习情境"列实为 情境+分值权重
            if role == "考核评价表" and "学习情境" in pair["label"] and field_name == "课程内容与学时":
                field_name = "学习情境分值权重"
            if (table_index, field_name) in seen_table_fields:
                continue
            seen_table_fields.add((table_index, field_name))
            row_index, value_col = pair["row"], pair["value_col"]
            value_cell = table.rows[row_index].cells[value_col]
            fmt = _format(value_cell.paragraphs[0]) if value_cell.paragraphs else {}
            locator = f"table:{table_index}/row:{row_index}/col:{value_col}/{pair['direction']}"
            add_slot(
                f"table:{table_index}:{field_name}:{pair['direction']}", locator, role,
                field_name, kind, repeat_scope, sources, pair["label"], fmt, "高",
            )
        if role in ("授课计划明细表", "单元教学设计表", "课程内容结构表", "课程内容与课时分配表") and len(table.rows) >= 2:
            for col_index, cell in _row_unique_cells_with_indices(table.rows[0]):
                header_text = _cell_text(cell)
                if not header_text:
                    continue
                matched = _match_field(header_text, mode="label")
                if not matched:
                    continue
                field_name, kind, sources = matched
                if (table_index, field_name) in seen_table_fields:
                    continue
                seen_table_fields.add((table_index, field_name))
                add_slot(
                    f"table:{table_index}:header:{field_name}",
                    f"table:{table_index}/col:{col_index}/header", role,
                    field_name, kind, repeat_scope, sources, header_text,
                    _format(cell.paragraphs[0]) if cell.paragraphs else {}, "高",
                )

    para_seq, table_seq = _document_order(document)
    ordered = sorted(slots.values(), key=lambda row: _locator_sort_key(row[2], para_seq, table_seq))
    for order, row in enumerate(ordered):
        db.execute(
            """INSERT INTO contract_slots
            (contract_id,slot_key,locator,section_title,field_name,classification,
             structure_json,format_json,content_req,confidence,required,manual_override,approval_status,sort_order)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (*row, order),
        )
    return len(ordered)


def parse_contract(library_id: int) -> dict:
    """解析库模板：生成/覆盖契约版本（已确认的版本保留，新解析开新版本）。"""
    item = store.rows("SELECT * FROM template_library WHERE id=?", (library_id,))
    if not item:
        raise ValueError("模板不存在")
    item = item[0]
    path = Path(item["file_path"])
    if not path.exists():
        raise ValueError("模板文件已不存在")

    if item["file_format"] != "docx":
        store.execute(
            "UPDATE template_library SET status='暂不支持',notes=? WHERE id=?",
            (f".{item['file_format']} 格式适配在阶段5接入，当前仅登记管理", library_id),
        )
        return {"status": "暂不支持", "message": f".{item['file_format']} 格式暂不支持解析"}

    document = Document(path)
    tables = []
    for table_index, table in enumerate(document.tables):
        rows = []
        merged = []
        for row_index, row in enumerate(table.rows):
            values = []
            for cell_index, cell in enumerate(row.cells):
                values.append(_cell_text(cell))
                merge = _cell_merge(cell)
                if merge["column_span"] > 1 or merge["vertical_merge"]:
                    merged.append({"row": row_index, "column": cell_index, **merge})
            rows.append(values)
        header = " | ".join(rows[0]) if rows else ""
        body = "\n".join(" | ".join(r) for r in rows[1:])
        role = _table_role(header, body)
        tables.append({
            "index": table_index,
            "role": role,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "header_text": header[:200],
            "merged_cells": merged,
            "repeat_mode": {
                "课程内容与课时分配表": "按课程项目和子任务重复",
                "课程内容结构表": "按课程项目重复",
                "考核评价表": "按选定考核项目重复",
                "学习情境描述表": "按学习情境及子情境重复",
                "授课计划明细表": "按排课任务重复",
                "单元教学设计表": "按周次或教学单元重复",
                "教学设计基本信息表": "与教学组织表成对，按周次或教学单元重复",
                "教学设计教学组织表": "与基本信息表成对，按所选体例重复",
                "职业能力分析表": "按职业标准项重复",
                "教学内容要求表": "按课程项目重复",
                "能力训练项目表": "按能力训练项目重复",
            }.get(role, "固定结构"),
        })

    cover_fields = _cover_fields(document)
    title_fields = _course_title_fields(document)
    content = _extract_content(document)
    content["mandates"] = _mandates_for(item["doc_type"])

    structural = {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "section_count": len(document.sections),
        "tables": tables,
        "cover": {"fields": cover_fields, "title_fields": title_fields},
    }

    existing = store.rows(
        "SELECT version FROM template_contracts WHERE library_id=? ORDER BY version DESC LIMIT 1",
        (library_id,),
    )
    next_version = (existing[0]["version"] + 1) if existing else 1

    with store.connect() as db:
        cursor = db.execute(
            """INSERT INTO template_contracts
            (library_id,version,status,structural_json,content_json,slot_count,parse_message,parsed_at)
            VALUES (?,?,?,?,?,?,?,CURRENT_TIMESTAMP)""",
            (
                library_id, next_version, "草稿",
                json.dumps(structural, ensure_ascii=False),
                json.dumps(content, ensure_ascii=False),
                0, "",
            ),
        )
        contract_id = cursor.lastrowid
        slot_count = _build_contract_slots(db, contract_id, document, structural, content)
        db.execute("UPDATE template_contracts SET slot_count=? WHERE id=?", (slot_count, contract_id))
        db.execute(
            "UPDATE template_library SET status='已解析',updated_at=CURRENT_TIMESTAMP WHERE id=?",
            (library_id,),
        )
        db.commit()

    message = (
        f"表格 {len(tables)} 个（含角色识别）；封面字段 {len(cover_fields)} 项；标题占位 {len(title_fields)} 处；"
        f"内容指令 {len(content['instructions'])} 条；参考格式 {len(content['reference_formats'])} 处；"
        f"强制条款 {len(content['mandates'])} 条；槽位 {slot_count} 个"
    )
    return {"status": "已解析", "contract_id": contract_id, "version": next_version, "message": message, "slot_count": slot_count}
