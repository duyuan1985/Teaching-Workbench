from docx import Document
from docx.oxml.ns import qn

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

# 检查P109-P119的字体和字号
for pi in range(109, 120):
    p = doc.paragraphs[pi]
    txt = p.text.strip()[:30]
    for ri, r in enumerate(p.runs):
        rPr = r._element.find(qn('w:rPr'))
        sz = ''
        szCs = ''
        font = ''
        color = ''
        if rPr is not None:
            sz_elem = rPr.find(qn('w:sz'))
            if sz_elem is not None:
                sz = sz_elem.get(qn('w:val'))
            szCs_elem = rPr.find(qn('w:szCs'))
            if szCs_elem is not None:
                szCs = szCs_elem.get(qn('w:val'))
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                font = rFonts.get(qn('w:eastAsia')) or rFonts.get(qn('w:ascii')) or ''
            color_elem = rPr.find(qn('w:color'))
            if color_elem is not None:
                color = color_elem.get(qn('w:val'))
        print(f'P{pi} R{ri}: font={font} sz={sz} szCs={szCs} color={color} | {txt}')

# 对比正常显示的段落
print('\n--- 对比正常段落 ---')
for pi in [35, 38, 73]:
    p = doc.paragraphs[pi]
    txt = p.text.strip()[:30]
    for ri, r in enumerate(p.runs):
        rPr = r._element.find(qn('w:rPr'))
        sz = ''
        font = ''
        if rPr is not None:
            sz_elem = rPr.find(qn('w:sz'))
            if sz_elem is not None:
                sz = sz_elem.get(qn('w:val'))
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                font = rFonts.get(qn('w:eastAsia')) or ''
        print(f'P{pi} R{ri}: font={font} sz={sz} | {txt}')
