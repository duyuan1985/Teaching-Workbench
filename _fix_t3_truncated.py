"""
修复教学进度表（Table 3）中被截断的单元格内容
C3=ability_goal, C5=knowledge_goal, C6=ideological_goal, C7=quality_goal
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import store

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)
t3 = doc.tables[3]

# 获取任务数据
tasks = store.rows("SELECT * FROM tasks WHERE offering_id=20 ORDER BY seq")
print(f"任务数: {len(tasks)}")

# 获取模板格式
template_cell = t3.cell(3, 3)
template_p = template_cell.paragraphs[0] if template_cell.paragraphs else None
font_name = '仿宋'
font_size_val = '21'  # 10.5pt = 21 half-points
if template_p and template_p.runs:
    r = template_p.runs[0]
    if r.font.name:
        font_name = r.font.name
    if r.font.size:
        font_size_val = str(int(r.font.size.pt * 2))
print(f"格式: font={font_name}, size={font_size_val}")

def set_cell_multiline(tc, text, font_name, font_size_val):
    """设置单元格多行文本（支持\n换行）"""
    # 清除所有段落
    paras = tc.findall(qn('w:p'))
    # 保留第一个段落元素，删除其余
    if paras:
        first_p = paras[0]
        for extra in paras[1:]:
            tc.remove(extra)
        # 清除第一个段落的内容
        for r in first_p.findall(qn('w:r')):
            first_p.remove(r)
    else:
        first_p = OxmlElement('w:p')
        tc.append(first_p)

    # 按换行符分割
    lines = text.split('\n')
    for li, line in enumerate(lines):
        if li > 0:
            # 新段落
            p = OxmlElement('w:p')
            tc.append(p)
        else:
            p = first_p

        # 段落格式
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p.insert(0, pPr)

        # run
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), font_size_val)
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), font_size_val)
        rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = line
        r.append(t)
        p.append(r)

# 修复每个任务行（R3-R32对应任务1-30）
fixed_count = 0
for ri in range(3, min(3 + len(tasks), len(t3.rows))):
    task = tasks[ri - 3]

    # C3: ability_goal
    ability = task.get('ability_goal', '')
    if ability:
        tc = t3.cell(ri, 3)._element
        set_cell_multiline(tc, ability, font_name, font_size_val)
        fixed_count += 1

    # C5: knowledge_goal
    knowledge = task.get('knowledge_goal', '')
    if knowledge:
        tc = t3.cell(ri, 5)._element
        set_cell_multiline(tc, knowledge, font_name, font_size_val)

    # C6: ideological_goal
    ideological = task.get('ideological_goal', '')
    if ideological:
        tc = t3.cell(ri, 6)._element
        set_cell_multiline(tc, ideological, font_name, font_size_val)

    # C7: quality_goal
    quality = task.get('quality_goal', '')
    if quality:
        tc = t3.cell(ri, 7)._element
        set_cell_multiline(tc, quality, font_name, font_size_val)

doc.save(fp)
print(f"\n修复完成，共修复{fixed_count}行的4列内容")
print(f"保存到: {fp}")

# 验证
doc2 = Document(fp)
t3_2 = doc2.tables[3]
print("\n=== 验证 R3（任务1）完整文本 ===")
for ci in [3, 5, 6, 7]:
    txt = t3_2.cell(3, ci).text.strip()
    print(f"  C{ci} [{len(txt)} chars]: {txt[:200]}")
    if len(txt) > 200:
        print(f"    ... (total {len(txt)} chars)")

print("\n=== 验证 R15（任务13）完整文本 ===")
for ci in [3, 5, 6, 7]:
    txt = t3_2.cell(15, ci).text.strip()
    print(f"  C{ci} [{len(txt)} chars]: {txt[:200]}")
