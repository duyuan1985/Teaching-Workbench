"""
对比教学设计进度表(Table 3)与课程标准学习情境表(Table 1)的学习目标
"""
from docx import Document
import store

# 教学设计进度表
fp_d = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc_d = Document(fp_d)
t3 = doc_d.tables[3]

# 课程标准学习情境表
fp_s = r'生成结果\精修版\2023-2024-2《商务数据分析》课程标准 杜媛.docx'
doc_s = Document(fp_s)
t1_s = doc_s.tables[1]  # 序号/学习情境/子情境/课时

print('=== 课程标准学习情境表(Table 1)结构 ===')
for ri in range(len(t1_s.rows)):
    texts = [c.text.strip()[:30] for c in t1_s.rows[ri].cells]
    print(f'  R{ri}: {texts}')

# 课程标准是否有学习目标内容？
print('\n=== 课程标准 所有表格概览 ===')
for ti, t in enumerate(doc_s.tables):
    first_cell = t.rows[0].cells[0].text.strip() if t.rows else ''
    print(f'  Table {ti}: rows={len(t.rows)}, first_cell="{first_cell[:40]}"')

# 数据库中的任务数据
print('\n=== 数据库任务数据(前3个任务) ===')
tasks = store.rows("SELECT * FROM tasks WHERE offering_id=20 ORDER BY seq")
for t in tasks[:3]:
    print(f'\n--- 任务{t["seq"]}: {t["title"][:40]} ---')
    print(f'  knowledge_goal: {t["knowledge_goal"][:100]}')
    print(f'  ability_goal: {t["ability_goal"][:100]}')
    print(f'  ideological_goal: {t["ideological_goal"][:100]}')
    print(f'  quality_goal: {t["quality_goal"][:100]}')
