import sqlite3, json
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

conn = sqlite3.connect(r'e:\开发\AIGC\教学档案工作台\data\workbench.db')
conn.row_factory = sqlite3.Row
c = conn.cursor()

# Get template file path
tf = c.execute("SELECT template_path FROM template_files WHERE id=9").fetchone()
template_path = tf['template_path']
print(f"Template path: {template_path}")

# Open template and examine table structures
doc = Document(template_path)

# Table 0 (goals table)
print("\n=== Table 0 (课程目标表) ===")
t0 = doc.tables[0]
print(f"  Rows: {len(t0.rows)}")
for ri, row in enumerate(t0.rows):
    tcs = row._tr.findall(qn("w:tc"))
    print(f"  Row {ri}: {len(tcs)} tc elements, {len(row.cells)} grid cells")
    for ci, cell in enumerate(row.cells):
        txt = cell.text.strip()[:30]
        # Check for gridSpan
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        gridspan = None
        vmerge = None
        if tcPr is not None:
            gs = tcPr.find(qn("w:gridSpan"))
            if gs is not None:
                gridspan = gs.get(qn("w:val"))
            vm = tcPr.find(qn("w:vMerge"))
            if vm is not None:
                vmerge = vm.get(qn("w:val")) or "continue"
        print(f"    cell[{ci}] text='{txt}' gridSpan={gridspan} vMerge={vmerge}")

# Table 2 (assessment table)
print("\n=== Table 2 (考核评价表) ===")
t2 = doc.tables[2]
print(f"  Rows: {len(t2.rows)}")
for ri, row in enumerate(t2.rows):
    tcs = row._tr.findall(qn("w:tc"))
    print(f"  Row {ri}: {len(tcs)} tc elements, {len(row.cells)} grid cells")
    for ci, cell in enumerate(row.cells):
        txt = cell.text.strip()[:40]
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        gridspan = None
        vmerge = None
        if tcPr is not None:
            gs = tcPr.find(qn("w:gridSpan"))
            if gs is not None:
                gridspan = gs.get(qn("w:val"))
            vm = tcPr.find(qn("w:vMerge"))
            if vm is not None:
                vmerge = vm.get(qn("w:val")) or "continue"
        # Only print unique cells (by tc id)
        print(f"    cell[{ci}] text='{txt}' gridSpan={gridspan} vMerge={vmerge}")

# Table 1 (content/hours table) - check first few rows
print("\n=== Table 1 (课时分配表) ===")
t1 = doc.tables[1]
print(f"  Rows: {len(t1.rows)}")
for ri, row in enumerate(t1.rows[:4]):
    tcs = row._tr.findall(qn("w:tc"))
    print(f"  Row {ri}: {len(tcs)} tc elements, {len(row.cells)} grid cells")
    for ci, cell in enumerate(row.cells):
        txt = cell.text.strip()[:30]
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        gridspan = None
        vmerge = None
        if tcPr is not None:
            gs = tcPr.find(qn("w:gridSpan"))
            if gs is not None:
                gridspan = gs.get(qn("w:val"))
            vm = tcPr.find(qn("w:vMerge"))
            if vm is not None:
                vmerge = vm.get(qn("w:val")) or "continue"
        print(f"    cell[{ci}] text='{txt}' gridSpan={gridspan} vMerge={vmerge}")

# Check paragraph 39 and 44
print("\n=== Paragraphs 35-46 ===")
for pi in range(35, min(47, len(doc.paragraphs))):
    p = doc.paragraphs[pi]
    print(f"  para[{pi}]: '{p.text.strip()[:80]}'")

conn.close()
