"""
分析教学设计模板的完整结构
检查所有段落标题、表格位置、缺失内容
"""
from docx import Document
from docx.oxml.ns import qn

# 用原始模板分析
template_fp = r"原始资料\模板\模板5：教学设计\模板5：教学设计 模板（2023-2024）.docx"
current_fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"

print("=" * 70)
print("一、原始模板结构")
print("=" * 70)

doc = Document(template_fp)
print(f"段落: {len(doc.paragraphs)}, 表格: {len(doc.tables)}")

print("\n--- 段落 ---")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f"  P{i}: {text[:80]}")

print("\n--- 表格概览 ---")
for ti, t in enumerate(doc.tables):
    rows = len(t.rows)
    cols = len(t.rows[0].cells) if rows > 0 else 0
    # 提取第一行内容判断表格类型
    r0_text = ""
    if rows > 0:
        for c in t.rows[0].cells:
            ct = c.text.strip()
            if ct:
                r0_text = ct[:30]
                break
    print(f"  表{ti}: {rows}行x{cols}列 | 首行: {r0_text}")

# ============================================================
# 对比当前文档
# ============================================================
print(f"\n{'='*70}")
print("二、当前生成文档结构")
print("=" * 70)

doc2 = Document(current_fp)
print(f"段落: {len(doc2.paragraphs)}, 表格: {len(doc2.tables)}")

print("\n--- 段落 ---")
for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    if text:
        print(f"  P{i}: {text[:80]}")

print("\n--- 表格概览 ---")
for ti, t in enumerate(doc2.tables):
    rows = len(t.rows)
    cols = len(t.rows[0].cells) if rows > 0 else 0
    r0_text = ""
    if rows > 0:
        for c in t.rows[0].cells:
            ct = c.text.strip()
            if ct:
                r0_text = ct[:30]
                break
    print(f"  表{ti}: {rows}行x{cols}列 | 首行: {r0_text}")

# ============================================================
# 检查缺失内容
# ============================================================
print(f"\n{'='*70}")
print("三、缺失内容检查")
print("=" * 70)

current_paras = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
template_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

# 模板中有但生成文档中没有的段落标题
template_titles = [t for t in template_paras if any(k in t for k in [
    "一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、",
    "课程信息", "课程定位", "教学设计", "教学实施", "教学资源", "说明", "第一节课",
    "单元教学", "教学设计·", "模版体例", "体例"
])]

print("模板中的标题段落:")
for t in template_titles:
    found = any(t[:10] in cp for cp in current_paras)
    status = "✓存在" if found else "✗缺失"
    print(f"  {status}: {t[:60]}")
