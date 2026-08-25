import io
import json
import math
import re
import tempfile
from collections import Counter
from pathlib import Path

import pdfplumber
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

import store


ROOT = Path(__file__).parent
DEFAULT_TEMPLATE_PATH = ROOT / "原始资料" / "模板" / "模板6：成绩分析.docx"


def _undouble(text):
    chars = []
    index = 0
    while index < len(text):
        if index + 1 < len(text) and text[index] == text[index + 1]:
            chars.append(text[index])
            index += 2
        else:
            chars.append(text[index])
            index += 1
    return "".join(chars)


def _number(value):
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _text(value):
    return str(value or "").strip()


def parse_grade_pdf(payload):
    with pdfplumber.open(io.BytesIO(payload)) as pdf:
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    header = _undouble(lines[1]) if len(lines) > 1 else ""
    term_match = re.search(r"(20\d{2})-(20\d{2})学年第([一二])学期", header)
    course_match = re.search(r"课\s*程：\[([^]]+)](.+?)\s+考核方式：([^\s]+)\s+学分：([\d.]+)", text)
    teacher_match = re.search(r"任课教师：\[[^]]+](.+?)\s+上课班级：([^\s]+).*?上课人数：(\d+)", text)
    nature_match = re.search(r"课程类别：([^\s]+)", text)
    if not course_match or not teacher_match:
        raise ValueError("无法从PDF识别课程、教师或班级信息，请确认文件是教务系统导出的原始成绩表。")
    students = []
    pattern = re.compile(r"^\d+\s+(\S+)\s+(.+?)\s+(男|女)\s+(\S+)\s+初修\s+(.+)$")
    for line in lines:
        match = pattern.match(line)
        if not match:
            continue
        tail = match.group(5).split()
        if len(tail) < 2:
            continue
        usual = _number(tail[0])
        final = _number(tail[-2]) if len(tail) >= 3 else _number(tail[0])
        comprehensive_text = tail[-1]
        comprehensive = _number(comprehensive_text)
        absent = "缺考" in comprehensive_text
        students.append({
            "student_no": match.group(1), "student_name": match.group(2), "gender": match.group(3),
            "class_name": match.group(4), "usual_score": usual, "final_score": final or 0,
            "comprehensive_score": comprehensive, "absent": absent,
        })
    if not students:
        raise ValueError("PDF中没有识别到学生成绩明细。")
    class_name = Counter(item["class_name"] for item in students).most_common(1)[0][0]
    term = ""
    if term_match:
        term = f"{term_match.group(1)}-{term_match.group(2)}-{'1' if term_match.group(3)=='一' else '2'}"
    return {
        "term": term,
        "course_code": _text(course_match.group(1)), "course_name": _text(course_match.group(2)),
        "assessment_method": _text(course_match.group(3)), "credits": float(course_match.group(4)),
        "teacher": _text(teacher_match.group(1)), "teaching_section": _text(teacher_match.group(2)),
        "expected_count": int(teacher_match.group(3)), "course_nature_raw": nature_match.group(1) if nature_match else "",
        "class_name": class_name, "students": students, "raw_text": text,
    }


def _statistics(students):
    scores = [float(item["comprehensive_score"] or 0) for item in students]
    bands = ((90, 101), (80, 90), (70, 80), (60, 70), (50, 60), (0, 50))
    counts = [sum(1 for score in scores if low <= score < high) for low, high in bands]
    total = len(scores)
    average = sum(scores) / total if total else 0
    deviation = math.sqrt(sum((score - average) ** 2 for score in scores) / total) if total else 0
    return {
        "counts": counts, "percentages": [count * 100 / total if total else 0 for count in counts],
        "average": average, "deviation": deviation, "highest": max(scores, default=0),
        "lowest": min(scores, default=0), "pass_rate": sum(score >= 60 for score in scores) * 100 / total if total else 0,
    }


def _font(size, bold=False):
    candidates = [Path(r"C:\Windows\Fonts\simhei.ttf") if bold else Path(r"C:\Windows\Fonts\simsun.ttc"), Path(r"C:\Windows\Fonts\msyh.ttc")]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def _make_chart(stats, output):
    image = Image.new("RGB", (960, 430), "white")
    draw = ImageDraw.Draw(image)
    title_font, label_font = _font(30, True), _font(22)
    draw.text((300, 12), "综合成绩分布", fill="#202124", font=title_font)
    left, top, right, bottom = 85, 70, 920, 350
    draw.line((left, top, left, bottom), fill="#5f6368", width=3)
    draw.line((left, bottom, right, bottom), fill="#5f6368", width=3)
    labels = ("90-100", "80-89", "70-79", "60-69", "50-59", "0-49")
    maximum = max(max(stats["counts"], default=1), 1)
    points = []
    for index, (label, count) in enumerate(zip(labels, stats["counts"])):
        x = left + 65 + index * 140
        y = bottom - int(count / maximum * 235)
        points.append((x, y))
        draw.text((x - 35, bottom + 15), label, fill="#3c4043", font=label_font)
        draw.text((x - 8, y - 32), str(count), fill="#1769aa", font=label_font)
    if len(points) > 1:
        smooth = []
        extended = [points[0], *points, points[-1]]
        for index in range(1, len(extended) - 2):
            p0, p1, p2, p3 = extended[index - 1:index + 3]
            for step in range(16):
                t = step / 16
                t2, t3 = t * t, t * t * t
                x = 0.5 * ((2 * p1[0]) + (-p0[0] + p2[0]) * t + (2*p0[0] - 5*p1[0] + 4*p2[0] - p3[0]) * t2 + (-p0[0] + 3*p1[0] - 3*p2[0] + p3[0]) * t3)
                y = 0.5 * ((2 * p1[1]) + (-p0[1] + p2[1]) * t + (2*p0[1] - 5*p1[1] + 4*p2[1] - p3[1]) * t2 + (-p0[1] + 3*p1[1] - 3*p2[1] + p3[1]) * t3)
                smooth.append((round(x), round(y)))
        smooth.append(points[-1])
        draw.line(smooth, fill="#1769aa", width=5, joint="curve")
    for x, y in points:
        draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill="#d93025")
    image.save(output)


def _analysis_text(offering, parsed, stats):
    units = store.rows("SELECT project_title,source_skills FROM curriculum_units WHERE offering_id=? AND review_action<>'删除' ORDER BY seq", (offering["id"],))
    projects = [item["project_title"] for item in units[:8]]
    project_text = "、".join(projects[:4]) or "课程核心知识与技能项目"
    excellent = stats["counts"][0] + stats["counts"][1]
    below = stats["counts"][4] + stats["counts"][5]
    return {
        "coverage": f"考核依据《{offering['course_name']}》课程标准和授课计划，覆盖{project_text}等主要内容，兼顾基础知识、规范操作、任务实施和综合应用，与课程目标及实际教学内容基本一致。",
        "difficulty": f"综合成绩90分及以上{stats['counts'][0]}人，80至89分{stats['counts'][1]}人，60分以下{below}人；平均分{stats['average']:.2f}分，标准差{stats['deviation']:.2f}。成绩具有一定区分度，具体考核难度仍需结合过程记录、试题和评分材料复核。",
        "mastery": f"综合成绩80分及以上共{excellent}人，表明部分学生能较好完成课程核心任务。低分学生需继续巩固基础知识、操作步骤和成果检查能力；具体薄弱知识点应结合学生作品及评分表进一步分析。",
        "improvement": "后续教学应增加重点和易错环节的分层练习、现场反馈与问题复盘，对基础薄弱学生加强示范和阶段检查；考核中进一步细化评分标准、保留过程证据，并依据成绩分布持续优化难度梯度。",
    }


def _set_cell_font(cell, size=9, bold=False):
    """设置单元格字体为宋体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "宋体"
            run.font.size = Pt(size)
            run.font.bold = bold
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.insert(0, rfonts)
            rfonts.set(qn("w:eastAsia"), "宋体")


def _fill_term_and_checkboxes(paragraph, term_label, is_exam):
    """填充学期文本和复选框
    模板原文: "20  --20  学 年第   学期          课程考核方式(考试     考查  )"
    .doc转.docx后表单控件复选框丢失，直接插入字符复选框
    """
    # 构建新的学期行文本
    # 格式: "20XX-20XX学年 第□期中 ☑期末 学期    课程考核方式(☑考试 □考查)"
    midterm_box = '□'  # 期中（成绩分析默认是期末）
    final_box = '☑'    # 期末
    exam_box = '☑' if is_exam else '□'
    assess_box = '□' if is_exam else '☑'

    new_text = (
        f"{term_label}（{midterm_box}期中 {final_box}期末）"
        f"          "
        f"课程考核方式（{exam_box}考试  {assess_box}考查）"
    )

    # 写回段落（保留第一个run的格式）
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(new_text)


def _get_template_for_offering(offering_id):
    """获取课程实例对应的成绩分析模板路径"""
    rows = store.rows(
        "SELECT template_path FROM template_files WHERE offering_id=? AND document_type='成绩分析' LIMIT 1",
        (offering_id,),
    )
    if rows and Path(rows[0]["template_path"]).exists():
        return rows[0]["template_path"]
    return str(DEFAULT_TEMPLATE_PATH)


def generate_grade_analysis(offering_id, pdf_payload, source_filename, exam_date="", question_source="自命题", exam_mode="其他方式", marking_mode="教师本人自阅"):
    import shutil

    offerings = store.rows("SELECT * FROM offerings WHERE id=?", (offering_id,))
    if not offerings:
        raise ValueError("课程实例不存在。")
    offering = offerings[0]
    parsed = parse_grade_pdf(pdf_payload)
    if parsed["course_name"] != offering["course_name"]:
        raise ValueError(f"PDF课程为《{parsed['course_name']}》，当前选择的是《{offering['course_name']}》，请重新选择。")
    stats = _statistics(parsed["students"])
    absent_count = sum(1 for item in parsed["students"] if item["absent"])
    failed = [item for item in parsed["students"] if item["absent"] or item["comprehensive_score"] is None or item["comprehensive_score"] < 60]
    analysis = _analysis_text(offering, parsed, stats)
    term_parts = offering["term"].split("-")
    term_label = f"{term_parts[0]}-{term_parts[1]}学年 第{'一' if term_parts[-1]=='1' else '二'}学期"
    assessment_method = parsed["assessment_method"]  # 考试 / 考查
    is_exam = assessment_method == "考试"

    output_dir = Path(store.get_setting("output_root", ROOT / "生成结果")) / offering["term"] / offering["course_name"] / "成绩分析"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{offering['course_name']} {parsed['class_name']} 试卷分析.docx"

    # 获取模板
    template_path = _get_template_for_offering(offering_id)

    with tempfile.TemporaryDirectory(prefix="grade-analysis-") as temp_dir:
        temp = Path(temp_dir)
        chart_path = temp / "chart.png"
        _make_chart(stats, chart_path)

        # 复制模板
        import shutil
        shutil.copy2(template_path, str(output_path))
        doc = Document(str(output_path))
        table = doc.tables[0]

        # --- 学期和考核方式复选框 ---
        _fill_term_and_checkboxes(doc.paragraphs[2], term_label, is_exam)

        # --- 基本信息表（行0-4）---
        # 标签在偶数列，值在奇数列（0,2,4,6,8,10,12,14,16...）
        # 实际结构：每行8对标签-值（但有合并单元格）
        course_nature = offering.get("course_nature") or ("必修课" if "必修" in parsed["course_nature_raw"] else "选修课")
        teacher_name = offering.get("teacher_name", "") or store.get_setting("teacher_name", parsed["teacher"])
        department = offering.get("department", "") or store.get_setting("department", "经济贸易系")

        # (行索引, 值列索引, 值)
        fields = [
            (0, 2, department),       # 所在系
            (0, 6, offering["major"]),  # 专业
            (0, 10, parsed["class_name"]),  # 班级
            (0, 14, offering["course_name"]),  # 课程名称
            (1, 2, teacher_name),     # 任课教师
            (1, 6, offering["total_hours"]),  # 学时数
            (1, 10, offering["credits"]),  # 学分数
            (1, 14, course_nature),    # 课程性质
            (2, 2, parsed["expected_count"]),  # 应考人数
            (2, 6, len(parsed["students"])),   # 实考人数
            (2, 10, absent_count),    # 缺考人数
            (2, 14, question_source),  # 出题方式
            (3, 2, 100),               # 卷面总分
            (3, 6, exam_date),         # 考试日期
            (3, 10, exam_mode),        # 考试方式
            (3, 14, marking_mode),     # 阅卷方式
            (4, 2, f"{stats['pass_rate']:.2f}%"),  # 及格率
            (4, 6, f"{stats['highest']:g}"),       # 最高分
            (4, 10, f"{stats['lowest']:g}"),       # 最低分
            (4, 14, f"{stats['average']:.2f}"),    # 平均分
        ]
        for row, col, value in fields:
            cell = table.rows[row].cells[col]
            cell.text = str(value)
            _set_cell_font(cell, size=9)

        # --- 成绩分布表（行6-11）---
        for i in range(6):
            row_idx = 6 + i
            count_cell = table.rows[row_idx].cells[14]
            count_cell.text = str(stats["counts"][i])
            _set_cell_font(count_cell, size=9)
            pct_cell = table.rows[row_idx].cells[16]
            pct_cell.text = f"{stats['percentages'][i]:.2f}%"
            _set_cell_font(pct_cell, size=9)

        # --- 插入成绩分布曲线图 ---
        chart_cell = table.rows[5].cells[1]
        last_para = chart_cell.paragraphs[-1]
        run = last_para.add_run()
        run.add_picture(str(chart_path), width=Inches(4.5))
        chart_cell.paragraphs[-1].alignment = 1  # 居中

        # --- 考核内容分析（行12列1）---
        analysis_cell = table.rows[12].cells[1]
        analysis_sections = [
            ("一、考核内容覆盖面分析", analysis["coverage"]),
            ("二、考核内容难易程度分析", analysis["difficulty"]),
            ("三、学生对知识点/技能点掌握情况分析", analysis["mastery"]),
            ("四、教师在今后教学中应注意的问题、改进的思路或建议", analysis["improvement"]),
        ]
        paragraphs = list(analysis_cell.paragraphs)
        for heading, content in reversed(analysis_sections):
            found = False
            for pi, para in enumerate(paragraphs):
                if para.text.strip() == heading:
                    if pi + 1 < len(paragraphs) and paragraphs[pi + 1].text.strip():
                        new_para = analysis_cell.add_paragraph(content)
                        new_para._p.getparent().remove(new_para._p)
                        para._p.addnext(new_para._p)
                    else:
                        if pi + 1 < len(paragraphs):
                            paragraphs[pi + 1].text = content
                    found = True
                    break
            if not found:
                analysis_cell.add_paragraph(heading)
                analysis_cell.add_paragraph(content)

        # 设置分析内容字体
        for para in analysis_cell.paragraphs:
            for run in para.runs:
                run.font.name = "宋体"
                run.font.size = Pt(9)
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = rpr.makeelement(qn("w:rFonts"), {})
                    rpr.insert(0, rfonts)
                rfonts.set(qn("w:eastAsia"), "宋体")

        # --- 不及格/缺考名单（行13-15）---
        def _get_student_cells(row_idx, start_col=2):
            seen = set()
            cells = []
            for ci in range(start_col, len(table.rows[row_idx].cells)):
                cell = table.rows[row_idx].cells[ci]
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                cells.append(cell)
            return cells

        number_cells = _get_student_cells(13, 2)
        name_cells = _get_student_cells(14, 2)
        score_cells = _get_student_cells(15, 2)
        capacity = min(len(number_cells), len(name_cells), len(score_cells))
        if len(failed) > capacity:
            raise ValueError(f"不及格或缺考学生共{len(failed)}人，当前模板名单区域只能填写{capacity}人。")

        for i in range(capacity):
            if i < len(failed):
                item = failed[i]
                number_cells[i].text = item["student_no"]
                name_cells[i].text = item["student_name"]
                score_cells[i].text = "缺考" if item["absent"] else f"{(item['comprehensive_score'] or 0):g}"
            else:
                number_cells[i].text = ""
                name_cells[i].text = ""
                score_cells[i].text = ""
            for cell in [number_cells[i], name_cells[i], score_cells[i]]:
                _set_cell_font(cell, size=9)

        # --- 删除表后的说明段落 ---
        tbl = table._tbl
        parent = tbl.getparent()
        elements_to_remove = []
        remove_started = False
        for child in list(parent):
            if child is tbl:
                remove_started = True
                continue
            if remove_started:
                elements_to_remove.append(child)
        for elem in elements_to_remove:
            parent.remove(elem)

        doc.save(str(output_path))

    store.execute(
        "INSERT INTO grade_analysis_documents(offering_id,class_name,source_filename,output_path) VALUES (?,?,?,?)",
        (offering_id, parsed["class_name"], source_filename, str(output_path)),
    )
    return output_path, {
        "term_label": term_label,
        "assessment_method": assessment_method,
        "stats": stats,
        "failed_count": len(failed),
    }
