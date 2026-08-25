import hashlib
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

import store


HEADING_RE = re.compile(r"^[一二三四五六七八九十]+、|^（[一二三四五六七八九十]+）|^\d+[、.]\s*")

# 字段定义：(关键词列表, 字段名, 内容类型, 数据来源, 匹配模式)
# 匹配模式: label=必须是标签格式(短文本+冒号/空白), heading=标题级, any=任意位置
FIELD_SPECS = (
    # === 基本信息类 ===
    (("课程名称及课程编号", "课程名称", "课程编号"), "课程基本信息", "事实填写", "课程实例、教学安排表", "label"),
    (("课程类型",), "课程类型", "事实填写", "课程实例、教学安排表", "label"),
    (("课程性质",), "课程性质", "事实填写", "课程实例、教学安排表", "label"),
    (("学时学分",), "学时学分", "事实填写", "课程实例、教学安排表", "label"),
    (("总学时",), "总学时", "事实填写", "课程实例、教学安排表", "label"),
    (("理论学时", "讲授学时"), "理论学时", "事实填写", "课程实例、教学安排表", "label"),
    (("实践学时",), "实践学时", "事实填写", "课程实例、教学安排表", "label"),
    (("实验学时",), "实验学时", "事实填写", "课程实例、教学安排表", "label"),
    (("学分",), "学分", "事实填写", "课程实例、教学安排表", "label"),
    (("考核类型",), "考核类型", "事实填写", "课程实例、教学安排表", "label"),
    (("考核方式",), "考核方式", "事实填写", "课程实例、教学安排表", "label"),
    (("开设学期", "授课学期"), "开设学期", "事实填写", "课程实例", "label"),
    (("授课班级",), "授课班级", "按班事实填写", "排课记录", "label"),
    (("授课日期", "日期"), "授课日期", "按班事实填写", "排课记录、进程表、校历", "label"),
    (("教学环境", "授课地点", "实训地点"), "教学环境", "按班事实填写", "排课记录", "label"),
    (("授课教师", "主讲教师"), "授课教师", "事实填写", "系统设置", "label"),
    (("实训教师",), "实训教师", "可选事实填写", "系统设置", "label"),
    (("周次",), "周次", "按排课填写", "排课记录、进程表、校历", "label"),
    (("学时",), "任务学时", "按排课填写", "课程任务、排课记录", "label"),
    (("内容提要",), "教学任务", "结构生成", "课程任务、课程蓝本", "label"),
    (("教学任务", "章节、项目、情景"), "教学任务", "结构生成", "课程任务", "label"),
    (("适用专业",), "适用专业", "事实填写", "课程实例", "label"),
    (("编制单位",), "编制单位", "事实填写", "教研室设置", "label"),
    (("合作单位",), "合作单位", "事实填写", "企业信息", "label"),
    (("课程代码",), "课程代码", "事实填写", "教学安排表", "label"),
    (("所属系部", "所属系部（教研室）"), "所属系部", "事实填写", "学校设置", "label"),

    # === 分析撰写类 ===
    (("课程性质",), "课程性质", "分析撰写", "课程实例、课程语义模型、先导后续课程", "heading"),
    (("知识目标", "认知目标"), "知识目标", "分析撰写", "课程语义模型、教材知识体系", "heading"),
    (("能力目标",), "能力目标", "分析撰写", "课程语义模型、项目成果与岗位任务", "heading"),
    (("思政目标", "课程思政", "思想政治素质", "德育目标"), "思政目标", "分析撰写", "课程语义模型、项目风险与社会价值", "heading"),
    (("素质目标",), "素质目标", "分析撰写", "课程语义模型、工作过程与职业要求", "heading"),
    (("总体思路", "设计理念", "设计思路"), "课程设计总体思路", "分析撰写", "课程语义模型、四新、模板要求", "heading"),
    (("考核评价", "考核方案", "考试与评价", "考核与评价"), "考核评价", "分析撰写", "课程任务、成果类型、评价规则", "heading"),
    (("教师知识", "教师素质", "教师要求", "教师基本要求", "专任教师"), "教师知识能力要求", "分析撰写", "课程语义模型、工具与安全要求", "heading"),
    (("教材学情", "学情分析"), "教材学情分析及教育理念", "分析撰写", "教材难度、任务先后关系、班级情况", "heading"),
    (("教学场景设计",), "教学场景设计", "分析撰写", "教学环境、任务类型、教学方法", "heading"),
    (("教学步骤与内容", "教学过程", "教学活动"), "教学活动流程", "按任务撰写", "课程任务、项目资源、源码与PPT", "heading"),
    (("教法学法", "教法", "学法"), "教法学法", "按任务撰写", "教学活动、学生基础与任务类型", "heading"),
    (("达成目标", "目标达成"), "达成目标", "按任务撰写", "知识、能力、思政和素质目标", "heading"),
    (("课堂小结",), "课堂小结", "按任务撰写", "本次任务目标与验收标准", "heading"),
    (("课后作业",), "课后作业", "按任务撰写", "项目成果、评价标准", "heading"),
    (("教学反思",), "教学反思", "按任务撰写", "目标、活动、评价观察点", "heading"),

    # === 结构生成类 ===
    (("课程内容划分", "课时分配", "学习情境", "课程结构与内容", "课程内容及要求", "课程结构"), "课程内容与学时", "结构生成", "课程蓝本、课程任务", "heading"),
    (("教学时间", "课时"), "教学时间", "按任务填写", "任务学时与课堂活动分配", "label"),

    # === 资源类 ===
    (("教材编写与选用", "教材选用", "教材"), "教材选用", "事实与分析", "课程实例、教材资源", "heading"),
    (("课程资源", "资源准备", "教学资源准备", "教学资源"), "教学资源", "资源映射", "项目资源、PPT、源码、实训文档", "heading"),

    # === 成绩分析专用 ===
    (("所在系", "系部", "系（部）"), "所属系部", "事实填写", "教学安排表", "label"),
    (("专   业", "专业"), "适用专业", "事实填写", "课程实例", "label"),
    (("班   级", "班级"), "授课班级", "事实填写", "教学安排表、成绩表", "label"),
    (("任课教师",), "授课教师", "事实填写", "教学安排表、系统设置", "label"),
    (("学时数",), "总学时", "事实填写", "教学安排表", "label"),
    (("学分数",), "学分", "事实填写", "教学安排表", "label"),
    (("课程性质",), "课程性质", "事实填写", "教学安排表", "label"),
    (("应考人数",), "应考人数", "事实填写", "成绩表", "label"),
    (("实考人数",), "实考人数", "事实填写", "成绩表", "label"),
    (("缺考人数",), "缺考人数", "事实填写", "成绩表", "label"),
    (("出题方式",), "出题方式", "事实填写", "人工填写", "label"),
    (("卷面总分",), "卷面总分", "事实填写", "默认100", "label"),
    (("考试日期",), "考试日期", "事实填写", "人工填写", "label"),
    (("考试方式",), "考试方式", "事实填写", "人工填写", "label"),
    (("阅卷方式",), "阅卷方式", "事实填写", "人工填写", "label"),
    (("及格率", "及格率%"), "及格率", "事实填写", "成绩统计", "label"),
    (("最 高 分", "最高分"), "最高分", "事实填写", "成绩统计", "label"),
    (("最 低 分", "最低分"), "最低分", "事实填写", "成绩统计", "label"),
    (("平 均 分", "平均分"), "平均分", "事实填写", "成绩统计", "label"),
    (("标准分数段", "分数段"), "成绩分布", "结构生成", "成绩统计", "heading"),
    (("考核内容分析", "考核内容"), "考核内容分析", "分析撰写", "成绩分析AI生成", "heading"),
    (("成绩不及格", "不及格名单", "缺考名单"), "不及格缺考名单", "结构生成", "成绩表", "heading"),
)


def _format(paragraph):
    run = next((item for item in paragraph.runs if item.text.strip()), None)
    font = run.font if run else None
    paragraph_format = paragraph.paragraph_format
    east_asia = None
    if run is not None:
        r_fonts = run._element.find(qn("w:rPr"))
        if r_fonts is not None:
            fonts_el = r_fonts.find(qn("w:rFonts"))
            if fonts_el is not None and fonts_el.get(qn("w:eastAsia")):
                east_asia = fonts_el.get(qn("w:eastAsia"))
    return {
        "font_name": font.name if font else None,
        "east_asia_font": east_asia,
        "font_size_pt": font.size.pt if font and font.size else None,
        "bold": font.bold if font else None,
        "underline": bool(font.underline) if font is not None and font.underline else False,
        "alignment": int(paragraph.alignment) if paragraph.alignment is not None else None,
        "line_spacing": str(paragraph_format.line_spacing) if paragraph_format.line_spacing else None,
        "space_before_pt": paragraph_format.space_before.pt if paragraph_format.space_before else None,
        "space_after_pt": paragraph_format.space_after.pt if paragraph_format.space_after else None,
        "left_indent_pt": paragraph_format.left_indent.pt if paragraph_format.left_indent else None,
        "first_line_indent_pt": paragraph_format.first_line_indent.pt if paragraph_format.first_line_indent else None,
    }


def _requirements(text):
    items = re.findall(r"（(\d+)）(.*?)(?=（(\d+)）|$)", text)
    if items:
        return "\n".join(f"{index}. {content.strip('。； ')}" for index, content, _ in items)
    cleaned = re.sub(r"^（注[:：]", "", text).rstrip("） ")
    return cleaned


def _sources(text):
    sources = []
    mapping = (
        ("人才培养方案", "课程信息、专业定位与教师确认"),
        ("培养目标", "课程信息、专业定位与教材内容"),
        ("培养规格", "专业岗位要求、教材内容与教师确认"),
        ("先导", "课程衔接关系与教师确认"),
        ("后继", "课程衔接关系与教师确认"),
        ("岗位", "岗位标准与行业资料"),
        ("教材", "指定教材包"),
        ("学时", "教学安排表与课程主表"),
    )
    for keyword, source in mapping:
        if keyword in text and source not in sources:
            sources.append(source)
    return "；".join(sources)


def _cell_merge(cell):
    properties = cell._tc.tcPr
    span = properties.find(qn("w:gridSpan"))
    merge = properties.find(qn("w:vMerge"))
    return {
        "column_span": int(span.get(qn("w:val"))) if span is not None else 1,
        "vertical_merge": merge.get(qn("w:val"), "continue") if merge is not None else "",
    }


def _cell_text(cell):
    """获取单元格纯文本"""
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _is_label_text(text):
    """判断文本是否像标签（短文本、可能带冒号、不含过多内容）"""
    compact = re.sub(r"\s+", "", text)
    # 允许标签后带括号注释（如 教学任务（章节、项目、情景）），检查时剥离
    compact = re.sub(r"[（(][^）)]*[）)]$", "", compact)
    if len(compact) > 25:
        return False
    # 排除以"注"、"说明"、"示例"开头的文本
    if re.match(r"^(注|说明|示例|要求)[：:]", compact):
        return False
    # 排除包含谓语动词的句子（是、为、有、在、将等）
    if re.search(r"[是为有在将应需可]", compact):
        # 但如果只是"为XX"格式且很短，可能还是标签
        if len(compact) <= 6 and re.match(r"^[为应需可]", compact):
            pass  # 太短可能是标签
        else:
            return False
    # 标签通常以关键词开头，后跟冒号、空白或直接结束
    if re.match(r"^[\u4e00-\u9fa5A-Za-z0-9]+[：: ]*$", compact):
        return True
    # 标签也可能只是纯文本，不带冒号（表格左列）
    if len(compact) <= 12 and not re.search(r"[，。；！？、]", compact):
        # 额外检查：纯名词性短语，不包含句子结构
        return True
    return False


def _match_field(text, mode="any"):
    """
    匹配字段。
    mode: 
      - 'label': 仅匹配标签格式（短文本，标签样式）
      - 'heading': 仅匹配标题级别（以标题编号开头）
      - 'any': 任意匹配（但仍有基本长度限制）
    """
    compact = re.sub(r"\s+", "", text)
    if not compact:
        return None
    
    for keywords, field_name, kind, sources, match_mode in FIELD_SPECS:
        # 检查是否有关键词匹配
        matched_keyword = None
        for keyword in keywords:
            kw_compact = re.sub(r"\s+", "", keyword)
            if kw_compact in compact:
                matched_keyword = kw_compact
                break
        if not matched_keyword:
            continue
        rest = compact[len(matched_keyword):len(matched_keyword) + 1]

        # 根据模式验证
        if mode == "label":
            # "关键词：值"开头（如 学时学分：Ｘ学时（Ｘ学分））视为标签，不受占比限制
            starts_as_label = compact.startswith(matched_keyword) and rest in "：:"
            if not starts_as_label and not _is_label_text(text):
                continue
            # 标签模式：文本应该主要由关键词构成，关键词占文本比例应较高
            # 剥离尾部括号注释后计算比例（如 教学任务（章节、项目、情景）→ 教学任务）
            compact_clean = re.sub(r"[（(][^）)]*[）)]$", "", compact)
            ratio = len(matched_keyword) / len(compact_clean) if compact_clean else 1.0
            if ratio < 0.4 and not starts_as_label:
                continue
        
        elif mode == "heading":
            # 标题模式：必须以标题编号开头（一、二、..., 1. 2. ...）
            if not HEADING_RE.match(text.strip()):
                continue
            # 标题文本也不应过长
            if len(compact) > 30:
                continue
        
        # 排除常见的误匹配情况
        # 1. 如果文本包含"注"、"说明"、"示例"等，且关键词不在开头，跳过
        if re.search(r"^(注|说明|示例|要求)", compact) and matched_keyword not in compact[:4]:
            continue
        # 2. 如果文本以"（注"或"(注"开头，跳过
        if compact.startswith("（注") or compact.startswith("(注"):
            continue
        # 3. 如果文本包含"填写"、"待填"等占位词，跳过
        if "填写" in compact or "待填" in compact:
            continue
        
        return field_name, kind, sources
    
    return None


def _table_role(header_text, body_text):
    """识别表格角色"""
    text = f"{header_text}\n{body_text}"
    compact = re.sub(r"\s+", "", text)
    compact_header = re.sub(r"\s+", "", header_text)

    def has_any(keywords):
        return any(k in compact for k in keywords)

    def has_all(keywords):
        return all(k in compact for k in keywords)

    # 教学设计基本信息表
    if has_all(("周次", "课时", "授课班级")):
        return "教学设计基本信息表"
    # 授课计划明细表
    if has_all(("周次", "日期", "课堂教学")) or (has_all(("教学目标和主要内容", "作业、考核"))):
        return "授课计划明细表"
    # 教学设计教学组织表
    if has_any(("教学场景设计",)) and has_any(("参照体例", "教学过程", "教学步骤")):
        return "教学设计教学组织表"
    # 单元教学设计表
    if has_any(("教学步骤与内容",)) and has_any(("达成目标",)):
        return "单元教学设计表"
    # 学习情境描述表
    if compact_header.startswith("课程：") and "学习情境" in compact:
        return "学习情境描述表"
    if "教学方法和建议" in text and "考核与评价" in text:
        return "学习情境描述表"
    # 课程目标表
    if has_all(("知识目标", "能力目标")) and (has_any(("思政目标", "素质目标"))):
        return "课程目标表"
    # 考核评价表（须在课时分配表之前判断：其表头含"学习情境"，但主体是考核方式/权重）
    if has_any(("过程性评价", "终结性考核", "过程性考核", "结果评价", "增值评价", "综合评价")):
        return "考核评价表"
    # 课程内容与课时分配表
    if has_any(("学习情境", "课时分配")):
        return "课程内容与课时分配表"
    # 课程内容结构表
    if has_any(("序号", "模块")) and has_any(("学时分配", "学时")):
        return "课程内容结构表"
    if has_all(("评价类型", "评价内容")):
        return "考核评价表"
    # 职业能力分析表
    if has_any(("职业标准",)) and has_any(("工作任务", "能力元素", "相关知识")):
        return "职业能力分析表"
    # 教学内容要求表
    if has_any(("生成性成果",)) and has_any(("教学要求", "教学内容", "思政要点")):
        return "教学内容要求表"
    # 能力训练项目表
    if has_any(("能力训练项目",)) and has_any(("能力训练子目标", "训练方式")):
        return "能力训练项目表"
    # 课程基本信息表（封面表格）
    if has_any(("课程名称",)) and has_any(("课程类型", "课程性质", "课程代码", "所属系部", "授课学期")):
        return "课程基本信息表"
    return "普通表格"


def _template_contract(document):
    """生成模板结构契约"""
    tables = []
    for table_index, table in enumerate(document.tables):
        rows = []
        merged = []
        for row_index, row in enumerate(table.rows):
            values = []
            for cell_index, cell in enumerate(row.cells):
                value = _cell_text(cell)
                values.append(value)
                merge = _cell_merge(cell)
                if merge["column_span"] > 1 or merge["vertical_merge"]:
                    merged.append({"row": row_index, "column": cell_index, **merge})
            rows.append(values)
        header = " | ".join(rows[0]) if rows else ""
        body = "\n".join(" | ".join(row) for row in rows[1:])
        role = _table_role(header, body)
        tables.append({
            "index": table_index,
            "role": role,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "header_text": header,
            "merged_cells": merged,
            "repeat_mode": {
                "课程内容与课时分配表": "按课程项目和子任务重复",
                "课程内容结构表": "按课程项目重复",
                "考核评价表": "按选定考核项目重复",
                "学习情境描述表": "按学习情境及子情境重复",
                "授课计划明细表": "按排课任务重复",
                "单元教学设计表": "按周次或教学单元重复",
                "教学设计基本信息表": "与教学组织表成对，按周次或教学单元重复",
                "教学设计教学组织表": "与基本信息表成对，按所选体例重复",
                "职业能力分析表": "按职业标准项重复",
                "教学内容要求表": "按课程项目重复",
                "能力训练项目表": "按能力训练项目重复",
            }.get(role, "固定结构"),
        })
    return {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "section_count": len(document.sections),
        "tables": tables,
    }


def _row_unique_cells_with_indices(row):
    """获取一行中的唯一单元格及其原始列索引（跳过合并重复引用）"""
    cells = []
    seen = set()
    for col_idx, cell in enumerate(row.cells):
        marker = id(cell._tc)
        if marker in seen:
            continue
        seen.add(marker)
        cells.append((col_idx, cell))
    return cells


def _identify_label_value_pairs(table):
    """
    识别表格中的标签-值对。
    返回: list of (label_text, row_index, label_col, value_col)
    """
    pairs = []
    rows = list(table.rows)
    if len(rows) < 1:
        return pairs

    # 策略1：水平标签-值对（标签在左列，值在右列）
    for row_index, row in enumerate(rows):
        cells = _row_unique_cells_with_indices(row)
        if len(cells) < 2:
            continue
        for ci in range(len(cells) - 1):
            label_col, label_cell = cells[ci]
            value_col, value_cell = cells[ci + 1]
            label_text = _cell_text(label_cell)
            if not label_text or not _is_label_text(label_text) or len(label_text) > 20:
                continue
            value_text = _cell_text(value_cell)
            if value_text and not _is_label_text(value_text):
                continue
            pairs.append({
                "label": label_text,
                "row": row_index,
                "label_col": label_col,
                "value_col": value_col,
                "direction": "horizontal",
            })

    # 策略2：垂直标签-值对（标签在首行，值在下方行）
    if len(rows) >= 2:
        header_cells = _row_unique_cells_with_indices(rows[0])
        for col_idx, cell in header_cells:
            label_text = _cell_text(cell)
            if not label_text or not _is_label_text(label_text):
                continue
            pairs.append({
                "label": label_text,
                "row": 0,
                "label_col": col_idx,
                "value_col": col_idx,
                "direction": "vertical",
            })

    return pairs


def _build_slots(db, template_file_id, document, contract):
    """构建模板槽位"""
    db.execute("DELETE FROM template_slots WHERE template_file_id=?", (template_file_id,))
    slots = []
    seen_fields = {}  # (section, field_name) -> slot index，用于去重
    
    def add_slot(slot_key, locator, section, field_name, kind, repeat_scope, sources, label_text, fmt):
        """添加槽位，自动去重，优先保留表格中的槽位"""
        key = (section, field_name)
        if key in seen_fields:
            # 已存在同字段同章节的槽位，比较优先级
            existing_idx = seen_fields[key]
            existing_is_table = slots[existing_idx][1].startswith("table:")
            new_is_table = locator.startswith("table:")
            # 表格槽位优先级高于段落槽位
            if new_is_table and not existing_is_table:
                # 替换为表格槽位
                slots[existing_idx] = (slot_key, locator, section, field_name, kind, repeat_scope, sources, label_text, fmt)
            return
        seen_fields[key] = len(slots)
        slots.append((slot_key, locator, section, field_name, kind, repeat_scope, sources, label_text, fmt))
    
    # === 段落级槽位 ===
    current_heading = ""
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        
        # 更新当前章节标题（标题段落也可能是可填写的字段，如"（四）教学资源"）
        if HEADING_RE.search(text) and not text.startswith("（注"):
            current_heading = text.split("(")[0].split("（注")[0].strip()
        
        # 跳过注释、说明、示例类段落
        if (text.startswith("（注") or text.startswith("(注") or 
            text.startswith("注：") or text.startswith("注:") or
            text.startswith("说明：") or text.startswith("示例") or
            text.startswith("（说明") or text.startswith("(说明")):
            continue
        
        # 尝试匹配字段（先尝试 label 模式匹配"标签：值"格式，再尝试 heading 模式匹配标题级字段）
        matched = _match_field(text, mode="label")
        if not matched:
            matched = _match_field(text, mode="heading")
        if matched:
            field_name, kind, sources = matched
            add_slot(
                f"paragraph:{index}:{field_name}",
                f"paragraph:{index}",
                current_heading,
                field_name, kind, "单次", sources, text,
                _format(paragraph)
            )
    
    # === 表格级槽位 ===
    for table_index, table in enumerate(document.tables):
        role = contract["tables"][table_index]["role"]
        repeat_scope = contract["tables"][table_index]["repeat_mode"]
        
        # 识别表格中的标签-值对
        pairs = _identify_label_value_pairs(table)
        
        for pair in pairs:
            label_text = pair["label"]
            matched = _match_field(label_text, mode="label")
            if not matched:
                continue
            
            field_name, kind, sources = matched
            row_index = pair["row"]
            
            # 获取标签单元格的格式（用于定位）
            label_cell = table.rows[row_index].cells[pair["label_col"]]
            fmt = _format(label_cell.paragraphs[0]) if label_cell.paragraphs else {}
            
            # 构建定位器：包含表格索引、行、列和方向
            locator = f"table:{table_index}/row:{row_index}/col:{pair['value_col']}/{pair['direction']}"
            slot_key = f"table:{table_index}:{field_name}:{pair['direction']}"
            
            add_slot(
                slot_key,
                locator,
                role,
                field_name, kind, repeat_scope, sources, label_text,
                fmt
            )
        
        # 对于明细表（重复结构的表），额外识别表头字段
        if role in ("授课计划明细表", "单元教学设计表", "课程内容结构表", "课程内容与课时分配表"):
            # 表头行通常是第一行
            if len(table.rows) >= 2:
                header_row = table.rows[0]
                header_cells = []
                seen = set()
                for cell in header_row.cells:
                    marker = id(cell._tc)
                    if marker in seen:
                        continue
                    seen.add(marker)
                    header_cells.append(cell)
                
                for col_index, cell in enumerate(header_cells):
                    header_text = _cell_text(cell)
                    if not header_text:
                        continue
                    matched = _match_field(header_text, mode="label")
                    if matched:
                        field_name, kind, sources = matched
                        fmt = _format(cell.paragraphs[0]) if cell.paragraphs else {}
                        locator = f"table:{table_index}/col:{col_index}/header"
                        slot_key = f"table:{table_index}:header:{field_name}"
                        add_slot(
                            slot_key,
                            locator,
                            role,
                            field_name, kind, repeat_scope, sources, header_text,
                            fmt
                        )
    
    # 写入数据库
    for slot in slots:
        db.execute(
            "INSERT INTO template_slots (template_file_id,slot_key,locator,section_title,field_name,content_kind,repeat_scope,source_priority,instruction_text,format_json) VALUES (?,?,?,?,?,?,?,?,?,?)",
            (template_file_id, *slot[:-1], json.dumps(slot[-1], ensure_ascii=False)),
        )
    return len(slots)


def analyze_template(template_file_id):
    """分析模板文件，提取规则和槽位"""
    item = store.rows("SELECT * FROM template_files WHERE id=?", (template_file_id,))
    if not item:
        raise ValueError("模板文件不存在。")
    path = Path(item[0]["template_path"])
    if not path.exists() or path.suffix.lower() != ".docx":
        raise ValueError("模板必须是存在的.docx文件。")
    
    document = Document(path)
    contract = _template_contract(document)
    rules = []
    current_heading = ""
    
    # === 提取规则（注释、说明、要求类内容）===
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if HEADING_RE.search(text) and not text.startswith("（注"):
            current_heading = text.split("(")[0].split("（注")[0].strip()
        if (text.startswith("（注") or text.startswith("(注") or
            text.startswith("注：") or text.startswith("注:") or
            "要求：" in text or "（注：" in text):
            rules.append(("段落", f"paragraph:{index}", current_heading, text, _format(paragraph)))
    
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                text = _cell_text(cell)
                if not text:
                    continue
                if "注：" in text or "要求：" in text or "……" in text or "填写" in text:
                    fmt = _format(next((p for p in cell.paragraphs if p.text.strip()), cell.paragraphs[0]))
                    rules.append(("表格单元格", f"table:{table_index}/row:{row_index}/cell:{cell_index}", current_heading, text, fmt))
    
    source_hash = hashlib.sha256(path.read_bytes()).hexdigest()
    
    with store.connect() as db:
        # 保存规则
        db.execute("DELETE FROM template_rules WHERE template_file_id=?", (template_file_id,))
        for seq, (location_type, location_ref, section, instruction, fmt) in enumerate(rules, 1):
            db.execute(
                """INSERT INTO template_rules
                (template_file_id,seq,location_type,location_ref,section_title,instruction_text,content_requirements,data_sources,format_json)
                VALUES (?,?,?,?,?,?,?,?,?)""",
                (template_file_id, seq, location_type, location_ref, section, instruction,
                 _requirements(instruction), _sources(instruction), json.dumps(fmt, ensure_ascii=False)),
            )
        
        # 保存分析结果
        db.execute(
            """INSERT INTO template_analyses (template_file_id,source_hash,rule_count,contract_json,analysis_status,analyzed_at)
            VALUES (?,?,?,?,'待确认',CURRENT_TIMESTAMP)
            ON CONFLICT(template_file_id) DO UPDATE SET source_hash=excluded.source_hash,rule_count=excluded.rule_count,
            contract_json=excluded.contract_json,analysis_status='待确认',analyzed_at=CURRENT_TIMESTAMP""",
            (template_file_id, source_hash, len(rules), json.dumps(contract, ensure_ascii=False)),
        )
        
        # 构建槽位
        slot_count = _build_slots(db, template_file_id, document, contract)
        db.commit()
    
    return len(rules)
