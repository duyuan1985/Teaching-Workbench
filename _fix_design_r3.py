"""
第三轮修复：修复表5和表6
表5（教学设计·基本信息）结构：
  R0: 周次(col0-1合并) | 值(col2-4合并) | 授课班级(col5) | 值(col6-9合并) | col10-19空白列
  R1: 授课教师(col0-1) | 值(col2-4) | 授课日期(col5) | 值(col6-9)
  R2: 课程类型(col0-1) | 值(col2-4) | 教学环境(col5) | 值(col6-9)
  R3: 教学任务(col0-1) | 值(col2-9合并)
  R4: 教学目标 | 知识目标(col1) | 值(col2-9)
  R5: 教学目标 | 能力目标(col1) | 值(col2-9)
  R6: 教学目标 | 思政目标(col1) | 值(col2-9)
  R7: 教学目标 | 素质目标(col1) | 值(col2-9)
  R8: 教材学情分析(col0-1) | 值(col2-9)
  R9-R14: 教学要求(col0-1) | 重点(col2-6) | 难点(col7) | col8-9

表6（教学设计·教学组织）结构：
  R0: 教学场景设计(col0-1) | 值(col2-5合并)
  R1: 教学资源准备(col0-1) | 值(col2-5合并)
  R2: 教学过程(col0-1) | 教学步骤与内容(col2) | 教法学法(col3) | 达成目标(col4) | 时间(col5)
  R3-R4: 教学活动流程(col0) | 教学导入(col1) | 内容(col2) | 教法学法(col3) | 达成目标(col4) | 时间(col5)
  R5-R7: 教学活动流程(col0) | 任务1-3(col1) | 内容(col2) | 教法学法(col3) | 达成目标(col4) | 时间(col5)
  R8: 课堂小结(col0-1) | 内容(col2) | 教法学法(col3) | 达成目标(col4) | 时间(col5)  -- col2-5合并
  R9: 课后作业(col0-1) | 内容(col2-5合并)
  R10: 教学反思(col0-1) | 内容(col2-5合并)

问题：
1. 表5内容重复到所有列（合并单元格写入问题）
2. 表5教学任务应该写"第X章 情境名\n子情景名"
3. 表5R9-R14教学要求内容错误（应是重点难点，不是模板残留）
4. 表5R9-R14打√列不正确
5. 表6内容重复到col2-5（合并单元格问题）
6. 表6格式应为仿宋五号
7. 表5格式应为仿宋五号
"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from copy import deepcopy
import store

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)

# 获取任务数据
tasks = store.rows("SELECT * FROM tasks WHERE offering_id=20 ORDER BY seq")
sessions = store.rows("SELECT * FROM sessions WHERE offering_id=20 ORDER BY week_no, lesson_date")

# ============================================================
# 工具函数
# ============================================================
def clear_cell(cell):
    """清空单元格内容"""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''

def set_cell_text(cell, text, font_name='仿宋', size=10.5, bold=False):
    """设置单元格文本（仿宋五号=10.5pt），支持多行"""
    clear_cell(cell)
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    
    lines = text.split('\n')
    for li, line in enumerate(lines):
        if li > 0:
            run = p.add_run()
            run.add_break()
        else:
            run = p.add_run()
        run.text = line
        run.font.name = font_name
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        run.font.size = Pt(size)
        run.font.bold = bold

def get_chapter_num(seq):
    nums = ['一','二','三','四','五','六','七','八','九']
    return f"第{nums[seq-1] if seq <= 9 else seq}章"

# ============================================================
# 为每个任务生成表5和表6内容
# ============================================================

def gen_t5_content(task, idx, session):
    """生成表5内容"""
    title = task.get('title', f'任务{idx+1}')
    seq = task.get('seq', idx+1)
    
    # 解析章节和子任务
    parts = title.split('：')
    if len(parts) >= 2:
        chapter_name = parts[0]
        sub_tasks = [s.strip() for s in parts[1].split('、')]
    else:
        chapter_name = title
        sub_tasks = [title]
    
    chapter_num = get_chapter_num(seq)
    
    # 教学任务
    task_text = f"{chapter_num} {chapter_name}\n"
    for st in sub_tasks:
        task_text += f"子情景：{st}\n"
    
    # 日期和周次
    lesson_date = session.get('lesson_date', '') if session else ''
    week_no = session.get('week_no', '') if session else ''
    
    # 教学目标
    if seq == 1:
        knowledge = "1. 理解数据分析的基本概念、分类和适用场景\n2. 掌握电商核心数据分析指标体系\n3. 掌握常用数据分析方法理论\n4. 熟悉数据分析的基本流程和规范"
        ability = "1. 能够识别电商业务中的数据分析需求\n2. 能够运用对比分析法分析电商销售数据\n3. 能够选择合适的数据分析方法解决实际问题"
        sizheng = "1. 建立数据安全意识，遵守《个人信息保护法》\n2. 树立诚信分析精神，坚持真实客观的数据呈现"
        quality = "1. 培养数据驱动的分析思维\n2. 养成按流程规范操作的习惯"
    elif seq == 2:
        knowledge = "1. 陈述Excel数据分析工具的功能模块\n2. 掌握SUMIF、VLOOKUP等常用函数的语法\n3. 解释数据透视表的创建步骤和布局设置\n4. 总结Excel在数据分析中的优势与局限"
        ability = "1. 能够使用Excel函数进行数据计算\n2. 能够创建数据透视表进行多维分析\n3. 能够制作数据透视图展示分析结果"
        sizheng = "1. 树立精益求精的工匠精神，确保数据计算准确\n2. 增强规范操作意识，遵循数据处理标准"
        quality = "1. 培养独立思考和问题解决能力\n2. 养成检查和验证计算结果的习惯"
    else:
        knowledge = f"1. 陈述{chapter_name}的基本概念和核心功能\n2. 掌握{chapter_name}的主要API和常用方法\n3. 解释{chapter_name}在数据分析中的应用场景\n4. 总结{chapter_name}的操作流程和注意事项"
        ability = f"1. 能够安装和配置{chapter_name}环境\n2. 能够使用{chapter_name}进行数据处理和分析\n3. 能够编写代码完成{chapter_name}相关实训任务"
        sizheng = "1. 增强数据安全与隐私保护意识\n2. 培养科技报国情怀，关注技术服务乡村振兴"
        quality = "1. 培养代码规范编写和测试的习惯\n2. 提升团队协作和成果展示能力"
    
    # 教材学情分析
    analysis = (
        f"教材分析：本任务{chapter_name}是商务数据分析课程的重要组成部分，"
        f"在课程体系中起着{'基础铺垫' if seq <= 3 else '技能提升' if seq <= 6 else '综合应用' if seq <= 9 else '总结提升'}的作用。"
        f"本任务内部各部分逻辑关系为：先讲解基本概念，再演示操作步骤，最后通过实训巩固。\n"
        f"学情分析：学生已具备{'Python编程和电商基础知识' if seq <= 3 else '前序章节内容基础'}，"
        f"对实操环节兴趣较高，但对{'抽象概念理解' if seq <= 3 else '代码编写和调试'}存在畏难情绪，"
        f"需通过案例引导和分步演示降低难度。\n"
        f"教师教育理念：坚持以学生为中心、做中学，注重培养实践能力和创新思维。"
    )
    
    # 教学要求（重点、难点）
    if seq == 1:
        key_points = "1. 数据分析的基本概念和分类\n2. 电商核心数据分析指标\n3. 常用分析方法的应用"
        difficult = "1. 分析方法的选择和应用\n2. 业务需求与分析方法的对应\n3. 从数据中提炼有效结论"
    elif seq == 2:
        key_points = "1. Excel常用函数的语法和参数\n2. 数据透视表的创建步骤\n3. 数据透视图的制作"
        difficult = "1. 复杂函数的参数理解\n2. 多维交叉分析\n3. 大数据量下的性能优化"
    else:
        key_points = f"1. {chapter_name}的基本概念和核心原理\n2. 核心API的使用方法和参数\n3. 操作流程和注意事项"
        difficult = f"1. 代码调试和错误处理\n2. API参数组合使用\n3. 异常情况的处理"
    
    # 教学要求中的子任务列表
    req_items = []
    for st in sub_tasks:
        req_items.append(st)
    
    return {
        'week': str(week_no),
        'hours': '2',
        'class': '2022电商教学班',
        'teacher': '杜媛',
        'date': lesson_date,
        'course_type': '理实一体课程',
        'env': '801教室',
        'task': task_text,
        'knowledge': knowledge,
        'ability': ability,
        'sizheng': sizheng,
        'quality': quality,
        'analysis': analysis,
        'key_points': key_points,
        'difficult': difficult,
        'req_items': req_items,
    }

def gen_t6_content(task, idx, session):
    """生成表6内容"""
    title = task.get('title', f'任务{idx+1}')
    seq = task.get('seq', idx+1)
    
    parts = title.split('：')
    if len(parts) >= 2:
        chapter_name = parts[0]
        sub_tasks = [s.strip() for s in parts[1].split('、')]
    else:
        chapter_name = title
        sub_tasks = [title]
    
    chapter_num = get_chapter_num(seq)
    
    # 教学场景设计
    scene = f"801教室，配备多媒体教学设备、Python开发环境（Anaconda+Jupyter Notebook）、Excel 2019，电商实训数据集"
    
    # 教学资源准备
    resources = (
        f"1. 多媒体课件：{chapter_name}PPT课件\n"
        f"2. 微课视频：{sub_tasks[0]}\n"
        f"3. 实训数据集：电商{'交易' if seq <= 4 else '用户行为'}数据\n"
        f"4. Python代码模板和示例\n"
        f"5. 超星学习通在线资源"
    )
    
    # 导入
    import_methods = [
        ("案例展示法", f"展示电商数据分析典型案例——某农产品电商平台通过用户购买行为数据分析发现用户偏好，引导思考数据分析的价值。学生观看案例、思考问题、小组讨论。"),
        ("复习导入法", f"回顾上节课内容，提问关键知识点，自然引入{chapter_name}的学习。学生回答问题、查漏补缺、明确学习方向。"),
        ("任务引入法", f"提出本节课实训任务，明确学习目标和成果要求。学生了解任务、分析需求、制定计划。"),
        ("情境导入法", f"创设电商企业数据分析工作情境，代入数据分析师角色。学生进入角色、分析场景、明确职责。"),
    ]
    import_method = import_methods[idx % 4]
    
    # 任务1-3内容
    tasks_content = []
    for si, sub_task in enumerate(sub_tasks[:3]):
        sub_task = sub_task.strip()
        task_detail = (
            f"任务{si+1}：{sub_task}（{'知识讲解' if si == 0 else '技术演示' if si == 1 else '实操练习'}）\n"
            f"一、{'基本概念' if si == 0 else '环境准备' if si == 1 else '实训目标'}\n"
            f"  {'1. ' + sub_task + '的定义与内涵' if si == 0 else '1. 启动开发环境' if si == 1 else '1. 独立完成' + sub_task + '实训'}\n"
            f"  {'2. 分类和适用场景' if si == 0 else '2. 导入所需库和数据' if si == 1 else '2. 检查结果并优化'}\n"
            f"二、{'核心知识点' if si == 0 else '操作演示' if si == 1 else '实训步骤'}\n"
            f"  {'1. 基本原理与流程' if si == 0 else '1. 教师演示完整操作流程' if si == 1 else '1. 分析任务需求'}\n"
            f"  {'2. 关键参数和配置' if si == 0 else '2. 学生跟随操作' if si == 1 else '2. 编写代码并调试'}\n"
            f"  {'3. 代码示例和运行结果' if si == 0 else '3. 巡回指导' if si == 1 else '3. 成果展示与点评'}\n"
            f"德育渗透：{'数据安全意识' if si == 0 else '规范操作意识' if si == 1 else '团队协作精神'}"
        )
        teach_method = (
            f"{'讲授法、案例教学法、启发式提问' if si == 0 else '操作演示法、逐步讲解法、巡回指导法' if si == 1 else '实操练习法、任务驱动法、巡回指导法、成果展示法'}；\n"
            f"学生{'听讲、思考、做笔记' if si == 0 else '跟随操作、提问讨论、记录笔记' if si == 1 else '独立操作、小组互助、展示分享'}"
        )
        achieve = (
            f"{'理解' + sub_task + '的基本概念和原理' if si == 0 else '掌握' + sub_task + '的操作步骤和代码实现' if si == 1 else '能独立完成' + sub_task + '实训，培养问题解决能力'}"
        )
        time = f"{'25' if si == 0 else '30' if si == 1 else '25'}分钟"
        
        tasks_content.append({
            'content': task_detail,
            'method': teach_method,
            'achieve': achieve,
            'time': time,
        })
    
    # 如果只有2个子任务，加综合练习
    if len(sub_tasks) <= 2:
        tasks_content.append({
            'content': f"任务{len(sub_tasks)+1}：综合练习（实操训练）\n一、实训目标\n  1. 综合运用本节课知识解决实际问题\n  2. 培养独立分析和解决问题的能力\n二、实训步骤\n  1. 分析任务需求\n  2. 编写代码并调试\n  3. 检查结果并优化\n德育渗透：精益求精的工匠精神",
            'method': '实操练习法、任务驱动法；学生独立操作、小组互助、展示分享',
            'achieve': '能综合运用所学知识解决实际问题',
            'time': '20分钟',
        })
    
    # 课堂小结
    summary_types = [
        f"本节课围绕{chapter_name}展开，核心知识点：①基本概念与分类；②核心原理与方法；③操作流程与注意事项。技能要点：能独立完成{chapter_name}相关操作。课堂参与积极，实操练习认真。",
        f"本节课收获最大的三点：①掌握了{chapter_name}的基本概念；②学会了核心操作步骤；③理解了实际应用场景。需注意：①概念理解要准确；②操作步骤要规范；③结果检查要仔细。",
        f"Q1：{chapter_name}的核心功能是什么？A1：数据处理与分析。Q2：操作的关键步骤？A2：环境准备→参数设置→执行操作→检查结果。Q3：常见问题如何解决？A3：参考错误提示，检查参数和代码。",
        f"本节课{chapter_name}与前序内容的联系：都是数据分析工具的应用。区别：{chapter_name}更注重代码实现。核心要点：1.概念→2.原理→3.操作→4.实训。学习建议：多练习、多思考、多总结。",
    ]
    summary = summary_types[idx % 4]
    
    # 课后作业
    homework_types = [
        f"1. 整理本节课知识点笔记\n2. 完成{chapter_name}实训任务，输出分析结果\n3. 思考：{chapter_name}在电商业务中有哪些应用场景？\n4. 预习下一章内容",
        f"基础层：1. 梳理{chapter_name}知识框架图 2. 完成3道基础练习题\n提升层：3. 用{chapter_name}分析电商数据 4. 对比与其他工具的异同\n拓展层：5. 查阅{chapter_name}最新文档 6. 尝试解决实际问题",
        f"项目任务：基于{chapter_name}完成电商数据分析小项目\n要求：1. 确定分析目标 2. 使用{chapter_name}处理数据 3. 输出报告（含图表）4. 下节课展示",
        f"1. 回顾本节课学习过程，填写学习反思表\n2. 用自己的话解释{chapter_name}核心概念\n3. 修改完善课堂实训代码，添加注释\n4. 在学习通平台完成本节测验",
        f"问题情境：某电商平台需要分析用户行为数据\n任务：1. 明确分析目标 2. 选择合适方法（含{chapter_name}）3. 编写代码 4. 撰写结论和建议\n要求：2-3人一组，下周展示",
    ]
    homework = homework_types[idx % 5]
    
    # 教学反思
    reflection_types = [
        f"本节课教学效果良好，{chapter_name}概念讲解清晰，学生理解到位。实训环节参与积极，大部分学生能独立完成任务。成功经验：案例导入有效激发兴趣；分步演示降低难度。改进方向：概念讲解可压缩5分钟留给实操；部分学生代码基础薄弱需加强个别指导。",
        f"学生反馈积极，多数表示案例教学帮助理解抽象概念，实操练习增强了动手能力。存在问题：少数学生对{'代码调试' if seq <= 8 else '综合应用'}感到困难，需在后续课程中加强练习。教学调整：增加同伴互助环节，让掌握较快的学生帮助其他同学。",
        f"教学方法和策略整体有效。任务驱动法让学生有明确目标，演示法让操作步骤可视化。需改进：导入环节可更有趣味性；小组讨论时间可适当延长；德育渗透可更自然融入。下节课将尝试引入更多互动元素。",
    ]
    reflection = reflection_types[idx % 3]
    
    return {
        'scene': scene,
        'resources': resources,
        'import_content': import_method[1],
        'import_method': '案例展示法、提问引导法、分组讨论法',
        'import_achieve': '激发学习兴趣，了解数据分析的实际应用价值',
        'import_time': '5分钟',
        'tasks': tasks_content,
        'summary': summary,
        'homework': homework,
        'reflection': reflection,
    }

# ============================================================
# 修复所有表5和表6
# ============================================================
print("修复表5和表6...")

for ti in range(5, len(doc.tables), 2):
    t5 = doc.tables[ti]
    t6 = doc.tables[ti+1] if ti+1 < len(doc.tables) else None
    
    task_idx = (ti - 5) // 2
    if task_idx >= len(tasks):
        continue
    
    task = tasks[task_idx]
    session = sessions[task_idx] if task_idx < len(sessions) else None
    
    c5 = gen_t5_content(task, task_idx, session)
    c6 = gen_t6_content(task, task_idx, session)
    
    # ============================================================
    # 修复表5
    # ============================================================
    # R0: 周次(col0-1) | 值(col2) | 课时(col3) | 值(col4) | 授课班级(col5) | 值(col6)
    set_cell_text(t5.cell(0, 2), c5['week'])
    set_cell_text(t5.cell(0, 4), c5['hours'])
    set_cell_text(t5.cell(0, 6), c5['class'])
    
    # R1: 授课教师(col0-1) | 值(col2) | 授课日期(col5) | 值(col6)
    set_cell_text(t5.cell(1, 2), c5['teacher'])
    set_cell_text(t5.cell(1, 6), c5['date'])
    
    # R2: 课程类型(col0-1) | 值(col2) | 教学环境(col5) | 值(col6)
    set_cell_text(t5.cell(2, 2), c5['course_type'])
    set_cell_text(t5.cell(2, 6), c5['env'])
    
    # R3: 教学任务(col0-1) | 值(col2)
    set_cell_text(t5.cell(3, 2), c5['task'])
    
    # R4-R7: 教学目标
    set_cell_text(t5.cell(4, 2), c5['knowledge'])
    set_cell_text(t5.cell(5, 2), c5['ability'])
    set_cell_text(t5.cell(6, 2), c5['sizheng'])
    set_cell_text(t5.cell(7, 2), c5['quality'])
    
    # R8: 教材学情分析
    set_cell_text(t5.cell(8, 2), c5['analysis'])
    
    # R9: 教学要求标题行（重点）
    set_cell_text(t5.cell(9, 2), c5['key_points'])
    set_cell_text(t5.cell(9, 7), "重点")
    
    # R10: 教学要求（难点）
    set_cell_text(t5.cell(10, 2), c5['difficult'])
    set_cell_text(t5.cell(10, 7), "难点")
    
    # R11-R14: 教学要求子任务打√
    for ri in range(11, min(15, len(t5.rows))):
        if ri - 11 < len(c5['req_items']):
            set_cell_text(t5.cell(ri, 2), c5['req_items'][ri - 11])
            # 打√列
            if len(t5.rows[ri].cells) > 7:
                set_cell_text(t5.cell(ri, 7), "√")
        else:
            # 空行
            if len(t5.rows[ri].cells) > 2:
                set_cell_text(t5.cell(ri, 2), "")
    
    # ============================================================
    # 修复表6
    # ============================================================
    if t6:
        # R0: 教学场景设计(col2)
        set_cell_text(t6.cell(0, 2), c6['scene'])
        
        # R1: 教学资源准备(col2)
        set_cell_text(t6.cell(1, 2), c6['resources'])
        
        # R3: 教学导入 - 内容(col2)
        set_cell_text(t6.cell(3, 2), c6['import_content'])
        set_cell_text(t6.cell(3, 3), c6['import_method'])
        set_cell_text(t6.cell(3, 4), c6['import_achieve'])
        set_cell_text(t6.cell(3, 5), c6['import_time'])
        
        # R4: 导入续
        set_cell_text(t6.cell(4, 2), f"教师归纳讨论结果，明确学习目标。学生明确目标、调整状态。")
        set_cell_text(t6.cell(4, 3), "总结归纳法、讲授法")
        set_cell_text(t6.cell(4, 4), "明确学习目标，导入新课")
        set_cell_text(t6.cell(4, 5), "5分钟")
        
        # R5-R7: 任务1-3
        for ti2 in range(min(3, len(c6['tasks']))):
            ri = 5 + ti2
            if ri < len(t6.rows):
                tc = c6['tasks'][ti2]
                set_cell_text(t6.cell(ri, 2), tc['content'])
                set_cell_text(t6.cell(ri, 3), tc['method'])
                set_cell_text(t6.cell(ri, 4), tc['achieve'])
                set_cell_text(t6.cell(ri, 5), tc['time'])
        
        # R8: 课堂小结
        set_cell_text(t6.cell(8, 2), c6['summary'])
        
        # R9: 课后作业
        set_cell_text(t6.cell(9, 2), c6['homework'])
        
        # R10: 教学反思
        set_cell_text(t6.cell(10, 2), c6['reflection'])
    
    if (task_idx + 1) % 5 == 0:
        print(f"  已修复{(task_idx + 1)}/30对表5/表6")

print(f"\n总共修复了30对表5/表6")
doc.save(fp)
print(f"保存完成: {fp}")
