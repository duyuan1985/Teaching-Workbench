"""契约预览：真实开课数据（offering 20 商务数据分析 2023-2024-2）+ 兜底内容生成对照文档。"""

import json
import re
import time
import urllib.request
from pathlib import Path

import contract_filler
import store
from docx import Document

OUT = Path(f"data/tmp_verify/课程标准_契约预览_商务数据分析_{time.strftime('%H%M%S')}.docx")
OLLAMA = "http://127.0.0.1:11434/api/chat"
MODEL = "qwen3:8b"
USE_AI = False  # 本地8B模型8项内容耗时17分钟且5项超时；速度方案留待阶段2

store.initialize()
offering = store.rows("SELECT * FROM offerings WHERE id=20")[0]
lib = store.rows("SELECT * FROM template_library WHERE doc_type='课程标准' AND version_label='2023-2024'")[0]
contract = store.rows(
    "SELECT * FROM template_contracts WHERE library_id=? ORDER BY version DESC LIMIT 1", (lib["id"],))[0]
slots = store.rows(
    "SELECT * FROM contract_slots WHERE contract_id=? ORDER BY sort_order", (contract["id"],))


def ask(prompt, json_mode=False):
    body = json.dumps({
        "model": MODEL,
        "messages": [{"role": "user", "content": prompt + "\n/no_think"}],
        "stream": False,
        "options": {"num_predict": 800, "temperature": 0.6},
    }).encode()
    req = urllib.request.Request(OLLAMA, data=body, headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req, timeout=180) as resp:
        text = json.loads(resp.read())["message"]["content"]
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.S).strip()
    if json_mode:
        m = re.search(r"\{.*\}", text, flags=re.S)
        if not m:
            raise ValueError("no json")
        return json.loads(m.group(0))
    return text


COURSE = offering["course_name"]
MAJOR = offering["major"]
CREDITS = int(offering["credits"])
HOURS = offering["total_hours"]
CN_NUM = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六", 7: "七", 8: "八"}
entry_year = int(re.search(r"(20\d{2})", offering["teaching_class"]).group(1))
term_start = int(offering["term"].split("-")[0])
term_sem = int(offering["term"].rsplit("-", 1)[1])
semester_no = (term_start - entry_year) * 2 + term_sem
CTX = {
    "course": COURSE, "major": MAJOR, "office": f"{MAJOR}教研室",
    "partner": store.rows("SELECT setting_value v FROM settings WHERE setting_key='partner_company'")[0]["v"],
    "credit": CREDITS, "hours": HOURS,
    "basic_lines": {
        "课程名称及课程编号": f"{COURSE}（课程代码：{offering['course_code']}）",
        "课程类型": "专业核心课（必修）",
        "学时学分": f"{HOURS}学时（{CREDITS}学分）",
        "开设学期": f"第{CN_NUM[semester_no]}学期",
    },
}

FB = {
    "nature": (
        f"《{COURSE}》是{MAJOR}专业的专业核心课（必修），在第{CN_NUM[semester_no]}学期开设，共{HOURS}学时、{CREDITS}学分。"
        "本专业面向农村电商运营、农产品网络营销等岗位群，培养能够运用数据工具开展业务分析、支撑经营决策的高素质技术技能人才，"
        "而农产品电商业务分散、数据来源多、决策链条长，客观上要求从业者具备从数据中提炼经营线索的能力。"
        "本课程正是围绕这一培养规格，将数据指标认知、采集清洗、分析可视化等岗位典型任务转化为学习内容，"
        "使学生在真实业务情境中形成“用数据说话、靠数据决策”的职业思维方式，直接支撑专业人才培养目标中数据分析能力的落地。"
        "课程以《电子商务基础》《网店运营》等课程为先导，为后续《农产品网络营销》《顶岗实习》等课程及岗位实践提供数据分析方法与工具支撑，"
        "在课程体系中起着承上启下的枢纽作用。"
    ),
    "goals": {
        "knowledge": "1.了解电商数据分析的基本流程与指标体系\n2.掌握数据采集、清洗与整理的常用方法\n3.熟悉描述性统计与对比分析的适用场景\n4.掌握常用可视化图表的制作规范",
        "ability": "1.具有从电商平台采集业务数据的能力\n2.具有完成数据清洗与预处理的能力\n3.具有运用统计方法开展对比与趋势分析的能力\n4.具有制作数据分析报告与数据看板的能力",
        "ideology": "马列主义方面，坚持实事求是的数据观，用数据客观反映经营实际\n理想信念方面，树立服务乡村振兴的职业理想，认识数据赋能农村电商的价值\n核心价值观方面，诚信采集与解读数据，不篡改、不造假\n传统文化方面，结合农产品电商案例增强文化自信与乡土情怀\n职业道德方面，严守数据安全与个人信息保护底线\n个人素养方面，养成严谨细致、精益求精的分析习惯",
        "quality": "创新意识方面，鼓励从多角度探索数据规律，提出改进建议\n安全意识方面，规范数据存储与传输，防范信息泄露\n团队协作方面，在小组项目中合理分工、协同完成分析任务\n独立自主方面，能够独立完成从数据到结论的完整分析\n个人自律方面，按时交付任务成果，主动复盘改进",
    },
    "design": (
        "本课程以“数据驱动乡村振兴”为核心理念，立足农产品电商真实业务需求，遵循“岗位任务引领、项目载体驱动、理实一体实施”的总体思路，"
        "按照电商数据分析的实际工作过程，将课程内容重构为数据认知与采集、数据清洗与预处理、数据分析与可视化三个学习情境，"
        "构建“理论讲授—工具实操—项目综合”三位一体的教学框架，"
        "使学生在完成完整业务分析项目的过程中习得知识、训练技能、养成素养，实现教、学、做合一。"
    ),
    "method": (
        "根据本课程的教学目标要求和课程特点以及有关学情，选择适合于本课程的最优化教学法。综合考虑教学效果和教学可操作性等因素，"
        "本课程选用项目教学法、案例教学法和任务教学法。"
        "项目教学法是以完整的电商数据分析项目为载体，学生在项目实施过程中经历“明确任务—采集数据—分析处理—形成报告”的完整工作过程，"
        "在做中学、学中做；案例教学法选用农产品电商真实经营案例，引导学生剖析数据背后的业务问题，训练分析思维；"
        "任务教学法将每个学习情境分解为可交付的工作任务，以任务驱动学习，便于过程考核与成果检验。"
    ),
    "assess": (
        "本课程采用过程性考核与终结性考核相结合的评价方式，过程性考核占40%，终结性考核占60%。"
        "过程性考核贯穿三个学习情境，以签到、课堂表现、作业、章节测试等方式评价学生的任务完成情况与学习态度；"
        "终结性考核以综合分析报告（作品）形式考查学生综合运用数据分析方法解决实际业务问题的能力，"
        "重视教学过程评价与成果评价相结合，突出对学生动手能力和职业素养的考查。"
    ),
    "teacher": (
        "在思想政治方面，任课教师应强化课程思政意识，提升自身思想政治水平，将思想政治教育、数据诚信与职业道德教育融入专业教学全过程。"
        "在专业方面，一是要有扎实的电子商务与数据分析专业基础，系统掌握统计方法、分析工具与可视化技术；"
        "二是要实时学习行业新业态、新方法、新技能，关注电商数据化运营的最新实践并将其贯穿于授课之中；"
        "三是要有企业实践经历，能够引入真实业务数据与项目案例开展教学。"
        "在其他方面，要加强自身对数据安全、个人信息保护等方面的认知，具备改革创新意识，能够持续改进教学方法与手段。"
    ),
}

AI_OK = []
start = time.time()
if USE_AI:
    try:
        FB["nature"] = ask(
            f"为高职《{COURSE}》课程（{MAJOR}专业，第{CN_NUM[semester_no]}学期，{HOURS}学时{CREDITS}学分）撰写课程标准中“一、课程性质”段落（250字左右，纯文本）。"
            "必须按三方面展开：（1）课程类型（专业核心课、必修）；（2）围绕人才培养目标和培养规格写本课程在人才培养中的地位、作用和功能；"
            "（3）与先导、后继课程的关系。")
        AI_OK.append("课程性质")
    except Exception as e:
        print("AI失败 课程性质:", type(e).__name__)
    for key, title, extra in (
        ("design", "1、本课程设计的总体思路", "只写总体思路（理念+框架），不要写课时分配和考核细节，150字左右"),
        ("method", "3、教学方法描述", "按格式：根据本课程的教学目标要求…本课程选用XX教学法。XX教学法是：…（说明内涵特点），150-200字"),
        ("assess", "4、考试与评价方式", "按学校统一要求：过程性考核40%+终结性考核60%，150字左右"),
        ("teacher", "（一）教师知识素质要求", "按三方面写：思想政治（课程思政融入）；专业（扎实基础、实时学习新技能、实践经历）；其他（安全与创新认知），200字左右"),
    ):
        try:
            FB[key] = ask(f"为高职《{COURSE}》课程（{MAJOR}专业）撰写课程标准“{title}”：{extra}，纯文本。")
            AI_OK.append(title)
        except Exception as e:
            print(f"AI失败 {title}:", type(e).__name__)

print(f"AI完成 {len(AI_OK)} 项，耗时 {time.time()-start:.0f}s")

CTX["sections"] = {
    "一、课程性质": FB["nature"],
    "1、本课程设计的总体思路": FB["design"],
    "3、教学方法描述": FB["method"],
    "4.考试与评价方式（或方案）": FB["assess"],
    "（一）教师知识素质要求": FB["teacher"],
    "（二）教材编写与选用":
        f"本课程教材依据《{offering['textbook_version']}》编写，体现“任务驱动、项目载体”的校企合作开发特色，"
        "内容覆盖数据认知与采集、数据清洗与预处理、数据分析与可视化等模块，"
        "随电商行业新业态、新技术发展动态修订，确保教学内容与岗位要求同步。",
    "（三）课程资源的开发与利用":
        "课程资源包括：配套教学课件与教学大纲、实训案例数据集与源代码、"
        "在线数据分析平台账号、校企合作提供的农产品电商脱敏业务数据，"
        "以及课程网站发布的拓展学习资料，支持学生课前预习、课中实操与课后巩固。",
    "（四）推荐参考书、期刊、网站": "推荐参考书、期刊与网站如下。",
    "1.参考书": "1.《电商数据分析与数据化运营》，人民邮电出版社\n2.《商务数据分析基础》，电子工业出版社",
    "2.期刊": "1.《电子商务世界》\n2.《统计与决策》",
    "3.网站": "1.中国电子商务研究中心（www.100ec.cn）\n2.国家统计局数据平台（data.stats.gov.cn）",
}

# 学时分配：60 = 16 + 14 + 30；理论实践各半
CTX["scenarios"] = [
    {"name": "数据认知与采集",
     "tasks": [("电商数据指标体系认知", 6, 3, 3), ("平台数据采集实操", 6, 3, 3), ("问卷设计与数据获取", 4, 2, 2)],
     "weight": 24, "final": "数据分析方案", "final_weight": 10,
     "assess": [("电商数据指标体系认知", "作业", 5), ("平台数据采集实操", "课堂表现", 5), ("问卷设计与数据获取", "作业", 4)]},
    {"name": "数据清洗与预处理",
     "tasks": [("数据质量检查", 6, 3, 3), ("缺失值与异常值处理", 8, 4, 4)],
     "weight": 20, "final": "清洗报告", "final_weight": 10,
     "assess": [("数据质量检查", "作品", 6), ("缺失值与异常值处理", "作业", 4)]},
    {"name": "数据分析与可视化",
     "tasks": [("描述性统计分析", 8, 4, 4), ("对比与趋势分析", 6, 3, 3), ("可视化图表制作", 8, 4, 4), ("数据看板搭建", 8, 4, 4)],
     "weight": 56, "final": "综合分析报告", "final_weight": 40,
     "assess": [("描述性统计分析", "章节测试", 4), ("对比与趋势分析", "作业", 4), ("可视化图表制作", "作品", 4), ("数据看板搭建", "成果", 4)]},
]
CTX["scenario_detail"] = {
    "name": CTX["scenarios"][0]["name"], "hours": 16,
    "goals": "知识目标：\n1.了解电商数据分析流程\n2.掌握核心数据指标含义\n能力目标：\n1.能采集平台业务数据\n2.能设计调研问卷",
    "content": "以农产品电商平台为对象，认识流量、转化、商品、服务四大类核心数据指标，完成店铺与商品数据采集，设计并发放调研问卷获取一手数据，形成规范的数据集。",
    "methods": "教学方法：\n1.案例教学法\n2.任务驱动法\n建议：\n1.结合真实店铺后台演示\n2.分组采集、交叉核验数据",
    "materials": "·多媒体教学设备\n·数字课件与指标手册\n·平台后台演示账号\n·问卷调研平台",
    "assessment": "评价内容：指标体系理解、数据采集完成率与规范性、问卷设计质量。",
    "remark": "通过本学习情境使学生建立电商数据意识，掌握规范的数据采集方法，为后续清洗与分析奠定基础。",
}
CTX["subscenario_detail"] = {
    "name": "电商数据指标体系认知", "hours": 6,
    "goals": "教学目标：\n1.能说出流量、转化、客单价等核心指标的含义与计算口径\n2.能解释指标之间的业务关系",
    "content": "流量指标、转化指标、商品指标、服务指标四大类指标的含义、计算方法与业务含义，指标体系的整体框架。",
    "methods": "教学方法：\n1.讲授法\n2.案例教学法\n建议：\n1.用真实店铺后台数据演示\n2.开展指标连线与计算练习",
    "materials": "·多媒体教学设备\n·店铺后台演示账号\n·指标体系手册",
    "assessment": "评价内容：指标术语解释准确性、指标计算练习正确率。",
    "remark": "为后续数据采集与分析任务奠定指标基础。",
}

filler = contract_filler.ContractFiller(lib["file_path"], slots)
filled, skipped = filler.fill(CTX)
filler.fill_goal_table(FB["goals"])
filler.fill_hour_table(CTX["scenarios"])
filler.fill_assessment_table(CTX["scenarios"])
filler.fill_scenario_table(CTX)
removed, replaced = filler.cleanup(CTX)
OUT.parent.mkdir(parents=True, exist_ok=True)
filler.save(OUT)

doc = Document(OUT)
non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
print(f"\n开课: {offering['term']} {COURSE}（{MAJOR}）{CREDITS}学分{HOURS}学时 第{CN_NUM[semester_no]}学期")
print(f"输出: {OUT}")
print(f"槽位填充 {len(filled)} 项，跳过 {len(skipped)} 项: {skipped}")
print(f"清理指令段落 {removed} 个，替换残留占位 {replaced} 处")
print(f"段落 {len(doc.paragraphs)}（非空 {non_empty}），表格 {len(doc.tables)}")
