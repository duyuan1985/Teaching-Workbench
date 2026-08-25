"""
填充考核方案表第二列（学习内容）的内容
R1-R3: 第一部分（过程性考核），合并
R4: 第二部分（终结性考核）
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)
t4 = doc.tables[4]
tbl = t4._tbl
rows = tbl.findall(qn('w:tr'))

def set_cell_text(tc, text, font_name='仿宋', font_size='21', bold=False):
    """清除tc内容并设置新文本"""
    for p in tc.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            p.remove(r)
    paras = tc.findall(qn('w:p'))
    if not paras:
        p = OxmlElement('w:p')
        tc.append(p)
    else:
        p = paras[0]
        for extra in paras[1:]:
            tc.remove(extra)
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p.insert(0, pPr)
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'center')
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), 'auto')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), font_size)
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), font_size)
    rPr.append(szCs)
    if bold:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)

# R1C1: 第一部分（vMerge=restart已存在）
tc_r1c1 = rows[1].findall(qn('w:tc'))[1]
set_cell_text(tc_r1c1, '第一部分', '仿宋', '21')

# R2C1, R3C1: 合并continue，清除文字（已经是空的）
# R4C1: 第二部分（无合并）
tc_r4c1 = rows[4].findall(qn('w:tc'))[1]
# 确保R4C1没有vMerge
tcPr = tc_r4c1.find(qn('w:tcPr'))
if tcPr is not None:
    for vm in tcPr.findall(qn('w:vMerge')):
        tcPr.remove(vm)
set_cell_text(tc_r4c1, '第二部分', '仿宋', '21')

doc.save(fp)
print('已填充第二列内容')

# 验证
doc2 = Document(fp)
t4_2 = doc2.tables[4]
for ri in range(len(t4_2.rows)):
    texts = [c.text.strip()[:30] for c in t4_2.rows[ri].cells]
    print(f'  R{ri}: {texts}')
