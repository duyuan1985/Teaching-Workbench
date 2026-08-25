from docx import Document
from docx.oxml.ns import qn

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

# 检查Table 6的结构
t6 = doc.tables[6]
print(f'=== Table 6: {len(t6.rows)} rows, {len(t6.columns)} cols ===')
for ri in range(len(t6.rows)):
    row = t6.rows[ri]
    tcs = row._tr.findall(qn('w:tc'))
    texts = []
    for tc in tcs:
        txt = ''
        for p in tc.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                for elem in r:
                    if elem.tag == qn('w:t'):
                        txt += elem.text or ''
        # 检查gridSpan
        tcPr = tc.find(qn('w:tcPr'))
        gs = ''
        if tcPr is not None:
            gs_elem = tcPr.find(qn('w:gridSpan'))
            if gs_elem is not None:
                gs = f'(gs={gs_elem.get(qn("w:val"))})'
        texts.append(f'{txt[:40]}{gs}' if txt or gs else '(空)')
    print(f'R{ri}: {texts}')

# 检查模板Table 6
tmpl_fp = r'原始资料\模板\模板5：教学设计\模板5：教学设计 模板（2023-2024）.docx'
tmpl_doc = Document(tmpl_fp)
tmpl_t6 = tmpl_doc.tables[6]
print(f'\n=== 模板 Table 6: {len(tmpl_t6.rows)} rows ===')
for ri in range(len(tmpl_t6.rows)):
    texts = [c.text.strip()[:50] for c in tmpl_t6.rows[ri].cells]
    print(f'R{ri}: {texts}')
