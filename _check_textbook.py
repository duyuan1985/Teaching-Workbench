from docx import Document

docs = {
    '教学设计': r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx',
    '课程标准': r'生成结果\精修版\2023-2024-2《商务数据分析》课程标准 杜媛.docx',
    '授课计划': r'生成结果\精修版\2023-2024-2《商务数据分析》授课计划 杜媛.docx',
}

for name, fp in docs.items():
    doc = Document(fp)
    print(f'=== {name} ===')
    for pi, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if '教材' in txt or '天津大学' in txt or '大数据分析方法' in txt:
            print(f'  P{pi}: {txt[:300]}')
    # 也检查表格中的教材信息
    for ti, t in enumerate(doc.tables):
        for ri in range(len(t.rows)):
            for ci in range(len(t.columns)):
                txt = t.cell(ri, ci).text.strip()
                if '教材' in txt or '天津大学' in txt or '大数据分析方法' in txt:
                    print(f'  Table{ti} R{ri}C{ci}: {txt[:300]}')
    print()
