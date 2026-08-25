"""
扩展第一节课设计梗概内容，增加时间分配、教学步骤、师生活动
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

# 新内容
new_lines = [
    "九、第一节课设计梗概",
    "第一节课为课程导入课（2学时），主要内容包括：",
    "（1）课程介绍（15分钟）：介绍课程性质——农村电子商务专业核心课，总学时60学时（理论30+实践30），3学分。讲解课程教学目标（知识目标、能力目标、思政目标、素质目标），说明过程性考核（40%）与终结性考核（60%）的具体标准和权重分配，介绍超星学习通平台使用方法、课程资源构成（微课视频36个、课件PPT 30套、题库500余道、实训数据集10套）及学习方法建议。",
    "（2）行业认知（20分钟）：通过招聘网站数据展示电商行业数据分析岗位需求量、薪资水平（初级8-12K、中级12-20K、高级20K+）和职业发展路径（数据分析师→数据产品经理→数据总监），分析岗位核心技能要求，对比传统电商运营与数据驱动运营的差异，激发学生学习兴趣和职业认同感。",
    "（3）案例分析（25分钟）：展示农产品电商平台真实数据分析案例——某农产品电商平台的用户购买行为分析，包含数据采集（订单数据、用户行为数据）、数据清洗（缺失值处理、异常值检测）、数据分析（RFM模型用户分群、关联规则挖掘）、数据可视化（销售趋势图、用户画像图）和报告撰写全流程，让学生直观感受数据分析在商务决策中的实际价值。",
    "（4）工具介绍（20分钟）：现场演示Python数据分析环境搭建，包括Anaconda安装、Jupyter Notebook启动和基本操作、Numpy/Pandas库导入测试；同时演示Excel高级功能（数据透视表、Power Query、图表制作），让学生初步了解课程将使用的两大工具体系，明确Python与Excel在数据分析中的互补关系。",
    "（5）思政融入（10分钟）：讲解《中华人民共和国数据安全法》《个人信息保护法》对数据分析的合规要求，通过\u201c大数据杀熟\u201d反面案例引导学生树立数据伦理意识，强调数据分析师应坚持真实分析、诚信报告的职业操守，培养学生数据安全意识和法治观念。",
    "（6）分组与任务分配（10分钟）：按4-5人一组将学生分组，每组分配本学期综合项目的初始方向（农产品电商、生鲜配送、农村物流等不同场景），发放项目任务书，明确团队分工要求（数据采集、数据清洗、数据分析、可视化、报告撰写），建立小组沟通群和协作机制。",
    "（7）课堂互动（10分钟）：通过超星学习通发起随堂投票和抢答，检验学生对课程框架和考核方式的理解程度，收集学生对课程的期望和疑问，现场解答共性问题，为后续教学节奏调整提供依据。",
    "（8）课后任务布置（10分钟）：要求学生完成以下任务：①注册超星学习通账号并加入课程班级；②预习第一章微课视频（数据分析概述、常用分析方法）；③安装Anaconda环境并测试Jupyter Notebook运行；④阅读教材第一章节内容；⑤各小组讨论确定综合项目的数据集方向，下次课提交项目选题报告。",
]

# 找到P109和P111的位置
p109 = doc.paragraphs[109]
p110 = doc.paragraphs[110]

# 获取格式
font_name = '仿宋_GB2312'
font_size_val = '21'
if p110.runs:
    r = p110.runs[0]
    if r.font.name:
        font_name = r.font.name
    if r.font.size:
        font_size_val = str(int(r.font.size.pt * 2))

print(f"格式: font={font_name}, size={font_size_val}")

# 替换P110的内容并添加新段落
# P109保持标题不变
# P110替换为第一个内容段
body = doc.element.body
all_p_elems = body.findall(qn('w:p'))

# 找到P110的元素
p110_elem = p110._element

def make_p(text, font_name, font_size_val):
    """创建段落元素"""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    # 段落格式：首行缩进
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLineChars'), '200')
    pPr.append(ind)
    # 行距auto
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), 'auto')
    pPr.append(spacing)
    p.append(pPr)
    
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), font_size_val)
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), font_size_val)
    rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p

# 替换P110
# 清除P110原有内容
for r in p110_elem.findall(qn('w:r')):
    p110_elem.remove(r)
# 设置新内容
pPr = p110_elem.find(qn('w:pPr'))
if pPr is None:
    pPr = OxmlElement('w:pPr')
    p110_elem.insert(0, pPr)
# 首行缩进
ind = pPr.find(qn('w:ind'))
if ind is None:
    ind = OxmlElement('w:ind')
    pPr.append(ind)
ind.set(qn('w:firstLineChars'), '200')
# 行距
spacing = pPr.find(qn('w:spacing'))
if spacing is None:
    spacing = OxmlElement('w:spacing')
    pPr.append(spacing)
spacing.set(qn('w:line'), 'auto')

# 添加run到P110
r = OxmlElement('w:r')
rPr = OxmlElement('w:rPr')
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'), font_name)
rFonts.set(qn('w:eastAsia'), font_name)
rFonts.set(qn('w:hAnsi'), font_name)
rFonts.set(qn('w:cs'), font_name)
rPr.append(rFonts)
sz = OxmlElement('w:sz')
sz.set(qn('w:val'), font_size_val)
rPr.append(sz)
szCs = OxmlElement('w:szCs')
szCs.set(qn('w:val'), font_size_val)
rPr.append(szCs)
r.append(rPr)
t = OxmlElement('w:t')
t.set(qn('xml:space'), 'preserve')
t.text = new_lines[1]  # 第一节课为课程导入课...
r.append(t)
p110_elem.append(r)

# 在P110后面插入新段落（new_lines[2:]）
prev_elem = p110_elem
for line in new_lines[2:]:
    new_p = make_p(line, font_name, font_size_val)
    prev_elem.addnext(new_p)
    prev_elem = new_p

doc.save(fp)
print("保存完成")

# 验证
doc2 = Document(fp)
in_section = False
for pi, p in enumerate(doc2.paragraphs):
    txt = p.text.strip()
    if '九、第一节课设计梗概' in txt:
        in_section = True
    if in_section:
        print(f'P{pi} [{len(txt)}]: {txt[:150]}')
    if in_section and '十、' in txt:
        break
