"""
修复所有29个单元教学设计表格：
1. R3教学任务章节号修正（按9章分组，非任务序号）
2. R9/R10标题行恢复"知识＆技能"
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import store
import re

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

# 章节映射：seq → (章节号, 章节名)
chapter_map = {
    1: ('第一章', '初识数据分析'),
    2: ('第一章', '初识数据分析'),
    3: ('第一章', '初识数据分析'),
    4: ('第一章', '初识数据分析'),
    5: ('第二章', 'Excel数据分析工具'),
    6: ('第二章', 'Excel数据分析工具'),
    7: ('第二章', 'Excel数据分析工具'),
    8: ('第二章', 'Excel数据分析工具'),
    9: ('第三章', 'Numpy数学运算库'),
    10: ('第三章', 'Numpy数学运算库'),
    11: ('第三章', 'Numpy数学运算库'),
    12: ('第三章', 'Numpy数学运算库'),
    13: ('第四章', 'Pandas数据分析库'),
    14: ('第四章', 'Pandas数据分析库'),
    15: ('第四章', 'Pandas数据分析库'),
    16: ('第四章', 'Pandas数据分析库'),
    17: ('第五章', 'SciPy科学计算库'),
    18: ('第五章', 'SciPy科学计算库'),
    19: ('第五章', 'SciPy科学计算库'),
    20: ('第六章', 'Sklearn数据统计基础'),
    21: ('第六章', 'Sklearn数据统计基础'),
    22: ('第六章', 'Sklearn数据统计基础'),
    23: ('第七章', 'Sklearn数据统计进阶'),
    24: ('第七章', 'Sklearn数据统计进阶'),
    25: ('第七章', 'Sklearn数据统计进阶'),
    26: ('第八章', 'Seaborn可视化分析库'),
    27: ('第八章', 'Seaborn可视化分析库'),
    28: ('第八章', 'Seaborn可视化分析库'),
    29: ('第九章', '综合评价与课程总结'),
    30: ('第九章', '综合评价与课程总结'),
}

# 获取数据库任务信息
tasks = store.rows("SELECT seq, title FROM tasks WHERE offering_id=20 ORDER BY seq")
task_by_seq = {t['seq']: t for t in tasks}

def clear_and_set_cell(tc, text, bold=True):
    """清除单元格内容并设置新文本"""
    paras = tc.findall(qn('w:p'))
    if not paras:
        p = OxmlElement('w:p')
        tc.append(p)
        paras = [p]
    p = paras[0]
    for r in p.findall(qn('w:r')):
        p.remove(r)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), '仿宋')
    rFonts.set(qn('w:eastAsia'), '仿宋')
    rFonts.set(qn('w:hAnsi'), '仿宋')
    rFonts.set(qn('w:cs'), '仿宋')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '21')
    rPr.append(szCs)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)
    r.append(rPr)
    t_elem = OxmlElement('w:t')
    t_elem.set(qn('xml:space'), 'preserve')
    t_elem.text = text
    r.append(t_elem)
    p.append(r)
    for extra_p in paras[1:]:
        tc.remove(extra_p)

def set_multiline_cell(tc, lines, bold=False):
    """设置多行文本到单元格"""
    paras = tc.findall(qn('w:p'))
    if not paras:
        p = OxmlElement('w:p')
        tc.append(p)
        paras = [p]
    p = paras[0]
    for r in p.findall(qn('w:r')):
        p.remove(r)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), '仿宋')
    rFonts.set(qn('w:eastAsia'), '仿宋')
    rFonts.set(qn('w:hAnsi'), '仿宋')
    rFonts.set(qn('w:cs'), '仿宋')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '21')
    rPr.append(szCs)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    r.append(rPr)
    for li, line in enumerate(lines):
        if li > 0:
            br = OxmlElement('w:br')
            r.append(br)
        t_elem = OxmlElement('w:t')
        t_elem.set(qn('xml:space'), 'preserve')
        t_elem.text = line
        r.append(t_elem)
    p.append(r)
    for extra_p in paras[1:]:
        tc.remove(extra_p)

def get_tc_text(tc):
    """获取单元格文本（包括换行）"""
    txt = ''
    for p in tc.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            for elem in r:
                if elem.tag == qn('w:t'):
                    txt += elem.text or ''
                elif elem.tag == qn('w:br'):
                    txt += '\n'
    return txt

# 处理所有29个表格
table_indices = list(range(5, 63, 2))
print(f'处理 {len(table_indices)} 个表格')

for idx, ti in enumerate(table_indices):
    seq = idx + 1  # seq 1-30
    t = doc.tables[ti]
    
    chapter_num, chapter_name = chapter_map[seq]
    task = task_by_seq.get(seq, {})
    task_title = task.get('title', '')
    
    # === 1. 修复R3教学任务 ===
    # 解析task_title: "初识数据分析：认识数据分析、常用数据分析方法"
    # → "第一章 初识数据分析\n子情景：认识数据分析\n子情景：常用数据分析方法"
    if '：' in task_title:
        parts = task_title.split('：', 1)
        scene_name = parts[0]  # 初识数据分析
        sub_scenes = parts[1].split('、')  # [认识数据分析, 常用数据分析方法]
    else:
        scene_name = task_title
        sub_scenes = [task_title]
    
    r3_lines = [f"{chapter_num} {scene_name}"]
    for ss in sub_scenes:
        r3_lines.append(f"子情景：{ss.strip()}")
    
    # 设置R3的C2单元格（内容单元格）
    r3 = t.rows[3]
    r3_tcs = r3._tr.findall(qn('w:tc'))
    # C2是第三个tc（前两个是标签）
    # 找到内容tc（通常是有gridSpan的较大单元格）
    for tc in r3_tcs:
        txt = get_tc_text(tc)
        if txt and '第' in txt:
            set_multiline_cell(tc, r3_lines, bold=False)
            break
    
    # === 2. 修复R9标题行 ===
    r9 = t.rows[9]
    r9_tcs = r9._tr.findall(qn('w:tc'))
    
    for ci, tc in enumerate(r9_tcs):
        txt = get_tc_text(tc).strip()
        # 如果C1位置包含数字开头的内容（被错误替换的知识＆技能）
        if re.match(r'^[1-9]', txt) and len(txt) > 5:
            clear_and_set_cell(tc, '知识＆技能', bold=True)
    
    # === 3. 修复R10标题行 ===
    r10 = t.rows[10]
    r10_tcs = r10._tr.findall(qn('w:tc'))
    
    for ci, tc in enumerate(r10_tcs):
        txt = get_tc_text(tc).strip()
        if re.match(r'^[1-9]', txt) and len(txt) > 5:
            clear_and_set_cell(tc, '知识＆技能', bold=True)
    
    print(f'Table {ti} (seq={seq}): {chapter_num} {chapter_name} | R3={r3_lines[0]} + {len(r3_lines)-1}子情景')

doc.save(fp)
print(f'\n保存完成')

# 验证
doc2 = Document(fp)
for ti in [5, 9, 13, 21, 35, 61]:
    t = doc2.tables[ti]
    # R3
    r3_txt = ''
    for tc in t.rows[3]._tr.findall(qn('w:tc')):
        txt = get_tc_text(tc)
        if '第' in txt:
            r3_txt = txt.replace('\n', ' | ')
            break
    # R9
    r9_tcs = t.rows[9]._tr.findall(qn('w:tc'))
    r9_texts = [get_tc_text(tc)[:15] for tc in r9_tcs]
    print(f'Table {ti}: R3=[{r3_txt}] R9={r9_texts}')
