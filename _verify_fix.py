from docx import Document

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)
print(f'总表格数: {len(doc.tables)}')

# 验证Table 5
t = doc.tables[5]
print(f'\n=== Table 5 R3 ===')
print(t.rows[3].cells[2].text)

print(f'\n=== Table 5 R9-R14 ===')
for ri in range(9, 15):
    texts = [c.text.strip()[:30] for c in t.rows[ri].cells]
    print(f'  R{ri}: C2={texts[2][:30]} | C7={texts[7]} | C8={texts[8]} | C11={texts[11]}')

# 验证Table 7
t7 = doc.tables[7]
print(f'\n=== Table 7 R3 ===')
print(t7.rows[3].cells[2].text)

print(f'\n=== Table 7 R9-R14 ===')
for ri in range(9, 15):
    texts = [c.text.strip()[:30] for c in t7.rows[ri].cells]
    print(f'  R{ri}: C2={texts[2][:30]} | C7={texts[7]} | C8={texts[8]} | C11={texts[11]}')

# 检查最后一个表格
last_t = doc.tables[-1]
print(f'\n=== Last Table {len(doc.tables)-1} R3 ===')
print(last_t.rows[3].cells[2].text[:100])
