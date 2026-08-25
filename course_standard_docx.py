import shutil
from pathlib import Path

from docx import Document

import store
from course_standard_model import build_course_standard_model


def _set_cell(cell, text):
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = str(text)
        return
    paragraph = paragraphs[0]
    runs = paragraph.runs
    if runs:
        runs[0].text = str(text)
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(str(text))
    for extra in paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def _set_paragraph(paragraph, text):
    runs = paragraph.runs
    if runs:
        runs[0].text = str(text)
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(str(text))


def _replace_first(document, old, new):
    for paragraph in document.paragraphs:
        if old in paragraph.text:
            _set_paragraph(paragraph, paragraph.text.replace(old, new))
            return True
    return False


def _set_paragraph_by_prefix(document, prefix, text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            _set_paragraph(paragraph, text)
            return True
    return False


def _fill_metadata(document, model):
    offering = model["offering"]
    _set_paragraph_by_prefix(document, "课程名称及课程编号：", f"课程名称及课程编号：{offering['course_name']}（{offering['course_code']}）")
    _set_paragraph_by_prefix(document, "课程类型：", f"课程类型：{offering['course_type']}")
    credits = float(offering["credits"])
    credits_text = str(int(credits)) if credits.is_integer() else str(credits)
    _set_paragraph_by_prefix(document, "学时学分：", f"学时学分：{offering['total_hours']}学时（{credits_text}学分）")
    semester = "第四学期" if offering["term"] == "2023-2024-2" else offering["term"]
    _set_paragraph_by_prefix(document, "开设学期：", f"开设学期：{semester}")
    _set_paragraph_by_prefix(document, "适用专业：", f"适用专业：                 {offering['major']}")


def _fill_course_nature(document, model):
    heading_index = next(
        (index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "一、课程性质"),
        None,
    )
    if heading_index is None:
        raise ValueError("课程标准中未找到“一、课程性质”位置。")
    nature = model["course_nature"]
    for offset, text in enumerate(nature, 1):
        if heading_index + offset >= len(document.paragraphs):
            raise ValueError("课程性质正文位置不足。")
        _set_paragraph(document.paragraphs[heading_index + offset], text)


def _fill_course_goals(document, model):
    table = document.tables[0]
    goals = model["course_goals"]
    keys = ("knowledge", "ability", "ideological", "quality")
    for column, key in enumerate(keys):
        value = "\n".join(f"{index}. {text}" for index, text in enumerate(goals[key], 1))
        _set_cell(table.rows[1].cells[column], value)


def _fill_course_design(document, model):
    heading_index = next(
        (index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "1、本课程设计的总体思路"),
        None,
    )
    if heading_index is None:
        raise ValueError("课程标准中未找到课程设计总体思路位置。")
    for offset, text in enumerate(model["course_design"], 1):
        _set_paragraph(document.paragraphs[heading_index + offset], text)


def _fill_content_table(document, model):
    table = document.tables[1]
    rows = [row for row in table.rows[1:] if row.cells and row.cells[0].text.strip() != "合计"]
    expected = [sub for scenario in model["scenarios"] for sub in [(scenario, item) for item in scenario["subscenarios"]]]
    if len(rows) != len(expected):
        raise ValueError(f"课程标准课时表需要{len(expected)}行，模板当前只有{len(rows)}行；需要先执行模板行扩展。")
    for row, (scenario, sub) in zip(rows, expected):
        _set_cell(row.cells[0], scenario["seq"])
        scenario_label = (
            scenario["title"] if scenario["title"] == "综合评价与课程总结"
            else f"项目{scenario['number']} {scenario['title']}"
        )
        _set_cell(row.cells[1], scenario_label)
        _set_cell(row.cells[2], sub["seq"])
        _set_cell(row.cells[3], sub["title"])
        _set_cell(row.cells[4], sub["hours"])
        _set_cell(row.cells[5], f"{scenario['hours']}（理论{scenario['theory_hours']}+实践{scenario['practice_hours']}）")
    total = table.rows[-1]
    _set_cell(total.cells[4], model["offering"]["total_hours"])
    _set_cell(total.cells[5], f"理论{sum(s['theory_hours'] for s in model['scenarios'])}+实践{sum(s['practice_hours'] for s in model['scenarios'])}")


def generate_course_standard(offering_id, template_path, output_path):
    model = build_course_standard_model(offering_id)
    source = Path(template_path)
    destination = Path(output_path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)
    document = Document(destination)
    _fill_metadata(document, model)
    _fill_course_nature(document, model)
    _fill_course_goals(document, model)
    _fill_course_design(document, model)
    _fill_content_table(document, model)
    document.save(destination)
    return destination
