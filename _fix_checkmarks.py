"""
修复R11-R14的打勾问题：
- 只保留与知识＆技能内容行数对应的行有勾
- 清除空行的勾
- 重点、难点、识记、理解、应用、评价 都对应知识＆技能行打勾
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import store

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

def get_tc_text(tc):
    txt = ''
    for p in tc.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            for elem in r:
                if elem.tag == qn('w:t'):
                    txt += elem.text or ''
    return txt.strip()

def clear_cell(tc):
    """清除单元格内容"""
    for p in tc.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            p.remove(r)

def set_check(tc, check=True):
    """设置或清除勾"""
    clear_cell(tc)
    if check:
        p = tc.findall(qn('w:p'))
        if not p:
            p = [OxmlElement('w:p')]
            tc.append(p[0])
        p = p[0]
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), '仿宋')
        rFonts.set(qn('w:eastAsia'), '仿宋')
        rFonts.set(qn('w:hAnsi'), '仿宋')
        rFonts.set(qn('w:cs'), '仿宋')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '21')
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '21')
        rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = '√'
        r.append(t)
        p.append(r)

# 处理所有29个表格
table_indices = list(range(5, 63, 2))
tasks = store.rows("SELECT seq, title FROM tasks WHERE offering_id=20 ORDER BY seq")
task_by_seq = {t['seq']: t for t in tasks}

for idx, ti in enumerate(table_indices):
    seq = idx + 1
    t = doc.tables[ti]
    task = task_by_seq.get(seq, {})
    title = task.get('title', '')
    
    # 解析子情景数量
    if '：' in title:
        parts = title.split('：', 1)
        sub_scenes = parts[1].split('、')
    else:
        sub_scenes = [title]
    
    num_items = len(sub_scenes)  # 知识＆技能行数
    
    # R11-R14: 对应知识＆技能的4行
    # 模板R11-R14的tc结构：
    # TC0: 标签(vMerge=continue)
    # TC1: 知识＆技能内容(gridSpan=5)
    # TC2: 重点(gridSpan=2)
    # TC3: 难点(gridSpan=3)
    # TC4: 识记(gridSpan=2)  -- 实际gridSpan可能不同
    # TC5: 理解(gridSpan=3)
    # TC6: 应用(gridSpan=2)
    # TC7: 评价
    
    for ri_offset in range(4):
        ri = 11 + ri_offset
        row = t.rows[ri]
        tcs = row._tr.findall(qn('w:tc'))
        
        if ri_offset < num_items:
            # 这行有内容，保留勾（已经在之前的修复中设置了内容）
            # 确保知识＆技能列有内容
            has_content = False
            if len(tcs) > 1:
                tc1_text = get_tc_text(tcs[1])
                if tc1_text:
                    has_content = True
            
            if not has_content:
                # 需要设置子情景名称
                if ri_offset < len(sub_scenes):
                    scene_name = sub_scenes[ri_offset].strip()
                    # 分割"任务实施：xxx"的情况
                    if '：' in scene_name:
                        parts2 = scene_name.split('：')
                        scene_name = parts2[-1]
                    # 设置到TC1
                    if len(tcs) > 1:
                        clear_cell(tcs[1])
                        p = tcs[1].findall(qn('w:p'))
                        if not p:
                            p = [OxmlElement('w:p')]
                            tcs[1].append(p[0])
                        p = p[0]
                        r = OxmlElement('w:r')
                        rPr = OxmlElement('w:rPr')
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), '仿宋')
                        rFonts.set(qn('w:eastAsia'), '仿宋')
                        rFonts.set(qn('w:hAnsi'), '仿宋')
                        rFonts.set(qn('w:cs'), '仿宋')
                        rPr.append(rFonts)
                        sz = OxmlElement('w:sz')
                        sz.set(qn('w:val'), '21')
                        rPr.append(sz)
                        szCs = OxmlElement('w:szCs')
                        szCs.set(qn('w:val'), '21')
                        rPr.append(szCs)
                        r.append(rPr)
                        t_elem = OxmlElement('w:t')
                        t_elem.set(qn('xml:space'), 'preserve')
                        t_elem.text = scene_name
                        r.append(t_elem)
                        p.append(r)
        else:
            # 这行不应该有内容和勾，清除
            for tc in tcs:
                tc1_text = get_tc_text(tc)
                if tc1_text:
                    clear_cell(tc)
    
    print(f'Table {ti} (seq={seq}): {num_items}个子情景，清除R{11+num_items}-R14的空行内容')

doc.save(fp)
print(f'\n保存完成')

# 验证Table 5
doc2 = Document(fp)
t = doc2.tables[5]
print(f'\n=== 验证 Table 5 ===')
for ri in range(9, 15):
    row = t.rows[ri]
    tcs = row._tr.findall(qn('w:tc'))
    texts = []
    for tc in tcs:
        txt = get_tc_text(tc)
        texts.append(txt[:15] if txt else '(空)')
    print(f'R{ri}: {texts}')
