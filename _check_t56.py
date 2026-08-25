"""
分析表5和表6的完整结构
"""
from docx import Document
from docx.oxml.ns import qn

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)

# 找到第一对表5/表6（索引5和6）
t5 = doc.tables[5]
t6 = doc.tables[6]

print(f"=== 表5（教学设计·基本信息）===")
print(f"{len(t5.rows)}行 x {len(t5.rows[0].cells)}列")
for ri in range(len(t5.rows)):
    cells = t5.rows[ri].cells
    texts = []
    for ci in range(min(10, len(cells))):
        t = cells[ci].text.strip()[:30]
        texts.append(f"[{ci}]{t}")
    print(f"  R{ri}: {' | '.join(texts)}")

print(f"\n=== 表6（教学设计·教学组织）===")
print(f"{len(t6.rows)}行 x {len(t6.rows[0].cells)}列")
for ri in range(len(t6.rows)):
    cells = t6.rows[ri].cells
    texts = []
    for ci in range(len(cells)):
        t = cells[ci].text.strip()[:40]
        texts.append(f"[{ci}]{t}")
    print(f"  R{ri}: {' | '.join(texts)}")

# 检查表5的所有列（20列）
print(f"\n=== 表5完整列数 ===")
print(f"R0 cells: {len(t5.rows[0].cells)}")
for ci in range(len(t5.rows[0].cells)):
    t = t5.rows[0].cells[ci].text.strip()
    print(f"  col{ci}: {t}")
