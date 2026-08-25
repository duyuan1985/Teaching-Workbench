"""
对比课程标准和授课计划的内容一致性
检查：章节任务、学时分配、考核方式、课程信息
"""
from docx import Document

std_fp = r"生成结果\精修版\2023-2024-2《商务数据分析》课程标准 杜媛.docx"
plan_fp = r"生成结果\精修版\2023-2024-2《商务数据分析》授课计划 杜媛.docx"

std = Document(std_fp)
plan = Document(plan_fp)

print("=" * 70)
print("课程标准 vs 授课计划 一致性检查")
print("=" * 70)

# ============================================================
# 1. 课程信息对比
# ============================================================
print("\n一、课程信息")
print("-" * 50)

# 课程标准：从段落中提取
std_info = {}
for p in std.paragraphs:
    t = p.text.strip()
    if "课程名称" in t and "课程编号" in t:
        std_info["课程名"] = t
    elif "课程类型" in t:
        std_info["类型"] = t
    elif "学时学分" in t:
        std_info["学时学分"] = t
    elif "开设学期" in t:
        std_info["学期"] = t

# 授课计划：从段落中提取
plan_info = {}
for p in plan.paragraphs:
    t = p.text.strip()
    if "总学时" in t:
        plan_info["学时"] = t
    elif "课程教学日程表" in t:
        plan_info["日程表标题"] = t

print(f"课程标准:")
for k, v in std_info.items():
    print(f"  {k}: {v[:60]}")
print(f"授课计划:")
for k, v in plan_info.items():
    print(f"  {k}: {v[:60]}")

# ============================================================
# 2. 章节任务对比（表1 vs 表0）
# ============================================================
print("\n二、章节任务对比")
print("-" * 50)

# 课程标准表1：课程内容划分及课时分配
# 找到表1（含"学习情境"标题的表）
std_table1 = None
for t in std.tables:
    if len(t.rows) > 1:
        header = t.rows[0].cells[0].text.strip()
        if "序号" in header or "学习情境" in header:
            std_table1 = t
            break

std_chapters = []
if std_table1:
    for ri in range(1, len(std_table1.rows)):
        cells = std_table1.rows[ri].cells
        seq = cells[0].text.strip()
        name = cells[1].text.strip() if len(cells) > 1 else ""
        hours = cells[-1].text.strip() if cells else ""
        if seq and seq != "合计" and name:
            std_chapters.append({"seq": seq, "name": name, "hours": hours})

# 授课计划表0：日程表
plan_table0 = plan.tables[0] if plan.tables else None
plan_tasks = []
if plan_table0:
    current_seq = ""
    current_chapter = ""
    for ri in range(2, len(plan_table0.rows)):
        cells = plan_table0.rows[ri].cells
        seq = cells[0].text.strip()
        theory = cells[3].text.strip()
        if seq and seq != current_seq:
            current_seq = seq
            if theory and ("第" in theory and "章" in theory):
                current_chapter = theory
                plan_tasks.append({"seq": seq, "chapter": theory, "tasks": []})
            elif theory:
                plan_tasks.append({"seq": seq, "chapter": theory, "tasks": [theory]})
            else:
                # 可能是实践行
                pass
        elif theory and plan_tasks:
            plan_tasks[-1]["tasks"].append(theory)

print(f"课程标准章节数: {len(std_chapters)}")
for ch in std_chapters:
    print(f"  {ch['seq']}. {ch['name'][:40]} 学时:{ch['hours'][:20]}")

print(f"\n授课计划任务数: {len(plan_tasks)}")
for t in plan_tasks:
    print(f"  {t['seq']}. {t['chapter'][:40]} 子任务:{len(t['tasks'])}个")

# ============================================================
# 3. 考核方式对比
# ============================================================
print("\n三、考核方式对比")
print("-" * 50)

# 课程标准表2：考核评价
std_table2 = None
for t in std.tables:
    for ri in range(min(3, len(t.rows))):
        cell_text = t.rows[ri].cells[0].text.strip()
        if "过程性考核" in cell_text or "终结性考核" in cell_text or "考核" in cell_text:
            std_table2 = t
            break
    if std_table2:
        break

if std_table2:
    print("课程标准考核表:")
    for ri in range(min(10, len(std_table2.rows))):
        cells = std_table2.rows[ri].cells
        texts = [c.text.strip()[:15] for c in cells]
        if any(texts):
            print(f"  R{ri}: {texts}")

# 授课计划补充说明中的考核
plan_t1 = plan.tables[1] if len(plan.tables) > 1 else None
if plan_t1:
    cell_text = plan_t1.rows[0].cells[0].text
    # 找考核段落
    for line in cell_text.split("\n"):
        if "考核" in line or "40%" in line or "60%" in line:
            print(f"  授课计划: {line.strip()[:80]}")

# ============================================================
# 4. 学时总数对比
# ============================================================
print("\n四、学时对比")
print("-" * 50)

# 课程标准表1合计行
if std_table1:
    for ri in range(len(std_table1.rows)):
        if "合计" in std_table1.rows[ri].cells[0].text:
            total_cell = std_table1.rows[ri].cells[-1].text.strip()
            print(f"课程标准总学时: {total_cell[:40]}")
            break

# 授课计划表0计算总行数
if plan_table0:
    data_rows = len(plan_table0.rows) - 2  # 减去表头
    print(f"授课计划数据行数: {data_rows}")
    # 统计理论学时和实践学时
    theory_hours = 0
    practice_hours = 0
    for ri in range(2, len(plan_table0.rows)):
        cells = plan_table0.rows[ri].cells
        if cells[4].text.strip() == "1":
            theory_hours += 1
        if len(cells) > 7 and cells[7].text.strip() == "1":
            practice_hours += 1
    print(f"授课计划理论学时: {theory_hours}, 实践学时: {practice_hours}, 总计: {theory_hours + practice_hours}")

# ============================================================
# 5. 任务列表对比
# ============================================================
print("\n五、任务列表一致性")
print("-" * 50)

# 从课程标准提取任务
std_ch_names = set()
if std_table1:
    for ch in std_chapters:
        std_ch_names.add(ch["name"][:6])

# 从授课计划提取章节
plan_ch_names = set()
for t in plan_tasks:
    # 去掉"第X章 "前缀
    name = t["chapter"]
    if "章 " in name:
        name = name.split("章 ", 1)[1]
    plan_ch_names.add(name[:6])

print(f"课程标准章节: {std_ch_names}")
print(f"授课计划章节: {plan_ch_names}")
common = std_ch_names & plan_ch_names
only_std = std_ch_names - plan_ch_names
only_plan = plan_ch_names - std_ch_names
print(f"共同章节: {common}")
print(f"仅课程标准有: {only_std}")
print(f"仅授课计划有: {only_plan}")

if not only_std and not only_plan:
    print("\n✓ 章节完全一致")
else:
    print("\n✗ 章节不一致")
