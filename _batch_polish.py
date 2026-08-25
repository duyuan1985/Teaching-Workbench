"""
步骤7：批量精修
1. 重新生成被覆盖的4门课程（ID=22,23,29,30）—带专业后缀
2. 清除2024-2025-2课标模板残留
3. 统一格式：行高自适应、清除空段落、行距auto
"""
import os, sys, re, time
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import store
import generate

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "生成结果", "精修版")

# ============================================================
# 第1部分：重新生成被覆盖的4门课程
# ============================================================
def regen_overwritten_courses():
    """重新生成ID=22,23,29,30（同名课程带专业后缀）"""
    print("=" * 60)
    print("第1部分：重新生成被覆盖的课程（带专业后缀）")
    print("=" * 60)

    overwrite_ids = [22, 23, 29, 30]
    for oid in overwrite_ids:
        o = store.rows('SELECT * FROM offerings WHERE id=?', [oid])
        if not o:
            continue
        o = o[0]
        print(f"\nID={oid} {o['course_name']} ({o['term']}) 专业={o.get('major','')}")
        try:
            results = generate.generate_all(oid)
            for dt, r in results.items():
                if 'path' in r:
                    print(f"  {dt}: {os.path.basename(r['path'])}")
                elif 'error' in r:
                    print(f"  {dt}: ERROR - {r['error']}")
        except Exception as e:
            print(f"  失败: {e}")
    print()


# ============================================================
# 第2部分：清除2024-2025-2课标模板残留
# ============================================================
def clean_template_residuals():
    """清除课标中的模板说明文字"""
    print("=" * 60)
    print("第2部分：清除模板残留文字")
    print("=" * 60)

    # 模板残留模式（新模板特有）
    residual_patterns = [
        "××××课程是××××专业",
        "××××（如专业核心）课程",
        "体现本课程在课程体系中的特色",
        "说明课程在专业人才培养中的作用",
        "服务××岗位",
        "在课程整体设计思路基础上，依据……",
        "确定XX个模块/项目及其学时",
        "其中包括理论课时XX个",
        "以学生为中心，注重学生的学习兴趣",
        "引导学生运用XXX",
    ]

    cleaned_count = 0
    for fname in os.listdir(OUTPUT_DIR):
        if not fname.endswith('.docx') or fname.startswith('~$'):
            continue

        fpath = os.path.join(OUTPUT_DIR, fname)
        try:
            doc = Document(fpath)
            changed = False

            for p in doc.paragraphs:
                text = p.text.strip()
                for pat in residual_patterns:
                    if pat in text:
                        # 找到包含模板说明的段落，清空内容
                        for r in p.runs:
                            r.text = ""
                        changed = True
                        break

            # 也清理表格中的残留
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            text = p.text.strip()
                            for pat in residual_patterns:
                                if pat in text:
                                    for r in p.runs:
                                        r.text = ""
                                    changed = True
                                    break

            if changed:
                doc.save(fpath)
                cleaned_count += 1
                print(f"  清理: {fname}")

        except Exception as e:
            print(f"  错误: {fname} - {e}")

    print(f"\n共清理 {cleaned_count} 个文档\n")


# ============================================================
# 第3部分：统一格式（行高、空段落、行距）
# ============================================================
def fix_all_heights(doc):
    """将所有表格行高设为自动"""
    for t in doc.tables:
        for row in t.rows:
            trPr = row._tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                row._tr.insert(0, trPr)
            # 移除固定行高，设为auto
            trHeight = trPr.find(qn('w:trHeight'))
            if trHeight is not None:
                trHeight.set(qn('w:hRule'), 'auto')
                # 保留最小高度但设为auto


def fix_empty_paragraphs(doc):
    """清除表格单元格中的空段落（保留至少1个）"""
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                paras = cell.paragraphs
                if len(paras) <= 1:
                    continue
                # 从后往前删除空段落（保留至少1个）
                for i in range(len(paras) - 1, 0, -1):
                    p = paras[i]
                    if not p.text.strip() and not p.runs:
                        p._element.getparent().remove(p._element)


def fix_line_spacing(doc):
    """设置段落行距为auto（防止文字截断）"""
    for p in doc.paragraphs:
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                line = spacing.get(qn('w:line'))
                lineRule = spacing.get(qn('w:lineRule'))
                # 如果行距是固定值且小于字号，改为auto
                if lineRule == 'exact' or (line and lineRule == 'atLeast'):
                    spacing.set(qn('w:lineRule'), 'auto')
                    spacing.set(qn('w:line'), '360')  # 1.5倍行距

    # 表格内段落也处理
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pPr = p._element.find(qn('w:pPr'))
                    if pPr is not None:
                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is not None:
                            lineRule = spacing.get(qn('w:lineRule'))
                            if lineRule == 'exact':
                                spacing.set(qn('w:lineRule'), 'auto')
                                spacing.set(qn('w:line'), '360')


def fix_font_consistency(doc):
    """统一sz和szCs"""
    for p in doc.paragraphs:
        for r in p.runs:
            rpr = r._element.find(qn('w:rPr'))
            if rpr is not None:
                sz = rpr.find(qn('w:sz'))
                szCs = rpr.find(qn('w:szCs'))
                if sz is not None and szCs is not None:
                    sz_val = sz.get(qn('w:val'))
                    szCs_val = szCs.get(qn('w:val'))
                    if sz_val != szCs_val:
                        szCs.set(qn('w:val'), sz_val)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        rpr = r._element.find(qn('w:rPr'))
                        if rpr is not None:
                            sz = rpr.find(qn('w:sz'))
                            szCs = rpr.find(qn('w:szCs'))
                            if sz is not None and szCs is not None:
                                sz_val = sz.get(qn('w:val'))
                                szCs_val = szCs.get(qn('w:val'))
                                if sz_val != szCs_val:
                                    szCs.set(qn('w:val'), sz_val)


def fix_empty_para_indent(doc):
    """清除空段落的大缩进"""
    for p in doc.paragraphs:
        if not p.text.strip():
            pPr = p._element.find(qn('w:pPr'))
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    flc = ind.get(qn('w:firstLineChars'))
                    if flc and int(flc) > 10:
                        pPr.remove(ind)


def apply_all_format_fixes():
    """对所有文档应用格式修复"""
    print("=" * 60)
    print("第3部分：统一格式修复（行高/空段落/行距/字体）")
    print("=" * 60)

    files = [f for f in os.listdir(OUTPUT_DIR)
             if f.endswith('.docx') and not f.startswith('~$')]
    total = len(files)
    print(f"共 {total} 个文档待处理\n")

    for i, fname in enumerate(sorted(files), 1):
        fpath = os.path.join(OUTPUT_DIR, fname)
        try:
            doc = Document(fpath)

            fix_all_heights(doc)
            fix_empty_paragraphs(doc)
            fix_line_spacing(doc)
            fix_font_consistency(doc)
            fix_empty_para_indent(doc)

            doc.save(fpath)
            print(f"  [{i}/{total}] {fname[:40]}... OK")
        except Exception as e:
            print(f"  [{i}/{total}] {fname[:40]}... ERROR: {e}")

    print(f"\n格式修复完成！共处理 {total} 个文档\n")


# ============================================================
# 主流程
# ============================================================
if __name__ == '__main__':
    t0 = time.time()

    # 第1部分：重新生成被覆盖的课程
    regen_overwritten_courses()

    # 第2部分：清除模板残留
    clean_template_residuals()

    # 第3部分：统一格式修复
    apply_all_format_fixes()

    elapsed = time.time() - t0
    print("=" * 60)
    print(f"全部精修完成！总耗时: {elapsed:.0f}秒 ({elapsed/60:.1f}分钟)")
    print("=" * 60)

    # 统计最终文件
    files = [f for f in os.listdir(OUTPUT_DIR) if f.endswith('.docx') and not f.startswith('~$')]
    print(f"最终文档数: {len(files)} 份")
