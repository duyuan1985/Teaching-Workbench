"""
课程标准生成器
"""
import os, shutil
from copy import deepcopy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from generate import (
    safe_set_text, clear_indent, set_indent_chars, fix_spacing_auto,
    set_alignment, set_page_break_before, set_space_before, set_line_spacing,
    write_cell, write_cell_lines, _set_font_name, _set_font_size,
    merge_v, merge_h, get_vmerge_val, get_cell_text_raw, CN_NUMS, get_opening_semester
)


def generate_standard(offering, tasks, sessions, units, output_dir=None):
    """生成课程标准"""
    from generate import get_template_path, get_output_path

    template_path = get_template_path("standard", offering)
    if not template_path or not os.path.exists(template_path):
        raise FileNotFoundError(f"课程标准模板不存在: {template_path}")

    fp = get_output_path("standard", offering, output_dir)
    shutil.copy2(template_path, fp)
    doc = Document(fp)
    P = doc.paragraphs

    cn = offering["course_name"]
    cc = offering.get("course_code", "")
    tn = offering.get("teacher_name") or "杜媛"
    th = int(offering.get("total_hours", 60))
    cr = int(offering.get("credits", 3))
    mj = offering.get("major", "")
    dp = offering.get("department", "经济贸易系")
    ct = offering.get("course_type", "专业核心课")
    term = offering["term"]

    # ============================================================
    # 1. 封面
    # ============================================================
    _fill_cover(doc, offering, P)

    # ============================================================
    # 2. 第二页
    # ============================================================
    _fill_page2(doc, offering, P)

    # ============================================================
    # 3. 正文段落（课程性质、设计思路等）
    # ============================================================
    _fill_body_paragraphs(doc, offering, P)

    # ============================================================
    # 检测模板版本
    # ============================================================
    is_new_template = len(doc.tables) >= 6

    if is_new_template:
        # 新模板（2024-2025-2、2025-2026）：
        # 表0(9列): 课程信息
        # 表1(4列): 职业标准/工作任务/能力元素/相关知识
        # 表2(3列): 素质目标/知识目标/能力目标
        # 表3(4列): 序号/模块名称/任务名称/学时分配 → 内容表
        # 表4(6列): 模块/任务/教学内容/教学要求/思政要点/生成性成果
        # 表5(4列): 评价类型/评价内容/评价主体/评价方式 → 考核表

        # 表0：课程信息
        if len(doc.tables) > 0:
            _fill_course_info_table_new(doc.tables[0], offering)

        # 表2：课程目标
        if len(doc.tables) > 2:
            _fill_objectives_table_new(doc.tables[2], offering)

        # 表3：课程内容划分及课时分配
        if len(doc.tables) > 3:
            _fill_content_table_new(doc.tables[3], offering, tasks, units)

        # 表5：考核评价
        if len(doc.tables) > 5:
            _fill_assessment_table_new(doc.tables[5], offering, tasks, units)
    else:
        # 旧模板（2023-2024、2024-2025-1）：
        # 表0(4列): 知识目标/能力目标/思政目标/素质目标
        # 表1(6列): 序号/学习情境/子情境/子情境/课时/课时 → 内容表
        # 表2(5列): 考核方案
        # 表3(6列): 学习情境描述

        # 表0：课程目标
        if len(doc.tables) > 0:
            _fill_objectives_table(doc.tables[0], offering)

        # 表1：课程内容划分及课时分配
        if len(doc.tables) > 1:
            _fill_content_table(doc.tables[1], offering, tasks, units)

        # 表2：考核评价
        if len(doc.tables) > 2:
            _fill_assessment_table(doc.tables[2], offering, tasks, units)

        # 表3：学习情境描述
        if len(doc.tables) > 3:
            _fill_scenario_table(doc.tables[3], offering, tasks, units)

    doc.save(fp)
    return fp


def _fill_cover(doc, offering, P):
    """封面"""
    cn = offering["course_name"]
    dp = offering.get("department", "经济贸易系")
    mj = offering.get("major", "")

    for p in P:
        text = p.text.strip()
        # 适用专业
        if "适用专业" in text:
            _set_cover_line(p, "适用专业：", mj)
            _align_cover_line(p)
        # 编制单位
        elif "编制单位" in text:
            dept = "农村电子商务教研室"
            _set_cover_line(p, "编制单位：", dept)
            _align_cover_line(p)
        # 合作单位
        elif "合作单位" in text:
            co = "天津滨海迅腾科技集团有限公司"
            _set_cover_line(p, "合作单位：", co)
            _align_cover_line(p)
        # 课程名称（封面标题）
        elif (
            "《" in text
            and "》" in text
            and "课程标准" not in text
            and "课程考核评价" not in text
        ):
            clear_indent(p)
            fix_spacing_auto(p)
            set_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
            safe_set_text(p, f"《{cn}》", font="黑体", size=22, bold=True)
        # 日期
        elif "年" in text and "月" in text and len(text) < 20:
            clear_indent(p)
            safe_set_text(p, "2024年2月", font="黑体", size=18, bold=False)

    # 清除封面空段落的大缩进
    for p in P[:20]:
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                flc = ind.get(qn('w:firstLineChars'))
                if flc and int(flc) > 10:
                    pPr.remove(ind)


def _set_cover_line(p, label, value):
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    label_run = p.add_run(label)
    _set_font_name(label_run, "宋体")
    _set_font_size(label_run, 16)
    value_run = p.add_run(_pad_underline(value))
    _set_font_name(value_run, "宋体")
    _set_font_size(value_run, 16)
    rpr = value_run._element.get_or_add_rPr()
    u = rpr.find(qn("w:u"))
    if u is None:
        u = OxmlElement("w:u")
        rpr.append(u)
    u.set(qn("w:val"), "single")


def _align_cover_line(p):
    clear_indent(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.right_indent = Cm(0)

def _pad_underline(text, target_len=14):
    """补全角空格使下划线等长"""
    if len(text) < target_len:
        pad = target_len - len(text)
        # 前后各补一半
        front = pad // 2
        back = pad - front
        return "\u3000" * front + text + "\u3000" * back
    return text


def _add_underline_to_content(p, label_len, content_len):
    """给内容部分加下划线（标签不加）"""
    # 目前简单实现：整个run加下划线
    # 后续可精确控制
    if p.runs:
        for r in p.runs:
            rpr = r._element.get_or_add_rPr()
            u = rpr.find(qn('w:u'))
            if u is None:
                u = OxmlElement('w:u')
                rpr.append(u)
            u.set(qn('w:val'), 'single')


def _fill_page2(doc, offering, P):
    """第二页"""
    cn = offering["course_name"]
    cc = offering.get("course_code", "")
    th = int(offering.get("total_hours", 60))
    cr = int(offering.get("credits", 3))
    ct = offering.get("course_type", "专业核心课")

    for p in P:
        text = p.text.strip()
        # 课程标准标题
        if "课程标准" in text and "《" in text:
            set_page_break_before(p)
            # 段前空一行
            set_space_before(p, 24)
            clear_indent(p)
            set_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
            safe_set_text(p, f"《{cn}》课程标准", font="黑体", size=18, bold=False)

        # 课程信息4行
        elif "课程名称" in text and "课程编号" in text:
            safe_set_text(p, f"课程名称及课程编号：{cn}（{cc}）", font="黑体", size=14, bold=True)
            clear_indent(p)
        elif text.startswith("课程类型") and len(text) < 20:
            safe_set_text(p, f"课程类型：{ct}", font="黑体", size=14, bold=True)
            clear_indent(p)
        elif "学时学分" in text:
            safe_set_text(p, f"学时学分：{th}学时（{cr}学分）", font="黑体", size=14, bold=True)
            clear_indent(p)
        elif "开设学期" in text:
            safe_set_text(p, f"开设学期：{get_opening_semester(offering) or offering.get('term', '')}", font="黑体", size=14, bold=True)
            clear_indent(p)


def _clear_paragraph(paragraph):
    """清空段落文字但保留模板段落格式。"""
    for run in list(paragraph.runs):
        run.text = ""


def _write_body_paragraph(paragraph, text):
    """按课程标准正文规范写入仿宋小四内容。"""
    safe_set_text(paragraph, text, font="仿宋_GB2312", size=12, bold=False)
    set_indent_chars(paragraph, 2)


def _write_body_heading(paragraph, text):
    """按模板写入正文标题，避免模板说明和标题混在一起。"""
    safe_set_text(paragraph, text, font="黑体", size=12, bold=False)
    clear_indent(paragraph)


def _clone_paragraph_after(anchor, source=None):
    """克隆一个模板段落并插入到 anchor 后。"""
    source = source or anchor
    paragraph_xml = deepcopy(source._p)
    anchor._p.addnext(paragraph_xml)
    return Paragraph(paragraph_xml, anchor._parent)


def _find_paragraph(doc, predicate):
    for paragraph in doc.paragraphs:
        if predicate(paragraph.text.strip()):
            return paragraph
    return None


def _fill_paragraph_block(doc, anchor, contents, stop_prefixes, source=None):
    """在标题后的原有槽位中填充正文，清除剩余说明，不覆盖下一个标题。"""
    paragraphs = list(doc.paragraphs)
    try:
        anchor_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is anchor._p)
    except StopIteration:
        return
    stop_index = len(paragraphs)
    for index in range(anchor_index + 1, len(paragraphs)):
        text = paragraphs[index].text.strip()
        if any(text.startswith(prefix) for prefix in stop_prefixes):
            stop_index = index
            break
    slots = paragraphs[anchor_index + 1:stop_index]
    cursor = anchor
    body_source = source or (slots[0] if slots else anchor)
    for index, content in enumerate(contents):
        if index < len(slots):
            paragraph = slots[index]
        else:
            paragraph = _clone_paragraph_after(cursor, body_source)
        _write_body_paragraph(paragraph, content)
        cursor = paragraph
    for paragraph in slots[len(contents):]:
        _clear_paragraph(paragraph)


def _fill_body_paragraphs(doc, offering, P):
    """填充课程标准正文并清理所有模板说明。"""
    cn = offering["course_name"]
    major = offering.get("major", "本专业")
    course_type = offering.get("course_type", "专业核心课")
    textbook = offering.get("textbook_version", "")

    heading_rules = (
        ("一、课程性质", "一、课程性质"),
        ("二、课程目标", "二、课程目标"),
        ("三、课程设计", "三、课程设计"),
        ("1、本课程设计的总体思路", "1、本课程设计的总体思路"),
        ("2、课程内容划分及课时分配", "2、课程内容划分及课时分配"),
        ("3、教学方法描述", "3、教学方法描述"),
        ("4.考试与评价方式", "4.考试与评价方式（或方案）"),
        ("四、学习情境描述", "四、学习情境描述"),
        ("五、课程实施", "五、课程实施"),
        ("（一）教师知识素质要求", "（一）教师知识素质要求"),
        ("（二）教材编写与选用", "（二）教材编写与选用"),
        ("（三）课程资源的开发与利用", "（三）课程资源的开发与利用"),
        ("（四）推荐参考书、期刊、网站", "（四）推荐参考书、期刊、网站"),
    )
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        for prefix, heading in heading_rules:
            if text.startswith(prefix):
                _write_body_heading(paragraph, heading)
                break
        if "课程考核评价" in text and "《" in text:
            _write_body_heading(paragraph, f"《{cn}》课程考核评价")

    body_template = _find_paragraph(
        doc,
        lambda text: text.startswith("　　") and not text.startswith("　　课程："),
    )

    nature = _find_paragraph(doc, lambda text: text.startswith("一、课程性质"))
    if nature:
        _fill_paragraph_block(
            doc, nature,
            [f"　　《{cn}》是{major}人才培养方案中的{course_type}，承担培养学生运用数字化工具解决业务问题的知识、技能和职业素养。课程面向数据采集、处理、分析、可视化和成果表达等岗位任务，落实规范操作、数据安全、诚信分析、团队协作和持续学习等培养规格。"],
            ["二、课程目标"], body_template,
        )

    design = _find_paragraph(doc, lambda text: text.startswith("1、本课程设计的总体思路"))
    if design:
        _fill_paragraph_block(
            doc, design,
            [
                f"　　课程设计依据{major}专业人才培养方案和岗位能力要求，以真实或仿真的商务数据分析项目为载体，将知识学习、技能训练、成果检查和职业规范贯穿教学全过程，按照“任务分析—方案设计—实施操作—检查评价—迭代改进”的路径组织{int(offering.get('total_hours', 60))}学时教学。",
                "　　课程内容主动对接行业新标准、新技术、新工艺和新方法，融入数据安全、个人信息保护、知识产权和诚信分析要求；依托学习通、雨课堂等平台开展线上线下混合式教学，依据学习过程数据持续改进教学。",
            ],
            ["2、课程内容划分及课时分配"], body_template,
        )

    content_heading = _find_paragraph(doc, lambda text: text.startswith("2、课程内容划分及课时分配"))
    if content_heading:
        _fill_paragraph_block(doc, content_heading, [], ["3、教学方法描述"])

    methods = _find_paragraph(doc, lambda text: text.startswith("3、教学方法描述"))
    if methods:
        _fill_paragraph_block(
            doc, methods,
            [
                "　　根据本课程的教学目标、课程内容和学生特点，综合考虑教学效果与教学可操作性，选用项目教学法、任务驱动法、案例教学法、演示教学法和小组合作法。",
                "　　项目教学法以企业真实项目为载体，将知识学习与技能训练融入项目实施全过程；任务驱动法以具体任务为线索推进教学，明确任务要求、操作步骤和成果标准；案例教学法通过分析电商行业真实数据案例，引导学生理解方法的应用场景。",
                "　　教学实施依托数字化平台开展课前预习、课中互动测验、分组实操、成果展示和课后反馈，依据学习过程数据进行分层指导，促进学生主动思考、乐于实践并持续改进。",
            ],
            ["4.考试与评价方式"], body_template,
        )

    assessment_heading = _find_paragraph(doc, lambda text: text.startswith("4.考试与评价方式"))
    if assessment_heading:
        _fill_paragraph_block(
            doc, assessment_heading,
            ["　　本课程采用过程性考核与终结性考核相结合的多元评价方式。过程性考核占40%，包括签到10%、课堂表现10%、作业20%；终结性考核占60%，采用综合作品或课程项目成果评价，重点考查知识掌握、技能操作、数据安全、职业素养和创新实践能力。"],
            ["《", "四、学习情境描述"], body_template,
        )

    evaluation_heading = _find_paragraph(doc, lambda text: "课程考核评价" in text and "《" in text)
    if evaluation_heading:
        _write_body_heading(evaluation_heading, f"《{cn}》课程考核评价")
        _fill_paragraph_block(doc, evaluation_heading, [], ["四、学习情境描述"])

    teacher = _find_paragraph(doc, lambda text: text.startswith("（一）教师知识素质要求"))
    if teacher:
        _fill_paragraph_block(
            doc, teacher,
            ["　　任课教师应具备电子商务、大数据或数据科学相关专业背景，熟悉电商运营与数据分析业务流程，掌握Excel、Python及相关数据分析工具，了解人工智能新技术；具备课程思政意识、理实一体化教学设计能力和数字化资源开发能力，能够依据学情改进教学。"],
            ["（二）教材编写与选用", "（三）课程资源的开发与利用", "（四）推荐参考书、期刊、网站", "六、"], body_template,
        )

    textbook_heading = _find_paragraph(doc, lambda text: text.startswith("（二）教材编写与选用"))
    if textbook_heading:
        textbook_name = textbook or "课程配套教材"
        _fill_paragraph_block(
            doc, textbook_heading,
            [
                f"　　本课程选用《{textbook_name}》作为主要教材，并配套PPT课件、源代码、数据集和任务单。",
                "　　教材编写与选用应以职业能力培养为核心，以项目和工作过程组织内容，及时更新行业新技术，并融入数据安全、知识产权和职业道德要求。",
            ],
            ["（三）课程资源的开发与利用", "（四）推荐参考书、期刊、网站", "六、"], body_template,
        )

    resources_heading = _find_paragraph(doc, lambda text: text.startswith("（三）课程资源的开发与利用"))
    if resources_heading:
        _fill_paragraph_block(
            doc, resources_heading,
            [
                "　　充分利用教材配套PPT、源代码、实训数据集和学习通平台资源，建设在线课程、企业真实案例库和分层练习，结合Kaggle、天池等数据平台拓展学习。",
                "　　与合作企业共同开发教学项目和实训案例，建立资源来源、适用项目、版本和版权信息记录，定期更新教学内容并检查个人信息保护和运行安全。",
            ],
            ["（四）推荐参考书、期刊、网站", "六、"], body_template,
        )

    reference_heading = _find_paragraph(doc, lambda text: text.startswith("（四）推荐参考书、期刊、网站"))
    if reference_heading:
        import re
        book_match = re.match(r"(.+?)[（(](.+?)[)）]$", str(textbook).strip())
        book_title = book_match.group(1) if book_match else str(textbook).strip()
        _fill_paragraph_block(doc, reference_heading, [], ["1.参考书"], body_template)
        book_heading = _find_paragraph(doc, lambda text: text.startswith("1.参考书"))
        if book_heading:
            _fill_paragraph_block(
                doc, book_heading,
                [f"（1）《{book_title}》；", "（2）《Python数据科学手册》；", f"（3）《{cn}与数据化运营》或同类行业案例教材。"],
                ["2.期刊", "3.网站", "六、"], body_template,
            )
        journal_heading = _find_paragraph(doc, lambda text: text.startswith("2.期刊"))
        if journal_heading:
            _fill_paragraph_block(
                doc, journal_heading,
                ["（1）《数据分析与知识发现》；（2）《大数据》；（3）《电子商务》。"],
                ["3.网站", "六、"], body_template,
            )
        site_heading = _find_paragraph(doc, lambda text: text.startswith("3.网站"))
        if site_heading:
            _fill_paragraph_block(
                doc, site_heading,
                ["（1）Python官方文档；（2）Pandas官方文档；（3）Kaggle数据科学社区；（4）国家统计局网站。"],
                ["六、"], body_template,
            )

    residue_prefixes = ("（注：", "(注：", "注：", "（正文", "（要求：", "（表格中内容", "(表格中内容", "参考格式如下", "XXXX教学法是")
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(residue_prefixes) or "顶格，行距" in text:
            _clear_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip().startswith(residue_prefixes):
                        _clear_paragraph(paragraph)


def _fill_objectives_table(t, offering):
    """表0：按知识、能力、思政、素质四列完整填充。"""
    while len(t.rows) > 2:
        t.rows[-1]._element.getparent().remove(t.rows[-1]._element)
    if len(t.rows) < 2:
        t.add_row()
    columns = [
        [
            "① 理解数据分析的基本概念、类型、工作流程和适用场景",
            "② 掌握数据分析方法理论和常用指标体系",
            "③ 掌握Excel数据处理、函数、数据透视和图表可视化方法",
            "④ 掌握Python数据分析库进行数据清洗、统计分析和可视化的方法",
            "⑤ 了解机器学习与人工智能在商务数据分析中的应用",
        ],
        [
            "① 能够运用Excel完成数据透视、统计分析和图表可视化",
            "② 能够使用NumPy和Pandas完成数据运算、清洗、聚合和透视分析",
            "③ 能够使用SciPy和scikit-learn完成统计分析、模型构建与效果评价",
            "④ 能够检查运行结果、定位常见问题并持续优化分析成果",
            "⑤ 能够独立完成商务数据分析项目并规范撰写分析报告",
        ],
        [
            "马列主义方面，运用唯物辩证法分析数据现象和商业问题，用发展的眼光认识数字化变革。",
            "理想信念方面，树立科技报国理想，关注数字技术服务产业发展和乡村振兴的实践价值。",
            "核心价值观方面，坚持诚信、公正和法治，不篡改数据、不误导决策。",
            "传统文化方面，践行实事求是、严谨细致的治学态度，传承工匠精神。",
            "职业道德方面，遵守个人信息保护和数据安全法规，严守商业秘密。",
            "个人素养方面，培养批判性思维、问题解决能力和持续改进习惯。",
        ],
        [
            "创新意识方面，培养数据驱动的创新思维，善于运用新技术探索分析方法。",
            "安全意识方面，规范数据处理流程，防范数据泄露风险。",
            "团队协作方面，提升任务分工、沟通反馈、成果检查和按时交付能力。",
            "独立自主方面，能够自主查阅资料、调试代码并解决问题。",
            "个人自律方面，养成按步骤实施、及时测试、记录问题和持续改进的习惯。",
        ],
    ]
    row = t.rows[1]
    for index, lines in enumerate(columns):
        if index < len(row.cells):
            write_cell_lines(row.cells[index], lines, "仿宋_GB2312", 9, {})

def _fill_content_table(t, offering, tasks, units):
    """表1：课程内容划分及课时分配"""
    # 按学习情境分组
    chapter_map = {}
    chapter_order = []
    for task in tasks:
        ch = task["chapter"]
        if ch not in chapter_map:
            chapter_map[ch] = []
            chapter_order.append(ch)
        chapter_map[ch].append(task)

    # 保留表头，清空数据行
    while len(t.rows) > 1:
        t.rows[-1]._element.getparent().remove(t.rows[-1]._element)

    total_hours = 0
    cn_idx = 0

    sub_idx = 0
    for chapter in chapter_order:
        ch_tasks = chapter_map[chapter]
        cn_idx += 1
        sub_idx = 0

        for task in ch_tasks:
            sub_idx += 1
            row = t.add_row()
            ri = len(t.rows) - 1

            # col0: 序号
            write_cell(row.cells[0], str(cn_idx), "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)

            # col1: 学习情境名（第一章XX格式，第一章的第一个任务才加章节号）
            if task == ch_tasks[0]:
                ch_text = f"第{CN_NUMS[min(cn_idx, 9) - 1]}章 {chapter}"
            else:
                ch_text = chapter
            write_cell(row.cells[1], ch_text, "仿宋_GB2312", 9, False)

            # col2: 子情境序号
            write_cell(row.cells[2], str(sub_idx), "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)

            # col3: 子情境名称
            title = task["title"].split("：", 1)[1] if "：" in task["title"] else task["title"]
            write_cell(row.cells[3], title, "仿宋_GB2312", 9, False)

            # col4: 子情境课时
            write_cell(row.cells[4], "2（理论1+实践1）", "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)

            # col5: 总课时（按学习情境合并）
            ch_hours = len(ch_tasks) * 2
            write_cell(row.cells[5], f"{ch_hours}（理论{ch_hours//2}+实践{ch_hours//2}）",
                      "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)

            total_hours += 2

    # 合计行
    row = t.add_row()
    ri = len(t.rows) - 1
    merge_h(t, ri, 0, 4)  # 合并前5列
    write_cell(row.cells[0], "合计", "仿宋_GB2312", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(row.cells[5], f"{total_hours}学时（理论{total_hours//2}学时+实践{total_hours//2}学时）",
              "仿宋_GB2312", 9, True, WD_ALIGN_PARAGRAPH.CENTER)

    # 合并同章节的序号、学习情境、总课时列
    ch_start = 1
    for chapter_index, chapter in enumerate(chapter_order, 1):
        ch_tasks = chapter_map[chapter]
        ch_end = ch_start + len(ch_tasks) - 1
        # 清空将被纵向合并的下方单元格，避免合并后出现“1/1/1/1”等重复文字。
        for lower in range(ch_start + 1, ch_end + 1):
            for col in (0, 1, 5):
                write_cell(t.rows[lower].cells[col], "", "仿宋_GB2312", 9, False)
        merge_v(t, ch_start, ch_end, 0)
        merge_v(t, ch_start, ch_end, 1)
        merge_v(t, ch_start, ch_end, 5)
        ch_text = f"第{CN_NUMS[min(chapter_index, 9) - 1]}章 {chapter}"
        ch_hours = len(ch_tasks) * 2
        write_cell(t.rows[ch_start].cells[0], str(chapter_index), "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(t.rows[ch_start].cells[1], ch_text, "仿宋_GB2312", 9, False)
        write_cell(t.rows[ch_start].cells[5], f"{ch_hours}（理论{ch_hours//2}+实践{ch_hours//2}）", "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)
        ch_start = ch_end + 1

def _fill_assessment_table(t, offering, tasks, units):
    """表2：考核评价"""
    chapter_map = {}
    chapter_order = []
    for task in tasks:
        ch = task["chapter"]
        if ch not in chapter_map:
            chapter_map[ch] = []
            chapter_order.append(ch)
        chapter_map[ch].append(task)

    # 保留表头，清空数据行
    while len(t.rows) > 2:
        t.rows[-1]._element.getparent().remove(t.rows[-1]._element)

    # 跳过最后一个章节（综合评价与课程总结 = 终结性考核）
    assess_chapters = chapter_order[:-1] if len(chapter_order) > 1 else chapter_order

    cn_idx = 0
    data_start = 2  # 数据从第2行开始
    row_cursor = data_start

    for chapter in assess_chapters:
        ch_tasks = chapter_map[chapter]
        cn_idx += 1

        # 每个任务3行：签到/课堂表现/作业
        for task in ch_tasks:
            title = task["title"].split("：", 1)[1] if "：" in task["title"] else task["title"]

            methods = [
                ("签到", 2),
                ("课堂表现", 2),
                ("作业", 1),
            ]

            start_row = row_cursor
            for mi, (method, weight) in enumerate(methods):
                if row_cursor >= len(t.rows):
                    t.add_row()
                row = t.rows[row_cursor]

                # 单元
                if mi == 0:
                    write_cell(row.cells[0], f"单元{cn_idx}（情境{cn_idx}）：{chapter}",
                              "仿宋_GB2312", 9, False)
                # 任务
                if mi == 0:
                    write_cell(row.cells[1], title, "仿宋_GB2312", 9, False)
                # 考核方式
                write_cell(row.cells[2], method, "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)
                # 权重
                write_cell(row.cells[3], f"{weight}%", "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)

                row_cursor += 1

            # 合并单元和任务列
            merge_v(t, start_row, start_row + 2, 0)
            merge_v(t, start_row, start_row + 2, 1)

    # 综合作品行：col4合并所有数据行
    comp_start = data_start
    comp_end = row_cursor - 1
    if comp_end >= comp_start:
        merge_v(t, comp_start, comp_end, 4)
        write_cell(t.rows[comp_start].cells[4], "综合作品\n60%", "仿宋_GB2312", 9, True, WD_ALIGN_PARAGRAPH.CENTER)

    # 表头居中
    for ri in range(2):
        for ci in range(len(t.rows[ri].cells)):
            for p in t.rows[ri].cells[ci].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _fill_scenario_table(t, offering, tasks, units):
    """表3：课程信息一次；情境和子情境均为标题、三列表头、内容、资源表头、内容五行。"""
    groups = {}
    order = []
    for task in tasks:
        chapter = task["chapter"]
        if chapter not in groups:
            groups[chapter] = []
            order.append(chapter)
        groups[chapter].append(task)
    scenarios = order[:-1] if len(order) > 1 else order
    source_rows = [deepcopy(t.rows[i]._tr) for i in range(min(11, len(t.rows)))]
    if len(source_rows) < 6:
        raise ValueError("课程标准学习情境模板至少需要6行结构")
    while len(t.rows):
        t.rows[0]._element.getparent().remove(t.rows[0]._element)

    def append_rows(rows):
        for row_xml in rows:
            t._element.append(deepcopy(row_xml))

    def fill_five_rows(base, title, hours, target, content, method, materials, assessment, note):
        merge_h(t, base, 0, 4)
        write_cell(t.rows[base].cells[0], title, "仿宋_GB2312", 9, True)
        write_cell(t.rows[base].cells[5], f"学时：{hours}", "仿宋_GB2312", 9, True)
        write_cell(t.rows[base+1].cells[0], "学习目标", "仿宋_GB2312", 9, True)
        merge_h(t, base+1, 1, 3)
        write_cell(t.rows[base+1].cells[1], "主要内容", "仿宋_GB2312", 9, True)
        merge_h(t, base+1, 4, 5)
        write_cell(t.rows[base+1].cells[4], "教学方法和建议", "仿宋_GB2312", 9, True)
        write_cell(t.rows[base+2].cells[0], target, "仿宋_GB2312", 9, False)
        merge_h(t, base+2, 1, 3)
        write_cell(t.rows[base+2].cells[1], content, "仿宋_GB2312", 9, False)
        merge_h(t, base+2, 4, 5)
        write_cell(t.rows[base+2].cells[4], method, "仿宋_GB2312", 9, False)
        write_cell(t.rows[base+3].cells[0], "教学材料、工具、媒体与参考资料", "仿宋_GB2312", 9, True)
        merge_h(t, base+3, 1, 3)
        write_cell(t.rows[base+3].cells[1], "考核与评价", "仿宋_GB2312", 9, True)
        merge_h(t, base+3, 4, 5)
        write_cell(t.rows[base+3].cells[4], "备注", "仿宋_GB2312", 9, True)
        write_cell(t.rows[base+4].cells[0], materials, "仿宋_GB2312", 9, False)
        merge_h(t, base+4, 1, 3)
        write_cell(t.rows[base+4].cells[1], assessment, "仿宋_GB2312", 9, False)
        merge_h(t, base+4, 4, 5)
        write_cell(t.rows[base+4].cells[4], note, "仿宋_GB2312", 9, False)

    append_rows(source_rows[:1])
    merge_h(t, 0, 0, 1)
    write_cell(t.rows[0].cells[0], f"课程：{offering['course_name']}", "仿宋_GB2312", 9, True)
    write_cell(t.rows[0].cells[2], "学分：", "仿宋_GB2312", 9, True)
    write_cell(t.rows[0].cells[3], str(int(offering.get("credits", 3))), "仿宋_GB2312", 9, True)
    write_cell(t.rows[0].cells[5], f"总学时：{int(offering.get('total_hours', 60))}", "仿宋_GB2312", 9, True)

    scenario_pattern = source_rows[1:6]
    sub_pattern = source_rows[6:11] if len(source_rows) >= 11 else source_rows[1:6]
    for ci, chapter in enumerate(scenarios, 1):
        chapter_tasks = groups[chapter]
        unit = _find_unit(chapter, units)
        skills = unit.get("source_skills", "") if unit else ""
        append_rows(scenario_pattern)
        base = len(t.rows) - 5
        target = f"知识目标：\n1. 掌握{chapter}核心概念和方法\n2. 理解任务实施流程\n能力目标：\n1. 能独立完成实训操作\n2. 能调试并解决常见问题\n思政目标：\n1. 遵守数据安全和个人信息保护要求\n素质目标：\n1. 形成规范操作、团队协作和持续改进习惯"
        fill_five_rows(base, f"学习情境{ci}：{chapter}", len(chapter_tasks)*2, target,
                       skills or f"{chapter}的理论讲解、案例分析与实操训练",
                       "教学方法：项目教学法、任务驱动法、案例教学法、操作演示法\n建议：理实一体化、案例导入、分层任务",
                       "教材、PPT课件、源代码、Jupyter Notebook、实训数据集和学习通资源",
                       "知识、技能、思政和素质目标综合评价；签到、课堂表现、作业与成果评价相结合",
                       "根据学情实施分层指导并及时反馈")
        for si, task in enumerate(chapter_tasks, 1):
            title = task["title"].split("：", 1)[1] if "：" in task["title"] else task["title"]
            append_rows(sub_pattern)
            base = len(t.rows) - 5
            target = "教学目标：\n" + (task.get("knowledge_goal", "") or f"掌握{title}的基本概念和操作方法") + "\n思政目标：\n" + (task.get("ideological_goal", "") or "遵守数据安全规范，坚持真实分析")
            fill_five_rows(base, f"学习子情境{ci}.{si}：{title}", 2, target,
                           title + "：基本概念与原理、操作方法与步骤、代码实现与演示、常见错误与调试、实训任务",
                           "教学方法：案例教学法、讲授法、实操练习法\n建议：理论联系实际、讲练结合、巡回指导",
                           "多媒体设备、PPT课件、源代码、Jupyter Notebook、实训数据集和学习通资源",
                           "知识理解、实训操作、职业素养；课堂提问+随堂练习+课后作业",
                           f"通过本子情境掌握{title}，为后续学习奠定基础。")

def _find_unit(chapter, units):
    for u in units:
        if u["project_title"] == chapter:
            return u
    return units[0] if units else {}


# ============================================================
# 新模板函数（2024-2025-2、2025-2026）
# ============================================================

def _fill_course_info_table_new(t, offering):
    """新模板表0：课程基本信息（9列）"""
    cn = offering["course_name"]
    dp = offering.get("department", "经济贸易系")
    mj = offering.get("major", "")
    th = int(offering.get("total_hours", 60))
    cr = int(offering.get("credits", 3))
    ct = offering.get("course_type", "专业核心课")
    term = offering["term"]

    rows = t.rows
    if len(rows) >= 2:
        # 行0: 课程名称(0-1) | 值(2-3) | 所属系部(4-6) | 值(7-8)
        if len(rows[0].cells) >= 9:
            write_cell(rows[0].cells[2], cn, "宋体", 10, False)
            write_cell(rows[0].cells[7], dp, "宋体", 10, False)
        if len(rows) >= 3:
            # 行2: 适用专业
            if len(rows[2].cells) >= 9:
                write_cell(rows[2].cells[2], mj, "宋体", 10, False)
                write_cell(rows[2].cells[7], ct, "宋体", 10, False)
        if len(rows) >= 4:
            # 行3: 学时/学分
            if len(rows[3].cells) >= 9:
                write_cell(rows[3].cells[2], f"{th}学时", "宋体", 10, False)
                write_cell(rows[3].cells[7], f"{cr}学分", "宋体", 10, False)


def _fill_objectives_table_new(t, offering):
    """新模板表2：课程目标（3列：素质目标/知识目标/能力目标）"""
    cn = offering["course_name"]

    # 清空多余行，保留表头
    while len(t.rows) > 2:
        t.rows[-1]._element.getparent().remove(t.rows[-1]._element)

    if len(t.rows) < 2:
        t.add_row()
    content_row = t.rows[1]

    # 列0: 素质目标
    quality_lines = [
        ("【素质目标】", True),
        ("1. 培养数据驱动的创新思维和问题解决意识", False),
        ("2. 树立数据安全意识和隐私保护观念", False),
        ("3. 提升沟通协作和项目合作能力", False),
        ("4. 恪守数据分析师职业操守和行为规范", False),
    ]
    write_cell_lines(content_row.cells[0], [t for t, _ in quality_lines],
                    "仿宋_GB2312", 9, {i: b for i, (_, b) in enumerate(quality_lines)})

    # 列1: 知识目标
    knowledge_lines = [
        ("【知识目标】", True),
        ("1. 掌握数据分析的基本概念、流程和方法体系", False),
        ("2. 理解数据采集、清洗、分析和可视化的技术原理", False),
        ("3. 熟悉Excel、Python数据分析库的功能和应用场景", False),
        ("4. 了解大数据、人工智能等新技术在商务分析中的应用", False),
    ]
    write_cell_lines(content_row.cells[1], [t for t, _ in knowledge_lines],
                    "仿宋_GB2312", 9, {i: b for i, (_, b) in enumerate(knowledge_lines)})

    # 列2: 能力目标
    ability_lines = [
        ("【能力目标】", True),
        ("1. 能够运用Excel进行数据透视分析和图表制作", False),
        ("2. 能够运用Python进行数据清洗和处理", False),
        ("3. 能够运用机器学习进行统计分析和建模", False),
        ("4. 能够完成完整的数据分析项目，撰写分析报告", False),
    ]
    write_cell_lines(content_row.cells[2], [t for t, _ in ability_lines],
                    "仿宋_GB2312", 9, {i: b for i, (_, b) in enumerate(ability_lines)})


def _fill_content_table_new(t, offering, tasks, units):
    """新模板表3：课程内容划分及课时分配（4列）"""
    # 按学习情境分组
    chapter_map = {}
    chapter_order = []
    for task in tasks:
        ch = task["chapter"]
        if ch not in chapter_map:
            chapter_map[ch] = []
            chapter_order.append(ch)
        chapter_map[ch].append(task)

    # 保留表头，清空数据行
    while len(t.rows) > 1:
        t.rows[-1]._element.getparent().remove(t.rows[-1]._element)

    total_hours = 0
    cn_idx = 0

    for chapter in chapter_order:
        ch_tasks = chapter_map[chapter]
        cn_idx += 1
        ch_hours = len(ch_tasks) * 2

        for sub_idx, task in enumerate(ch_tasks, 1):
            row = t.add_row()

            # col0: 序号
            write_cell(row.cells[0], str(cn_idx), "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)

            # col1: 模块（或项目）名称 → 第一章XX格式，第一章第一个任务才加章节号
            if sub_idx == 1:
                ch_text = f"第{CN_NUMS[min(cn_idx, 9) - 1]}章 {chapter}"
            else:
                ch_text = chapter
            write_cell(row.cells[1], ch_text, "仿宋_GB2312", 9, False)

            # col2: 任务名称
            title = task["title"].split("：", 1)[1] if "：" in task["title"] else task["title"]
            write_cell(row.cells[2], title, "仿宋_GB2312", 9, False)

            # col3: 学时分配
            write_cell(row.cells[3], "2（理论1+实践1）", "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)

            total_hours += 2

        # 合并同章节的序号、模块名称列；先清空下方单元格，避免合并后出现重复序号。
        data_start = len(t.rows) - len(ch_tasks)
        data_end = len(t.rows) - 1
        if data_end > data_start:
            for lower in range(data_start + 1, data_end + 1):
                write_cell(t.rows[lower].cells[0], "", "仿宋_GB2312", 9, False)
                write_cell(t.rows[lower].cells[1], "", "仿宋_GB2312", 9, False)
            merge_v(t, data_start, data_end, 0)
            merge_v(t, data_start, data_end, 1)

    # 合计行
    row = t.add_row()
    merge_h(t, len(t.rows) - 1, 0, 2)
    write_cell(row.cells[0], "合计", "仿宋_GB2312", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(row.cells[3],
              f"{total_hours}学时（理论{total_hours//2}学时+实践{total_hours//2}学时）",
              "仿宋_GB2312", 9, True, WD_ALIGN_PARAGRAPH.CENTER)


def _fill_assessment_table_new(t, offering, tasks, units):
    """新模板表5：考核评价（4列：评价类型/评价内容/评价主体/评价方式）"""
    chapter_map = {}
    chapter_order = []
    for task in tasks:
        ch = task["chapter"]
        if ch not in chapter_map:
            chapter_map[ch] = []
            chapter_order.append(ch)
        chapter_map[ch].append(task)

    # 保留表头行，清空数据行
    while len(t.rows) > 1:
        t.rows[-1]._element.getparent().remove(t.rows[-1]._element)

    # 跳过最后一个章节（综合评价与课程总结 = 终结性考核）
    assess_chapters = chapter_order[:-1] if len(chapter_order) > 1 else chapter_order

    # 过程性考核行
    row_proc = t.add_row()
    write_cell(row_proc.cells[0], "过程性考核\n（40%）", "仿宋_GB2312", 9, True, WD_ALIGN_PARAGRAPH.CENTER)

    # 收集所有任务的评价内容
    proc_items = []
    cn_idx = 0
    for chapter in assess_chapters:
        ch_tasks = chapter_map[chapter]
        cn_idx += 1
        for task in ch_tasks:
            title = task["title"].split("：", 1)[1] if "：" in task["title"] else task["title"]
            proc_items.append((f"单元{cn_idx}：{title}", cn_idx))

    # 评价内容
    proc_content = "\n".join([
        f"签到（10%）：每单元考勤",
        f"课堂表现（10%）：互动问答、实操表现",
        f"作业（20%）：每单元课后作业",
    ])
    write_cell(row_proc.cells[1], proc_content, "仿宋_GB2312", 9, False)

    write_cell(row_proc.cells[2], "教师评价\n学生互评", "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(row_proc.cells[3], "签到记录+课堂观察+作业批改", "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)

    # 终结性考核行
    row_final = t.add_row()
    write_cell(row_final.cells[0], "终结性考核\n（60%）", "仿宋_GB2312", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(row_final.cells[1], "综合作品：完整的数据分析项目报告\n（含数据采集、清洗、分析、可视化、结论建议）",
              "仿宋_GB2312", 9, False)
    write_cell(row_final.cells[2], "教师评价", "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(row_final.cells[3], "作品评分量表", "仿宋_GB2312", 9, False, WD_ALIGN_PARAGRAPH.CENTER)

    # 合并评价主体和评价方式列
    if len(t.rows) >= 3:
        merge_v(t, 1, 2, 2)
        merge_v(t, 1, 2, 3)

    # 表头居中
    for ci in range(len(t.rows[0].cells)):
        for p in t.rows[0].cells[ci].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
