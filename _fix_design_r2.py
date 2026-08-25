"""
第二轮修复：在每对表5/表6前插入完整的文字描述段落
模板结构（P116-P177）：
  P116: 《课程名》课程单元教学设计（标题）
  P145: 【教学任务】
  P146: 【授课教师】
  P147: 【授课时间/班级】
  P148: 【课时】
  P149: 【课程类型】
  P150: 【教学场景设计】
  P151: 【教学资源准备】
  P152: 【教学目标】
  P153: 知识目标
  P154: 能力目标
  P155: 思政目标
  P156: 素质目标
  P157: 【教材、学情分析及教育理念】
  P158: 教材分析
  P159: 学情分析
  P160: 教师教育理念
  P161: 【教学重点】
  P162: 【教学难点】
  P163: 【教法学法】
  P164: 【教学活动流程】
  P165: 导入新课
  P166: 讲授新课
  P167-173: 任务1-4
  P174: 【板书设计】
  P175: 【课堂小结】
  P176: 【课后作业】
  P177: 【教学反思】
"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from copy import deepcopy
import store
import os
import re

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
template_fp = r"原始资料\模板\模板5：教学设计\模板5：教学设计 模板（2023-2024）.docx"

# 读取模板中的段落XML作为模板
tmpl_doc = Document(template_fp)
tmpl_paras = tmpl_doc.paragraphs

# 提取P116-P177的段落元素作为模板（深拷贝）
text_template_elements = []
for i in range(116, 178):
    if i < len(tmpl_paras):
        text_template_elements.append(deepcopy(tmpl_paras[i]._element))

print(f"提取了{len(text_template_elements)}个模板段落")

# 获取任务数据
tasks = store.rows("SELECT * FROM tasks WHERE offering_id=20 ORDER BY seq")
sessions = store.rows("SELECT * FROM sessions WHERE offering_id=20 ORDER BY week_no, lesson_date")

print(f"任务数: {len(tasks)}")
print(f"课程数: {len(sessions)}")

# ============================================================
# 为每个任务生成文字描述内容
# ============================================================

def get_task_content(task, idx, session):
    """为每个任务生成详细的文字描述内容"""
    title = task.get('title', f'任务{idx+1}')
    seq = task.get('seq', idx+1)
    
    # 解析章节和子任务
    parts = title.split('：')
    if len(parts) >= 2:
        chapter_name = parts[0]
        sub_tasks = parts[1].split('、')
    else:
        chapter_name = title
        sub_tasks = [title]
    
    # 章节号
    chapter_num = f"第{['一','二','三','四','五','六','七','八','九'][seq-1] if seq <= 9 else seq}章"
    
    # 上课日期
    lesson_date = session.get('lesson_date', '') if session else ''
    week_no = session.get('week_no', '') if session else ''
    
    # ============================================================
    # 教学任务
    # ============================================================
    teach_task = f"{chapter_num} {chapter_name}\n"
    for st in sub_tasks:
        teach_task += f"子情景：{st.strip()}\n"
    
    # ============================================================
    # 教学目标
    # ============================================================
    # 根据章节生成不同目标
    knowledge_goals = []
    ability_goals = []
    sizheng_goals = []
    quality_goals = []
    
    if seq == 1:
        knowledge_goals = [
            "陈述数据分析的基本概念、分类和适用场景",
            "概述电商核心数据分析指标体系（GMV、转化率、客单价等）",
            "解释对比分析、分组分析、交叉分析等常用分析方法",
            "总结数据分析的基本流程和规范要求"
        ]
        ability_goals = [
            "能够识别电商业务中的数据分析需求",
            "能够运用对比分析法分析电商销售数据",
            "能够选择合适的数据分析方法解决实际问题"
        ]
        sizheng_goals = [
            "建立数据安全意识，遵守《个人信息保护法》",
            "树立诚信分析精神，坚持真实客观的数据呈现"
        ]
        quality_goals = [
            "培养数据驱动的分析思维",
            "养成按流程规范操作的习惯"
        ]
    elif seq == 2:
        knowledge_goals = [
            "陈述Excel数据分析工具的功能模块",
            "概述SUMIF、VLOOKUP等常用函数的语法和参数",
            "解释数据透视表的创建步骤和布局设置",
            "总结Excel在数据分析中的优势与局限"
        ]
        ability_goals = [
            "能够使用Excel函数进行数据计算",
            "能够创建数据透视表进行多维分析",
            "能够制作数据透视图展示分析结果"
        ]
        sizheng_goals = [
            "树立精益求精的工匠精神，确保数据计算准确",
            "增强规范操作意识，遵循数据处理标准"
        ]
        quality_goals = [
            "培养独立思考和问题解决能力",
            "养成检查和验证计算结果的习惯"
        ]
    elif seq <= 9:
        # 通用模板（根据章节调整）
        tool_name = chapter_name
        knowledge_goals = [
            f"陈述{tool_name}的基本概念和核心功能",
            f"概述{tool_name}的主要API和常用方法",
            f"解释{tool_name}在数据分析中的应用场景",
            f"总结{tool_name}的操作流程和注意事项"
        ]
        ability_goals = [
            f"能够安装和配置{tool_name}环境",
            f"能够使用{tool_name}进行数据处理和分析",
            f"能够编写代码完成{tool_name}相关实训任务"
        ]
        sizheng_goals = [
            "增强数据安全与隐私保护意识",
            "培养科技报国情怀，关注技术服务乡村振兴"
        ]
        quality_goals = [
            "培养代码规范编写和测试的习惯",
            "提升团队协作和成果展示能力"
        ]
    else:
        knowledge_goals = [
            "梳理课程完整知识体系",
            "总结各模块核心知识点和技能",
            "回顾行业新技术和新方法",
            "明确1+X证书和技能大赛要点"
        ]
        ability_goals = [
            "能够独立完成综合数据分析项目",
            "能够撰写规范的数据分析报告",
            "能够进行项目成果展示和答辩"
        ]
        sizheng_goals = [
            "坚定科技报国的理想信念",
            "树立终身学习和持续发展的意识"
        ]
        quality_goals = [
            "培养项目管理和时间管理能力",
            "养成总结反思和持续改进的习惯"
        ]
    
    # ============================================================
    # 教材分析
    # ============================================================
    material_analysis = (
        f"本单元{chapter_name}是商务数据分析课程的重要组成部分，"
        f"在课程体系中起着{ '基础铺垫' if seq <= 3 else '技能提升' if seq <= 6 else '综合应用' if seq <= 9 else '总结提升'}的作用。"
        f"本单元内部各部分逻辑关系为：先讲解基本概念和原理，再演示操作步骤，最后通过实训巩固，"
        f"遵循从理论到实践、从认识到应用的教学逻辑。"
    )
    
    # 学情分析
    student_analysis = (
        f"学生在本单元之前已学习{ 'Python程序设计和电子商务基础' if seq <= 3 else '前序章节内容'}，"
        f"具备{ '基本的编程思维和电商知识' if seq <= 3 else '一定的数据分析基础'}。"
        f"学习情感方面，学生对{chapter_name}的实操环节兴趣较高，"
        f"但对{ '抽象概念的理解' if seq <= 3 else '代码编写和调试'}存在畏难情绪，"
        f"需要通过案例引导和分步演示降低学习难度。"
    )
    
    # 教师教育理念
    teacher_concept = (
        "坚持以学生为中心、做中学的教育理念，注重培养学生的实践能力和创新思维，"
        "通过任务驱动和项目教学法，让学生在解决实际问题中掌握知识和技能。"
    )
    
    # ============================================================
    # 教学重点难点
    # ============================================================
    key_points = (
        f"1. {chapter_name}的基本概念和核心原理\n"
        f"2. {'常用函数和数据透视表的创建' if seq == 2 else '核心API的使用方法和参数设置' if seq <= 8 else '知识体系的梳理和综合应用'}\n"
        f"3. {'分析方法的实际应用' if seq == 1 else '操作步骤和注意事项' if seq <= 8 else '项目成果的完整呈现'}"
    )
    
    difficult_points = (
        f"1. {'数据分析方法的选择和应用' if seq == 1 else '复杂函数的参数理解' if seq == 2 else '代码调试和错误处理' if seq <= 8 else '综合项目的完整实施'}\n"
        f"2. {'业务需求与分析方法的对应关系' if seq == 1 else '数据透视表的多维交叉分析' if seq == 2 else 'API参数组合使用' if seq <= 8 else '项目时间管理和质量把控'}\n"
        f"3. {'从数据中提炼有效结论' if seq == 1 else '大数据量下的性能优化' if seq == 2 else '异常情况的处理' if seq <= 8 else '答辩中的问题应答'}"
    )
    
    # 教法学法
    teach_method = (
        f"教法：案例教学法（分析电商真实案例）、任务驱动法（以实训任务为载体）、"
        f"演示法（{ '演示Excel操作' if seq == 2 else '演示代码编写和运行' if seq <= 8 else '演示项目流程'}）、"
        f"启发式教学法（引导学生思考）。\n"
        f"学法：自主学习（课前预习微课）、小组讨论（讨论解决方案）、"
        f"实践操作（跟随教师操作）、互评法（交换检查成果）。"
    )
    
    # ============================================================
    # 教学活动流程
    # ============================================================
    import_methods = [
        ("案例导入法", "展示电商数据分析典型案例，引发思考"),
        ("复习导入法", "回顾上节课内容，自然引入新知识"),
        ("任务引入法", "提出实训任务，明确学习目标"),
        ("情境导入法", "创设电商工作情境，代入角色")
    ]
    import_method = import_methods[idx % 4]
    
    # 导入新课
    intro = (
        f"【{import_method[0]}】{import_method[1]}。"
        f"学生活动：思考、讨论、回答问题。"
        f"教师归纳：引入{chapter_name}的学习内容。"
        f"（5分钟）\n"
        f"德育渗透：引导学生关注数据安全与诚信分析的重要性。"
    )
    
    # 讲授新课 - 任务1-3
    tasks_content = []
    for si, sub_task in enumerate(sub_tasks[:3]):
        sub_task = sub_task.strip()
        task_content = (
            f"任务{si+1}：{sub_task}（知识讲解）\n"
            f"一、基本概念\n"
            f"  1. {sub_task}的定义与内涵\n"
            f"  2. {sub_task}的分类和适用场景\n"
            f"二、核心知识点\n"
            f"  1. {'基本原理与流程' if si == 0 else '常用方法与工具' if si == 1 else '实际应用与注意事项'}\n"
            f"  2. {'关键参数和配置' if seq <= 8 else '分析维度和指标'}\n"
            f"  3. {'代码示例和运行结果' if seq <= 8 else '案例分析步骤'}\n"
            f"三、操作演示\n"
            f"  1. 教师演示{sub_task}的完整操作流程\n"
            f"  2. 学生跟随操作，教师巡回指导\n"
            f"  3. 学生独立完成练习任务\n"
            f"德育渗透：培养学生的{ '数据安全意识' if si == 0 else '规范操作意识' if si == 1 else '团队协作精神'}\n"
            f"（{'20' if si == 0 else '20' if si == 1 else '15'}分钟）"
        )
        tasks_content.append(task_content)
    
    # 如果只有2个子任务，加一个综合练习
    if len(sub_tasks) <= 2:
        tasks_content.append(
            f"任务{len(sub_tasks)+1}：综合练习（实操训练）\n"
            f"一、实训目标\n"
            f"  1. 综合运用本节课所学知识解决实际问题\n"
            f"  2. 培养独立分析和解决问题的能力\n"
            f"二、实训步骤\n"
            f"  1. 分析任务需求，确定解决方案\n"
            f"  2. 编写{'代码' if seq <= 8 else '分析方案'}并运行调试\n"
            f"  3. 检查结果并优化\n"
            f"三、成果展示\n"
            f"  学生展示成果，教师点评\n"
            f"德育渗透：培养精益求精的工匠精神\n"
            f"（20分钟）"
        )
    
    # 板书设计
    blackboard = (
        f"【板书设计】\n"
        f"{'='*30}\n"
        f"{chapter_num} {chapter_name}\n"
        f"{'='*30}\n"
        f"一、基本概念\n"
        f"  · 定义\n"
        f"  · 分类\n"
        f"二、核心知识点\n"
        f"  · 要点1\n"
        f"  · 要点2\n"
        f"  · 要点3\n"
        f"三、操作流程\n"
        f"  1. 步骤1 → 2. 步骤2 → 3. 步骤3\n"
        f"四、注意事项\n"
        f"  · 注意1\n"
        f"  · 注意2\n"
        f"{'='*30}"
    )
    
    # 课堂小结
    summary_types = [
        ("【知识梳理】", f"本节课围绕{chapter_name}展开，核心知识点包括：①基本概念与分类；②核心原理与方法；③操作流程与注意事项。技能要点：能够独立完成{chapter_name}的相关操作。学习态度点评：课堂参与积极，实操练习认真。"),
        ("【三点总结】", f"本节课收获最大的三点：①掌握了{chapter_name}的基本概念；②学会了核心操作步骤；③理解了实际应用场景。需要注意的三点：①概念理解要准确；②操作步骤要规范；③结果检查要仔细。下一步计划：课后复习巩固，完成作业练习。"),
        ("【问答回顾】", f"Q1：{chapter_name}的核心功能是什么？A1：{sub_tasks[0] if sub_tasks else chapter_name}。Q2：操作的关键步骤有哪些？A2：环境准备→参数设置→执行操作→检查结果。Q3：常见问题如何解决？A3：参考错误提示，检查参数和代码。课堂表现点评：发言积极，讨论热烈。"),
        ("【对比总结】", f"本节课{chapter_name}与前序内容的联系与区别：联系在于都是数据分析工具/方法的应用；区别在于{chapter_name}{'更注重代码实现' if seq <= 8 else '更注重综合应用'}。核心要点回顾：1.概念→2.原理→3.操作→4.实训。学习建议：多练习、多思考、多总结。"),
    ]
    summary_type = summary_types[idx % 4]
    summary = f"{summary_type[0]}\n{summary_type[1]}\n（5分钟）"
    
    # 课后作业
    homework_types = [
        ("【基础巩固+实操提升】", f"1. 整理本节课知识点笔记，重点记录{chapter_name}的概念和操作流程。\n2. 完成实训任务：使用{chapter_name}处理电商数据集，输出分析结果。\n3. 思考题：{chapter_name}在实际电商业务中有哪些应用场景？举例说明。\n4. 预习下一章内容，观看微课视频。"),
        ("【分层作业】", f"基础层：1. 梳理{chapter_name}的知识框架图。2. 完成3道基础练习题。\n提升层：3. 使用{chapter_name}分析给定电商数据，撰写分析报告。4. 对比{chapter_name}与其他工具的异同。\n拓展层：5. 查阅{chapter_name}最新技术文档，了解新特性。6. 尝试用{chapter_name}解决一个实际问题。"),
        ("【项目式作业】", f"项目任务：基于{chapter_name}完成一个电商数据分析小项目。\n要求：1. 确定分析目标和数据来源。2. 使用{chapter_name}进行数据处理和分析。3. 输出分析报告（含图表）。4. 下节课展示分享。\n提交方式：电子版报告+代码文件。"),
        ("【复盘式作业】", f"1. 回顾本节课学习过程，填写学习反思表（学到了什么/还不会什么/下一步计划）。\n2. 用自己的话解释{chapter_name}的核心概念，录制1分钟语音。\n3. 修改完善课堂实训代码，添加注释。\n4. 在学习通平台完成本节测验。"),
        ("【PBL式作业】", f"问题情境：某电商平台需要分析用户购买行为数据，请你设计分析方案。\n任务：1. 明确分析目标和数据需求。2. 选择合适的方法（含{chapter_name}）。3. 编写分析代码。4. 撰写分析结论和建议。\n合作要求：2-3人一组，下周展示。"),
    ]
    homework_type = homework_types[idx % 5]
    homework = f"{homework_type[0]}\n{homework_type[1]}\n（30分钟）"
    
    # 教学反思
    reflection_types = [
        ("【教学效果反思】", f"本节课整体教学效果良好，{chapter_name}的概念讲解清晰，学生理解到位。实训环节学生参与积极，大部分学生能独立完成任务。成功经验：案例导入有效激发了学习兴趣；分步演示降低了操作难度。改进方向：时间分配需优化，概念讲解部分可压缩5分钟留给实操练习；部分学生代码基础薄弱，需加强个别指导。"),
        ("【学生反馈反思】", f"通过课堂观察和课后交流，学生对{chapter_name}的学习反馈积极。多数学生表示案例教学帮助理解抽象概念，实操练习增强了动手能力。存在问题：少数学生对{'代码调试' if seq <= 8 else '综合应用'}感到困难，需要在后续课程中加强练习。教学调整：增加同伴互助环节，让掌握较快的学生帮助其他同学。"),
        ("【方法策略反思】", f"本节课采用的教学方法和策略整体有效。任务驱动法让学生有明确目标，演示法让操作步骤可视化。需要改进的地方：导入环节可以更有趣味性；小组讨论时间可以适当延长；德育渗透可以更自然地融入教学内容而非单独提出。下节课将尝试引入更多互动元素。"),
    ]
    reflection_type = reflection_types[idx % 3]
    reflection = f"{reflection_type[0]}\n{reflection_type[1]}"
    
    # 教学场景设计
    scene = f"801教室（理实一体化教室，配有计算机、投影设备）"
    
    # 教学资源准备
    resources = (
        f"1. 多媒体课件PPT（{chapter_name}）\n"
        f"2. 微课视频（{sub_tasks[0] if sub_tasks else chapter_name}）\n"
        f"3. 实训数据集（电商{'交易' if seq <= 4 else '用户行为'}数据）\n"
        f"4. {'Python代码模板和示例' if seq <= 8 else '综合项目需求文档'}\n"
        f"5. 超星学习通在线资源"
    )
    
    # 组装所有段落内容
    content = {
        'title': f"《商务数据分析》课程单元教学设计",
        'task': teach_task,
        'teacher': '杜媛',
        'time_class': f"{lesson_date} / 2022电商教学班",
        'hours': '2学时',
        'course_type': '理实一体课程',
        'scene': scene,
        'resources': resources,
        'knowledge_goal': '；'.join(knowledge_goals),
        'ability_goal': '；'.join(ability_goals),
        'zhengzhi_goal': '；'.join(sizheng_goals),
        'quality_goal': '；'.join(quality_goals),
        'material_analysis': material_analysis,
        'student_analysis': student_analysis,
        'teacher_concept': teacher_concept,
        'key_points': key_points,
        'difficult_points': difficult_points,
        'teach_method': teach_method,
        'intro': intro,
        'task1': tasks_content[0] if len(tasks_content) > 0 else '',
        'task2': tasks_content[1] if len(tasks_content) > 1 else '',
        'task3': tasks_content[2] if len(tasks_content) > 2 else '',
        'task4': '',
        'blackboard': blackboard,
        'summary': summary,
        'homework': homework,
        'reflection': reflection,
    }
    
    return content

# ============================================================
# 在每对表5/表6前插入文字描述段落
# ============================================================
print("\n插入文字描述段落...")

doc = Document(fp)
body = doc.element.body

# 找到所有"教学设计·基本信息"标题段落
# 这些标题是在第一轮插入的，每对表5前面有一个
inserted = 0
for ti in range(5, len(doc.tables), 2):
    t5 = doc.tables[ti]
    t5_elem = t5._element
    
    # 找到表5前面的"教学设计·基本信息"标题
    prev = t5_elem.getprevious()
    title_elem = None
    while prev is not None:
        if prev.tag.endswith('}p'):
            text = ""
            for r in prev.findall(qn('w:r')):
                for t in r.findall(qn('w:t')):
                    text += t.text or ''
            if '教学设计·基本信息' in text:
                title_elem = prev
                break
        prev = prev.getprevious()
    
    if title_elem is None:
        print(f"  表{ti}: 未找到标题，跳过")
        continue
    
    # 获取任务索引
    task_idx = (ti - 5) // 2
    if task_idx >= len(tasks):
        print(f"  表{ti}: 任务索引{task_idx}超出范围")
        continue
    
    task = tasks[task_idx]
    session = sessions[task_idx] if task_idx < len(sessions) else None
    content = get_task_content(task, task_idx, session)
    
    # 在标题前插入文字描述段落（从模板段落复制结构，替换内容）
    # 使用模板P116-P177的段落结构
    prev_elem = title_elem.getprevious()  # 在标题前一个元素后插入
    
    # 实际上应该在标题前面插入所有文字描述段落
    # 先找到插入点：标题元素的前面
    insert_before = title_elem
    
    # 从模板段落创建新的段落
    new_elements = []
    for tmpl_elem in text_template_elements:
        new_elem = deepcopy(tmpl_elem)
        new_elements.append(new_elem)
    
    # 替换内容
    # P116: 标题 《课程名》课程单元教学设计
    # P145: 【教学任务】
    # P146: 【授课教师】
    # ...等
    
    # 段落索引映射（相对于text_template_elements）
    # 0: P116 标题
    # 29: P145 教学任务
    # 30: P146 授课教师
    # 31: P147 授课时间/班级
    # 32: P148 课时
    # 33: P149 课程类型
    # 34: P150 教学场景设计
    # 35: P151 教学资源准备
    # 36: P152 教学目标（标题）
    # 37: P153 知识目标
    # 38: P154 能力目标
    # 39: P155 思政目标
    # 40: P156 素质目标
    # 41: P157 教材学情分析（标题）
    # 42: P158 教材分析
    # 43: P159 学情分析
    # 44: P160 教师教育理念
    # 45: P161 教学重点
    # 46: P162 教学难点
    # 47: P163 教法学法
    # 48: P164 教学活动流程（标题）
    # 49: P165 导入新课
    # 50: P166 讲授新课
    # 51: P167 任务1
    # 52: P168 任务1内容
    # 53: P169 任务2
    # 54: P170 任务2内容
    # 55: P171 任务3
    # 56: P172 任务3内容
    # 57: P173 任务4
    # 58: P174 板书设计
    # 59: P175 课堂小结
    # 60: P176 课后作业
    # 61: P177 教学反思
    
    def set_elem_text(elem, text):
        """设置段落元素的内容，保留格式"""
        # 清除现有run内容
        for r in elem.findall(qn('w:r')):
            elem.remove(r)
        # 添加新run
        from docx.oxml import OxmlElement
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), '仿宋')
        rFonts.set(qn('w:eastAsia'), '仿宋')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '24')
        rPr.append(szCs)
        r.append(rPr)
        
        # 处理多行文本
        lines = text.split('\n')
        for li, line in enumerate(lines):
            if li > 0:
                br = OxmlElement('w:br')
                r.append(br)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = line
            r.append(t)
        
        elem.append(r)
    
    def set_elem_text_bold(elem, text):
        """设置段落元素的内容，加粗"""
        for r in elem.findall(qn('w:r')):
            elem.remove(r)
        from docx.oxml import OxmlElement
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), '仿宋')
        rFonts.set(qn('w:eastAsia'), '仿宋')
        rPr.append(rFonts)
        b = OxmlElement('w:b')
        rPr.append(b)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '24')
        rPr.append(szCs)
        r.append(rPr)
        
        lines = text.split('\n')
        for li, line in enumerate(lines):
            if li > 0:
                br = OxmlElement('w:br')
                r.append(br)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = line
            r.append(t)
        
        elem.append(r)
    
    # 替换各段落内容
    replacements = {
        0: (content['title'], True, 18),    # P116 标题
        29: (f"【教学任务】{content['task']}", False, 12),  # P145
        30: (f"【授课教师】{content['teacher']}", False, 12),  # P146
        31: (f"【授课时间/班级】{content['time_class']}", False, 12),  # P147
        32: (f"【课    时】{content['hours']}", False, 12),  # P148
        33: (f"【课程类型】{content['course_type']}", False, 12),  # P149
        34: (f"【教学场景设计】{content['scene']}", False, 12),  # P150
        35: (f"【教学资源准备】\n{content['resources']}", False, 12),  # P151
        36: ("【教学目标】", True, 12),  # P152
        37: (f"知识目标：{content['knowledge_goal']}", True, 12),  # P153
        38: (f"能力目标：{content['ability_goal']}", True, 12),  # P154
        39: (f"思政目标：{content['zhengzhi_goal']}", True, 12),  # P155
        40: (f"素质目标：{content['quality_goal']}", True, 12),  # P156
        41: ("【教材、学情分析及教育理念】", True, 12),  # P157
        42: (f"教材分析：{content['material_analysis']}", True, 12),  # P158
        43: (f"学情分析：{content['student_analysis']}", True, 12),  # P159
        44: (f"教师教育理念：{content['teacher_concept']}", True, 12),  # P160
        45: (f"【教学重点】\n{content['key_points']}", False, 12),  # P161
        46: (f"【教学难点】\n{content['difficult_points']}", False, 12),  # P162
        47: (f"【教法学法】{content['teach_method']}", False, 12),  # P163
        48: ("【教学活动流程】", True, 12),  # P164
        49: (f"导入新课：{content['intro']}", True, 12),  # P165
        50: ("讲授新课：", True, 12),  # P166
        51: (f"任务1：\n{content['task1']}", False, 12),  # P167+168合并
        52: (f"任务2：\n{content['task2']}", False, 12),  # P169+170合并
        53: (f"任务3：\n{content['task3']}", False, 12),  # P171+172合并
        54: ("", False, 12),  # P173 任务4（空）
        55: (f"【板书设计】\n{content['blackboard']}", False, 12),  # P174
        56: (f"【课堂小结】{content['summary']}", False, 12),  # P175
        57: (f"【课后作业】{content['homework']}", False, 12),  # P176
        58: (f"【教学反思】{content['reflection']}", False, 12),  # P177
    }
    
    for idx, (text, bold, size) in replacements.items():
        if idx < len(new_elements) and text:
            if bold:
                set_elem_text_bold(new_elements[idx], text)
            else:
                set_elem_text(new_elements[idx], text)
    
    # 插入到标题前面
    for new_elem in new_elements:
        insert_before.addprevious(new_elem)
    
    inserted += 1
    if inserted % 5 == 0:
        print(f"  已插入{inserted}/{len(tasks)}个任务的文字描述")

print(f"\n总共插入了{inserted}个任务的文字描述")

doc.save(fp)
print(f"保存完成: {fp}")
