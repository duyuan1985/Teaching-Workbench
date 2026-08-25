"""
授课计划生成器
"""
import os, shutil
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from generate import (
    safe_set_text, clear_indent, set_indent_chars, fix_spacing_auto,
    set_alignment, set_page_break_before, set_space_before, set_line_spacing,
    write_cell, write_cell_lines, _set_font_name, _set_font_size,
    merge_v, merge_h, get_vmerge_val, CN_NUMS
)

def generate_plan(offering, tasks, sessions, units, output_dir=None):
    """生成授课计划"""
    from generate import get_template_path, get_output_path

    template_path = get_template_path("plan", offering)
    if not template_path or not os.path.exists(template_path):
        raise FileNotFoundError(f"授课计划模板不存在: {template_path}")

    fp = get_output_path("plan", offering, output_dir)
    shutil.copy2(template_path, fp)
    doc = Document(fp)
    P = doc.paragraphs

    # ============================================================
    # 1. 封面
    # ============================================================
    term = offering["term"]
    cn = offering["course_name"]
    tn = offering.get("teacher_name") or "杜媛"
    cls = offering.get("teaching_class") or ""

    # 转换学期为中文格式
    term_cn = _term_to_chinese(term)

    # 找到关键段落（按内容匹配，不按索引）
    teacher_filled = False
    for p in P:
        text = p.text.strip()
        # 学期
        if "学年" in text and "学期" in text and "商务" not in text and "课程" not in text:
            clear_indent(p)
            fix_spacing_auto(p)
            set_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
            set_space_before(p, 0)
            safe_set_text(p, term_cn, font="宋体", size=24, bold=True)
        # 课程名（教案标题）
        elif "教案" in text or ("《" in text and "课程" not in text):
            if "商务" not in text and "课程" not in text and "日程" not in text:
                clear_indent(p)
                fix_spacing_auto(p)
                set_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
                set_space_before(p, 18)  # 段前空一行
                safe_set_text(p, f"《{cn}》教案", font="宋体", size=24, bold=True)
        # 主讲教师（封面，只填一次）
        elif "主讲教师" in text and "教研室" not in text and not teacher_filled:
            clear_indent(p)
            set_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
            safe_set_text(p, f"主讲教师：{tn}        授课班级：{cls}", font="宋体", size=16, bold=True)
            teacher_filled = True
        elif "主讲教师" in text and "教研室" not in text and teacher_filled:
            # 删除重复的教师段落
            for r in p.runs:
                r.text = ""
        # 日期
        elif "年" in text and "月" in text and "日" in text and len(text) < 30 and "总学时" not in text:
            clear_indent(p)
            set_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
            safe_set_text(p, _get_date_string(offering), font="宋体", size=16, bold=True)
        # 山西林业职业技术学院（封面标题）
        elif text == "山西林业职业技术学院":
            set_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)

    # 清除封面空段落的大缩进
    for p in P[:12]:
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                flc = ind.get(qn('w:firstLineChars'))
                if flc and int(flc) > 10:
                    pPr.remove(ind)

    # ============================================================
    # 2. 第二页：日程表标题
    # ============================================================
    for p in P:
        text = p.text.strip()
        if "课程教学日程表" in text:
            set_page_break_before(p)
            clear_indent(p)
            safe_set_text(p, f"山西林业职业技术学院{term_cn}    {cn}    课程教学日程表",
                         font="宋体", size=16, bold=True)
        # 专业/班级/学时
        elif "专业" in text and "班级" in text and "总学时" in text:
            th = int(offering.get("total_hours", 60))
            lh = int(offering.get("lecture_hours", th // 2))
            eh = int(offering.get("experiment_hours", th - lh))
            # 确保理论和实践各一半
            if eh == 0 and lh == th:
                lh = th // 2
                eh = th - lh
            mj = offering.get("major", "")
            date_str = _get_date_string(offering)
            clear_indent(p)
            safe_set_text(
                p,
                f"专业：{mj}    班级：{cls}    总学时：{th}学时（其中：课堂教学：{lh}    实验实习：{eh}）    {date_str}",
                font="宋体", size=14, bold=False
            )
        # 签字栏
        elif "主讲教师" in text and "教研室主任" in text:
            safe_set_text(p, f"主讲教师：{tn}                    教研室主任：                              系部主任：",
                         font="宋体", size=10.5, bold=False)

    # ============================================================
    # 3. 表0：日程表
    # ============================================================
    if len(doc.tables) > 0:
        _fill_schedule_table(doc.tables[0], offering, tasks, sessions, units)

    # ============================================================
    # 4. 表1：补充说明
    # ============================================================
    if len(doc.tables) > 1:
        _fill_supplement(doc.tables[1], offering, tasks)

    # ============================================================
    # 5. 德育教育
    # ============================================================
    # 找到德育标题段落
    for p in P:
        if "德育教育" in p.text and "总提纲" in p.text:
            set_page_break_before(p)
            clear_indent(p)
            set_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
            safe_set_text(p, "山西林业职业技术学院德育教育教学总提纲",
                         font="宋体", size=18, bold=True)
        elif "课程：" in p.text and "任课教师" in p.text and "德育" not in p.text:
            set_space_before(p, 18)
            clear_indent(p)
            safe_set_text(p, f"{term_cn}    课程：{cn}    任课教师：{tn}",
                         font="宋体", size=12, bold=False)

    # 表2：德育教育
    if len(doc.tables) > 2:
        _fill_deyu(doc.tables[2], offering)

    doc.save(fp)
    return fp


def _get_date_string(offering):
    """从offering获取日期字符串"""
    term = offering.get("term", "")
    # 从term提取年份
    if "2023" in term:
        return "2024年2月20日"
    elif "2024-2025-1" in term:
        return "2024年9月2日"
    elif "2024-2025-2" in term:
        return "2025年2月24日"
    elif "2025-2026-1" in term:
        return "2025年9月1日"
    elif "2025-2026-2" in term:
        return "2026年2月23日"
    elif "2026-2027-1" in term:
        return "2026年9月7日"
    return "2024年2月20日"


def _term_to_chinese(term):
    """将学期代码转为中文格式，如 2023-2024-2 → 2023——2024学年第二学期"""
    import re
    m = re.match(r"(\d{4})-(\d{4})-(\d)", term)
    if m:
        y1, y2, s = m.groups()
        semester = "第一学期" if s == "1" else "第二学期"
        return f"{y1}——{y2}学年{semester}"
    return term


def _fill_schedule_table(t, offering, tasks, sessions, units):
    """填充日程表"""
    # 保留表头行，清空数据行
    while len(t.rows) > 2:
        t.rows[-1]._element.getparent().remove(t.rows[-1]._element)

    # 按学习情境分组
    chapter_map = {}
    chapter_order = []
    for task in tasks:
        ch = task["chapter"]
        if ch not in chapter_map:
            chapter_map[ch] = []
            chapter_order.append(ch)
        chapter_map[ch].append(task)

    cn_num_idx = 0
    row_cursor = 2  # 从第3行开始（0-based=2）
    global_seq = 0  # 全局序号

    for ch_idx, chapter in enumerate(chapter_order):
        ch_tasks = chapter_map[chapter]
        cn_num_idx += 1

        for ti, task in enumerate(ch_tasks):
            global_seq += 1
            sess = _find_sess(task, sessions)
            week = task.get("week_no", "")
            date = task.get("lesson_date", "") or (sess.get("lesson_date", "") if sess else "")
            classroom = (sess.get("classroom", "801教室") if sess else "801教室")
            exp_room = classroom  # 实训地点同授课地点

            title = task["title"].split("：", 1)[1] if "：" in task["title"] else task["title"]

            # 获取子任务列表
            sub_tasks = _get_sub_tasks(task, units)
            if not sub_tasks:
                sub_tasks = [title]

            # 理论内容行数 = 1(情境名) + 子任务数
            theory_lines = 1 + len(sub_tasks)
            # 实践内容1行
            total_rows = theory_lines + 1

            # 记录起始行
            start_row = row_cursor

            # 添加行
            for ri in range(total_rows):
                if row_cursor >= len(t.rows):
                    t.add_row()
                row_cursor += 1

            # 填充内容
            # 序号（全局序号，按日期排序）
            write_cell(t.rows[start_row].cells[0], str(global_seq), "宋体", 12, False, WD_ALIGN_PARAGRAPH.CENTER)
            # 周次
            write_cell(t.rows[start_row].cells[1], str(week), "宋体", 12, False, WD_ALIGN_PARAGRAPH.CENTER)
            # 日期
            write_cell(t.rows[start_row].cells[2], date, "宋体", 12, False, WD_ALIGN_PARAGRAPH.CENTER)

            # 理论内容：每个子任务占独立一行
            # 第0行：情境名（加章节号，加粗）
            ch_title = f"第{CN_NUMS[min(cn_num_idx, 9) - 1]}章 {chapter}" if ti == 0 else chapter
            write_cell(t.rows[start_row].cells[3], ch_title, "宋体", 12, ti == 0)
            # 第1~N行：每个子任务占一行
            for si, st in enumerate(sub_tasks):
                write_cell(t.rows[start_row + 1 + si].cells[3], st, "宋体", 12, False)

            # 实践内容
            practice_text = _get_practice_content(task, chapter)
            write_cell(t.rows[start_row + theory_lines].cells[6], practice_text, "宋体", 12, False)

            # 理论学时
            write_cell(t.rows[start_row].cells[4], "1", "宋体", 12, False, WD_ALIGN_PARAGRAPH.CENTER)
            # 理论授课地点
            write_cell(t.rows[start_row].cells[5], classroom, "宋体", 12, False, WD_ALIGN_PARAGRAPH.CENTER)
            # 实践学时
            write_cell(t.rows[start_row + theory_lines].cells[7], "1", "宋体", 12, False, WD_ALIGN_PARAGRAPH.CENTER)
            # 实验地点（同授课地点）
            write_cell(t.rows[start_row + theory_lines].cells[8], exp_room, "宋体", 12, False, WD_ALIGN_PARAGRAPH.CENTER)
            # 实训教师
            write_cell(t.rows[start_row].cells[9], offering.get("teacher_name") or "杜媛", "宋体", 12, False, WD_ALIGN_PARAGRAPH.CENTER)

            # 合并公共列：序号、周次、日期、实训教师 跨所有行
            for col in [0, 1, 2, 9]:
                merge_v(t, start_row, start_row + total_rows - 1, col)
            # 理论学时、理论授课地点 跨理论行
            for col in [4, 5]:
                merge_v(t, start_row, start_row + theory_lines - 1, col)

            # 表格内容宋体小四
            for ri in range(start_row, start_row + total_rows):
                for ci in range(len(t.rows[ri].cells)):
                    for p in t.rows[ri].cells[ci].paragraphs:
                        clear_indent(p)


def _find_sess(task, sessions):
    """查找排课"""
    for s in sessions:
        if s.get("week_no") == task.get("week_no") and s.get("lesson_date") == task.get("lesson_date"):
            return s
    ws = [s for s in sessions if s.get("week_no") == task.get("week_no")]
    return ws[0] if ws else {}


def _get_sub_tasks(task, units):
    """从任务标题中提取子任务列表"""
    title = task["title"].split("：", 1)[1] if "：" in task["title"] else task["title"]
    # 任务标题格式："章节名：子任务1、子任务2"
    # 用顿号分隔子任务
    parts = [p.strip() for p in title.split("、") if p.strip()]
    # 如果有"任务实施"前缀，去掉
    parts = [p.replace("任务实施：", "") if p.startswith("任务实施：") else p for p in parts]
    return parts if parts else [title]


def _get_practice_content(task, chapter):
    """生成实践内容摘要"""
    title = task["title"].split("：", 1)[1] if "：" in task["title"] else task["title"]
    # 根据章节生成实践内容
    practice_map = {
        "初识": "分析电商数据分析典型案例",
        "Excel": "完成Excel数据透视表与图表分析实操",
        "Numpy": "NumPy数组运算与数据处理实操",
        "Pandas": "Pandas数据清洗与分析实操",
        "SciPy": "SciPy统计检验与假设分析实操",
        "Sklearn": "机器学习模型训练与评估实操",
        "Seaborn": "Seaborn数据可视化图表制作实操",
        "综合": "综合数据分析项目实战",
    }
    for key, val in practice_map.items():
        if key in chapter:
            return val
    return f"{title}实训操作"


def _fill_supplement(t, offering, tasks):
    """填充补充说明表"""
    cn = offering["course_name"]
    cc = offering.get("course_code", "")
    th = int(offering.get("total_hours", 60))
    cr = int(offering.get("credits", 3))
    mj = offering.get("major", "")
    tb = offering.get("textbook_version", "")

    sections = [
        ("一、计划依据",
         f"　　本课程授课计划依据《{cn}》课程标准（{cc}）及{offering.get('teaching_class','')}级{mj}专业人才培养方案制定，总学时{th}学时（理论{th//2}学时+实践{th//2}学时），学分{cr}分，开设于{offering.get('term','')}学期。"),
        ("二、教学目的",
         f"　　通过本课程学习，使学生掌握数据分析的基本概念、方法理论和常用工具，能够运用Excel、Python（NumPy、Pandas、SciPy、scikit-learn、Seaborn）进行电商数据采集、清洗、分析和可视化，培养数据化思维和实操能力，达到电商数据分析师岗位的核心能力要求。"),
        ("三、选用教材",
         f"　　《{tb}》，天津大学出版社，配套PPT课件、源代码和实训资源。"),
        ("四、教学方式方法",
         f"　　采用理实一体化教学，项目引领、任务驱动。依托超星学习通平台开展线上线下混合式教学，课前推送微课预习，课中讲授+实操+互动测验，课后作业+学情数据分析。指导学生运用AI大模型工具辅助代码调试与分析思路设计。"),
        ("五、教学安排",
         f"　　本课程共{len(tasks)}个教学任务，分{len(set(t['chapter'] for t in tasks))}个学习情境。每周2次课，每次2学时（理论1学时+实践1学时）。调课说明：根据校历安排，法定节假日课程顺延或调整至周末补课。"),
        ("六、本学期授课更新内容（新标准、新技术、新工艺、新方法）",
         f"　　1. 新标准：融入《中华人民共和国数据安全法》《个人信息保护法》对数据处理合规性要求；对接1+X数据采集与处理职业技能等级证书标准和全国职业院校技能大赛大数据分析与应用赛项标准。2. 新技术：引入大语言模型（LLM）辅助数据分析、AIGC智能图表生成、数据中台与实时数据分析等新技术。3. 新方法：采用项目引领、任务驱动教学法，依托学习通平台实施混合式教学，基于学情数据精准分析学习效果并动态调整教学策略。"),
        ("七、考核方式与成绩构成",
         f"　　采用过程性考核（40%）+终结性考核（60%）的方式。过程性考核（40%）：签到（10%）+课堂表现（10%）+作业（20%），每单元签到1.25%、课堂表现1.25%、作业2.5%，共8个学习情境。终结性考核（60%）：综合作品——基于真实电商数据集完成完整的数据分析项目，包含数据采集、清洗、分析、可视化和报告撰写。"),
    ]

    cell = t.rows[0].cells[0]
    # 清空
    while len(cell.paragraphs) > 1:
        cell.paragraphs[-1]._element.getparent().remove(cell.paragraphs[-1]._element)
    p0 = cell.paragraphs[0]
    for r in p0.runs:
        r.text = ""

    first = True
    for title, content in sections:
        if first:
            p = p0
            first = False
        else:
            p = cell.add_paragraph()

        r = p.add_run(title)
        _set_font_name(r, "仿宋_GB2312")
        _set_font_size(r, 12)
        r.font.bold = True
        clear_indent(p)
        set_line_spacing(p, 1.5)

        # 内容段
        p2 = cell.add_paragraph()
        r2 = p2.add_run(content)
        _set_font_name(r2, "仿宋_GB2312")
        _set_font_size(r2, 12)
        r2.font.bold = False
        clear_indent(p2)
        set_line_spacing(p2, 1.5)


def _fill_deyu(t, offering):
    """填充德育教育表"""
    cn = offering["course_name"]
    tn = offering.get("teacher_name", "杜媛")
    term = offering["term"]

    sections = [
        ("一、马列主义方面",
         "　　运用马克思主义唯物辩证法分析数据现象，理解量变与质变的关系，用发展的眼光看待数据驱动的商业变革。引导学生运用辩证唯物主义世界观认识数据分析中整体与局部、现象与本质的关系。"),
        ("二、理想信念方面",
         "　　树立科技报国理想，关注数字技术服务乡村振兴的实践价值，增强通过数据分析服务经济社会发展的使命感。结合农村电子商务数据分析案例，培养学生的家国情怀和社会责任感。"),
        ("三、核心价值观方面",
         "　　践行诚信、公正、法治，坚持真实分析和规范验证，不篡改数据、不误导决策。以数据诚信践行社会主义核心价值观，培养数据伦理意识。"),
        ("四、传统文化方面",
         "　　汲取\u201c实事求是\u201d传统文化精髓，在数据分析中做到求真务实、严谨细致，传承工匠精神。将中华优秀传统文化中\u201c格物致知\u201d的治学态度融入数据分析过程。"),
        ("五、职业道德方面",
         "　　严守数据保密协议，不泄露商业机密和用户隐私，遵守《个人信息保护法》《数据安全法》，恪守数据分析师职业操守。培养数据安全意识和职业规范意识。"),
        ("六、个人素养方面",
         "　　培养数据驱动的批判性思维和问题解决能力，养成规范操作、主动学习、持续改进的良好习惯。提升团队协作、沟通表达和抗压能力。"),
    ]

    cell = t.rows[0].cells[0]
    # 清空
    while len(cell.paragraphs) > 1:
        cell.paragraphs[-1]._element.getparent().remove(cell.paragraphs[-1]._element)
    p0 = cell.paragraphs[0]
    for r in p0.runs:
        r.text = ""

    first = True
    for title, content in sections:
        if first:
            p = p0
            first = False
        else:
            p = cell.add_paragraph()

        r = p.add_run(title)
        _set_font_name(r, "仿宋_GB2312")
        _set_font_size(r, 12)
        r.font.bold = True
        clear_indent(p)  # 标题顶格
        set_line_spacing(p, 1.5)

        # 内容段
        p2 = cell.add_paragraph()
        r2 = p2.add_run(content)
        _set_font_name(r2, "仿宋_GB2312")
        _set_font_size(r2, 12)
        r2.font.bold = False
        clear_indent(p2)  # 清除模板firstLineChars残留
        set_line_spacing(p2, 1.5)
