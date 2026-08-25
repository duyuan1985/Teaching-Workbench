"""
详细对比三个文档的关键信息
1. 课程基本信息（学时、学分、学期、先修课、后续课）
2. 课程目标（认知/知识、能力、思政、素质）
3. 教学方法
4. 考核方案
5. 教材信息
6. 学习情境划分
7. 师资/实训条件
"""
from docx import Document
from docx.oxml.ns import qn

fp_design = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
fp_standard = r'生成结果\精修版\2023-2024-2《商务数据分析》课程标准 杜媛.docx'
fp_plan = r'生成结果\精修版\2023-2024-2《商务数据分析》授课计划 杜媛.docx'

def get_all_text(fp):
    """获取文档全部文本（段落+表格）"""
    doc = Document(fp)
    lines = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            lines.append(txt)
    for t in doc.tables:
        for ri in range(len(t.rows)):
            for ci in range(len(t.columns)):
                txt = t.cell(ri, ci).text.strip()
                if txt:
                    lines.append(txt)
    return '\n'.join(lines)

# ============================================================
# 1. 课程基本信息对比
# ============================================================
print('='*80)
print('1. 课程基本信息对比')
print('='*80)

# 教学设计 - Table 0
doc_d = Document(fp_design)
t0_d = doc_d.tables[0]
print('\n--- 教学设计 Table 0 ---')
for ri in range(len(t0_d.rows)):
    texts = [c.text.strip()[:40] for c in t0_d.rows[ri].cells]
    print(f'  R{ri}: {texts}')

# 课程标准 - paragraphs
doc_s = Document(fp_standard)
print('\n--- 课程标准 基本信息 ---')
for p in doc_s.paragraphs:
    txt = p.text.strip()
    for kw in ['课程名称', '课程类型', '学时', '学分', '开设学期', '先导', '后续']:
        if kw in txt:
            print(f'  {txt[:120]}')
            break

# 授课计划
doc_p = Document(fp_plan)
print('\n--- 授课计划 基本信息 ---')
for p in doc_p.paragraphs:
    txt = p.text.strip()
    for kw in ['专业', '班级', '总学时', '主讲教师', '学期']:
        if kw in txt:
            print(f'  {txt[:120]}')
            break

# ============================================================
# 2. 课程目标对比
# ============================================================
print('\n' + '='*80)
print('2. 课程目标对比')
print('='*80)

# 教学设计 目标
print('\n--- 教学设计 课程目标 ---')
in_goal = False
for p in doc_d.paragraphs:
    txt = p.text.strip()
    if '认知目标' in txt or '知识目标' in txt:
        in_goal = True
        print(f'  {txt[:120]}')
    elif '能力目标' in txt:
        in_goal = True
        print(f'  {txt[:120]}')
    elif '思政目标' in txt:
        in_goal = True
        print(f'  {txt[:120]}')
    elif '素质目标' in txt:
        in_goal = True
        print(f'  {txt[:120]}')
    elif in_goal and txt.startswith('①') or txt.startswith('②') or txt.startswith('③') or txt.startswith('④'):
        print(f'  {txt[:120]}')
    elif in_goal and ('（二）' in txt or '（三）' in txt or '（四）' in txt):
        in_goal = False

# 课程标准 目标
print('\n--- 课程标准 课程目标 ---')
# Find Table 1 in course standard (should be objectives)
for ti, t in enumerate(doc_s.tables):
    first_cell = t.rows[0].cells[0].text.strip() if t.rows else ''
    if '目标' in first_cell or '描述' in first_cell or '知识' in first_cell:
        print(f'  Table {ti}: first_cell="{first_cell}"')
        for ri in range(len(t.rows)):
            texts = [c.text.strip()[:50] for c in t.rows[ri].cells]
            print(f'    R{ri}: {texts}')

# ============================================================
# 3. 教学方法对比
# ============================================================
print('\n' + '='*80)
print('3. 教学方法对比')
print('='*80)

print('\n--- 教学设计 教学方法 ---')
for p in doc_d.paragraphs:
    txt = p.text.strip()
    if '教学法' in txt or '教学模式' in txt or '教学方法' in txt:
        print(f'  {txt[:150]}')

print('\n--- 课程标准 教学方法 ---')
for p in doc_s.paragraphs:
    txt = p.text.strip()
    if '教学法' in txt or '教学方法' in txt:
        print(f'  {txt[:150]}')

# ============================================================
# 4. 考核方案对比
# ============================================================
print('\n' + '='*80)
print('4. 考核方案对比')
print('='*80)

print('\n--- 教学设计 考核方案 ---')
for p in doc_d.paragraphs:
    txt = p.text.strip()
    if '考核' in txt or '过程性' in txt or '终结性' in txt:
        print(f'  {txt[:150]}')

print('\n--- 教学设计 考核表(Table 4) ---')
t4_d = doc_d.tables[4]
for ri in range(len(t4_d.rows)):
    texts = [c.text.strip()[:50] for c in t4_d.rows[ri].cells]
    print(f'  R{ri}: {texts}')

print('\n--- 课程标准 考核方案 ---')
for p in doc_s.paragraphs:
    txt = p.text.strip()
    if '考核' in txt or '过程性' in txt or '终结性' in txt:
        print(f'  {txt[:150]}')

print('\n--- 课程标准 考核表 ---')
for ti, t in enumerate(doc_s.tables):
    first_cell = t.rows[0].cells[0].text.strip() if t.rows else ''
    if '学习情境' in first_cell and '分值' in first_cell:
        print(f'  Table {ti}:')
        for ri in range(min(5, len(t.rows))):
            texts = [c.text.strip()[:30] for c in t.rows[ri].cells]
            print(f'    R{ri}: {texts}')
        print(f'    ... ({len(t.rows)} rows total)')

# ============================================================
# 5. 教材信息对比
# ============================================================
print('\n' + '='*80)
print('5. 教材信息对比')
print('='*80)

print('\n--- 教学设计 教材 ---')
for p in doc_d.paragraphs:
    txt = p.text.strip()
    if '教材' in txt or '天津大学' in txt:
        print(f'  {txt[:200]}')

print('\n--- 课程标准 教材 ---')
for p in doc_s.paragraphs:
    txt = p.text.strip()
    if '教材' in txt or '天津大学' in txt:
        print(f'  {txt[:200]}')

print('\n--- 授课计划 教材 ---')
for p in doc_p.paragraphs:
    txt = p.text.strip()
    if '教材' in txt or '天津大学' in txt:
        print(f'  {txt[:200]}')

# ============================================================
# 6. 学习情境划分对比
# ============================================================
print('\n' + '='*80)
print('6. 学习情境划分对比')
print('='*80)

print('\n--- 教学设计 学习情境(Table 1) ---')
t1_d = doc_d.tables[1]
for ri in range(len(t1_d.rows)):
    texts = [c.text.strip()[:50] for c in t1_d.rows[ri].cells]
    print(f'  R{ri}: {texts}')

print('\n--- 课程标准 学习情境 ---')
for ti, t in enumerate(doc_s.tables):
    first_cell = t.rows[0].cells[0].text.strip() if t.rows else ''
    if '学习情境' in first_cell and ('序' in first_cell or '名' in first_cell or '情' in first_cell):
        print(f'  Table {ti}: first="{first_cell[:30]}"')
        for ri in range(min(10, len(t.rows))):
            texts = [c.text.strip()[:40] for c in t.rows[ri].cells]
            print(f'    R{ri}: {texts}')
        if len(t.rows) > 10:
            print(f'    ... ({len(t.rows)} rows total)')
