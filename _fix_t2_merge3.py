"""
修复表2合并单元格 - 彻底清除旧文字
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)
t2 = doc.tables[2]

merge_groups = [
    (1, 3, '1', '初识数据分析'),
    (4, 6, '2', 'Excel数据分析工具'),
    (7, 9, '3', 'Numpy数学运算库'),
    (10, 12, '4', 'Pandas数据分析库'),
    (13, 15, '5', 'SciPy科学计算库'),
    (16, 18, '6', 'Sklearn数据统计基础'),
    (19, 21, '7', 'Sklearn数据统计进阶'),
    (22, 24, '8', 'Seaborn可视化分析库'),
    (25, 26, '9', '综合评价与课程总结'),
]

for start, end, seq_val, name_val in merge_groups:
    for ci in [0, 1]:
        val = seq_val if ci == 0 else name_val
        
        # 第一行：vMerge=restart
        tr = t2.rows[start]._tr
        tcs = tr.findall(qn('w:tc'))
        if ci >= len(tcs):
            continue
        tc = tcs[ci]
        
        # 设置tcPr
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        for vm in tcPr.findall(qn('w:vMerge')):
            tcPr.remove(vm)
        vm_new = OxmlElement('w:vMerge')
        vm_new.set(qn('w:val'), 'restart')
        tcPr.append(vm_new)
        
        # 删除所有旧段落
        for p in tc.findall(qn('w:p')):
            tc.remove(p)
        
        # 创建新段落
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), '仿宋')
        rFonts.set(qn('w:eastAsia'), '仿宋')
        rFonts.set(qn('w:hAnsi'), '仿宋')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '21')
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = val
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)
        tc.append(p)
    
    # 后续行：vMerge=continue，清空所有段落
    for ri in range(start + 1, end + 1):
        tr = t2.rows[ri]._tr
        tcs = tr.findall(qn('w:tc'))
        for ci in [0, 1]:
            if ci >= len(tcs):
                continue
            tc = tcs[ci]
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            for vm in tcPr.findall(qn('w:vMerge')):
                tcPr.remove(vm)
            vm_new = OxmlElement('w:vMerge')
            tcPr.append(vm_new)
            
            # 删除所有旧段落，只保留一个空段落
            ps = tc.findall(qn('w:p'))
            for p in ps[1:]:  # 保留第一个段落
                tc.remove(p)
            # 清空第一个段落的run
            if ps:
                for r in ps[0].findall(qn('w:r')):
                    ps[0].remove(r)

print("修复完成")

doc.save(fp)

# 验证
doc2 = Document(fp)
t2 = doc2.tables[2]
for ri in range(min(10, len(t2.rows))):
    cells = t2.rows[ri].cells
    texts = [c.text.strip()[:20] for c in cells[:3]]
    print(f"  R{ri}: {texts}")
