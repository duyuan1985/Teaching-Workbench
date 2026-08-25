"""
AI增强模板分析器 - 基于大语言模型的深度模板解读

使用智谱GLM等大模型，深度理解模板中的文字说明、表格结构和填写要求，
生成更精准的槽位定义和内容生成规则。

工作方式：
1. 先由基础分析器提取结构（段落、表格、合并关系）
2. 将结构信息分批送给AI进行语义理解
3. AI输出结构化的槽位定义JSON
4. 与基础分析结果合并，存入数据库
"""

import json
import re
from pathlib import Path
from docx import Document

import store
import template_analyzer
from ai.ai_router import ask_result


# 系统提示词：告诉AI它的角色和输出格式
SYSTEM_PROMPT = """你是一个专业的教学文档模板分析专家。你的任务是分析职业教育课程文档模板，识别模板中需要填写的区域（槽位），并给出每个槽位的精确定义。

请严格按照以下要求工作：
1. 仔细阅读模板内容，理解每个填写区域的含义
2. 识别哪些是"标签-值"格式的简单字段，哪些是需要大段撰写的分析类内容
3. 注意模板中的"注"、"说明"、"要求"等文字，这些是填写指导
4. 输出严格的JSON格式，不要有任何额外文字

JSON格式说明：
{
  "slots": [
    {
      "field_name": "字段的标准名称（如：课程名称、课程性质、知识目标等）",
      "content_kind": "内容类型，必须是以下之一：事实填写/分析撰写/结构生成/按任务填写/按班事实填写",
      "fill_requirement": "根据模板说明提炼的填写要求",
      "suggested_source": "建议的数据来源（如：课程实例、课程任务、教材内容、课程语义模型等）",
      "location_hint": "位置描述，帮助定位这个字段在模板的哪里"
    }
  ]
}

内容类型说明：
- 事实填写：简单的事实信息，如课程名称、学时、教师姓名等
- 分析撰写：需要分析、总结、撰写的段落内容，如课程性质、课程设计思路等
- 结构生成：需要按一定结构生成的表格或列表，如课程内容与学时分配
- 按任务填写：每个教学任务都有一份的内容，如单元教学设计
- 按班事实填写：按班级不同而变化的事实信息，如授课班级、授课日期

请只输出JSON，不要输出任何其他文字或解释。"""


def _table_to_text(table, max_rows=10):
    """将表格转换为文本描述，送给AI分析"""
    lines = []
    rows = list(table.rows)
    
    # 处理合并单元格：用id去重
    for row_idx, row in enumerate(rows[:max_rows]):
        cells = []
        seen = set()
        for cell in row.cells:
            marker = id(cell._tc)
            if marker in seen:
                continue
            seen.add(marker)
            text = template_analyzer._cell_text(cell)
            # 截断过长的内容
            if len(text) > 100:
                text = text[:100] + "..."
            cells.append(text)
        lines.append(f"行{row_idx}: {' | '.join(cells)}")
    
    if len(rows) > max_rows:
        lines.append(f"... 还有 {len(rows) - max_rows} 行")
    
    return "\n".join(lines)


def _paragraphs_to_text(document, start_idx=0, count=30):
    """提取段落文本"""
    lines = []
    for i in range(start_idx, min(start_idx + count, len(document.paragraphs))):
        text = document.paragraphs[i].text.strip()
        if text:
            if len(text) > 200:
                text = text[:200] + "..."
            lines.append(f"段落{i}: {text}")
    return "\n".join(lines)


def analyze_template_with_ai(template_file_id, force_online=True):
    """
    使用AI深度分析模板。
    
    Args:
        template_file_id: 模板文件ID
        force_online: 是否强制使用在线AI（智谱）
    
    Returns:
        dict: AI分析结果，包含 slots 列表
    """
    item = store.rows("SELECT * FROM template_files WHERE id=?", (template_file_id,))
    if not item:
        raise ValueError("模板文件不存在。")
    
    path = Path(item[0]["template_path"])
    document_type = item[0]["document_type"]
    document = Document(path)
    
    print(f"\n[AI分析] 开始分析模板: {document_type}")
    print(f"  段落数: {len(document.paragraphs)}")
    print(f"  表格数: {len(document.tables)}")
    
    all_slots = []
    
    # === 第一步：分析段落 ===
    print(f"\n  [1/3] 分析段落内容...")
    para_text = _paragraphs_to_text(document, 0, min(50, len(document.paragraphs)))
    para_prompt = f"""这是一个职业教育{document_type}模板的段落内容。请识别其中需要填写的区域（槽位）。

注意：
- 带有"（注"、"注："、"说明："、"（要求"等开头的段落是填写说明，不是需要填写的区域
- 带有"********"、"……"、"XXX"、"20XX"等占位符的地方通常是需要填写的
- 标题性段落（如"一、课程性质"）通常意味着下方需要填写内容
- 标签格式（如"课程名称："）后面的值需要填写

段落内容：
{para_text}

请输出需要填写的槽位的JSON。"""

    result = ask_result(para_prompt, SYSTEM_PROMPT, force_online=force_online, show_details=True)
    if result["success"]:
        try:
            # 尝试提取JSON
            content = result["content"].strip()
            # 有时AI会用 ```json 包裹
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0].strip()
            elif "```" in content:
                content = content.split("```")[1].split("```")[0].strip()
            
            data = json.loads(content)
            if not isinstance(data, dict):
                data = {}
            slots = data.get("slots", [])
            print(f"    AI识别出 {len(slots)} 个段落级槽位")
            for s in slots:
                s["_source"] = "paragraph"
            all_slots.extend(slots)
        except (json.JSONDecodeError, KeyError, AttributeError) as e:
            print(f"    ⚠️ AI返回格式解析失败: {e}")
            print(f"    原始内容前200字: {result['content'][:200]}")
    else:
        print(f"    ⚠️ AI调用失败: {result.get('error')}")
    
    # === 第二步：逐个分析表格 ===
    print(f"\n  [2/3] 分析表格内容...")
    
    # 先获取基础分析的表格角色
    contract = template_analyzer._template_contract(document)
    
    for table_idx, table in enumerate(document.tables):
        role = contract["tables"][table_idx]["role"]
        rows_count = len(table.rows)
        cols_count = len(table.columns)
        
        # 跳过过大的表格（可能是示例内容）
        if rows_count > 30:
            print(f"    表{table_idx} ({role}): 行数过多，跳过AI分析")
            continue
        
        table_text = _table_to_text(table, max_rows=min(15, rows_count))
        
        table_prompt = f"""这是一个职业教育{document_type}模板中的第{table_idx}个表格。
表格角色: {role}
表格大小: {rows_count}行 x {cols_count}列

请分析这个表格中需要填写的单元格（槽位）。

注意：
- 表格第一行通常是表头（列名）
- 左列如果是短文本+冒号，通常是标签，右侧是值单元格
- 带有"……"、"XXX"等占位符的单元格是需要填写的
- 带有"注"、"说明"的单元格是填写要求，不是槽位
- 合并单元格需要特别注意其实际含义

表格内容：
{table_text}

请输出需要填写的槽位的JSON。对于表格中的槽位，location_hint请描述清楚在第几行第几列。"""

        result = ask_result(table_prompt, SYSTEM_PROMPT, force_online=force_online, show_details=False)
        if result["success"]:
            try:
                content = result["content"].strip()
                if "```json" in content:
                    content = content.split("```json")[1].split("```")[0].strip()
                elif "```" in content:
                    content = content.split("```")[1].split("```")[0].strip()
                
                data = json.loads(content)
                if not isinstance(data, dict):
                    data = {}
                slots = data.get("slots", [])
                print(f"    表{table_idx} ({role}): AI识别出 {len(slots)} 个槽位")
                for s in slots:
                    s["_source"] = f"table:{table_idx}"
                    s["_table_role"] = role
                all_slots.extend(slots)
            except (json.JSONDecodeError, KeyError, AttributeError) as e:
                print(f"    表{table_idx}: ⚠️ 解析失败: {e}")
        else:
            print(f"    表{table_idx}: ⚠️ AI调用失败")
    
    # === 第三步：汇总并保存 ===
    print(f"\n  [3/3] 汇总分析结果...")
    print(f"    AI共识别出 {len(all_slots)} 个槽位")
    
    # 保存到数据库（ai_analysis 字段）
    analysis_json = json.dumps({
        "ai_slots": all_slots,
        "ai_source": result.get("source", ""),
        "ai_model": result.get("model", ""),
    }, ensure_ascii=False)
    
    with store.connect() as db:
        # 检查是否已有记录
        existing = db.execute(
            "SELECT template_file_id FROM template_analyses WHERE template_file_id=?",
            (template_file_id,)
        ).fetchone()
        
        if existing:
            db.execute(
                "UPDATE template_analyses SET ai_analysis_json=?, ai_analysis_status='已完成' WHERE template_file_id=?",
                (analysis_json, template_file_id)
            )
        else:
            db.execute(
                "INSERT INTO template_analyses (template_file_id, ai_analysis_json, ai_analysis_status) VALUES (?, ?, '已完成')",
                (template_file_id, analysis_json)
            )
        db.commit()
    
    print(f"    ✓ 结果已保存")
    return all_slots


def compare_analysis(template_file_id):
    """
    对比基础分析和AI分析的结果。
    
    Returns:
        dict: 对比结果
    """
    # 获取基础分析结果
    base_slots = store.rows(
        "SELECT field_name, content_kind, locator, section_title FROM template_slots WHERE template_file_id=? ORDER BY field_name",
        (template_file_id,)
    )
    
    # 获取AI分析结果
    ai_row = store.rows(
        "SELECT ai_analysis_json, ai_analysis_status FROM template_analyses WHERE template_file_id=?",
        (template_file_id,)
    )
    
    result = {
        "base_count": len(base_slots),
        "ai_count": 0,
        "base_only": [],
        "ai_only": [],
        "both": [],
    }
    
    if not ai_row or not ai_row[0]["ai_analysis_json"]:
        result["ai_status"] = "未分析"
        return result
    
    result["ai_status"] = ai_row[0]["ai_analysis_status"]
    
    try:
        ai_data = json.loads(ai_row[0]["ai_analysis_json"])
        if not isinstance(ai_data, dict):
            ai_data = {}
        ai_slots = ai_data.get("ai_slots", [])
        result["ai_count"] = len(ai_slots)

        # 标准化字段名用于对比
        base_fields = {s["field_name"] for s in base_slots}
        ai_fields = {s.get("field_name", "") for s in ai_slots}

        result["base_only"] = sorted(base_fields - ai_fields)
        result["ai_only"] = sorted(ai_fields - base_fields)
        result["both"] = sorted(base_fields & ai_fields)

    except (json.JSONDecodeError, KeyError, AttributeError):
        pass
    
    return result


if __name__ == "__main__":
    # 测试：分析课程标准模板
    offering_id = 27
    templates = store.rows(
        "SELECT id, document_type FROM template_files WHERE offering_id=? ORDER BY id",
        (offering_id,)
    )
    
    for t in templates:
        print(f"\n{'='*60}")
        print(f"分析模板: {t['document_type']} (ID: {t['id']})")
        print('='*60)
        
        try:
            slots = analyze_template_with_ai(t["id"])
            
            print(f"\n--- AI分析的槽位列表 ---")
            for i, s in enumerate(slots):
                print(f"{i+1}. {s.get('field_name', '?')} ({s.get('content_kind', '?')})")
                print(f"   来源: {s.get('_source', '?')}")
                if s.get('fill_requirement'):
                    req = s['fill_requirement'][:80]
                    print(f"   要求: {req}...")
        except Exception as e:
            print(f"分析失败: {e}")
            import traceback
            traceback.print_exc()
