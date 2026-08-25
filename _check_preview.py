"""检查预览文档：字体/格式/内容要点逐项核对用户反馈。"""
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

path = sorted(Path("data/tmp_verify").glob("课程标准_契约预览_商务数据分析_*.docx"))[-1]
doc = Document(path)


def run_font(run):
    rpr = run._element.find(qn("w:rPr"))
    ea = ""
    if rpr is not None:
        rf = rpr.find(qn("w:rFonts"))
        if rf is not None:
            ea = rf.get(qn("w:eastAsia")) or ""
    return f"{ea}/{run.font.size.pt if run.font.size else '?'}pt"


print("=== 段落（前22段含字体） ===")
for i, p in enumerate(doc.paragraphs[:22]):
    if not p.text.strip():
        continue
    fonts = {run_font(r) for r in p.runs if r.text.strip()}
    print(f"[{i}] {fonts} {p.text[:60]}")

print("\n=== 其余段落文本 ===")
for i, p in enumerate(doc.paragraphs):
    if i < 22 or not p.text.strip():
        continue
    print(f"[{i}] {p.text[:80]}")

print("\n=== 表格 ===")
for ti, t in enumerate(doc.tables):
    print(f"--- 表{ti} ({len(t.rows)}行) ---")
    for ri, row in enumerate(t.rows[:20]):
        seen, texts = set(), []
        for c in row.cells:
            if id(c._tc) in seen:
                continue
            seen.add(id(c._tc))
            texts.append(c.text.replace("\n", "|")[:42])
        print(f"  r{ri}: {' ‖ '.join(texts)}")
