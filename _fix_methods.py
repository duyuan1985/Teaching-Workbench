"""
更新教学设计的教学方法，与课程标准一致
课程标准：项目教学法、任务驱动法、案例教学法 + 数字技术应用
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from copy import deepcopy

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)

# 找到教学方法段落
paras = doc.paragraphs
start_idx = None
end_idx = None
for pi, p in enumerate(paras):
    txt = p.text.strip()
    if '（四）课程教学模式和教学方法设计' in txt:
        start_idx = pi
    if start_idx and '（五）' in txt:
        end_idx = pi
        break

print(f"教学方法段落: P{start_idx}-P{end_idx-1}")
for pi in range(start_idx, end_idx):
    print(f"  P{pi}: {paras[pi].text.strip()[:150]}")

# 新内容
new_lines = [
    "（四）课程教学模式和教学方法设计",
    "教学模式：理实一体化教学模式，理论讲解与实操训练交替进行，每个学习情境安排理论1学时和实践1学时。",
    "教学方法：本课程选用项目教学法、任务驱动法、案例教学法。",
    "项目教学法是以企业真实项目为载体，将知识学习和技能训练融入项目实施全过程的教学方法。",
    "任务驱动法以具体任务为线索推进教学，激发学生学习动力。",
    "案例教学法通过分析电商行业真实数据案例，引导学生理解方法的应用场景和操作规范。",
    "在数字技术应用方面：本课程依托超星学习通平台实施线上线下混合式教学，课前推送微课视频与预习测验，课中运用随堂投票、抢答、限时测验等功能实时采集学情，课后基于平台学习行为数据精准分析学习效果并动态调整教学策略。同时，指导学生规范运用AI大模型工具（如通义千问、Kimi等）辅助数据探索思路设计、Python代码调试与图表美化，培养学生在人工智能时代会用工具、善用工具、慎用工具的数字素养与信息伦理意识。",
]

# 获取格式模板
template_p = paras[start_idx + 1]
font_name = '仿宋_GB2312'
font_size = Pt(10.5)
if template_p.runs:
    r = template_p.runs[0]
    font_name = r.font.name or '仿宋_GB2312'
    if r.font.size:
        font_size = r.font.size
print(f"格式: font={font_name}, size={font_size}")

# 替换段落内容
body = doc.element.body
all_p_elems = body.findall(qn('w:p'))

# 收集旧段落元素
old_elems = []
for pi in range(start_idx, end_idx):
    old_elems.append(paras[pi]._element)

# 替换已有段落
for i in range(min(len(old_elems), len(new_lines))):
    p_elem = old_elems[i]
    # 清除run
    for r in p_elem.findall(qn('w:r')):
        p_elem.remove(r)
    # 添加新run
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)
    sz_val = str(int(font_size.pt * 2))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), sz_val)
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), sz_val)
    rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = new_lines[i]
    r.append(t)
    p_elem.append(r)

# 如果新内容多于旧段落，添加新段落
if len(new_lines) > len(old_elems):
    template_elem = old_elems[-1]
    for i in range(len(old_elems), len(new_lines)):
        new_p = deepcopy(template_elem)
        for r in new_p.findall(qn('w:r')):
            new_p.remove(r)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), sz_val)
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), sz_val)
        rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = new_lines[i]
        r.append(t)
        new_p.append(r)
        template_elem.addnext(new_p)
        template_elem = new_p

# 如果旧段落多于新内容，删除多余
elif len(old_elems) > len(new_lines):
    for i in range(len(new_lines), len(old_elems)):
        old_elems[i].getparent().remove(old_elems[i])

doc.save(fp)
print("\n保存完成")

# 验证
doc2 = Document(fp)
in_section = False
for pi, p in enumerate(doc2.paragraphs):
    txt = p.text.strip()
    if '（四）课程教学模式和教学方法设计' in txt:
        in_section = True
    if in_section:
        print(f"P{pi}: {txt[:200]}")
    if in_section and '（五）' in txt:
        break
