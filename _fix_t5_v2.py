"""
修复v2：直接操作XML修改标题行和教学任务格式
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import re

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

def clear_and_set_cell(tc, text, bold=True):
    """清除单元格内容并设置新文本"""
    # 删除所有段落的内容（保留第一个段落）
    paras = tc.findall(qn('w:p'))
    if not paras:
        p = OxmlElement('w:p')
        tc.append(p)
        paras = [p]
    
    p = paras[0]
    # 删除所有run
    for r in p.findall(qn('w:r')):
        p.remove(r)
    
    # 创建新run
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), '仿宋')
    rFonts.set(qn('w:eastAsia'), '仿宋')
    rFonts.set(qn('w:hAnsi'), '仿宋')
    rFonts.set(qn('w:cs'), '仿宋')
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '21')
    rPr.append(szCs)
    
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)
    
    r.append(rPr)
    
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    
    # 删除多余的段落
    for extra_p in paras[1:]:
        tc.remove(extra_p)

def set_multiline_cell(tc, lines, bold=False):
    """设置多行文本到单元格"""
    paras = tc.findall(qn('w:p'))
    if not paras:
        p = OxmlElement('w:p')
        tc.append(p)
        paras = [p]
    
    p = paras[0]
    for r in p.findall(qn('w:r')):
        p.remove(r)
    
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), '仿宋')
    rFonts.set(qn('w:eastAsia'), '仿宋')
    rFonts.set(qn('w:hAnsi'), '仿宋')
    rFonts.set(qn('w:cs'), '仿宋')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '21')
    rPr.append(szCs)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    r.append(rPr)
    
    for li, line in enumerate(lines):
        if li > 0:
            br = OxmlElement('w:br')
            r.append(br)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = line
        r.append(t)
    p.append(r)
    
    for extra_p in paras[1:]:
        tc.remove(extra_p)

# 处理所有30个表格
table_indices = list(range(5, 63, 2))  # 5,7,9,...,61
print(f'处理 {len(table_indices)} 个表格')

for ti in table_indices:
    t = doc.tables[ti]
    
    # === 1. 修复R9标题行 ===
    # 获取R9的单元格元素
    r9 = t.rows[9]
    r9_tcs = r9._tr.findall(qn('w:tc'))
    
    # 打印当前状态（仅第一个表格）
    if ti == 5:
        print(f'\nTable 5 R9 has {len(r9_tcs)} tc elements')
        for ci, tc in enumerate(r9_tcs):
            txt = ''
            for p in tc.findall(qn('w:p')):
                for r in p.findall(qn('w:r')):
                    for t_elem in r.findall(qn('w:t')):
                        txt += t_elem.text or ''
            # 检查gridSpan和vMerge
            tcPr = tc.find(qn('w:tcPr'))
            gs = ''
            vm = ''
            if tcPr is not None:
                gs_elem = tcPr.find(qn('w:gridSpan'))
                if gs_elem is not None:
                    gs = f' gridSpan={gs_elem.get(qn("w:val"))}'
                vm_elem = tcPr.find(qn('w:vMerge'))
                if vm_elem is not None:
                    vm_val = vm_elem.get(qn('w:val'))
                    vm = f' vMerge={vm_val or "continue"}'
            print(f'  TC{ci}: [{txt[:40]}]{gs}{vm}')
    
    # R9结构应该是：
    # TC0: 教学要求 (vMerge=continue)
    # TC1: 知识＆技能 (gridSpan=5)
    # TC2: 重点
    # TC3: 难点 (gridSpan=3)
    # TC4: 目标达成度 (gridSpan=9)
    
    # 但当前可能结构不同，需要找到实际位置
    # 通过文本内容找到需要修改的单元格
    for ci, tc in enumerate(r9_tcs):
        txt = ''
        for p in tc.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                for t_elem in r.findall(qn('w:t')):
                    txt += t_elem.text or ''
        
        # 如果C2位置包含数字开头的内容（被错误替换的知识＆技能）
        if re.match(r'^[1-9]', txt) and '分析' in txt:
            clear_and_set_cell(tc, '知识＆技能', bold=True)
        # 如果有"难点"但应该在"重点"位置
        elif txt == '难点' and ci <= 2:
            # 这个位置应该是"重点"
            # 检查下一个TC是否已经有"难点"
            if ci + 1 < len(r9_tcs):
                next_txt = ''
                for p in r9_tcs[ci+1].findall(qn('w:p')):
                    for r in p.findall(qn('w:r')):
                        for t_elem in r.findall(qn('w:t')):
                            next_txt += t_elem.text or ''
                if next_txt == '难点' or next_txt == '':
                    # 当前位置改为"重点"
                    clear_and_set_cell(tc, '重点', bold=True)
                    # 如果下一个是空的，设为"难点"
                    if next_txt == '':
                        clear_and_set_cell(r9_tcs[ci+1], '难点', bold=True)
    
    # === 2. 修复R10标题行 ===
    r10 = t.rows[10]
    r10_tcs = r10._tr.findall(qn('w:tc'))
    
    for ci, tc in enumerate(r10_tcs):
        txt = ''
        for p in tc.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                for t_elem in r.findall(qn('w:t')):
                    txt += t_elem.text or ''
        
        if re.match(r'^[1-9]', txt) and '分析' in txt:
            clear_and_set_cell(tc, '知识＆技能', bold=True)
        elif txt == '难点' and ci <= 2:
            if ci + 1 < len(r10_tcs):
                next_txt = ''
                for p in r10_tcs[ci+1].findall(qn('w:p')):
                    for r in p.findall(qn('w:r')):
                        for t_elem in r.findall(qn('w:t')):
                            next_txt += t_elem.text or ''
                if next_txt == '难点' or next_txt == '':
                    clear_and_set_cell(tc, '重点', bold=True)
                    if next_txt == '':
                        clear_and_set_cell(r10_tcs[ci+1], '难点', bold=True)
    
    # === 3. 修复R3教学任务格式 ===
    r3 = t.rows[3]
    r3_tcs = r3._tr.findall(qn('w:tc'))
    
    for tc in r3_tcs:
        txt = ''
        for p in tc.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                for t_elem in r.findall(qn('w:t')):
                    txt += t_elem.text or ''
        
        if '子情景' in txt and '子情景' in txt.replace('子情景', '', 1):
            # 有两个"子情景"，需要分行
            lines = re.split(r'(子情景[：:])', txt)
            result = []
            for i in range(0, len(lines), 2):
                if i + 1 < len(lines):
                    result.append(lines[i] + lines[i+1])
                else:
                    if lines[i]:
                        result.append(lines[i])
            set_multiline_cell(tc, result, bold=False)

doc.save(fp)
print('\n保存完成')

# 验证
doc2 = Document(fp)
t = doc2.tables[5]
print('\n=== 验证 Table 5 ===')
print(f'R3: {t.rows[3].cells[2].text}')
r9 = t.rows[9]._tr.findall(qn('w:tc'))
for ci, tc in enumerate(r9):
    txt = ''
    for p in tc.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            for t_elem in r.findall(qn('w:t')):
                txt += t_elem.text or ''
    print(f'R9 TC{ci}: {txt[:30]}')
