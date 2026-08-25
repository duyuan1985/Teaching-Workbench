"""契约驱动填充引擎（阶段3原型）。

按已解析契约的槽位顺序，把课程上下文填入模板副本；
所有写入值加黄色高亮，用于人工对照"哪些位置随生成而变化"。
"""

import copy
import json
import re

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt

BODY_FONT = "仿宋_GB2312"
BODY_SIZE = 12


def _apply_font(run, name=BODY_FONT, size=BODY_SIZE, bold=None):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def _style_like(run, source_fmt):
    """按契约格式槽（模板实测值）设置 run 字体；无契约则用正文默认。"""
    name = source_fmt.get("font") or BODY_FONT
    size = source_fmt.get("font_size_pt") or BODY_SIZE
    _apply_font(run, name, size, source_fmt.get("bold"))


def _write_runs(paragraph, text, fmt=None, highlight=True, font_name=None):
    """整段重写：清空 runs，写入多行文本（\\n 转行），套格式并高亮。"""
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = str(text)
    merged = dict(fmt or {})
    if font_name:
        merged["font"] = font_name
    _style_like(run, merged)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _append_value(paragraph, value, fmt=None):
    """段末追加值（保留模板前缀），用于 标签：值 型正文行。"""
    run = paragraph.add_run(str(value))
    _style_like(run, fmt or {})
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _insert_body_after(heading_paragraph, text):
    """空节：在标题段后插入正文段（可含多行，每段首行缩进两字符）。"""
    from docx.text.paragraph import Paragraph
    anchor = heading_paragraph._p
    size = BODY_SIZE
    for line in str(text).split("\n"):
        if not line.strip():
            continue
        new_p = copy.deepcopy(heading_paragraph._p)
        anchor.addnext(new_p)
        anchor = new_p
        paragraph = Paragraph(new_p, heading_paragraph._parent)
        _write_runs(paragraph, line)
        paragraph.paragraph_format.line_spacing = Pt(23)
        paragraph.paragraph_format.first_line_indent = Pt(size * 2)
    return anchor


def _unique_cells(row):
    """按物理 tc 去重返回单元格（vMerge 续行解析到主格）。"""
    seen, cells = set(), []
    for cell in row.cells:
        if id(cell._tc) not in seen:
            seen.add(id(cell._tc))
            cells.append(cell)
    return cells


def _fill_cell(cell, text, size=9):
    paragraph = cell.paragraphs[0]
    _write_runs(paragraph, text, {"font": BODY_FONT, "font_size_pt": size})
    for extra in cell.paragraphs[1:]:
        _write_runs(extra, "", {})


class ContractFiller:
    def __init__(self, template_path, slots):
        self.doc = Document(template_path)
        self.slots = slots  # 已按 sort_order 排序的契约槽位列表

    # ---------- 槽位路由 ----------

    def fill(self, context):
        # 先把所有段落型槽位解析成 Paragraph 对象（避免后续插入/删除造成索引位移）
        targets = {}
        for slot in self.slots:
            key, structure = slot["slot_key"], json.loads(slot["structure_json"] or "{}")
            for role, source in (("anchor", key), ("heading", structure.get("heading_locator", ""))):
                m = re.search(r"paragraph:(\d+)", source)
                if m and key.split(":")[0] in ("paragraph", "cover", "section"):
                    n = int(m.group(1))
                    if n < len(self.doc.paragraphs):
                        targets[(key, role)] = self.doc.paragraphs[n]
        filled, skipped = [], []
        for slot in self.slots:
            key = slot["slot_key"]
            try:
                ok = self._dispatch(slot, context, targets)
                (filled if ok else skipped).append(key)
            except Exception as exc:  # 预览模式：单槽失败不中断
                skipped.append(f"{key}!{type(exc).__name__}")
        return filled, skipped

    def _dispatch(self, slot, context, targets):
        key, locator = slot["slot_key"], slot["locator"]
        structure = json.loads(slot["structure_json"]) if slot["structure_json"] else {}
        fmt = json.loads(slot["format_json"]) if slot["format_json"] else {}
        if key.startswith("paragraph:") and "课程名称" in key:
            return self._fill_title(targets.get((key, "anchor")), fmt, context)
        if key.startswith("cover:"):
            return self._fill_cover(targets.get((key, "anchor")), fmt, context)
        if key.startswith("table:"):
            return False  # 复杂表格由专门方法处理（见 fill_tables）
        if key.startswith("section:"):
            return self._fill_section(structure, context, targets.get((key, "anchor")), targets.get((key, "heading")))
        return self._fill_labeled_paragraph(targets.get((key, "anchor")), fmt, context)

    # ---------- 标题占位 / 封面 / 标签行 ----------

    def _fill_title(self, paragraph, fmt, context):
        if paragraph is None:
            return False
        course = context["course"]
        text = paragraph.text
        if "课程标准" in text:
            text = f"《{course}》课程标准"
        elif "考核评价" in text:
            text = f"《{course}》课程考核评价"
        else:
            text = f"《{course}》"
        _write_runs(paragraph, text, fmt, font_name="黑体")
        return True

    def _fill_cover(self, paragraph, fmt, context):
        if paragraph is None:
            return False
        label_map = {"适用专业": context["major"], "编制单位": context["office"], "合作单位": context["partner"]}
        label = next((k for k in label_map if k in paragraph.text), None)
        if not label:
            return False
        value = label_map[label]
        for run in paragraph.runs:
            run.text = ""
        label_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        label_run.text = f"{label}："
        remaining = max(0, 16 - len(str(value)))
        value_run = paragraph.add_run("　" * (remaining // 2) + str(value) + "　" * (remaining - remaining // 2))
        value_run.underline = True
        value_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        # 封面字段：宋体三号，下划线仅覆盖值区域
        _apply_font(value_run, "宋体", 16)
        _apply_font(label_run, "宋体", 16)
        return True

    def _fill_labeled_paragraph(self, paragraph, fmt, context):
        if paragraph is None:
            return False
        text = paragraph.text
        for label, value in context["basic_lines"].items():
            if text.startswith(label):
                base = text.split("：")[0] + "："
                _write_runs(paragraph, base, fmt, font_name="黑体")
                _append_value(paragraph, value, fmt)
                paragraph.runs[-1].font.name = "黑体"
                paragraph.runs[-1]._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")
                return True
        return False

    # ---------- 章节正文 ----------

    def _fill_section(self, structure, context, anchor, heading):
        heading_title = structure.get("heading", "")
        content = context["sections"].get(heading_title)
        if not content:
            return False
        if structure.get("body_type") == "empty" or anchor is None:
            if heading is None:
                return False
            _insert_body_after(heading, content)
        else:
            _write_runs(anchor, content, structure.get("format") or {})
            anchor.paragraph_format.line_spacing = Pt(23)
        return True

    # ---------- 课程目标表（表0） ----------

    def fill_goal_table(self, goals):
        table = self.doc.tables[0]
        row = table.rows[1]
        for cell, text in zip(row.cells, [goals["knowledge"], goals["ability"], goals["ideology"], goals["quality"]]):
            _fill_cell(cell, text, size=10)

    # ---------- 课时分配表（表1） ----------
    # 任务格式：(名称, 学时, 理论, 实践) 或 (名称, 学时)；课时列输出 "8（理论4+实践4）"

    def fill_hour_table(self, scenarios):
        table = self.doc.tables[1]
        data_rows = table.rows[1:-1]  # 去表头与合计行
        need = sum(len(s["tasks"]) for s in scenarios)
        pristine = copy.deepcopy(data_rows[0]._tr)  # 首次填充前留干净模板行
        while len(data_rows) < need:
            new_tr = copy.deepcopy(pristine)
            data_rows[-1]._tr.addnext(new_tr)
            data_rows = table.rows[1:-1]
        for extra in range(len(data_rows) - need):
            table._tbl.remove(data_rows[-1 - extra]._tr)
        data_rows = table.rows[1:-1]

        def hours_text(hours, theory, practice):
            if theory is None:
                return str(hours)
            return f"{hours}（理论{theory}+实践{practice}）"

        index = 0
        total = theory_total = practice_total = 0
        for s_no, scenario in enumerate(scenarios, start=1):
            for t_no, task in enumerate(scenario["tasks"]):
                name, hours = task[0], task[1]
                theory = task[2] if len(task) > 2 else None
                practice = task[3] if len(task) > 3 else None
                cells = _unique_cells(data_rows[index])
                texts = [f"{s_no}", scenario["name"] if t_no == 0 else "", f"{t_no + 1}", name,
                         hours_text(hours, theory, practice)]
                for cell, text in zip(cells, texts):
                    if text:
                        _fill_cell(cell, text)
                index += 1
                total += hours
                if theory is not None:
                    theory_total += theory
                    practice_total += practice
        total_cells = _unique_cells(table.rows[-1])
        _fill_cell(total_cells[0], "合计")
        _fill_cell(total_cells[-1], hours_text(total, theory_total, practice_total))

    # ---------- 考核评价表（表2） ----------

    def fill_assessment_table(self, scenarios):
        table = self.doc.tables[2]
        # 持有行对象：合并续行的 cells[0] 解析到主格，须用行对象本身删行
        rows_info = [(row, _unique_cells(row)) for row in table.rows[2:]]
        groups, current, prev_first = [], None, None
        for row, cells in rows_info:
            first_text = cells[0].text.strip()
            if first_text != prev_first:
                m = re.match(r"单元(\d+)", first_text)
                no = int(m.group(1)) if m else ((current["no"] + 1) if current else 1)
                current = {"no": no, "rows": []}
                groups.append(current)
                prev_first = first_text
            current["rows"].append((row, cells))
        remove_rows = []
        final_cell = None  # 末列通天柱（全表合并）：课程级终结考核说明，只写一次
        for group in groups:
            scenario = scenarios[group["no"] - 1] if group["no"] <= len(scenarios) else None
            if scenario is None:
                remove_rows.extend(row for row, _ in group["rows"])
                continue
            _, first_cells = group["rows"][0]
            _fill_cell(first_cells[0], f"单元{group['no']}（情境{group['no']}）：{scenario['name']}（{scenario['weight']}%）")
            if len(first_cells) >= 4:
                _fill_cell(first_cells[-2], f"{scenario['final']}（{scenario['final_weight']}%）")
            if final_cell is None:
                final_cell = first_cells[-1]
            used = {int(m.group(1)) for _, cells in group["rows"]
                    if (m := re.match(r"任务(\d+)", cells[1].text.strip()))}
            next_task = (max(used) + 1) if used else 1
            for row, cells in group["rows"]:
                task_text = cells[1].text.strip()
                m = re.match(r"任务(\d+)", task_text)
                if m and int(m.group(1)) <= len(scenario["assess"]):
                    t_no = int(m.group(1)) - 1
                elif "……" in task_text and next_task <= len(scenario["assess"]):
                    t_no = next_task - 1
                    next_task += 1
                elif "……" in task_text:
                    remove_rows.append(row)
                    continue
                else:
                    continue
                task_name, mode, weight_val = scenario["assess"][t_no]
                _fill_cell(cells[1], f"任务{t_no + 1}：{task_name}（{mode}）")
                _fill_cell(cells[2], f"{weight_val}%")
        for row in remove_rows:
            row._tr.getparent().remove(row._tr)
        if final_cell is not None:
            finals = "、".join(
                f"{s['final']}{s['final_weight']}%" for s in scenarios)
            _fill_cell(final_cell, f"考试、作品、成果类考核，合计60%\n（{finals}）")

    # ---------- 模板指令清理（填充完成后） ----------

    INSTRUCTION_RE = re.compile(
        r"^[（(]?\s*注[：:）)]|^[（(]要求|^[（(]说明|^参考格式如下|^XXXX教学法|^根据本课程的教学目标要求|"
        r"^[（(]正文行距|^[（(]表格中|^注[：:]")
    HEADING_FMT_RE = re.compile(r"[（(][^（）()]*?(行距|黑体|仿宋|宋体|顶格|下同|居中|加粗)[^（）()]*?[）)]")

    def cleanup(self, context):
        course = context["course"]
        removed, replaced = 0, 0
        to_delete = []
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            highlighted = any(r.font.highlight_color for r in paragraph.runs)
            if highlighted:
                continue
            if "《×××》" in text or "****" in text:
                new_text = text.replace("×××", course).replace("*" * 10, course)
                first = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                for run in paragraph.runs[1:]:
                    run.text = ""
                first.text = new_text
                first.font.highlight_color = WD_COLOR_INDEX.YELLOW
                replaced += 1
                continue
            if self.INSTRUCTION_RE.match(text) and len(text) < 260:
                to_delete.append(paragraph._p)
                removed += 1
                continue
            # 节标题内联格式说明（保留编号，删除括号说明）
            if self.HEADING_FMT_RE.search(text) and re.match(r"^([一二三四五六七八九十]+、|（[一二三四五六七八九十]+）|\d+[、.])", text):
                new_text = self.HEADING_FMT_RE.sub("", text).strip()
                if new_text != text:
                    first = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    for run in paragraph.runs[1:]:
                        run.text = ""
                    first.text = new_text
        for el in to_delete:
            el.getparent().remove(el)
        return removed, replaced

    # ---------- 学习情境描述表（表3，预览填模板自带的两块） ----------

    def fill_scenario_table(self, context):
        table = self.doc.tables[3]
        rows = table.rows
        _fill_cell(_unique_cells(rows[0])[0], f"课程：{context['course']}")
        cells0 = _unique_cells(rows[0])
        if len(cells0) >= 4:
            _fill_cell(cells0[1], f"学分：{context['credit']}")
            _fill_cell(cells0[3], f"总学时：{context['hours']}")
        s1 = context["scenario_detail"]
        _fill_cell(_unique_cells(rows[1])[0], f"学习情境1：{s1['name']}")
        _fill_cell(_unique_cells(rows[1])[1], f"学时：{s1['hours']}")
        b = _unique_cells(rows[3])
        _fill_cell(b[0], s1["goals"])
        _fill_cell(b[1], s1["content"])
        _fill_cell(b[2], s1["methods"])
        m1 = _unique_cells(rows[5])
        _fill_cell(m1[0], s1["materials"])
        _fill_cell(m1[1], s1["assessment"])
        _fill_cell(m1[2], s1["remark"])
        sub = context["subscenario_detail"]
        _fill_cell(_unique_cells(rows[6])[0], f"学习子情境1.1：{sub['name']}")
        _fill_cell(_unique_cells(rows[6])[1], f"学时：{sub['hours']}")
        b2 = _unique_cells(rows[8])
        _fill_cell(b2[0], sub["goals"])
        _fill_cell(b2[1], sub["content"])
        _fill_cell(b2[2], sub["methods"])
        m2 = _unique_cells(rows[10])
        _fill_cell(m2[0], sub["materials"])
        _fill_cell(m2[1], sub["assessment"])
        _fill_cell(m2[2], sub["remark"])

    def save(self, out_path):
        self.doc.save(out_path)
        return out_path
