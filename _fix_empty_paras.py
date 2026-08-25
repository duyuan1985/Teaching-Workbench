"""
清除所有表格单元格中的空段落
只保留有内容的段落，删除空段落
"""
from docx import Document
from docx.oxml.ns import qn

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

def has_text(p):
    """检查段落是否有文本内容"""
    for r in p.findall(qn('w:r')):
        for elem in r:
            if elem.tag == qn('w:t') and elem.text:
                return True
    return False

def is_empty_paragraph(p):
    """检查是否为空段落（无文本、无图片等有意义内容）"""
    # 有文本则非空
    if has_text(p):
        return False
    # 有图片则非空
    for r in p.findall(qn('w:r')):
        if r.find(qn('w:drawing')) is not None:
            return False
        if r.find(qn('w:pict')) is not None:
            return False
    return True

total_removed = 0

# 遍历所有表格
for ti in range(len(doc.tables)):
    t = doc.tables[ti]
    for ri in range(len(t.rows)):
        row = t.rows[ri]
        tcs = row._tr.findall(qn('w:tc'))
        for tc in tcs:
            paras = tc.findall(qn('w:p'))
            if len(paras) <= 1:
                continue
            # 保留至少一个段落
            for p in paras[1:]:  # 跳过第一个段落
                if is_empty_paragraph(p):
                    tc.remove(p)
                    total_removed += 1

doc.save(fp)
print(f'共删除 {total_removed} 个空段落')

# 验证
doc2 = Document(fp)
t = doc2.tables[6]
print('\n=== 验证 Table 6 ===')
for ri in range(11):
    row = t.rows[ri]
    tcs = row._tr.findall(qn('w:tc'))
    for ci, tc in enumerate(tcs):
        paras = tc.findall(qn('w:p'))
        empty_count = sum(1 for p in paras if is_empty_paragraph(p))
        if empty_count > 0 or len(paras) > 1:
            texts = []
            for p in paras:
                txt = ''
                for r in p.findall(qn('w:r')):
                    for elem in r:
                        if elem.tag == qn('w:t'):
                            txt += elem.text or ''
                texts.append(txt[:30] if txt else '(空)')
            print(f'R{ri} TC{ci}: {len(paras)}段落, {empty_count}空 | {texts[:3]}')
