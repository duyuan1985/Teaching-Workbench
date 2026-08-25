"""
对比教学设计、课程标准、授课计划三个文档的前九部分内容
"""
from docx import Document
from docx.oxml.ns import qn

# 三个文档路径
fp_design = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
fp_standard = r'生成结果\精修版\2023-2024-2《商务数据分析》课程标准 杜媛.docx'
fp_plan = r'生成结果\精修版\2023-2024-2《商务数据分析》授课计划 杜媛.docx'

def get_paras(fp):
    """获取文档所有段落文本"""
    doc = Document(fp)
    paras = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            paras.append((p.style.name if p.style else '', txt))
    return paras

def get_table_texts(fp):
    """获取文档所有表格文本"""
    doc = Document(fp)
    tables = []
    for ti, t in enumerate(doc.tables):
        rows_data = []
        for ri in range(len(t.rows)):
            texts = [c.text.strip() for c in t.rows[ri].cells]
            rows_data.append(texts)
        tables.append(rows_data)
    return tables

# ============================================================
# 教学设计 - 段落结构
# ============================================================
print('='*80)
print('教学设计 段落结构（前30段）')
print('='*80)
paras_d = get_paras(fp_design)
for i, (style, txt) in enumerate(paras_d[:80]):
    print(f'P{i:3d} [{style[:10]:10s}]: {txt[:100]}')

print()
print('='*80)
print('课程标准 段落结构（前40段）')
print('='*80)
paras_s = get_paras(fp_standard)
for i, (style, txt) in enumerate(paras_s[:60]):
    print(f'P{i:3d} [{style[:10]:10s}]: {txt[:100]}')

print()
print('='*80)
print('授课计划 段落结构（前40段）')
print('='*80)
paras_p = get_paras(fp_plan)
for i, (style, txt) in enumerate(paras_p[:60]):
    print(f'P{i:3d} [{style[:10]:10s}]: {txt[:100]}')
