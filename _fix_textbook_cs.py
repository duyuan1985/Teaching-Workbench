"""
修改课程标准中的教材描述，与教学设计统一
"""
from docx import Document
from docx.oxml.ns import qn

unified_textbook_desc = '《大数据分析方法项目实战》，天津大学出版社。由天津滨海迅腾科技集团有限公司编写，天津大学出版社出版。该教材为十三五职业教育国家规划教材，以项目引领、任务驱动方式组织教学内容，内容涵盖数据采集、清洗、分析、可视化全流程，配有丰富的电商案例和实训项目，符合高职高专学生的认知规律和理实一体化教学要求。'

fp_s = r'生成结果\精修版\2023-2024-2《商务数据分析》课程标准 杜媛.docx'
doc_s = Document(fp_s)

p48 = doc_s.paragraphs[48]
old_text = p48.text
print(f'修改前: {old_text[:150]}')

# 拆分：第一段(教材选用) + 第二段(教材编写要求)
split_marker = '（2）教材编写要求'
if split_marker in old_text:
    part2 = old_text[old_text.index(split_marker):]
else:
    part2 = ''

# 获取格式
font_name = '仿宋'
if p48.runs:
    r = p48.runs[0]
    if r.font.name:
        font_name = r.font.name

# 清除原有run
for r in p48.runs:
    r._element.getparent().remove(r._element)

# 新第一段
new_part1 = f'（1）教材选用：本课程选用{unified_textbook_desc}'
new_full_text = new_part1 + '\n' + part2 if part2 else new_part1

# 添加新run
run = p48.add_run(new_full_text)
run.font.name = font_name
run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

doc_s.save(fp_s)
print(f'修改后: {p48.text[:300]}')
print('\n保存成功')
