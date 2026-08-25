"""
检查模板封面结构
"""
from docx import Document
from docx.oxml.ns import qn

template_fp = r"原始资料\模板\模板5：教学设计\模板5：教学设计 模板（2023-2024）.docx"
doc = Document(template_fp)

print("=== 模板封面段落（P0-P40）===")
for i, p in enumerate(doc.paragraphs[:45]):
    text = p.text.strip()
    runs_info = ""
    if p.runs:
        r = p.runs[0]
        fn = r.font.name or ""
        sz = r.font.size.pt if r.font.size else ""
        bold = r.font.bold
        runs_info = f"[{fn},{sz},bold={bold}]"
    print(f"  P{i}: {runs_info} '{text[:80]}'")

print("\n=== 模板表0 ===")
t0 = doc.tables[0]
print(f"{len(t0.rows)}行 x {len(t0.rows[0].cells)}列")
for ri in range(len(t0.rows)):
    cells = t0.rows[ri].cells
    texts = [c.text.strip()[:30] for c in cells]
    print(f"  R{ri}: {texts}")

print("\n=== 模板表1 ===")
t1 = doc.tables[1]
print(f"{len(t1.rows)}行 x {len(t1.rows[0].cells)}列")
for ri in range(len(t1.rows)):
    cells = t1.rows[ri].cells
    texts = [c.text.strip()[:30] for c in cells]
    print(f"  R{ri}: {texts}")
