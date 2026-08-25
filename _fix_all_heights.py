from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

# 修复所有表格（基本信息+教学组织）的所有行高为auto
for ti in range(len(doc.tables)):
    t = doc.tables[ti]
    for ri in range(len(t.rows)):
        tr = t.rows[ri]._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            continue
        trHeight = trPr.find(qn('w:trHeight'))
        if trHeight is None:
            continue
        hRule = trHeight.get(qn('w:hRule'))
        if hRule == 'exact':
            trHeight.set(qn('w:val'), '0')
            trHeight.set(qn('w:hRule'), 'auto')

doc.save(fp)
print('所有表格的固定行高(exact)已全部改为auto')

# 验证
doc2 = Document(fp)
remaining = 0
for ti in range(len(doc2.tables)):
    t = doc2.tables[ti]
    for ri in range(len(t.rows)):
        tr = t.rows[ri]._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            continue
        trHeight = trPr.find(qn('w:trHeight'))
        if trHeight is None:
            continue
        if trHeight.get(qn('w:hRule')) == 'exact':
            remaining += 1
print(f'剩余exact行高: {remaining}')
