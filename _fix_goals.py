"""
更新教学设计课程目标内容，与课程标准一致
保留目标名称"认知目标"，不改为"知识目标"
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from copy import deepcopy

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)

# 找到目标段落
paras = doc.paragraphs
target_indices = {}
for pi, p in enumerate(paras):
    txt = p.text.strip()
    if txt == '1、认知目标：':
        target_indices['cognitive'] = pi
    elif txt == '2、能力目标：':
        target_indices['ability'] = pi
    elif txt == '3、思政目标：':
        target_indices['sizheng'] = pi
    elif txt == '4、素质目标：':
        target_indices['quality'] = pi
    elif txt == '（二）课程内容设计':
        target_indices['end'] = pi
        break

print(f"目标段落索引: {target_indices}")

# ============================================================
# 新内容（与课程标准一致）
# ============================================================
# 认知目标（保留名称，内容用课程标准的知识目标）
cognitive_items = [
    "① 理解数据分析的基本概念、分类（描述性、诊断性、预测性、指令性）和适用场景",
    "② 掌握数据分析方法理论（PEST、5W2H、SWOT等）和常用指标体系",
    "③ 掌握Excel数据分析工具（数据透视表、函数、图表）的使用方法",
    "④ 掌握Python数据分析库（Numpy、Pandas、SciPy）的使用方法",
    "⑤ 掌握Scikit-learn机器学习库的分类、回归、聚类模型构建方法",
]

# 能力目标
ability_items = [
    "① 能够运用Excel进行数据透视、统计分析和图表可视化",
    "② 能够使用Numpy进行数组运算、数学函数和统计函数计算",
    "③ 能够使用Pandas进行数据清洗、标准化、聚合和透视分析",
    "④ 能够使用Scikit-learn构建分类、回归和聚类模型并评估效果",
    "⑤ 能够独立完成商务数据分析项目并撰写完整的分析报告",
]

# 思政目标
sizheng_items = [
    "① 马列主义方面，运用马克思主义唯物辩证法分析数据现象，理解量变与质变关系，用发展的眼光看待数据驱动的商业变革",
    "② 理想信念方面，树立科技报国理想，关注数字技术服务乡村振兴的实践价值，增强通过数据分析服务经济社会发展的使命感",
    "③ 核心价值观方面，践行诚信、公正、法治，坚持真实分析和规范验证，不篡改数据、不误导决策，以数据诚信践行社会主义核心价值观",
    "④ 传统文化方面，汲取\u201c实事求是\u201d传统文化精髓，在数据分析中做到求真务实、严谨细致，传承工匠精神",
    "⑤ 职业道德方面，严守数据保密协议，不泄露商业机密和用户隐私，遵守《个人信息保护法》《数据安全法》，恪守数据分析师职业操守",
    "⑥ 个人素养方面，培养数据驱动的批判性思维和问题解决能力，养成规范操作、主动学习、持续改进的良好习惯",
]

# 素质目标
quality_items = [
    "① 创新意识方面，培养数据驱动的创新思维，善于从数据中发现商业价值和创新机会，运用AIGC等新技术探索数据分析新方法",
    "② 安全意识方面，建立数据安全意识，规范数据处理流程，防范数据泄露风险，确保分析过程合规合法",
    "③ 团队协作方面，提升任务分工、沟通反馈、成果检查和按时交付能力，在小组项目中发挥协作精神",
    "④ 独立自主方面，培养独立完成数据分析任务的能力，能够自主查阅文档、调试代码、解决问题",
    "⑤ 个人自律方面，养成按步骤实施、及时测试、记录问题和持续改进的习惯，增强抗压能力和自我调节能力",
]

# ============================================================
# 获取模板段落的格式（从现有①条目获取）
# ============================================================
# 找到现有的①段落作为格式模板
template_p = None
for pi in range(target_indices['cognitive'] + 1, target_indices['ability']):
    txt = paras[pi].text.strip()
    if txt.startswith('①'):
        template_p = paras[pi]
        break

def get_run_format(template_para):
    """从模板段落获取run格式信息"""
    if not template_p.runs:
        return {'font_name': '仿宋', 'font_size': Pt(10.5)}
    r = template_p.runs[0]
    return {
        'font_name': r.font.name or '仿宋',
        'font_size': r.font.size or Pt(10.5),
    }

fmt = get_run_format(template_p)
print(f"模板格式: font={fmt['font_name']}, size={fmt['font_size']}")

# ============================================================
# 收集需要替换的段落范围
# ============================================================
# 认知目标：cognitive+1 到 ability-1
# 能力目标：ability+1 到 sizheng-1 （跳过"通过任务引领"段）
# 思政目标：sizheng+1 到 quality-1
# 素质目标：quality+1 到 end-1

sections = [
    (target_indices['cognitive'] + 1, target_indices['ability'], cognitive_items),
    (target_indices['ability'] + 1, target_indices['sizheng'], ability_items),
    (target_indices['sizheng'] + 1, target_indices['quality'], sizheng_items),
    (target_indices['quality'] + 1, target_indices['end'], quality_items),
]

# ============================================================
# 执行替换
# ============================================================
def replace_para_text(para, text, font_name, font_size):
    """替换段落文本，保留格式"""
    # 清除所有run
    for r in para.runs:
        r._element.getparent().remove(r._element)
    # 添加新run
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

body = doc.element.body
all_paras = body.findall(qn('w:p'))

# 需要从后往前替换，避免索引变化
# 先收集所有要操作的段落元素和目标文本
operations = []
for start_idx, end_idx, items in sections:
    print(f"\n区间 P{start_idx}-P{end_idx-1}: {end_idx - start_idx}个旧段落 -> {len(items)}个新条目")

    # 收集旧段落元素（跳过非①开头的段落如"通过任务引领活动，学生能"）
    old_paras = []
    for pi in range(start_idx, end_idx):
        txt = paras[pi].text.strip()
        if txt.startswith('①') or txt.startswith('②') or txt.startswith('③') or txt.startswith('④') or txt.startswith('⑤') or txt.startswith('⑥'):
            old_paras.append(paras[pi]._element)
        elif txt and not txt.startswith('通过'):
            # 保留非条目段落
            pass

    print(f"  旧条目段落: {len(old_paras)}, 新条目: {len(items)}")

    # 替旧条目文本
    for i in range(min(len(old_paras), len(items))):
        p_elem = old_paras[i]
        # 清除run
        for r in p_elem.findall(qn('w:r')):
            p_elem.remove(r)
        # 添加新run
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), fmt['font_name'])
        rFonts.set(qn('w:eastAsia'), fmt['font_name'])
        rFonts.set(qn('w:hAnsi'), fmt['font_name'])
        rFonts.set(qn('w:cs'), fmt['font_name'])
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(fmt['font_size'].pt * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(fmt['font_size'].pt * 2)))
        rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = items[i]
        r.append(t)
        p_elem.append(r)

    # 如果新条目多于旧条目，需要添加新段落
    if len(items) > len(old_paras):
        # 用最后一个旧段落作为模板
        template_elem = old_paras[-1] if old_paras else paras[start_idx]._element
        for i in range(len(old_paras), len(items)):
            new_p = deepcopy(template_elem)
            # 清除run
            for r in new_p.findall(qn('w:r')):
                new_p.remove(r)
            # 添加新run
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), fmt['font_name'])
            rFonts.set(qn('w:eastAsia'), fmt['font_name'])
            rFonts.set(qn('w:hAnsi'), fmt['font_name'])
            rFonts.set(qn('w:cs'), fmt['font_name'])
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(fmt['font_size'].pt * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(fmt['font_size'].pt * 2)))
            rPr.append(szCs)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = items[i]
            r.append(t)
            new_p.append(r)
            # 插入到模板段落后面
            template_elem.addnext(new_p)
            template_elem = new_p  # 下一个接在后面

    # 如果旧条目多于新条目，删除多余段落
    elif len(old_paras) > len(items):
        for i in range(len(items), len(old_paras)):
            p_elem = old_paras[i]
            p_elem.getparent().remove(p_elem)

doc.save(fp)
print("\n保存完成")

# 验证
doc2 = Document(fp)
in_section = False
for pi, p in enumerate(doc2.paragraphs):
    txt = p.text.strip()
    if '课程目标设计' in txt:
        in_section = True
    if in_section:
        print(f"P{pi}: {txt[:150]}")
    if in_section and '（二）' in txt:
        break
