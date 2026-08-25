from docx import Document
from docx.oxml.ns import qn

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

# 检查3个问题
# 1. 教材学情分析(R8) - 检查前5个任务是否相同
print('=== 问题1: 教材学情分析 ===')
for ti in [5, 7, 9, 11, 13]:
    t = doc.tables[ti]
    r8_txt = t.rows[8].cells[2].text.strip()[:80]
    print(f'Table {ti}: {r8_txt}')

# 2. 知识＆技能难点列 - 检查Table 5的R11-R14
print('\n=== 问题2: 难点列打勾 ===')
t = doc.tables[5]
for ri in range(11, 15):
    row = t.rows[ri]
    tcs = row._tr.findall(qn('w:tc'))
    texts = []
    for ci, tc in enumerate(tcs):
        txt = ''
        for p in tc.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                for elem in r:
                    if elem.tag == qn('w:t'):
                        txt += elem.text or ''
        texts.append(f'TC{ci}={txt}' if txt else f'TC{ci}=(空)')
    print(f'R{ri}: {texts}')

# 3. 教学活动流程 - 检查Table 6的任务内容长度
print('\n=== 问题3: 教学活动流程内容长度 ===')
t6 = doc.tables[6]
for ri in range(3, 8):
    row = t6.rows[ri]
    c3_text = row.cells[2].text.strip()
    print(f'R{ri} (C3): {len(c3_text)}字符 | {c3_text[:100]}')
