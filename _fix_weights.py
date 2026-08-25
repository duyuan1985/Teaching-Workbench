"""
修改课程标准表2的考核权重：签到1.25%+课堂表现1.25%+作业2.5%
"""
from docx import Document
from docx.oxml.ns import qn

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》课程标准 杜媛.docx"
doc = Document(fp)

# 找到表2（考核评价表）
t2 = None
for t in doc.tables:
    if len(t.rows) > 2:
        r0 = t.rows[0].cells[0].text.strip()
        if "考核方式" in r0 or "过程性考核" in r0 or "单元" in r0:
            t2 = t
            break

if not t2:
    # 尝试找含"签到"的表
    for t in doc.tables:
        for ri in range(min(5, len(t.rows))):
            for ci in range(len(t.rows[ri].cells)):
                if "签到" in t.rows[ri].cells[ci].text:
                    t2 = t
                    break
            if t2:
                break
        if t2:
            break

if t2:
    print(f"找到考核表: {len(t2.rows)}行 x {len(t2.rows[0].cells)}列")
    print("修改前:")
    for ri in range(min(12, len(t2.rows))):
        cells = t2.rows[ri].cells
        texts = [c.text.strip()[:15] for c in cells]
        print(f"  R{ri}: {texts}")
    
    # 权重在col3（索引3），考核方式名称在col2（索引2）
    changed = 0
    for ri in range(2, len(t2.rows)):
        cells = t2.rows[ri].cells
        if len(cells) < 4:
            continue
        method_cell = cells[2].text.strip()
        weight_cell = cells[3]
        
        if "签到" in method_cell:
            for p in weight_cell.paragraphs:
                for r in p.runs:
                    r.text = "1.25%"
            changed += 1
        elif "课堂表现" in method_cell:
            for p in weight_cell.paragraphs:
                for r in p.runs:
                    r.text = "1.25%"
            changed += 1
        elif "作业" in method_cell:
            for p in weight_cell.paragraphs:
                for r in p.runs:
                    r.text = "2.5%"
            changed += 1
    
    print(f"\n修改了{changed}个权重格子")
    print("修改后:")
    for ri in range(min(12, len(t2.rows))):
        cells = t2.rows[ri].cells
        texts = [c.text.strip()[:15] for c in cells]
        print(f"  R{ri}: {texts}")
    
    doc.save(fp)
    print(f"\n保存完成: {fp}")
else:
    print("未找到考核评价表")
