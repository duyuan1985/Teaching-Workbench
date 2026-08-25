"""公共 Word 文档操作工具函数。

从 document_generator.py 中提取的可复用 Word 操作函数，
供 document_generator.py 及其他模块共享使用。
"""

from docx.oxml.ns import qn
from docx.shared import Pt

__all__ = [
    "_set_paragraph",
    "_set_cell",
    "_set_underlined_field",
    "_set_body_font",
    "_set_cell_body_font",
    "_set_cell_small_font",
    "_set_table_font",
    "_replace_prefix",
    "_term_label",
    "_plan_term_label",
]


def _set_paragraph(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = str(text)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(str(text))


def _set_cell(cell, text):
    paragraph = cell.paragraphs[0]
    _set_paragraph(paragraph, text)
    for extra in cell.paragraphs[1:]:
        _set_paragraph(extra, "")


def _set_underlined_field(paragraph, label, value, field_width=18):
    """Write a fixed-width, centered East Asian cover field after its label."""
    source_run = paragraph.runs[0] if paragraph.runs else None
    for run in paragraph.runs:
        run.text = ""
    label_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    label_run.text = label
    value = str(value)
    remaining = max(0, field_width - len(value))
    left = remaining // 2
    right = remaining - left
    field_run = paragraph.add_run("　" * left + value + "　" * right)
    field_run.underline = True
    if source_run is not None:
        field_run.bold = source_run.bold
        field_run.italic = source_run.italic
        field_run.font.name = source_run.font.name
        field_run.font.size = source_run.font.size


def _set_body_font(paragraph):
    for run in paragraph.runs:
        run.font.name = "仿宋_GB2312"
        run.font.size = Pt(12)
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")


def _set_cell_body_font(cell):
    for paragraph in cell.paragraphs:
        _set_body_font(paragraph)


def _set_cell_small_font(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "仿宋_GB2312"
            run.font.size = Pt(9)
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")


def _set_table_font(table, name="仿宋_GB2312", size=9):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = name
                    run.font.size = Pt(size)
                    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def _replace_prefix(document, prefix, text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            _set_paragraph(paragraph, text)
            return


def _term_label(term):
    parts = str(term).split("-")
    if len(parts) >= 3:
        semester = "第一学期" if parts[-1] == "1" else "第二学期"
        return f"{parts[0]}—{parts[1]}学年{semester}"
    return str(term)


def _plan_term_label(term):
    parts = str(term).split("-")
    if len(parts) >= 3:
        semester = "第一学期" if parts[-1] == "1" else "第二学期"
        return f"{parts[0]}——{parts[1]}学年  {semester}"
    return str(term)
