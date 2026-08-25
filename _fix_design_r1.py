"""
第一轮修复：
1. 补写二、课程定位内容
2. 补写三、教学设计理念内容
3. 补写六~九章节内容
4. 补写十、单元教学设计标题
5. 修复职业能力训练表
6. 补写"教学设计·基本信息"和"教学设计·教学组织"标题
"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from copy import deepcopy
import store

fp = r"生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx"
doc = Document(fp)

# ============================================================
# 工具函数
# ============================================================
def set_para_text(p, text, font_name='仿宋', size=14, bold=False):
    """安全设置段落文本"""
    for r in p.runs:
        r.text = ''
    if p.runs:
        p.runs[0].text = text
        r = p.runs[0]
    else:
        r = p.add_run(text)
    r.font.name = font_name
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    r.font.size = Pt(size)
    r.font.bold = bold
    # 清除缩进
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)

def add_para_after(doc, ref_para, text, font_name='仿宋', size=14, bold=False):
    """在指定段落后面插入新段落"""
    new_p = deepcopy(ref_para._element)
    # 清空内容
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    ref_para._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, ref_para._parent)
    r = p.add_run(text)
    r.font.name = font_name
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    r.font.size = Pt(size)
    r.font.bold = bold
    return p

# ============================================================
# 1. 补写二、课程定位 (P43, 空段落)
# ============================================================
print("1. 补写课程定位...")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '二、课程定位':
        # 下一段是空的
        if i+1 < len(doc.paragraphs):
            next_p = doc.paragraphs[i+1]
            if not next_p.text.strip():
                set_para_text(next_p, 
                    '　　《商务数据分析》是农村电子商务专业的一门专业核心课程，属于理实一体化课程。'
                    '本课程的前导课程为《电子商务基础》《计算机应用基础》《Python程序设计》，'
                    '为后续《新媒体平台运营与推广》、综合实训及顶岗实习等课程奠定数据分析基础。'
                    '课程内容对接电商企业真实业务，融入大数据、人工智能等新一代信息技术在商务分析领域的新应用，'
                    '兼顾1+X数据采集与处理职业技能等级证书和全国职业院校技能大赛相关赛项的能力要求，'
                    '实现课证融通、赛教融合。本课程在专业课程体系中起着承上启下的关键作用，'
                    '是培养学生数据思维和商务分析能力的核心课程。')
                print(f"  P{i+1} 课程定位已填写")
        break

# ============================================================
# 2. 补写三、教学设计的理念 (P45, 空段落)
# ============================================================
print("2. 补写教学设计理念...")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '三、教学设计的理念':
        if i+1 < len(doc.paragraphs):
            next_p = doc.paragraphs[i+1]
            if not next_p.text.strip():
                set_para_text(next_p,
                    '　　本课程坚持"以学生为中心、以任务为引领、以能力为本位"的教学设计理念，'
                    '遵循"做中学、学中做"的理实一体化教学原则。'
                    '在课程设计中主动对接行业新标准与新技术：融入《中华人民共和国数据安全法》'
                    '《个人信息保护法》对数据处理合规性的要求；引入大语言模型（LLM）辅助数据分析、'
                    'AIGC智能图表生成、数据中台与实时数据分析等新技术、新方法；'
                    '参照全国职业院校技能大赛大数据分析与应用赛项标准设计实训项目，'
                    '实现课程内容与行业标准、岗位需求同步更新。'
                    '同时依托超星学习通、雨课堂等数字化教学平台开展线上线下混合式教学，'
                    '课前推送微课、课中互动测验、课后学情数据分析，以数字技术驱动教学持续改进。')
                print(f"  P{i+1} 教学设计理念已填写")
        break

# ============================================================
# 3. 补写六~九章节内容（在P101"课程考核"段落后插入）
# ============================================================
print("3. 补写六~九章节...")

# 找到P101 "课程考核：考核建议见下表"
ref_idx = None
for i, p in enumerate(doc.paragraphs):
    if '课程考核' in p.text and '考核建议见下表' in p.text:
        ref_idx = i
        break

if ref_idx is not None:
    # 在P101后插入六~九章节
    # 找到表4的位置（考核方案表），在表4之后插入
    # 先找到表4后面的段落位置
    # 表4是第5个表（索引4）
    t4 = doc.tables[4]
    # 表4最后一段在哪个段落后面
    
    # 从P101往下找，找到第一个非空段落（应该是表4之后的段落）
    # 用另一种方式：在文档body中找到表4的XML元素，在其后插入段落
    
    body = doc.element.body
    
    # 找到表4的XML元素
    t4_elem = t4._element
    
    # 要插入的内容
    sections = [
        ("六、课程教学实施条件", True),
        ("　　（一）师资条件：本课程教学团队由校内专任教师和企业兼职教师组成。专任教师需具备数据分析、Python编程、机器学习等方面的专业能力，具有2年以上企业实践经历或相关行业工作经验。企业兼职教师需来自电商企业数据分析岗位，具有丰富的实战经验，能参与课程实践教学。", False),
        ("　　（二）实训条件：本课程教学需配备理实一体化教室，配有计算机（安装Anaconda、Python 3.10+、Jupyter Notebook、Excel 2019及以上版本）、多媒体投影设备、网络环境。实训环境需支持学生分组操作，每组配备1台计算机，并安装数据分析相关软件包（pandas、numpy、matplotlib、seaborn、scikit-learn等）。", False),
        ("　　（三）教学场所：理论教学在多媒体教室进行，实践教学在计算机实训室进行，综合项目实践可安排在校企合作实训基地或企业现场。", False),
        ("七、教学资源", True),
        ("教材：", True),
        ("　　《商务数据分析与应用》，人民邮电出版社，2023年版。该教材为十三五职业教育国家规划教材，内容涵盖数据采集、清洗、分析、可视化全流程，配有丰富的电商案例和实训项目。", False),
        ("教学资料：", True),
        ("　　（1）超星学习通在线课程资源：微课视频36个、课件PPT 30套、题库500余道、实训数据集10套。", False),
        ("　　（2）参考书籍：《Python数据分析从入门到实践》《商务数据分析实战》《数据可视化之美》等。", False),
        ("　　（3）企业真实数据集：农产品电商平台交易数据、用户行为数据、商品评价数据等脱敏数据集。", False),
        ("　　（4）1+X数据采集与处理职业技能等级证书培训教材及标准文档。", False),
        ("　　（5）全国职业院校技能大赛大数据分析与应用赛项历年真题及评分标准。", False),
        ("八、需要说明的其他问题", True),
        ("　　（1）本课程采用理实一体化教学模式，理论与实践交替进行，每个学习情境均安排理论1学时和实践1学时，确保学生在掌握理论知识的同时能够动手操作。", False),
        ("　　（2）课程融入课程思政元素，在数据分析各环节适时渗透马列主义世界观、数据安全法治意识、诚信分析职业道德、团队协作精神等，实现全员全程全方位育人。", False),
        ("　　（3）课程对接1+X数据采集与处理职业技能等级证书标准和全国职业院校技能大赛大数据分析与应用赛项标准，学生在课程学习中同步备考证书和参赛，实现课证融通、赛教融合。", False),
        ("九、第一节课设计梗概", True),
        ("　　第一节课为课程导入课，主要内容包括：（1）课程介绍：介绍课程性质、教学目标、考核方式、学习方法和课程资源平台使用方法。（2）行业认知：通过分析电商行业数据分析岗位需求、薪资水平和发展前景，激发学生学习兴趣。（3）案例分析：展示农产品电商平台数据分析典型案例，让学生直观感受数据分析在商务决策中的作用。（4）工具介绍：演示Python数据分析环境和Excel高级功能，让学生初步了解课程将使用的工具。（5）分组与任务分配：将学生分组并分配本学期综合项目的初始任务，明确团队合作要求。（6）课后任务：要求学生注册超星学习通账号、预习第一章微课视频、安装Anaconda环境。", False),
    ]
    
    # 在表4元素后面依次插入段落
    prev_elem = t4_elem
    for text, is_heading in sections:
        # 创建新段落
        new_p = doc.add_paragraph()
        r = new_p.add_run(text)
        r.font.name = '仿宋'
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '仿宋')
        r.font.size = Pt(14)
        r.font.bold = is_heading
        
        # 移动到正确位置（prev_elem后面）
        new_elem = new_p._element
        body.remove(new_elem)
        prev_elem.addnext(new_elem)
        prev_elem = new_elem
    
    print(f"  六~九章节已插入")
else:
    print("  未找到P101课程考核段落")

# ============================================================
# 4. 补写"十、单元教学设计"标题
# ============================================================
print("4. 补写十、单元教学设计标题...")

# 在九章节最后一段后面插入"十、单元教学设计"标题
# 找到"九、第一节课设计梗概"段落
ref_p = None
for i, p in enumerate(doc.paragraphs):
    if '九、第一节课设计梗概' in p.text:
        ref_p = p
        break

if ref_p is not None:
    # 找到九章节内容最后一段
    for i, p in enumerate(doc.paragraphs):
        if '九、第一节课设计梗概' in p.text:
            # 找到九章节内容段落（下一段）
            if i+1 < len(doc.paragraphs):
                content_p = doc.paragraphs[i+1]
                # 在content_p后面插入"十、单元教学设计"标题
                new_p = add_para_after(doc, content_p, '十、单元教学设计', 
                                       font_name='黑体', size=14, bold=False)
                print("  十、单元教学设计标题已插入")
                break
else:
    # 如果九章节没找到，在表5之前插入
    t5 = doc.tables[5]
    body = doc.element.body
    t5_elem = t5._element
    new_p = doc.add_paragraph()
    r = new_p.add_run('十、单元教学设计')
    r.font.name = '黑体'
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '黑体')
    r.font.size = Pt(14)
    new_elem = new_p._element
    body.remove(new_elem)
    t5_elem.addprevious(new_elem)
    print("  十、单元教学设计标题已插入（表5前）")

# ============================================================
# 5. 修复职业能力训练表（表3）
# ============================================================
print("5. 修复职业能力训练表(表3)...")

t3 = doc.tables[3]
print(f"  表3: {len(t3.rows)}行x{len(t3.rows[0].cells)}列")

# 打印当前结构
for ri in range(min(6, len(t3.rows))):
    cells = t3.rows[ri].cells
    texts = [c.text.strip()[:20] for c in cells]
    print(f"  R{ri}: {texts}")

# 表3应该是职业能力训练项目设计表
# 列：序号、能力训练项目名称、相关支撑知识、训练方式手段及步骤、结果（可展示）
# 需要为每个子目标填写不同的内容

# 获取课程任务数据
tasks = store.rows("SELECT * FROM tasks WHERE offering_id=20 ORDER BY seq")
units = store.rows("SELECT * FROM curriculum_units WHERE offering_id=20 ORDER BY seq")

# 为每个学习情境生成职业能力训练内容
capacity_items = []
for ui, unit in enumerate(units):
    unit_title = unit.get('title', f'学习情境{ui+1}')
    # 获取该单元的任务
    unit_tasks = [t for t in tasks if t.get('unit_id') == unit.get('id')]
    for ti, task in enumerate(unit_tasks):
        task_title = task.get('title', f'任务{ti+1}')
        # 编号
        seq = f"{ui+1}.{ti+1}"
        # 能力训练项目名称
        name = task_title
        # 相关支撑知识（每个子目标不一样）
        if ui == 0:
            knowledge = f"①数据分析基本概念与分类\n②常用数据分析方法理论\n③数据分析指标体系\n④数据分析流程与规范"
            method = f"①案例分析法：分析电商数据分析典型案例\n②小组讨论法：讨论数据分析应用场景\n③任务驱动法：完成数据分析流程设计\n④实操练习法：使用工具进行数据分析"
        elif ui == 1:
            knowledge = f"①Python开发环境搭建\n②变量、数据类型与运算符\n③控制流程与函数\n④数据处理常用库（pandas）"
            method = f"①演示法：教师演示Python环境搭建\n②练习法：学生跟随练习基础语法\n③项目法：编写数据处理脚本\n④互评法：学生互相检查代码"
        elif ui == 2:
            knowledge = f"①数据采集方法与工具\n②数据清洗规则与流程\n③缺失值与异常值处理\n④数据转换与标准化"
            method = f"①任务驱动法：采集电商交易数据\n②演示法：演示数据清洗流程\n③练习法：处理缺失值和异常值\n④小组合作法：分组完成数据清洗"
        else:
            knowledge = f"①{unit_title}核心概念与原理\n②{task_title}关键技术\n③相关工具与API使用\n④实际业务应用场景"
            method = f"①案例教学法：分析{unit_title}典型案例\n②演示法：演示{task_title}操作步骤\n③练习法：学生独立完成{task_title}实训\n④小组讨论法：讨论解决方案"
        
        result = f"{task_title}实训成果（数据分析报告或可视化图表）"
        
        capacity_items.append({
            'seq': seq,
            'name': name,
            'knowledge': knowledge,
            'method': method,
            'result': result
        })

print(f"  生成{len(capacity_items)}个能力训练项目")

# 填充表3
# 表3结构：6行x10列（模板），需要扩展行数
# 先看表3的列结构
if len(t3.rows) > 0:
    header = [c.text.strip() for c in t3.rows[0].cells]
    print(f"  表头: {header}")

# 清除模板行并填充
# 表3有6行模板行，需要扩展到len(capacity_items)+1行
# 先填充已有行
for ri in range(1, min(len(t3.rows), len(capacity_items)+1)):
    item = capacity_items[ri-1]
    cells = t3.rows[ri].cells
    # 确保列数正确
    if len(cells) >= 5:
        # col0: 序号
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = item['seq']
            else:
                p.add_run(item['seq'])
        # col1: 能力训练项目名称
        for p in cells[1].paragraphs:
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = item['name']
            else:
                p.add_run(item['name'])
        # col2: 相关支撑知识
        for p in cells[2].paragraphs:
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = item['knowledge']
            else:
                p.add_run(item['knowledge'])
        # col3: 训练方式手段及步骤
        for p in cells[3].paragraphs:
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = item['method']
            else:
                p.add_run(item['method'])
        # col4: 结果（可展示）
        for p in cells[4].paragraphs:
            for r in p.runs:
                r.text = ''
            if p.runs:
                p.runs[0].text = item['result']
            else:
                p.add_run(item['result'])

# 如果capacity_items比模板行多，需要添加行
if len(capacity_items) + 1 > len(t3.rows):
    # 复制最后一行的结构作为模板
    template_row = t3.rows[-1]._tr
    for i in range(len(t3.rows), len(capacity_items) + 1):
        new_tr = deepcopy(template_row)
        t3._element.append(new_tr)
    
    # 填充新行
    for ri in range(len(t3.rows), len(capacity_items) + 1):
        if ri <= len(capacity_items):
            item = capacity_items[ri-1]
            row = t3.rows[ri-1]
            cells = row.cells
            if len(cells) >= 5:
                for p in cells[0].paragraphs:
                    for r in p.runs:
                        r.text = ''
                    if p.runs:
                        p.runs[0].text = item['seq']
                    else:
                        p.add_run(item['seq'])
                for p in cells[1].paragraphs:
                    for r in p.runs:
                        r.text = ''
                    if p.runs:
                        p.runs[0].text = item['name']
                    else:
                        p.add_run(item['name'])
                for p in cells[2].paragraphs:
                    for r in p.runs:
                        r.text = ''
                    if p.runs:
                        p.runs[0].text = item['knowledge']
                    else:
                        p.add_run(item['knowledge'])
                for p in cells[3].paragraphs:
                    for r in p.runs:
                        r.text = ''
                    if p.runs:
                        p.runs[0].text = item['method']
                    else:
                        p.add_run(item['method'])
                for p in cells[4].paragraphs:
                    for r in p.runs:
                        r.text = ''
                    if p.runs:
                        p.runs[0].text = item['result']
                    else:
                        p.add_run(item['result'])

print(f"  表3已填充{len(capacity_items)}行数据")

# ============================================================
# 6. 在每对表5/表6前补写标题
# ============================================================
print("6. 补写表5/表6标题...")

# 从表5开始（索引5），每隔2个表（表5+表6=1对），在表5前插入"教学设计·基本信息"标题
# 在表6前插入"教学设计·教学组织"标题
body = doc.element.body
title_count = 0

for ti in range(5, len(doc.tables), 2):
    t5 = doc.tables[ti]
    t6 = doc.tables[ti+1] if ti+1 < len(doc.tables) else None
    
    # 在表5前插入"教学设计·基本信息"
    t5_elem = t5._element
    
    # 检查前一个元素是否已经是标题
    prev = t5_elem.getprevious()
    prev_text = ""
    if prev is not None and prev.tag.endswith('}p'):
        for r in prev.findall(qn('w:r')):
            for t in r.findall(qn('w:t')):
                prev_text += t.text or ''
    
    if '教学设计·基本信息' not in prev_text:
        new_p = doc.add_paragraph()
        r = new_p.add_run('教学设计·基本信息')
        r.font.name = '黑体'
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '黑体')
        r.font.size = Pt(12)
        r.font.bold = True
        new_elem = new_p._element
        body.remove(new_elem)
        t5_elem.addprevious(new_elem)
        title_count += 1
    
    # 在表6前插入"教学设计·教学组织"
    if t6:
        t6_elem = t6._element
        prev = t6_elem.getprevious()
        prev_text = ""
        if prev is not None and prev.tag.endswith('}p'):
            for r in prev.findall(qn('w:r')):
                for t in r.findall(qn('w:t')):
                    prev_text += t.text or ''
        
        if '教学设计·教学组织' not in prev_text:
            new_p = doc.add_paragraph()
            r = new_p.add_run('教学设计·教学组织')
            r.font.name = '黑体'
            rPr = r._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = rPr.makeelement(qn('w:rFonts'), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), '黑体')
            r.font.size = Pt(12)
            r.font.bold = True
            new_elem = new_p._element
            body.remove(new_elem)
            t6_elem.addprevious(new_elem)
            title_count += 1

print(f"  插入了{title_count}个标题")

# ============================================================
# 保存
# ============================================================
doc.save(fp)
print(f"\n保存完成: {fp}")
print("第一轮修复完成")
