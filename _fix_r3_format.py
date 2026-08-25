"""
修复R3教学任务的换行格式
当前：第一章 初识数据分析子情景：\n认识数据分析子情景：\n常用数据分析方法
应为：第一章 初识数据分析\n子情景：认识数据分析\n子情景：常用数据分析方法
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

fp = r'生成结果\精修版\2023-2024-2《商务数据分析》教学设计 杜媛.docx'
doc = Document(fp)

def set_multiline_cell(tc, lines, bold=False):
    """设置多行文本到单元格"""
    paras = tc.findall(qn('w:p'))
    if not paras:
        p = OxmlElement('w:p')
        tc.append(p)
        paras = [p]
    
    p = paras[0]
    for r in p.findall(qn('w:r')):
        p.remove(r)
    
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), '仿宋')
    rFonts.set(qn('w:eastAsia'), '仿宋')
    rFonts.set(qn('w:hAnsi'), '仿宋')
    rFonts.set(qn('w:cs'), '仿宋')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '21')
    rPr.append(szCs)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    r.append(rPr)
    
    for li, line in enumerate(lines):
        if li > 0:
            br = OxmlElement('w:br')
            r.append(br)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = line
        r.append(t)
    p.append(r)
    
    for extra_p in paras[1:]:
        tc.remove(extra_p)

# 处理所有29个表格
table_indices = list(range(5, 63, 2))
fixed_count = 0

for ti in table_indices:
    t = doc.tables[ti]
    r3 = t.rows[3]
    r3_tcs = r3._tr.findall(qn('w:tc'))
    
    for tc in r3_tcs:
        # 获取完整文本（包括换行）
        txt = ''
        for p in tc.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                for elem in r:
                    if elem.tag == qn('w:t'):
                        txt += elem.text or ''
                    elif elem.tag == qn('w:br'):
                        txt += '\n'
        
        # 去掉已有的换行，重新分割
        flat_txt = txt.replace('\n', '')
        
        if '子情景' in flat_txt:
            # 在"子情景"前插入换行
            parts = re.split(r'(子情景[：:])', flat_txt)
            # parts = ['第一章 初识数据分析', '子情景：', '认识数据分析', '子情景：', '常用数据分析方法']
            lines = []
            i = 0
            while i < len(parts):
                if parts[i].startswith('子情景'):
                    # 子情景部分，和后面的内容合并
                    if i + 1 < len(parts):
                        lines.append(parts[i] + parts[i+1])
                        i += 2
                    else:
                        lines.append(parts[i])
                        i += 1
                else:
                    if parts[i]:
                        lines.append(parts[i])
                    i += 1
            
            set_multiline_cell(tc, lines, bold=False)
            fixed_count += 1

doc.save(fp)
print(f'修复了 {fixed_count} 个表格的R3格式')

# 验证
doc2 = Document(fp)
for ti in [5, 7, 9]:
    t = doc2.tables[ti]
    r3 = t.rows[3]
    r3_tcs = r3._tr.findall(qn('w:tc'))
    for tc in r3_tcs:
        txt = ''
        for p in tc.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                for elem in r:
                    if elem.tag == qn('w:t'):
                        txt += elem.text or ''
                    elif elem.tag == qn('w:br'):
                        txt += ' | '
        if '子情景' in txt or '第一章' in txt or '第二章' in txt:
            print(f'Table {ti} R3: {txt}')
